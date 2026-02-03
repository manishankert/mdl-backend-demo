# services/document_editor.py
import os
import re
import json
import logging
from io import BytesIO
from typing import Optional

import requests
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from utils.docx_utils import clear_runs, add_hyperlink

logging.basicConfig(level=logging.INFO)


def para_text(p) -> str:
    return "".join(run.text for run in p.runs)


def force_paragraph_font_size(p, size_pt=12):
    for r in p.runs:
        r.font.size = Pt(size_pt)


def set_font_size_to_12(doc):
    """Set all text in document to 12pt font."""
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.size = Pt(12)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(12)

    for section in doc.sections:
        for container in (section.header, section.footer):
            for p in container.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(12)


def looks_like_optional_plural_text(s: str) -> bool:
    """Find paragraphs that still have '(s)' or '(es)' style tokens or subject-verb '(s)'."""
    s = (s or "")
    return any(t in s for t in ["(s)", "(es)", "violate(s)", "address(es)", "appear(s)"]) or " audit finding" in s.lower() or " corrective action" in s.lower()


def pluralize_with_openai(text: str, total_findings: int) -> Optional[str]:
    """
    Use OpenAI to convert optional-plural boilerplate into grammatically correct text.
    Returns rewritten string, or None on failure.
    """
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        return None

    singular = (total_findings == 1)
    style_hint = (
        "Use singular grammar (is/addresses/appears; finding, issue, CAP, date, corrective action)."
        if singular else
        "Use plural grammar (are/address/appear; findings, issues, CAPs, dates, corrective actions)."
    )

    system = (
        "You are revising boilerplate text in a U.S. government letter. "
        "Rewrite the provided text to be grammatically correct and natural, "
        "resolving any optional plural tokens like '(s)' or '(es)' and fixing subject-verb agreement. "
        "Preserve meaning and tone; do not add or remove content beyond grammar and number agreement. "
        "Return only the final sentence(s) with no quotes."
    )
    user = (
        f"{style_hint}\n\n"
        "Rewrite the text below to be grammatically correct. Resolve all '(s)' / '(es)' tokens and subject-verb forms. "
        "Keep the same information, formal tone, and punctuation.\n\n"
        f"Text:\n{text}"
    )

    try:
        resp = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            data=json.dumps({
                "model": "gpt-4o-mini",
                "messages": [
                    {"role": "system", "content": system},
                    {"role": "user", "content": user},
                ],
                "temperature": 0,
            }),
            timeout=15,
        )
        resp.raise_for_status()
        out = resp.json()
        rewritten = (out.get("choices", [{}])[0].get("message", {}).get("content") or "").strip()
        return rewritten or None
    except Exception:
        return None


def ai_fix_pluralization_in_doc(doc, total_findings: int):
    """
    Find paragraphs with '(s)/(es)' style text or affected phrases and fix them via OpenAI.
    Falls back silently if API not available.
    """
    candidates = []
    # Scan body paragraphs
    for p in doc.paragraphs:
        t = para_text(p)
        if looks_like_optional_plural_text(t):
            candidates.append(p)
    # Also scan header/footer just in case
    for sec in doc.sections:
        for container in (sec.header, sec.footer):
            for p in container.paragraphs:
                t = para_text(p)
                if looks_like_optional_plural_text(t):
                    candidates.append(p)

    # Rewrite each candidate via OpenAI; if it fails, leave as-is
    for p in candidates:
        original = para_text(p).strip()
        if not original:
            continue
        rewritten = pluralize_with_openai(original, total_findings)
        if rewritten and rewritten != original:
            clear_runs(p)
            p.add_run(rewritten)


def fix_mdl_grammar_text(text: str, n_findings: int) -> str:
    singular = (n_findings == 1)
    be = "is" if singular else "are"

    # NBSP -> space
    out = text.replace("\u00A0", " ")

    # If tokens are still present, resolve them
    out = re.sub(r"\[\s*is\s*/\s*are\s*\]", be, out, flags=re.IGNORECASE)
    out = re.sub(r"\[\s*The\s*\]", "The" if singular else "", out, flags=re.IGNORECASE)
    out = re.sub(r"\(s\)", "" if singular else "s", out)
    out = re.sub(r"\bviolate\s*\(s\)\b", "violates" if singular else "violate", out, flags=re.IGNORECASE)
    out = re.sub(r"\bappear\s*\(s\)\b", "appears" if singular else "appear", out, flags=re.IGNORECASE)
    out = re.sub(r"\baddress\s*\(es\)\b", "addresses" if singular else "address", out, flags=re.IGNORECASE)
    out = re.sub(r"\baddresses\s*\(es\)\b", "addresses", out, flags=re.IGNORECASE)
    out = re.sub(r"\(es\)", "", out)

    # FIX the remaining grammar
    # Insert missing "is/are" after these subjects if it's missing
    out = re.sub(r"\b(The audit finding(?:s)?)\s+(?=sustained\b)", rf"\1 {be} ", out, flags=re.IGNORECASE)
    out = re.sub(r"\b(The CAP(?:s)?)\s*,?\s*if implemented,\s+(?!is\b|are\b)", rf"\1, if implemented, {be} ", out, flags=re.IGNORECASE)
    out = re.sub(r"\b(the corrective action(?:s)?)\s+(?=subject\b)", rf"\1 {be} ", out, flags=re.IGNORECASE)

    # Fix singular verb agreement when subject is singular
    if singular:
        out = re.sub(r"\bissue\s+violate\b", "issue violates", out, flags=re.IGNORECASE)
        out = re.sub(r"\bfinding\s+appear\b", "finding appears", out, flags=re.IGNORECASE)
        out = re.sub(r"\bCAP,\s*if implemented,\s*is responsive to the audit finding,\s*address\b",
                     "CAP, if implemented, is responsive to the audit finding, addresses",
                     out, flags=re.IGNORECASE)

    # Cleanup spacing/punctuation
    out = re.sub(r"[ \t]{2,}", " ", out)
    out = re.sub(r"\s+([,.;:])", r"\1", out)
    return out


def apply_mdl_grammar(doc, n_findings: int):
    def rewrite_paragraph(p):
        old = p.text
        new = fix_mdl_grammar_text(old, n_findings)
        if new != old:
            for r in p.runs[::-1]:
                p._p.remove(r._r)
            p.add_run(new)

            # FORCE font size to 12pt for entire paragraph
            force_paragraph_font_size(p, 12)

    # body
    for p in doc.paragraphs:
        rewrite_paragraph(p)

    # tables
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    rewrite_paragraph(p)


def replace_email_with_mailto_link(p, email: str):
    """Replace occurrences of the email in paragraph text with a clickable mailto hyperlink."""
    if email not in p.text:
        return False

    full = p.text
    parts = full.split(email)

    clear_runs(p)

    # rebuild: text + hyperlink + text (+ possible repeats)
    for i, chunk in enumerate(parts):
        if chunk:
            p.add_run(chunk)
        if i < len(parts) - 1:
            add_hyperlink(p, email, f"mailto:{email}", font_pt=12)

    return True


def postprocess_docx(doc_bytes: bytes, model: dict) -> bytes:
    bio = BytesIO(doc_bytes)
    doc = Document(bio)

    email = model.get("treasury_contact_email", "ORP_SingleAudits@treasury.gov")

    correct_auditee = (model.get("auditee_name") or model.get("recipient_name") or "").strip()
    if correct_auditee.lower().startswith("the "):
        correct_auditee = correct_auditee[4:].strip()

    date_text = (model.get("fy_end_text") or model.get("fy_end_date") or model.get("fiscal_year_end") or "").strip()

    for p in doc.paragraphs:
        text = p.text
        if "Treasury has reviewed the single audit report for" not in text:
            continue

        # Remove leading "the" before the auditee in the sentence
        text = re.sub(
            r'(Treasury has reviewed the single audit report for )the\s+',
            r'\1',
            text,
            flags=re.IGNORECASE
        )
        # Find the date in the paragraph
        date_in_doc = None
        if date_text and date_text in text:
            date_in_doc = date_text
        else:
            m = re.search(r'([A-Za-z]+ \d{1,2}, \d{4})', text)
            if m:
                date_in_doc = m.group(1)

        clear_runs(p)

        # rebuild auditee and date with two bold runs
        if correct_auditee and date_in_doc and (correct_auditee in text) and (date_in_doc in text):
            # Split around auditee first
            pre_a, rest = text.split(correct_auditee, 1)
            # Then split the remaining text around date
            pre_d, post_d = rest.split(date_in_doc, 1)

            p.add_run(pre_a)
            r1 = p.add_run(correct_auditee)
            r1.bold = True

            p.add_run(pre_d)
            r2 = p.add_run(date_in_doc)
            r2.bold = True

            p.add_run(post_d)
        elif correct_auditee and (correct_auditee in text):
            # Fallback: bold only auditee
            pre, post = text.split(correct_auditee, 1)
            p.add_run(pre)
            r = p.add_run(correct_auditee)
            r.bold = True
            p.add_run(post)
        elif date_in_doc and (date_in_doc in text):
            # Fallback: bold only date
            pre, post = text.split(date_in_doc, 1)
            p.add_run(pre)
            r = p.add_run(date_in_doc)
            r.bold = True
            p.add_run(post)
        else:
            # Fallback: just keep text
            p.add_run(text)

        force_paragraph_font_size(p, 12)

        break

    # hyperlink treasury email everywhere in the doc
    for p in doc.paragraphs:
        changed = replace_email_with_mailto_link(p, email)
        if changed:
            force_paragraph_font_size(p, 12)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
