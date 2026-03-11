# services/template_processor.py
import os
import re
import logging
from io import BytesIO
from typing import Dict, List, Any, Optional

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.text.paragraph import Paragraph
from docx.table import Table
from fastapi import HTTPException

from utils.text_utils import title_case, norm_txt
from utils.docx_utils import (
    clear_runs,
    para_text,
    tight_paragraph,
    apply_grid_borders,
    insert_after,
    remove_paragraph,
    set_table_cell_margins,
    set_table_preferred_width_and_indent,
    set_row_height_and_allow_break,
    set_table_column_widths,
    apply_program_table_spacing,
    set_table_bold_borders,
)
from services.mdl_builder import format_letter_date
from services.document_editor import (
    set_font_size_to_12,
    apply_mdl_grammar,
    ai_fix_pluralization_in_doc,
)

from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logging.basicConfig(level=logging.INFO)

# Words that must remain fully uppercase
UPPERCASE_TOKENS = {
    "LLC", "LLP", "PLLC", "PC", "PA", "INC", "CO", "CORP",
    "CPA", "CPA'S", "CPA\u2019S", "CPAS", "CPAs", "CFA", "EA",
    "USA", "U.S.", "US",
    "CFO", "CEO", "COO", "CIO", "CAO", "VP", "HR", "IT",
    "PKF", "EFPR",
    # US state abbreviations
    "AL", "AK", "AZ", "AR", "CA", "CO", "CT", "DE", "FL", "GA",
    "HI", "ID", "IL", "IN", "IA", "KS", "KY", "LA", "ME", "MD",
    "MA", "MI", "MN", "MS", "MO", "MT", "NE", "NV", "NH", "NJ",
    "NM", "NY", "NC", "ND", "OH", "OK", "OR", "PA", "RI", "SC",
    "SD", "TN", "TX", "UT", "VT", "VA", "WA", "WV", "WI", "WY",
    "DC", "PR", "GU", "VI", "AS", "MP",
}

LOWERCASE_WORDS = {
    "a", "an", "the", "and", "or", "but", "nor", "for", "so", "yet",
    "at", "by", "in", "of", "on", "to", "up", "as", "is", "it",
    "with", "from", "into", "onto", "over", "than", "that",
}

def ensure_leading_the(name: str) -> str:
    name = (name or "").strip()
    if not name:
        return name
    return name if name.lower().startswith("the ") else f"The {name}"

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def render_questioned_cost_cell(cell, value: str):

    # Normalize input
    v = (value or "").strip()

    lines = [ln.strip() for ln in v.splitlines() if ln.strip()] if v else []

    if not lines:
        lines = ["Questioned Cost:", "None", "Disallowed Cost:", "None"]

    cell.text = ""
    paras = []

    # First paragraph
    p0 = cell.paragraphs[0]
    p0.text = lines[0]
    paras.append(p0)

    # Remaining paragraphs
    for ln in lines[1:]:
        paras.append(cell.add_paragraph(ln))

    for p in paras:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    paras[-1].paragraph_format.space_after = Pt(12)

def smart_title_case(text: str) -> str:
    if not text:
        return text

    words = re.split(r"(\s+)", text.strip())  # preserve spacing
    out = []
    first_word = True

    for w in words:
        if not w.strip():
            out.append(w)
            continue
        elif w.strip().rstrip(".,;:!?").upper() in UPPERCASE_TOKENS:
            clean = w.strip().rstrip(".,;:!?")
            suffix = w.strip()[len(clean):]
            result = clean.upper().replace("'S", "'s").replace("\u2019S", "\u2019s")
            out.append(result + suffix)
        elif w.strip().lower() in LOWERCASE_WORDS and not first_word:
            out.append(w.strip().lower())
        elif "/" in w:
            parts = w.split("/")
            cased = []
            for part in parts:
                if not part:
                    cased.append(part)
                elif part.strip().upper() in UPPERCASE_TOKENS:
                    result = part.strip().upper()
                    result = result.replace("'S", "'s").replace("\u2019S", "\u2019s")
                    cased.append(result)
                elif part.strip().lower() in LOWERCASE_WORDS:
                    cased.append(part.strip().lower())
                elif part.isupper():
                    cased.append(part.capitalize())
                else:
                    cased.append(part[0].upper() + part[1:] if part else part)
            out.append("/".join(cased))
        elif w.isupper():
            out.append(w.capitalize())
        else:
            out.append(w)
        #logging.info(f"smart_title_case token: {repr(w)} -> upper={repr(w.strip().upper())} in_tokens={w.strip().upper() in UPPERCASE_TOKENS}")
        first_word = False

    return "".join(out)


def iter_cells_in_table(tbl: Table):
    for row in tbl.rows:
        for cell in row.cells:
            yield cell


def iter_all_paragraphs_in_container(container) -> list:
    items = []
    if hasattr(container, "paragraphs"):
        items.extend(container.paragraphs)
    if hasattr(container, "tables"):
        for t in container.tables:
            for c in iter_cells_in_table(t):
                items.extend(c.paragraphs)
                for nt in c.tables:
                    for nc in iter_cells_in_table(nt):
                        items.extend(nc.paragraphs)
    return items


def iter_all_paragraphs_full(doc):
    # body
    for p in doc.paragraphs:
        yield p
    # tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    # header/footer
    for sec in doc.sections:
        for container in (sec.header, sec.footer):
            for p in container.paragraphs:
                yield p


def replace_in_paragraph_run_aware(p: Paragraph, mapping: Dict[str, str]) -> bool:
    original = para_text(p)
    if not original:
        return False
    new_text = original
    for k, v in mapping.items():
        if k in new_text:
            new_text = new_text.replace(k, v)
    if new_text != original:
        clear_runs(p)
        p.add_run(new_text)
        return True
    return False


def replace_placeholders_docwide(doc: Document, mapping: Dict[str, str]):
    for p in iter_all_paragraphs_in_container(doc):
        replace_in_paragraph_run_aware(p, mapping)
    for sec in doc.sections:
        for p in iter_all_paragraphs_in_container(sec.header):
            replace_in_paragraph_run_aware(p, mapping)
        for p in iter_all_paragraphs_in_container(sec.footer):
            replace_in_paragraph_run_aware(p, mapping)


def find_anchor_paragraph(doc: Document, anchor: str) -> Optional[Paragraph]:
    for p in iter_all_paragraphs_in_container(doc):
        if anchor in para_text(p):
            return p
    for sec in doc.sections:
        for p in iter_all_paragraphs_in_container(sec.header):
            if anchor in para_text(p):
                return p
        for p in iter_all_paragraphs_in_container(sec.footer):
            if anchor in para_text(p):
                return p
    return None


def delete_immediate_next_table(anchor_para: Paragraph):
    """If the template has a placeholder table immediately after the anchor paragraph, delete it."""
    p_el = anchor_para._p
    nxt = p_el.getnext()
    if nxt is not None and nxt.tag.endswith("tbl"):
        parent = nxt.getparent()
        parent.remove(nxt)


def pick_table_style(doc: Document) -> Optional[str]:
    if getattr(doc, "tables", None):
        for t in doc.tables:
            try:
                if t.style and t.style.name:
                    return t.style.name
            except Exception:
                pass
    try:
        _ = doc.styles["Table Grid"]
        return "Table Grid"
    except KeyError:
        pass
    for st in doc.styles:
        try:
            if st.type == WD_STYLE_TYPE.TABLE:
                return st.name
        except Exception:
            continue
    return None


def build_program_table(doc: Document, program: Dict[str, Any]) -> Table:
    findings = program.get("findings", []) or []
    rows = max(1, len(findings)) + 1

    #tbl = doc.add_table(rows=rows, cols=6)  # 6 columns (added Repeat Finding)
    tbl = doc.add_table(rows=rows, cols=5)  # 6 columns (added Repeat Finding)
    _style = pick_table_style(doc)
    if _style:
        try:
            tbl.style = _style
        except Exception:
            pass
    apply_grid_borders(tbl)

    headers = [
        "Audit\nFinding #",
        "Compliance Type -\nAudit Finding Summary",
        "Audit Finding\nDetermination",
        "Questioned Cost\nDetermination",
        "CAP\nDetermination",
        #"Repeat\nFinding",
    ]
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        clear_runs(cell.paragraphs[0])
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        tight_paragraph(cell.paragraphs[0])

    if findings:
        for r, f in enumerate(findings, start=1):
            #for c in range(6):
            for c in range(5):
                cell = tbl.cell(r, c)
                clear_runs(cell.paragraphs[0])

                # Column-specific formatting
                if c == 0:  # Finding ID
                    cell.paragraphs[0].add_run(f.get("finding_id", ""))
                    repeat_ref = (f.get("repeat_prior_reference") or "").strip()
                    if repeat_ref:
                        cell.paragraphs[0].add_run(f"\n\nRepeat of {repeat_ref}")
                    elif f.get("is_repeat_finding"):
                        cell.paragraphs[0].add_run("\n\nRepeat Finding")
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                elif c == 1:  # Compliance Type - Audit Finding (SPECIAL FORMATTING)
                    compliance_type = f.get("compliance_type", "")
                    summary = f.get("summary", "").strip()

                    # Add compliance type in BOLD
                    if compliance_type:
                        bold_run = cell.paragraphs[0].add_run(compliance_type)
                        bold_run.bold = True

                    # Add hyphen with spaces
                    if compliance_type and summary:
                        #cell.paragraphs[0].add_run(" - ")
                        cell.paragraphs[0].add_run(" \u2013")
                        cell.paragraphs[0].add_run("\n")

                    # Add summary (not bold)
                    if summary:
                        cell.paragraphs[0].add_run("\n")
                        cell.paragraphs[0].add_run(summary)

                    # Left align this column
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

                elif c == 2:  # Audit Finding Determination
                    cell.paragraphs[0].add_run(f.get("audit_determination", "Sustained"))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                elif c == 3:  # Questioned Cost Determination
                    #cell.paragraphs[0].add_run(f.get("questioned_cost_determination", "None"))
                    render_questioned_cost_cell(cell, f.get("questioned_cost_determination", ""))
                    #cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                elif c == 4:  # CAP Determination
                    cell.paragraphs[0].add_run(f.get("cap_determination", "Not Applicable"))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                #elif c == 5:  # Repeat Finding
                 #   cell.paragraphs[0].add_run("Yes" if f.get("is_repeat_finding") else "No")
                  #  cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    else:
        cell = tbl.cell(1, 0)
        clear_runs(cell.paragraphs[0])
        cell.paragraphs[0].add_run("-")

    # FORMAT START
    set_table_cell_margins(tbl, top_in=0.00, bottom_in=0.00, left_in=0.06, right_in=0.06)

    # ---- Program table formatting (ONLY if 5 columns) ----
    set_table_preferred_width_and_indent(tbl, width_in=6.25, indent_in=0.05)

    for r in tbl.rows:
        set_row_height_and_allow_break(r, height_in=0.49, allow_break_across_pages=True)

    #set_table_column_widths(tbl, [0.73, 1.39, 1.0, 1.24, 1.09, 0.80])
    set_table_column_widths(tbl, [0.83, 1.59, 1.2, 1.44, 1.19])
    # ---- end program table formatting ----

    # SPACING MUST BE LAST so nothing overwrites it
    apply_program_table_spacing(tbl)

    set_table_bold_borders(tbl, size=12)

    # END FORMAT

    return tbl

def dedupe_programs(programs: list[dict]) -> list[dict]:
    seen = set()
    out = []
    for p in programs or []:
        aln = (p.get("assistance_listing") or p.get("aln") or "").strip()
        pname = (p.get("program_name") or p.get("program_title") or "").strip()
        # If program name is blank, dedupe on ALN alone; otherwise ALN+name
        key = (aln, pname) if pname else (aln,)
        if key in seen:
            continue
        seen.add(key)
        out.append(p)
    return out

def dedupe_findings(findings: list[dict]) -> list[dict]:
    seen = set()
    out = []
    for f in findings or []:
        fid = (f.get("finding_id") or f.get("audit_finding") or "").strip()
        if fid and fid in seen:
            continue
        if fid:
            seen.add(fid)
        out.append(f)
    return out

def dedupe_model_programs_and_findings(model: Dict[str, Any]) -> int:
    """
    Mutates model['programs'] to remove duplicate programs and duplicate findings.
    Returns total_findings AFTER dedupe.
    """
    programs = model.get("programs") or []
    programs = dedupe_programs(programs)

    for prog in programs:
        prog["findings"] = dedupe_findings(prog.get("findings") or [])

    model["programs"] = programs

    return sum(len(p.get("findings") or []) for p in programs)

def insert_program_tables_at_anchor_no_headers(doc: Document, anchor_para: Paragraph, programs: List[Dict[str, Any]]):
    """
    Insert program tables without creating duplicate headers.
    The template already has the header paragraph, we just insert tables.
    """
    # Clean anchor text
    text = para_text(anchor_para).replace("[[PROGRAM_TABLES]]", "")
    clear_runs(anchor_para)
    if text.strip():
        anchor_para.add_run(text)

    # FIX: Remove extra space after anchor paragraph (space above Findings table)
    tight_paragraph(anchor_para)

    # Delete any placeholder table immediately following the anchor
    delete_immediate_next_table(anchor_para)

    # DEDUPE programs + findings BEFORE rendering tables
    '''programs = dedupe_programs(programs or [])
    for prog in programs:
        prog["findings"] = dedupe_findings(prog.get("findings") or [])'''

    # Order programs by ALN
    PROGRAM_PRIORITY = {
        "21.027": 0,
        "21.026": 1,
        "21.023": 2,
        "21.029": 3,
        "21.031": 4,
        "21.032": 5,
    }
    def _al_key(p):
        aln = (p.get("assistance_listing") or "99.999")
        return (PROGRAM_PRIORITY.get(aln, 99), aln)
    programs_sorted = sorted(programs or [], key=_al_key)

    last = anchor_para
    

    # For SINGLE program: just insert table (header already exists in template)
    # For MULTIPLE programs: insert header + table for 2nd, 3rd, etc.
    for idx, p in enumerate(programs_sorted):
        al = p.get("assistance_listing", "Unknown")
        name = p.get("program_name", "Unknown Program")

        # Only add header for 2nd+ programs (first uses the template header)
        if idx > 0:
            heading_para = doc.add_paragraph()
            clear_runs(heading_para)

            header_run = heading_para.add_run("Assistance Listing Number/Program Name:")
            header_run.bold = True

            heading_para.add_run("\n")
            heading_para.add_run(f"{al} / {name}")

            tight_paragraph(heading_para)
            heading_para.paragraph_format.space_before = Pt(12)
            heading_para.paragraph_format.space_after = Pt(8)

            # Splice heading after 'last'
            heading_el = heading_para._p
            heading_el.getparent().remove(heading_el)
            insert_after(last, heading_el)
            last = heading_el

  

        # Insert table
        tbl = build_program_table(doc, p)
        tbl_el = tbl._tbl
        tbl_el.getparent().remove(tbl_el)
        insert_after(last, tbl_el)
        last = tbl_el

        # Spacer between programs (if multiple)
        if idx < len(programs_sorted) - 1:
            spacer = doc.add_paragraph()
            spacer_el = spacer._p
            spacer_el.getparent().remove(spacer_el)
            insert_after(last, spacer_el)
            last = spacer_el


def find_para_by_contains(doc: Document, needle: str) -> Optional[Paragraph]:
    def _norm(s: str) -> str:
        s = (s or "").replace("\u00A0", " ").replace("\xa0", " ")
        s = s.replace("\u200b", "").replace("\u200c", "").replace("\u200d", "")
        return " ".join(s.split())
    N = _norm(needle)
    for p in iter_all_paragraphs_in_container(doc):
        if N in _norm(para_text(p)):
            return p
    for sec in doc.sections:
        for p in iter_all_paragraphs_in_container(sec.header):
            if N in _norm(para_text(p)):
                return p
        for p in iter_all_paragraphs_in_container(sec.footer):
            if N in _norm(para_text(p)):
                return p
    return None


def remove_duplicate_program_headers(doc: Document, first_label: Paragraph):
    """
    Remove any duplicate 'Assistance Listing Number/Program Name' paragraphs
    that appear after the first one (the template's original).
    """
    # Get all paragraphs
    all_paras = list(doc.paragraphs)

    # Find the index of the first label
    try:
        first_idx = all_paras.index(first_label)
    except ValueError:
        return  # Can't find it, give up

    # Look for duplicates after the first one (within the next 5 paragraphs)
    for i in range(first_idx + 1, min(first_idx + 6, len(all_paras))):
        p = all_paras[i]
        text = para_text(p)

        # If this paragraph also contains "Assistance Listing Number/Program Name"
        if "Assistance Listing Number/Program Name" in text:
            #logging.info(f"Removing duplicate header: {text[:80]}")
            remove_paragraph(p)
            break  # Only remove one duplicate


def cleanup_post_table_narrative(doc, model):
    """
    Remove the repeated narrative paragraphs that appear after the program table(s).
    """
    # Collect IDs and summaries to match
    finding_ids = set()
    summaries = set()
    combos = set()
    for prog in (model.get("programs") or []):
        for f in (prog.get("findings") or []):
            fid = (f.get("finding_id") or "").strip()
            summ = (f.get("summary") or "").strip()
            combo = (f.get("compliance_and_summary") or "").strip()
            if fid:
                finding_ids.add(fid)
            if summ:
                summaries.add(norm_txt(summ))
            if combo:
                combos.add(norm_txt(combo))

    # Regex patterns that match the repeated narrative blocks in the body
    starts = [
        r"^\d{4}-\d{3}\s*-\s*",
        r"^\d{4}-\d{3}\s*[\u2013\u2014]\s*",
        r"^Auditor\s+Description\s+of\s+Condition",
        r"^Auditor\s+Recommendation\.?",
        r"^Responsible\s+Person\s*:",
        r"^Corrective\s+Action\.?",
        r"^Anticipated\s+Completion\s+Date\s*:",
        r"^Federal\s+Agency\s*:",
        r"^Federal\s+Program\s+Title\s*:",
        r"^Federal\s+Award\s+Identification",
        r"^Compliance\s+Requirement\s+Affected\s*:",
        r"^Award\s+Period\s*:",
        r"^Type\s+of\s+Finding\s*:",
        r"^Recommendation\s*:",
        r"^Explanation\s+of\s+disagreement",
        r"^Action\s+taken\s+in\s+response",
        r"^Name\s+of\s+the\s+contact\s+person",
        r"^Planned\s+completion\s+date",
        r"^SUSPENSION\s+AND\s+DEBARMENT",
        r"^PROCUREMENT",
    ]
    patt = re.compile("|".join(starts), re.IGNORECASE)

    # Remove paragraphs that match any of the above
    removed_count = 0
    for p in list(doc.paragraphs):
        t = norm_txt("".join(r.text for r in p.runs))
        if not t:
            continue

        # CRITICAL FIX: Skip the program header paragraph
        if "Assistance Listing Number/Program Name:" in t:
            #logging.info(f"Skipping program header from cleanup: {t[:80]}")
            continue

        should_remove = False
        reason = ""

        # Exact/contains matches
        if any(fid in t for fid in finding_ids):
            should_remove = True
            reason = f"contains finding ID"

        elif patt.search(t):
            should_remove = True
            reason = "matches FAC narrative pattern"

        elif any(s and s.lower() in t.lower() for s in summaries):
            should_remove = True
            reason = "matches summary"

        elif any(c and c.lower() in t.lower() for c in combos):
            should_remove = True
            reason = "matches combo"

        # NEW: Also check for common ALN patterns (21.027, SLFRP, etc.)
        elif re.search(r'\b\d{2}\.\d{3}\b', t) and "Assistance Listing Number/Program Name" not in t:
            should_remove = True
            reason = "contains ALN pattern"

        elif re.search(r'\bSLFRP\d+\b', t, re.IGNORECASE):
            should_remove = True
            reason = "contains SLFRP award number"

        if should_remove:
            #logging.info(f"Removing ({reason}): {t[:100]}")
            remove_paragraph(p)
            removed_count += 1

    #logging.info(f"Cleanup removed {removed_count} duplicate narrative paragraphs")


def unset_all_caps_everywhere(doc):
    # body paragraphs
    for p in doc.paragraphs:
        for r in p.runs:
            r.font.all_caps = False
            r.font.small_caps = False
    # tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.all_caps = False
                        r.font.small_caps = False
    # headers/footers
    for sec in doc.sections:
        for container in (sec.header, sec.footer):
            for p in container.paragraphs:
                for r in p.runs:
                    r.font.all_caps = False
                    r.font.small_caps = False


def email_postfix_cleanup(doc, email):
    """
    Strip leading bracket/curly tokens at paragraph start; fix ".The" joins.
    Safe to run before hyperlink insertion.
    """
    pat_leading = re.compile(r"^\s*(\[\s*treasury_contact_email\s*\]|\$\{treasury_contact_email\})\.?\s*")
    for p in iter_all_paragraphs_full(doc):
        t = para_text(p)
        if not t:
            continue

        new = pat_leading.sub("", t)
        if email and f"{email}.The" in new:
            new = new.replace(f"{email}.The", f"{email}. The")

        if new != t:
            clear_runs(p)
            p.add_run(new)


def strip_leading_token_artifacts(doc):
    pat = re.compile(r"^\s*\$\{[^}]+\}\.?\s*")
    for p in doc.paragraphs:
        t = para_text(p)
        if not t:
            continue
        new = pat.sub("", t)
        if new != t:
            clear_runs(p)
            p.add_run(new)


def fix_questioned_costs_grammar(doc):
    """Fix 'No questioned cost is' to 'No questioned costs are'."""
    for p in iter_all_paragraphs_full(doc):
        text = para_text(p)
        if "No questioned cost is included" in text:
            new_text = text.replace(
                "No questioned cost is included in this single audit report",
                "No questioned costs are included in this single audit report"
            )
            if new_text != text:
                clear_runs(p)
                p.add_run(new_text)
                #logging.info("Fixed questioned costs grammar")
                break

def fix_state_abbrevs(s: str) -> str:
    """Re-uppercase any 2-letter state abbreviation that got title-cased."""
    if not s:
        return s
    return re.sub(r'\b([A-Z][a-z])\b', lambda m: m.group(1).upper(), s)

def fix_narrative_bold(data: bytes, model: Dict[str, Any]) -> bytes:
    from io import BytesIO
    doc = Document(BytesIO(data))
    
    correct_auditee = fix_state_abbrevs(ensure_leading_the(
        model.get("auditee_name") or model.get("recipient_name") or ""
    ))
    correct_auditor = re.sub(r'\s+([,;])', r'\1', smart_title_case(
        model.get("auditor_name") or ""
    ))

    for p in iter_all_paragraphs_in_container(doc):
        text = para_text(p)
        if "Treasury has reviewed the single audit report for" not in text:
            continue

        pattern = (
            r'(Treasury has reviewed the single audit report for )(?:the )?'
            r'(.+?)'
            r'(,? prepared by )(.+?)'
            r'( for the fiscal year.*)'
        )
        m = re.search(pattern, text, re.DOTALL)
        if not m:
            break

        clear_runs(p)
        p.add_run(m.group(1)).font.size = Pt(12)
        r = p.add_run(correct_auditee)
        r.bold = True
        r.font.size = Pt(12)
        p.add_run(m.group(3)).font.size = Pt(12)
        p.add_run(correct_auditor).font.size = Pt(12)

        # Split group 5 to bold the date
        tail = m.group(5)  # " for the fiscal year ending on June 30, 2023. No questioned..."
        date_match = re.search(r'(.*?ending on )([A-Za-z]+ \d+, \d{4})(.*)', tail, re.DOTALL)
        if date_match:
            run_pre = p.add_run(date_match.group(1))
            run_pre.font.size = Pt(12)
            run_date = p.add_run(date_match.group(2))
            run_date.bold = True
            run_date.font.size = Pt(12)
            run_post = p.add_run(date_match.group(3))
            run_post.font.size = Pt(12)
        else:
            p.add_run(tail).font.size = Pt(12)
        break

    set_font_size_to_12(doc)  # re-apply after postprocess

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def build_docx_from_template(model: Dict[str, Any], *, template_path: str) -> bytes:
    """
    Open a .docx template and:
      1) Replace placeholders across the whole document (headers/footers too)
      2) Insert program tables at the [[PROGRAM_TABLES]] anchor
    """
    if not os.path.isfile(template_path):
        raise HTTPException(400, f"Template not found: {template_path}")
    

    # Narrow view: just the findings fields you care about
    '''for pi, prog in enumerate(model.get("programs") or []):
        logging.info(f"Program[{pi}]: aln={prog.get('assistance_listing')} name={prog.get('program_name')}")
        for fi, f in enumerate(prog.get("findings") or []):
            logging.info(
                f"  Finding[{fi}]: id={f.get('finding_id')} | "
                f"compliance_type={f.get('compliance_type')!r} | "
                f"summary={f.get('summary')!r} | "
                f"keys={list(f.keys())}"
            )'''
    # ── END DEBUG ────────────────────────────────────────────────────────────

    doc = Document(template_path)

    # Dates
    _, letter_date_long = format_letter_date(model.get("letter_date_iso"))

    # Header fields (defaults -> empty so placeholders never leak through)
    auditee = smart_title_case((model.get("auditee_name")
               or model.get("recipient_name")
               or ""))
    ein = model.get("ein", "") or ""
    street = model.get("street_address", "") or ""
    city = model.get("city", "") or ""
    state = model.get("state", "") or ""
    zipc = model.get("zip_code", "") or ""
    poc = smart_title_case(model.get("poc_name", "") or "")
    poc_t = smart_title_case(model.get("poc_title", "") or "")
    auditor = smart_title_case(model.get("auditor_name", "") or "")
    #logging.info(f"Auditor: {auditor}")
    #logging.info(f"Auditee: {auditee}")
    #logging.info(f"POC: {poc} ({poc_t})")
    fy_end = (model.get("period_end_text")
               or str(model.get("audit_year", ""))) or ""
    # Treasury contact email
    treasury_contact_email = "ORP_SingleAudits@treasury.gov "

    # Map BOTH styles of placeholders used by the template
    mapping = {
        # date stub used in some templates
        "Date XX, 2025": letter_date_long,

        # [bracket] style
        "[Recipient Name]": auditee,
        "[EIN]": ein,
        "[Street Address]": street,
        "[City]": city,
        "[State]": state,
        "[Zip Code]": zipc,
        "[Point of Contact]": poc,
        "[Point of Contact Title]": poc_t,
        "[Auditor Name]": auditor,
        "[Fiscal Year End Date]": fy_end,
        "[The]": "The",
        "[the]": "the",

        # ${curly} style
        "${recipient_name}": auditee,
        "${ein}": ein,
        "${street_address}": street,
        "${city}": city,
        "${state}": state,
        "${zip_code}": zipc,
        "${poc_name}": poc,
        "${poc_title}": poc_t,
        "${auditor_name}": auditor,
        "${fy_end_text}": fy_end,
    }

    # Ensure no None values sneak in
    mapping = {k: (v if v is not None else "") for k, v in mapping.items()}
    email = (model.get("treasury_contact_email") or "ORP_SingleAudits@treasury.gov ").strip()

    mapping.update({
        # bracket style used by template
        "[treasury_contact_email]": f" {email} ",
        # curly style just in case
        "${treasury_contact_email}": f" {email} "
    })

    # After you load/prepare model + mapping replacements, before table insertion/plurals:
    total_findings = dedupe_model_programs_and_findings(model)
    programs = model.get("programs") or []

    # 1) Replace placeholders everywhere (body + headers/footers + nested tables)
    replace_placeholders_docwide(doc, mapping)
    # 2) Fix questioned costs grammar
    fix_questioned_costs_grammar(doc)
    # 3) Run email cleanup BEFORE hyperlink creation
    email_postfix_cleanup(doc, email)
    strip_leading_token_artifacts(doc)

    # 5) Final cleanups that don't touch text
    unset_all_caps_everywhere(doc)

    # 2) Insert program tables at the anchor
    anchor = find_anchor_paragraph(doc, "[[PROGRAM_TABLES]]")
    logging.info(f"Anchor found: {anchor is not None}")
    if anchor:
        logging.info(f"Anchor text: {repr(para_text(anchor))}")
    if not anchor:
        raise HTTPException(400, "Template does not contain the [[PROGRAM_TABLES]] anchor paragraph.")
    programs = model.get("programs", []) or []

    try:
        label_p = find_para_by_contains(doc, "Assistance Listing Number/Program Name")
        progs = model.get("programs") or []
        if label_p is not None and progs:
            PROGRAM_PRIORITY = {
                "21.027": 0,
                "21.026": 1,
                "21.023": 2,
                "21.029": 3,
                "21.031": 4,
                "21.032": 5,
            }
            progs_sorted = sorted(progs, key=lambda p: (PROGRAM_PRIORITY.get(p.get("assistance_listing") or "", 99), p.get("assistance_listing") or "99.999"))
            first = progs_sorted[0]
            aln = (first.get("assistance_listing") or "").strip()
            pname = (first.get("program_name") or "").strip()

            # Clear the paragraph and add formatted text
            clear_runs(label_p)

            # Add bold header text
            header_run = label_p.add_run("Assistance Listing Number/Program Name:")
            header_run.bold = True
            # Add a line break (not new paragraph)
            label_p.add_run("\n")
            # Add the ALN and program name (not bold)
            label_p.add_run(f"{aln} / {pname}")
            # KEY FIX: Set tight spacing - use tight_paragraph for consistent removal
            tight_paragraph(label_p)

            pf = label_p.paragraph_format
            #logging.info(f"Label para - space_before: {pf.space_before}, space_after: {pf.space_after}")

            # After creating table:
            #logging.info(f"Table spacing check")
            # Remove any duplicate headers that follow
            # Only remove duplicate headers if there's exactly one program
            # For multiple programs, each table needs its own header
            if len(progs) == 1:
                remove_duplicate_program_headers(doc, label_p)
    except Exception as e:
        logging.warning(f"Error handling program headers: {e}")
        pass

    insert_program_tables_at_anchor_no_headers(doc, anchor, programs)

    # Remove duplicate narrative blocks under the table
    try:
        cleanup_post_table_narrative(doc, model)
    except Exception:
        pass

    # Grammar-fix optional plurals via OpenAI (if key set)
    try:
        total_findings = sum(len(prog.get("findings") or []) for prog in (model.get("programs") or []))
        ai_fix_pluralization_in_doc(doc, total_findings)
    except Exception:
        pass

    # Deterministic grammar fix: resolve [is/are], [The], (s), (es) tokens
    # MUST run BEFORE _strip_leftovers_in_container which strips all [...] patterns
    total_findings = sum(len(prog.get("findings") or []) for prog in (model.get("programs") or []))
    apply_mdl_grammar(doc, total_findings)

    for p in iter_all_paragraphs_full(doc):
        text = para_text(p)
        if "identified issues violates" in text:
            new_text = text.replace(
                "identified issues violates",
                "identified issues violate" if total_findings > 1 else "identified issue violates"
            )
            if new_text != text:
                clear_runs(p)
                p.add_run(new_text)
                break

    # Final tidy: strip any *remaining* token patterns like ${...} or [...]
    def _strip_leftovers_in_container(container):
        for p in iter_all_paragraphs_in_container(container):
            t = para_text(p)
            if not t:
                continue
            new_t = t
            if "${" in new_t:
                new_t = re.sub(r"\$\{[^}]+\}", "", new_t)
            if "[" in new_t and "]" in new_t:
                new_t = re.sub(r"\[[^\]]+\]", "", new_t)
            if new_t != t:
                clear_runs(p)
                p.add_run(new_t)

    _strip_leftovers_in_container(doc)
    for sec in doc.sections:
        _strip_leftovers_in_container(sec.header)
        _strip_leftovers_in_container(sec.footer)

    set_font_size_to_12(doc)

    
    bio = BytesIO()

    # ========== FORCE FIX NARRATIVE PARAGRAPH (FINAL, BOLD-SAFE) ==========
    #correct_auditee = model.get("auditee_name") or model.get("recipient_name") or ""
    correct_auditee = ensure_leading_the(model.get("auditee_name") or model.get("recipient_name") or "")
    correct_auditor = smart_title_case(model.get("auditor_name") or "")

    # Strip leading "The "
    '''if correct_auditee.lower().startswith("the "):
        correct_auditee = correct_auditee[4:].strip()'''

    for p in iter_all_paragraphs_in_container(doc):
        text = para_text(p)

        if "Treasury has reviewed the single audit report for" not in text:
            continue

        pattern = (
            r'(Treasury has reviewed the single audit report for )(?:the )?'
            r'(.+?)'
            r'(, prepared by )(.+?)'
            r'( for the fiscal year)'
        )

        m = re.search(pattern, text)
        if not m:
            continue

        clear_runs(p)

        # Build paragraph with run-level formatting
        p.add_run(m.group(1))                 # fixed intro text

        r = p.add_run(correct_auditee)        # auditee
        r.bold = True                         # GUARANTEED bold

        p.add_run(m.group(3))                 # ", prepared by "
        p.add_run(correct_auditor)            # auditor
        p.add_run(m.group(5))                 # trailing text

        break

    doc.save(bio)
    return bio.getvalue()
