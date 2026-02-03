# services/html_converter.py
from io import BytesIO

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from bs4 import BeautifulSoup
from html2docx import HTML2Docx

from utils.docx_utils import apply_grid_borders, set_col_widths


def apply_inline_formatting(paragraph, node):
    def add_text(text, bold=False, italic=False, underline=False):
        if text is None:
            return
        r = paragraph.add_run(text)
        r.bold = bool(bold)
        r.italic = bool(italic)
        r.underline = bool(underline)

    for child in getattr(node, "children", []):
        name = getattr(child, "name", None)
        if name is None:
            add_text(str(child))
        elif name in ("b", "strong"):
            add_text(child.get_text(), bold=True)
        elif name in ("i", "em"):
            add_text(child.get_text(), italic=True)
        elif name == "u":
            add_text(child.get_text(), underline=True)
        elif name == "br":
            paragraph.add_run().add_break()
        else:
            if hasattr(child, "children"):
                apply_inline_formatting(paragraph, child)


def basic_html_to_docx(doc: Document, html_str: str):
    soup = BeautifulSoup(html_str, "html.parser")
    body = soup.body or soup

    for element in body.children:
        if getattr(element, "name", None) is None:
            txt = str(element).strip()
            if txt:
                doc.add_paragraph(txt)
            continue

        tag = element.name.lower()
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(tag[1])
            text = element.get_text(strip=False)
            doc.add_heading(text, level=min(max(level, 1), 6))
            continue

        if tag == "p":
            p = doc.add_paragraph()
            apply_inline_formatting(p, element)
            continue

        if tag in ("ul", "ol"):
            style = "List Bullet" if tag == "ul" else "List Number"
            for li in element.find_all("li", recursive=False):
                p = doc.add_paragraph(style=style)
                apply_inline_formatting(p, li)
            continue

        if tag == "table":
            rows = element.find_all("tr", recursive=False)
            if not rows:
                continue
            first_cells = rows[0].find_all(["th", "td"], recursive=False)
            cols = max(1, len(first_cells))
            first_is_header = any(c.name == "th" for c in first_cells)

            tbl = doc.add_table(rows=len(rows), cols=cols)

            try:
                tbl.style = "Table Grid"
            except Exception:
                pass
            apply_grid_borders(tbl)

            sect = doc.sections[0]
            content_width = sect.page_width - sect.left_margin - sect.right_margin
            col_w = int(content_width / cols)
            if cols != 5:
                set_col_widths(tbl, [col_w] * cols)

            for r_idx, tr in enumerate(rows):
                cells = tr.find_all(["th", "td"], recursive=False)
                for c_idx in range(cols):
                    cell = tbl.cell(r_idx, c_idx)
                    if not cell.paragraphs:
                        cell.add_paragraph()
                    p = cell.paragraphs[0]
                    if hasattr(p, "clear"):
                        p.clear()
                    if c_idx < len(cells):
                        apply_inline_formatting(p, cells[c_idx])
                    else:
                        p.text = ""
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            if first_is_header:
                for c in tbl.rows[0].cells:
                    for r in c.paragraphs[0].runs:
                        r.bold = True
                    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if tag in ("div", "section", "article"):
            p = doc.add_paragraph()
            apply_inline_formatting(p, element)
            continue

        txt = element.get_text(strip=True)
        if txt:
            doc.add_paragraph(txt)


def set_font_size_to_12(doc):
    """Set all text in document to 12pt font."""
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.size = Pt(12)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(12)

    for section in doc.sections:
        for container in (section.header, section.footer):
            for p in container.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(12)


def html_to_docx_bytes(html_str: str, *, force_basic: bool = False) -> bytes:
    doc = Document()
    if not force_basic:
        try:
            HTML2Docx().add_html_to_document(html_str or "", doc)
        except Exception:
            basic_html_to_docx(doc, html_str or "")
    else:
        basic_html_to_docx(doc, html_str or "")

    if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
        doc.add_paragraph("HTML result is empty.")

    set_font_size_to_12(doc)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
