# main.py
from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional
from io import BytesIO
import os, re, base64
from bs4 import BeautifulSoup
from docx.shared import Pt
import requests
from fastapi import FastAPI, HTTPException

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pydantic import BaseModel
# ============================================================
from urllib.parse import quote

# DOCX & HTML conversion
from docx import Document
from html2docx import HTML2Docx

# Azure (optional)
from azure.storage.blob import BlobServiceClient, BlobSasPermissions, generate_blob_sas
from fastapi.middleware.cors import CORSMiddleware


app = FastAPI(title="MDL DOCX Builder")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],   # or ["http://localhost:*","http://127.0.0.1:*"]
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
# ============================================================
# Helpers
# ============================================================
FAC_BASE = os.getenv("FAC_API_BASE", "https://api.fac.gov")
FAC_KEY  = os.getenv("FAC_API_KEY")

def _short(s: Optional[str], limit: int = 900) -> str:
    """
    Collapse whitespace and truncate to `limit` chars with an ellipsis.
    Safe for None.
    """
    if not s:
        return ""
    s = re.sub(r"\s+", " ", s.strip())
    return (s[: limit - 1] + "…") if len(s) > limit else s

def _fac_headers():
    key = os.getenv("FAC_API_KEY")
    if not key:
        raise HTTPException(500, "FAC_API_KEY not configured on the docx service")
    return {"X-Api-Key": key}

def _fac_get(path: str, params: Dict[str, Any]) -> Any:
    """GET to FAC with sane timeouts and clear errors."""
    try:
        r = requests.get(f"{FAC_BASE.rstrip('/')}/{path.lstrip('/')}",
                         headers=_fac_headers(), params=params, timeout=15)
        r.raise_for_status()
        return r.json()
    except requests.HTTPError as e:
        raise HTTPException(r.status_code if 'r' in locals() else 500,
                            f"FAC GET {path} failed: {getattr(r,'text','')}") from e
    except Exception as e:
        raise HTTPException(500, f"FAC GET {path} failed: {e}") from e

def _or_param(field: str, values: List[str]) -> str:
    """
    Build PostgREST OR param like:
    or=(finding_ref_number.eq.2024-001,finding_ref_number.eq.2024-002)
    requests will URL-encode it for us in params.
    """
    inner = ",".join([f"{field}.eq.{v}" for v in values])
    return f"({inner})"

def _shade_cell(cell, hex_fill="D9D9D9"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_fill)  # light gray
    tcPr.append(shd)

def _set_col_widths(table, widths):
    # widths are in EMU (docx internal units)
    for col_idx, w in enumerate(widths):
        for cell in table.columns[col_idx].cells:
            cell.width = w

def _tight_paragraph(p):
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)


def sanitize(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9._-]+", "_", name or "").strip("_")

def _parse_conn_str(conn: str) -> Dict[str, Optional[str]]:
    """
    Return dict with AccountName, AccountKey, BlobEndpoint.
    Supports both Azurite short form and full connection string.
    """
    parts = dict(p.split("=", 1) for p in conn.split(";") if "=" in p)

    if "UseDevelopmentStorage=true" in conn:
        return {
            "AccountName": "devstoreaccount1",
            "AccountKey": ("Eby8vdM02xNOcqFlqUwJPLlmEtlCDXJ1OUzFT50uSRZ6IFsu"
                           "Fq2UVErCz4I6tq/K1SZFPTOtr/KBHBeksoGMGw=="),
            "BlobEndpoint": "http://127.0.0.1:10000/devstoreaccount1"
        }

    return {
        "AccountName": parts.get("AccountName"),
        "AccountKey": parts.get("AccountKey"),
        "BlobEndpoint": parts.get("BlobEndpoint")  # present for Azurite full conn string
    }

def upload_and_sas(container: str, blob_name: str, data: bytes, ttl_minutes: int = 120) -> str:
    conn = os.getenv("AZURE_STORAGE_CONNECTION_STRING")
    if not conn:
        raise RuntimeError("AZURE_STORAGE_CONNECTION_STRING not set")

    info = _parse_conn_str(conn)
    account_name = info["AccountName"]
    account_key = info["AccountKey"]
    blob_endpoint = info.get("BlobEndpoint")

    bsc = BlobServiceClient.from_connection_string(conn)
    cc = bsc.get_container_client(container)
    try:
        cc.create_container()
    except Exception:
        pass

    cc.upload_blob(name=blob_name, data=data, overwrite=True)

    # 🔑 Key bits for Azurite
    sas = generate_blob_sas(
        account_name=account_name,
        account_key=account_key,
        container_name=container,
        blob_name=blob_name,
        permission=BlobSasPermissions(read=True),
        expiry=datetime.utcnow() + timedelta(minutes=ttl_minutes),
        version=os.getenv("AZURITE_SAS_VERSION", "2021-08-06"),  # Azurite-compatible
        protocol="http",  # Azurite is http
    )

    # Ensure query is safely encoded (avoid '+'/'/' issues)
    sas_q = quote(sas, safe="=&")

    base = blob_endpoint.rstrip("/") if blob_endpoint else f"https://{account_name}.blob.core.windows.net"
    return f"{base}/{container}/{blob_name}?{sas_q}"

# Local file fallback (when no Azure/Azurite configured)
LOCAL_SAVE_DIR = os.getenv("LOCAL_SAVE_DIR", "./_out")
PUBLIC_BASE_URL = os.getenv("PUBLIC_BASE_URL", "http://localhost:8000")

def save_local_and_url(blob_name: str, data: bytes) -> str:
    base_dir = LOCAL_SAVE_DIR
    full_path = os.path.join(base_dir, blob_name)
    os.makedirs(os.path.dirname(full_path), exist_ok=True)
    with open(full_path, "wb") as f:
        f.write(data)
    return f"{PUBLIC_BASE_URL}/local/{blob_name}"

@app.get("/local/{path:path}")
def get_local_file(path: str):
    full = os.path.join(LOCAL_SAVE_DIR, path)
    if not os.path.isfile(full):
        raise HTTPException(404, "Not found")
    return FileResponse(
        full,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

def html_to_docx_bytes(html: str) -> bytes:
    """
    Convert HTML to .docx. Try html2docx first; if it errors, fall back to a
    lightweight BeautifulSoup renderer that supports h1–h6, p, br, b/strong,
    i/em, u, ul/ol/li, and simple tables.
    """
    doc = Document()
    tried_html2docx = False

    # 1) Try the library converter
    try:
        tried_html2docx = True
        HTML2Docx().add_html_to_document(html or "", doc)
    except Exception:
        # 2) Fallback: simple manual renderer
        _basic_html_to_docx(doc, html or "")

    # Safety: ensure non-empty document
    if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
        if tried_html2docx:
            doc.add_paragraph("⚠️ html2docx failed; fallback also produced no content.")
        else:
            doc.add_paragraph("⚠️ HTML was empty.")

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def _apply_inline_formatting(paragraph, node):
    """
    Append text from a BeautifulSoup node to `paragraph`, applying inline
    formatting for <b>/<strong>, <i>/<em>, and <u>. Handles <br> as a line break.
    """
    from docx.text.run import Run
    def add_text(text, bold=False, italic=False, underline=False):
        if text is None:
            return
        r = paragraph.add_run(text)
        r.bold = bool(bold)
        r.italic = bool(italic)
        r.underline = bool(underline)

    for child in node.children:
        name = getattr(child, "name", None)
        if name is None:
            add_text(str(child))
        elif name in ("b", "strong"):
            add_text(child.get_text(), bold=True)
        elif name in ("i", "em"):
            add_text(child.get_text(), italic=True)
        elif name == "u":
            add_text(child.get_text(), underline=True)
        elif name == "br":
            paragraph.add_run().add_break()
        else:
            # Recurse for nested inline containers/spans
            if hasattr(child, "children"):
                _apply_inline_formatting(paragraph, child)

def _basic_html_to_docx(doc: Document, html: str):
    """
    Very simple HTML renderer: block tags (h1–h6, p, ul/ol/li, table) + inline formatting.
    """
    soup = BeautifulSoup(html, "html.parser")

    def add_paragraph_with_text(text, style=None):
        p = doc.add_paragraph()
        if style:
            p.style = style
        _apply_inline_formatting(p, BeautifulSoup(text, "html.parser"))
        return p

    body = soup.body or soup

    for element in body.children:
        if getattr(element, "name", None) is None:
            # Raw text nodes at top-level: wrap in paragraph
            txt = str(element).strip()
            if txt:
                doc.add_paragraph(txt)
            continue

        tag = element.name.lower()

        # Headings
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(tag[1])
            text = element.get_text(strip=False)
            doc.add_heading(text, level=min(max(level, 1), 6))
            continue

        # Paragraph
        if tag == "p":
            p = doc.add_paragraph()
            _apply_inline_formatting(p, element)
            continue

        # Unordered / Ordered lists
        if tag in ("ul", "ol"):
            list_style = "List Bullet" if tag == "ul" else "List Number"
            for li in element.find_all("li", recursive=False):
                p = doc.add_paragraph(style=list_style)
                _apply_inline_formatting(p, li)
            continue

        # Tables
        if tag == "table":
            rows = element.find_all("tr", recursive=False)
            if not rows:
                continue

            # Determine column count and whether first row is a header
            first_cells = rows[0].find_all(["th", "td"], recursive=False)
            cols = max(1, len(first_cells))
            first_is_header = any(c.name == "th" for c in first_cells)

            # Create table
            tbl = doc.add_table(rows=len(rows), cols=cols)
            tbl.style = "Table Grid"

            # Make the table span the content width and set equal column widths
            sect = doc.sections[0]
            content_width = sect.page_width - sect.left_margin - sect.right_margin
            col_w = int(content_width / cols)
            _set_col_widths(tbl, [col_w] * cols)

            # Fill rows
            for r_idx, tr in enumerate(rows):
                cells = tr.find_all(["th", "td"], recursive=False)
                for c_idx in range(cols):
                    cell = tbl.cell(r_idx, c_idx)
                    # ensure empty paragraph exists
                    if not cell.paragraphs:
                        cell.add_paragraph()
                    p = cell.paragraphs[0]
                    p.clear() if hasattr(p, "clear") else None  # for older python-docx this is no-op

                    # Render cell content
                    if c_idx < len(cells):
                        _apply_inline_formatting(p, cells[c_idx])
                    else:
                        p.text = ""

                    # Tight spacing and top vertical align
                    _tight_paragraph(p)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            # Style header row
            if first_is_header:
                for c in tbl.rows[0].cells:
                    _shade_cell(c, "E7E6E6")  # lighter gray
                    for r in c.paragraphs[0].runs:
                        r.bold = True
                    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        # Div / Section / Unknown container → render textual content as paragraphs
        if tag in ("div", "section", "article"):
            text = element.decode_contents()
            p = doc.add_paragraph()
            _apply_inline_formatting(p, element)
            continue

        # Default: dump text
        txt = element.get_text(strip=True)
        if txt:
            doc.add_paragraph(txt)

# ============================================================
# Models
# ============================================================

class BuildDocx(BaseModel):
    auditee_name: str
    ein: str
    audit_year: int
    body_html: Optional[str] = None
    body_html_b64: Optional[str] = None  # optional: if you prefer to send base64 html
    dest_path: Optional[str] = None
    filename: Optional[str] = None

class BuildFromFAC(BaseModel):
    auditee_name: str
    ein: str
    audit_year: int
    fac_general: List[Dict[str, Any]] = []
    fac_findings: List[Dict[str, Any]] = []
    fac_findings_text: List[Dict[str, Any]] = []
    fac_caps: List[Dict[str, Any]] = []
    federal_awards: List[Dict[str, Any]] = []
    dest_path: Optional[str] = None
    filename: Optional[str] = None

class BuildByReport(BaseModel):
    auditee_name: str
    ein: str
    audit_year: int
    report_id: str
    dest_path: Optional[str] = None
    only_flagged: bool = False
    max_refs: int = 15
    include_awards: bool = False  # default OFF to keep payload small
# ============================================================
# Routes
# ============================================================

@app.get("/healthz")
def healthz():
    return {"ok": True, "time": datetime.utcnow().isoformat()}

@app.post("/echo")
def echo(payload: Dict[str, Any]):
    return {"received": payload, "ts": datetime.utcnow().isoformat()}

@app.get("/debug/env")
def debug_env():
    import os
    key = os.getenv("FAC_API_KEY") or ""
    masked = (key[:4] + "…" + key[-2:]) if key else None
    return {"fac_api_key_present": bool(key), "fac_api_key_masked": masked}
@app.post("/build-docx-demo")
def build_docx_demo():
    document = Document()
    document.add_heading("Hello from the DOCX demo ✅", level=1)
    document.add_paragraph("If you can read this, your write/upload pipeline is good.")
    bio = BytesIO()
    document.save(bio)
    data = bio.getvalue()

    container = os.getenv("AZURE_BLOB_CONTAINER", "mdl-output")
    blob_name = "demo/hello.docx"

    if os.getenv("AZURE_STORAGE_CONNECTION_STRING"):
        url = upload_and_sas(container, blob_name, data)
    else:
        url = save_local_and_url(blob_name, data)

    return {"url": url, "size_bytes": len(data)}

# ---------- HTML -> DOCX (production route) ----------
@app.post("/build-docx")
def build_docx(req: BuildDocx):
    # Accept either plain HTML or base64-encoded HTML
    html = (req.body_html or "").strip()
    if (not html) and req.body_html_b64:
        try:
            html = base64.b64decode(req.body_html_b64).decode("utf-8", errors="ignore")
        except Exception:
            raise HTTPException(400, "Invalid base64 in body_html_b64")
    if not html:
        raise HTTPException(400, "body_html (or body_html_b64) is required")

    data = html_to_docx_bytes(html)

    container = os.getenv("AZURE_BLOB_CONTAINER", "mdl-output")
    folder = (req.dest_path or "").lstrip("/")
    base = req.filename or f"MDL-{sanitize(req.auditee_name)}-{sanitize(req.ein)}-{req.audit_year}.docx"
    blob_name = f"{folder}{base}" if folder else base

    url = (upload_and_sas(container, blob_name, data)
           if os.getenv("AZURE_STORAGE_CONNECTION_STRING")
           else save_local_and_url(blob_name, data))
    return {"url": url, "blob_path": f"{container}/{blob_name}", "size_bytes": len(data)}

# ---------- FAC arrays -> Preview MDL DOCX (no-LLM) ----------
def _compose_html_from_fac(req: "BuildFromFAC") -> str:
    g = (req.fac_general[0] if req.fac_general else {})
    rid = g.get("report_id") or "N/A"
    fac_date = g.get("fac_accepted_date") or "N/A"

    text_map = {
        (t.get("finding_ref_number") or t.get("reference_number")): (t.get("finding_text") or "")
        for t in (req.fac_findings_text or [])
    }
    cap_map = {
        (c.get("finding_ref_number") or c.get("reference_number")): (c.get("planned_action") or "")
        for c in (req.fac_caps or [])
    }
    prog_map = {
        a.get("award_reference"): a.get("federal_program_name")
        for a in (req.federal_awards or []) if a.get("award_reference")
    }

    def relevant(row: Dict[str, Any]) -> bool:
        flags = [
            row.get("is_material_weakness"),
            row.get("is_significant_deficiency"),
            row.get("is_questioned_costs"),
            row.get("is_modified_opinion"),
            row.get("is_other_findings"),
            row.get("is_other_matters"),
            row.get("is_repeat_finding"),
        ]
        return any(bool(x) for x in flags)

    findings = [f for f in (req.fac_findings or []) if relevant(f)]
    findings.sort(key=lambda x: str(x.get("reference_number") or ""))
    findings = findings[:50]  # cap for preview

    rows = []
    for f in findings:
        ref = str(f.get("reference_number") or "")
        program = prog_map.get(f.get("award_reference")) or "N/A"
        req_type = f.get("type_requirement") or ""
        sev = []
        if f.get("is_material_weakness"): sev.append("material_weakness")
        if f.get("is_significant_deficiency"): sev.append("significant_deficiency")
        if f.get("is_questioned_costs"): sev.append("questioned_costs")
        if f.get("is_modified_opinion"): sev.append("modified_opinion")
        if f.get("is_repeat_finding"): sev.append("repeat_finding")
        if f.get("is_other_matters"): sev.append("other_matters")
        if f.get("is_other_findings"): sev.append("other_findings")
        sev_str = ", ".join(sev) if sev else "—"
        summary = _short(text_map.get(ref, ""), 900)
        cap_txt = _short(cap_map.get(ref, ""), 400)
        rows.append(
            f"<tr><td>{ref}</td><td>{program}</td><td>{req_type}</td>"
            f"<td>{sev_str}</td><td>{summary}</td><td>{cap_txt}</td></tr>"
        )

    total = len(findings)
    exec_summary = (f"{total} MDL-relevant finding(s) identified."
                    if total else "No MDL-relevant findings identified per FAC records.")

    return f"""<!DOCTYPE html>
<html>
<head><meta charset="utf-8"><title>MDL Preview - {req.auditee_name} - {req.ein} - {req.audit_year}</title></head>
<body>
<h1>Master Decision Letter – Preview (No-LLM)</h1>
<p><b>Auditee:</b> {req.auditee_name}<br/>
<b>EIN:</b> {req.ein} &nbsp; <b>Audit Year:</b> {req.audit_year}<br/>
<b>FAC Report ID:</b> {rid} &nbsp; <b>FAC Accepted:</b> {fac_date}<br/>
<b>Date:</b> {datetime.utcnow().date().isoformat()}</p>

<h2>Executive Summary</h2>
<p>{exec_summary}</p>

<h2>Findings (first {total} shown)</h2>
<table border="1" cellspacing="0" cellpadding="6">
  <thead>
    <tr>
      <th>Ref #</th><th>Program</th><th>Requirement Type</th>
      <th>Severity</th><th>Summary (trimmed)</th><th>CAP (trimmed)</th>
    </tr>
  </thead>
  <tbody>
    {''.join(rows) if rows else '<tr><td colspan="6">None</td></tr>'}
  </tbody>
</table>

<div style="page-break-after: always;"></div>
<h2>Appendix A — Raw counts</h2>
<ul>
  <li>general rows: {len(req.fac_general or [])}</li>
  <li>findings rows (MDL-relevant): {total}</li>
  <li>findings_text rows: {len(req.fac_findings_text or [])}</li>
  <li>corrective_action_plans rows: {len(req.fac_caps or [])}</li>
  <li>federal_awards rows: {len(req.federal_awards or [])}</li>
</ul>
</body></html>"""

@app.post("/build-docx-from-fac")
def build_docx_from_fac(req: BuildFromFAC):
    html = _compose_html_from_fac(req)
    data = html_to_docx_bytes(html)

    container = os.getenv("AZURE_BLOB_CONTAINER", "mdl-output")
    folder = (req.dest_path or "").lstrip("/")
    base = req.filename or f"MDL-Preview-{sanitize(req.auditee_name)}-{sanitize(req.ein)}-{req.audit_year}.docx"
    blob_name = f"{folder}{base}" if folder else base

    url = (upload_and_sas(container, blob_name, data)
           if os.getenv("AZURE_STORAGE_CONNECTION_STRING")
           else save_local_and_url(blob_name, data))
    return {"url": url, "blob_path": f"{container}/{blob_name}", "size_bytes": len(data)}

import requests
from fastapi import HTTPException

FAC_BASE = os.getenv("FAC_API_BASE", "https://api.fac.gov")
FAC_KEY  = os.getenv("FAC_API_KEY")

def _fac_headers():
    if not FAC_KEY:
        raise HTTPException(500, "FAC_API_KEY not configured on the docx service")
    return {"X-Api-Key": FAC_KEY}

def _fac_get(path: str, params: Dict[str, Any]) -> Any:
    try:
        r = requests.get(f"{FAC_BASE.rstrip('/')}/{path.lstrip('/')}",
                         headers=_fac_headers(), params=params, timeout=15)
        r.raise_for_status()
        return r.json()
    except requests.HTTPError as e:
        raise HTTPException(r.status_code if 'r' in locals() else 500,
                            f"FAC GET {path} failed: {getattr(r,'text','')}") from e
    except Exception as e:
        raise HTTPException(500, f"FAC GET {path} failed: {e}") from e

def _or_param(field: str, values: List[str]) -> str:
    inner = ",".join([f"{field}.eq.{v}" for v in values])
    return f"({inner})"

@app.post("/build-docx-by-report")
def build_docx_by_report(req: BuildByReport):
    # 1) Minimal general (for accepted date)
    fac_general = _fac_get("general", {
        "report_id": f"eq.{req.report_id}",
        "select": "report_id,fac_accepted_date",
        "limit": 1
    })

    # 2) Findings (small; optionally only with flags)
    findings_params = {
        "report_id": f"eq.{req.report_id}",
        "select": "reference_number,award_reference,type_requirement,"
                  "is_material_weakness,is_significant_deficiency,is_questioned_costs,"
                  "is_modified_opinion,is_other_findings,is_other_matters,is_repeat_finding",
        "order": "reference_number.asc",
        "limit": str(req.max_refs)
    }
    if req.only_flagged:
        flagged = [
            "is_material_weakness","is_significant_deficiency","is_questioned_costs",
            "is_modified_opinion","is_other_findings","is_other_matters","is_repeat_finding"
        ]
        findings_params["or"] = "(" + ",".join([f"{f}.eq.true" for f in flagged]) + ")"

    fac_findings = _fac_get("findings", findings_params)

    # refs we actually fetched
    refs = [row.get("reference_number") for row in fac_findings if row.get("reference_number")]
    refs = refs[: req.max_refs]

    # 3) Finding text & CAP only for those refs
    if refs:
        fac_findings_text = _fac_get("findings_text", {
            "report_id": f"eq.{req.report_id}",
            "select": "finding_ref_number,finding_text",
            "order": "finding_ref_number.asc",
            "limit": str(len(refs)),
            "or": _or_param("finding_ref_number", refs)
        })
        fac_caps = _fac_get("corrective_action_plans", {
            "report_id": f"eq.{req.report_id}",
            "select": "finding_ref_number,planned_action",
            "order": "finding_ref_number.asc",
            "limit": str(len(refs)),
            "or": _or_param("finding_ref_number", refs)
        })
    else:
        fac_findings_text, fac_caps = [], []

    # 4) (Optional) awards kept small/off by default
    federal_awards = []
    if req.include_awards:
        federal_awards = _fac_get("federal_awards", {
            "report_id": f"eq.{req.report_id}",
            "select": "award_reference,federal_program_name",
            "order": "award_reference.asc",
            "limit": "50"
        })

    # 5) Reuse existing builder
    payload = BuildFromFAC(
        auditee_name=req.auditee_name,
        ein=req.ein,
        audit_year=req.audit_year,
        fac_general=fac_general,
        fac_findings=fac_findings,
        fac_findings_text=fac_findings_text,
        fac_caps=fac_caps,
        federal_awards=federal_awards,
        dest_path=req.dest_path or ""
    )
    return build_docx_from_fac(payload)

@app.get("/debug/storage")
def debug_storage():
    conn = os.getenv("AZURE_STORAGE_CONNECTION_STRING")
    info = _parse_conn_str(conn) if conn else {}
    return {
        "using_storage": bool(conn),
        "account": info.get("AccountName"),
        "blob_endpoint": info.get("BlobEndpoint"),
    }

@app.get("/debug/storage")
def debug_storage():
    conn = os.getenv("AZURE_STORAGE_CONNECTION_STRING")
    info = _parse_conn_str(conn) if conn else {}
    return {"using_storage": bool(conn), "account": info.get("AccountName"), "blob_endpoint": info.get("BlobEndpoint")}

@app.get("/debug/sas")
def debug_sas():
    url = upload_and_sas("mdl-output", "debug/hello.txt", b"hi", ttl_minutes=5)
    return {"url": url}


# ---------- DROP-IN: helpers to generate Treasury-style MDL HTML ----------

from datetime import datetime
import html, re
from typing import Any, Dict, List, Optional, Tuple

# If you already have _short in your file, remove this one.
def _short(s: Optional[str], limit: int) -> str:
    if not s:
        return ""
    s = s.strip()
    return (s[:limit] + "…") if len(s) > limit else s

def _norm_ref(s: Optional[str]) -> str:
    return re.sub(r"\s+", "", (s or "")).upper()

def summarize_finding_text(raw: str, max_chars: int = 1000) -> str:
    """
    Heuristic summarizer (no LLM). Keeps 1–3 short sentences that mention the control/compliance issue.
    """
    if not raw:
        return ""
    text = re.sub(r"\s+", " ", raw).strip()
    # Prefer lines after “Type of Finding / Criteria / Condition / Effect / Recommendation”
    # Basic heuristic: split on periods, keep first 2-3 concise sentences.
    parts = re.split(r"(?<=[.?!])\s+", text)
    picked = []
    for p in parts:
        if len(picked) >= 3:
            break
        # Skip boilerplate “Assistance Listing … Award Period …” lines
        if re.search(r"\b(Assistance Listing|Award Period|Federal Program|Identification Number|CLIN|CFDA)\b", p, re.I):
            continue
        picked.append(p)
    out = " ".join(picked) or text
    return _short(out, max_chars)

def format_letter_date(date_iso: Optional[str] = None) -> Tuple[str, str]:
    """
    Returns (yyyy-mm-dd, 'Month DD, YYYY'). If not provided, uses today UTC.
    """
    dt = datetime.fromisoformat(date_iso) if date_iso else datetime.utcnow()
    return dt.strftime("%Y-%m-%d"), dt.strftime("%B %d, %Y")

def render_mdl_html(model: Dict[str, Any]) -> str:
    """
    Render Treasury-style MDL HTML from normalized model (see build_mdl_model_from_fac).
    """
    # Required header bits
    letter_date_iso = model.get("letter_date_iso")
    letter_date_iso, letter_date_long = format_letter_date(letter_date_iso)

    auditee_name = model.get("auditee_name", "Recipient")
    ein = model.get("ein", "")
    address_lines = model.get("address_lines", [])
    attention_line = model.get("attention_line")
    period_end_text = model.get("period_end_text", str(model.get("audit_year", "")))

    # Build address/attention blocks
    address_block = "<br>".join(html.escape(x) for x in address_lines) if address_lines else ""
    attention_block = f"<p><strong>{html.escape(attention_line)}</strong></p>" if attention_line else ""

    programs = model.get("programs", [])
    not_sustained_notes = model.get("not_sustained_notes", [])

    def _render_program_table(p: Dict[str, Any]) -> str:
        rows_html = []
        for f in p.get("findings", []):
            rows_html.append(f"""
              <tr>
                <td>{html.escape(f.get('finding_id',''))}</td>
                <td>{html.escape(f.get('compliance_type',''))}</td>
                <td>{html.escape(f.get('summary',''))}</td>
                <td>{html.escape(f.get('audit_determination',''))}</td>
                <td>{html.escape(f.get('questioned_cost_determination',''))}</td>
                <td>{html.escape(f.get('cap_determination',''))}</td>
              </tr>
            """)
        if not rows_html:
            rows_html.append("""
              <tr>
                <td colspan="6"><em>No MDL-relevant findings identified for this program.</em></td>
              </tr>
            """)
        table = f"""
          <h3>Assistance Listing Number/Program Name: {html.escape(p.get('assistance_listing','Unknown'))} / {html.escape(p.get('program_name','Unknown'))}</h3>
          <table border="1" cellspacing="0" cellpadding="6" style="border-collapse:collapse; width:100%; font-size:10.5pt;">
            <tr>
              <th>Audit<br>Finding #</th>
              <th>Compliance Type -<br>Audit Finding</th>
              <th>Summary</th>
              <th>Audit Finding<br>Determination</th>
              <th>Questioned Cost<br>Determination</th>
              <th>CAP<br>Determination</th>
            </tr>
            {''.join(rows_html)}
          </table>
        """
        # Optional CAP blocks (only where cap_text is present)
        cap_blocks = []
        for f in p.get("findings", []):
            cap_text = f.get("cap_text")
            if cap_text:
                cap_blocks.append(f"""
                  <h4>Corrective Action Plan – {html.escape(f.get('finding_id',''))}</h4>
                  <p>{html.escape(cap_text)}</p>
                """)
        return table + ("\n".join(cap_blocks) if cap_blocks else "")

    programs_html = "\n".join(_render_program_table(p) for p in programs) if programs else "<p><em>No MDL-relevant findings identified per FAC records.</em></p>"

    not_sustained_html = ""
    if not_sustained_notes:
        notes_paras = "\n".join(f"<p>{html.escape(n)}</p>" for n in not_sustained_notes if n)
        not_sustained_html = f"<h3>FINDINGS NOT SUSTAINED</h3>\n{notes_paras}"

    # Full letter
    return f"""
<div style="font-family: Calibri, Arial, sans-serif; font-size:11pt; line-height:1.4;">
  <p style="text-align:right; margin:0 0 12pt 0;">{html.escape(letter_date_long)}</p>

  <p style="margin:0 0 12pt 0;">
    <strong>DEPARTMENT OF THE TREASURY</strong><br>
    WASHINGTON, D.C.
  </p>

  <p style="margin:0 0 12pt 0;">
    <strong>{html.escape(auditee_name)}</strong><br>
    EIN: {html.escape(ein)}<br>
    {address_block}
  </p>

  {attention_block}

  <p style="margin:12pt 0 12pt 0;">
    <strong>Subject:</strong> U.S. Department of the Treasury’s Management Decision Letter (MDL) for Single Audit Report for the period ending on {html.escape(period_end_text)}
  </p>

  <p>
    In accordance with 2 C.F.R. § 200.521(b), the U.S. Department of the Treasury (Treasury)
    is required to issue a management decision for single audit findings pertaining to awards under
    Treasury’s programs. Treasury’s review as part of its responsibilities under 2 C.F.R § 200.513(c)
    includes an assessment of Treasury’s award recipients’ single audit findings, corrective action plans (CAPs),
    and questioned costs, if any.
  </p>

  <p>
    Treasury has reviewed the single audit report for {html.escape(auditee_name)}. Treasury has made the following determinations
    regarding the audit finding(s) and CAP(s) listed below.
  </p>

  {programs_html}

  {not_sustained_html}

  <p>
    Please note, the corrective action(s) are subject to review during the recipient’s next annual single audit
    or program-specific audit, as applicable, to determine adequacy. If the same audit finding(s) appear in a future single
    audit report for this recipient, its current or future award funding under Treasury’s programs may be adversely impacted.
  </p>

  <p>
    The recipient may appeal Treasury’s decision for the audit finding(s) listed above. A written appeal must be submitted within
    30 calendar days of the date of this management decision letter to Treasury via email at
    <a href="mailto:ORP_SingleAudits@treasury.gov">ORP_SingleAudits@treasury.gov</a>.
    The appeal must include: (1) the specific reasons for disputing Treasury’s determination; (2) relevant documentation
    to support the recipient’s position; (3) an alternative course of action with an anticipated completion date of the action; and
    (4) the contact information of the managing official responsible for implementing the proposed alternative course of action.
  </p>

  <p>For questions regarding the audit finding(s), please email us at <a href="mailto:ORP_SingleAudits@treasury.gov">ORP_SingleAudits@treasury.gov</a>. Thank you.</p>

  <p style="margin-top:18pt;">Sincerely,<br><br>
  Audit and Compliance Resolution Team<br>
  Office of Capital Access<br>
  U.S. Department of the Treasury</p>
</div>
""".strip()

def build_mdl_model_from_fac(
    *,
    auditee_name: str,
    ein: str,
    audit_year: int,
    fac_general: List[Dict[str, Any]],
    fac_findings: List[Dict[str, Any]],
    fac_findings_text: List[Dict[str, Any]],
    fac_caps: List[Dict[str, Any]],
    federal_awards: List[Dict[str, Any]],
    # Optional enrichments / knobs
    period_end_text: Optional[str] = None,           # e.g. "June 30, 2024"
    address_lines: Optional[List[str]] = None,
    attention_line: Optional[str] = None,
    only_flagged: bool = False,
    max_refs: int = 10,
    auto_cap_determination: bool = True,
) -> Dict[str, Any]:
    """
    Normalizes FAC payloads into the MDL model that render_mdl_html() expects.
    - Joins /findings with /findings_text and /corrective_action_plans by normalized reference #
    - Groups findings by program using federal_awards (award_reference -> program name)
    - Fallbacks if flags or joins are missing
    """
    # Map award_reference -> program name (and try to derive assistance listing if present)
    award2name = {}
    for a in federal_awards or []:
        ref = a.get("award_reference")
        if ref:
            award2name[ref] = a.get("federal_program_name") or "Unknown Program"

    # Normalize text & CAP by finding ref
    text_by_ref = {_norm_ref(t.get("finding_ref_number")): (t.get("finding_text") or "").strip()
                   for t in (fac_findings_text or [])}
    cap_by_ref = {_norm_ref(c.get("finding_ref_number")): (c.get("planned_action") or "").strip()
                  for c in (fac_caps or [])}

    # Flag filter
    def _is_flagged(f: dict) -> bool:
        return any([
            f.get("is_material_weakness") is True,
            f.get("is_significant_deficiency") is True,
            f.get("is_questioned_costs") is True,
            f.get("is_modified_opinion") is True,
            f.get("is_other_findings") is True,
            f.get("is_other_matters") is True,
            f.get("is_repeat_finding") is True,
        ])

    # Build refs from /findings
    base_refs = []
    for f in fac_findings or []:
        if only_flagged and not _is_flagged(f):
            continue
        ref = f.get("reference_number")
        if ref:
            base_refs.append(ref)

    # Fallback: derive refs from /findings_text if none
    if not base_refs and fac_findings_text:
        base_refs = [t.get("finding_ref_number") for t in fac_findings_text if t.get("finding_ref_number")]

    # De-dupe, normalize, cap
    norm_refs: List[Tuple[str, str]] = []
    seen = set()
    for r in base_refs:
        k = _norm_ref(r)
        if k and k not in seen:
            norm_refs.append((r, k))
            seen.add(k)
    norm_refs = norm_refs[: max_refs or 10]

    # Group findings by award_reference → program name
    programs_map: Dict[str, Dict[str, Any]] = {}   # key: award_reference or 'UNKNOWN'
    for f in fac_findings or []:
        ref = f.get("reference_number")
        if not ref:
            continue
        kn = _norm_ref(ref)
        if kn not in {k for _, k in norm_refs}:
            # this finding wasn't selected
            continue

        award_ref = f.get("award_reference") or "UNKNOWN"
        group = programs_map.setdefault(award_ref, {
            "assistance_listing": f.get("assistance_listing") or "Unknown",  # FAC may not supply
            "program_name": award2name.get(award_ref, "Unknown Program"),
            "findings": []
        })

        summary = summarize_finding_text(text_by_ref.get(kn, ""))
        cap_text = cap_by_ref.get(kn)

        # basic determinations (can later be replaced with rules/LLM):
        audit_det = "Sustained"  # default conservative
        qcost_det = "None"
        dcost_det = "None"
        cap_det = "Accepted" if (auto_cap_determination and cap_text) else "Not Applicable"

        group["findings"].append({
            "finding_id": f.get("reference_number") or "",
            "compliance_type": f.get("type_requirement") or "",
            "summary": summary,
            "audit_determination": audit_det,
            "questioned_cost_determination": qcost_det,
            "disallowed_cost_determination": dcost_det,
            "cap_determination": cap_det,
            "cap_text": cap_text,
        })

    # If nothing grouped (e.g., /findings empty but /findings_text has narrative), create a single catchall section
    if not programs_map and norm_refs:
        catchall = {
            "assistance_listing": "Unknown",
            "program_name": "Unknown Program",
            "findings": []
        }
        for orig_ref, key in norm_refs:
            summary = summarize_finding_text(text_by_ref.get(key, ""))
            cap_text = cap_by_ref.get(key)
            cap_det = "Accepted" if (auto_cap_determination and cap_text) else "Not Applicable"
            catchall["findings"].append({
                "finding_id": orig_ref,
                "compliance_type": "",
                "summary": summary,
                "audit_determination": "Sustained",
                "questioned_cost_determination": "None",
                "disallowed_cost_determination": "None",
                "cap_determination": cap_det,
                "cap_text": cap_text,
            })
        programs_map["UNKNOWN"] = catchall

    # Build final model
    letter_date_iso, _ = format_letter_date(model.get("letter_date_iso") if (model := {}) else None)  # harmless default
    model = {
        "letter_date_iso": letter_date_iso,
        "auditee_name": auditee_name,
        "ein": f"{ein[:2]}-{ein[2:]}" if ein and ein.isdigit() and len(ein) == 9 else ein,
        "address_lines": address_lines or [],
        "attention_line": attention_line or "",
        "period_end_text": period_end_text or f"June 30, {audit_year}",   # TODO: replace when you fetch actual period end
        "audit_year": audit_year,
        "programs": list(programs_map.values()),
        "not_sustained_notes": [],  # you can inject notes later if needed
    }
    return model
