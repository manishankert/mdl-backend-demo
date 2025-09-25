from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional
from io import BytesIO
import os, re

# docx & html conversion
from docx import Document
from html2docx import html2docx

# Azure (optional—if AZURE_STORAGE_CONNECTION_STRING is set)
from azure.storage.blob import (
    BlobServiceClient, BlobSasPermissions, generate_blob_sas
)

app = FastAPI(title="MDL DOCX Builder (Local Test)")

# ------------------------------------------------------------
# Helpers
# ------------------------------------------------------------
def sanitize(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9._-]+", "_", name).strip("_")

def _parse_conn_str(conn: str):
    parts = dict(p.split("=", 1) for p in conn.split(";") if "=" in p)
    # Azurite short form support
    if "UseDevelopmentStorage=true" in conn:
        return {
            "AccountName": "devstoreaccount1",
            "AccountKey": ("Eby8vdM02xNOcqFlqUwJPLlmEtlCDXJ1OUzFT50uSRZ6IFsu"
                           "Fq2UVErCz4I6tq/K1SZFPTOtr/KBHBeksoGMGw=="),
            "BlobEndpoint": "http://127.0.0.1:10000/devstoreaccount1"
        }
    return {
        "AccountName": parts.get("AccountName"),
        "AccountKey": parts.get("AccountKey"),
        "BlobEndpoint": parts.get("BlobEndpoint")  # present for Azurite full conn str
    }


def upload_and_sas(container: str, blob_name: str, data: bytes, ttl_minutes: int = 120):
    conn = os.getenv("AZURE_STORAGE_CONNECTION_STRING")
    if not conn:
        raise RuntimeError("AZURE_STORAGE_CONNECTION_STRING not set")

    info = _parse_conn_str(conn)
    account_name = info["AccountName"]
    account_key = info["AccountKey"]
    blob_endpoint = info.get("BlobEndpoint")  # e.g., http://127.0.0.1:10000/devstoreaccount1

    from azure.storage.blob import BlobServiceClient, BlobSasPermissions, generate_blob_sas

    bsc = BlobServiceClient.from_connection_string(conn)
    cc = bsc.get_container_client(container)
    try:
        cc.create_container()
    except Exception:
        pass

    cc.upload_blob(name=blob_name, data=data, overwrite=True)

    sas = generate_blob_sas(
        account_name=account_name,
        account_key=account_key,
        container_name=container,
        blob_name=blob_name,
        permission=BlobSasPermissions(read=True),
        # Azurite is happier with an older version; optional but safe:
        version="2021-08-06",
        expiry=datetime.utcnow() + timedelta(minutes=ttl_minutes),
    )

    # Build the right base URL
    if blob_endpoint:
        base = blob_endpoint.rstrip("/")
    else:
        base = f"https://{account_name}.blob.core.windows.net"

    return f"{base}/{container}/{blob_name}?{sas}"


# Local file fallback (if you don’t want Azure/Azurite)
LOCAL_SAVE_DIR = os.getenv("LOCAL_SAVE_DIR", "./_out")
PUBLIC_BASE_URL = os.getenv("PUBLIC_BASE_URL", "http://localhost:8000")

def save_local_and_url(blob_name: str, data: bytes) -> str:
    base = LOCAL_SAVE_DIR
    os.makedirs(os.path.join(base, os.path.dirname(blob_name)), exist_ok=True)
    path = os.path.join(base, blob_name)
    with open(path, "wb") as f:
        f.write(data)
    return f"{PUBLIC_BASE_URL}/local/{blob_name}"


def html_to_docx_bytes(html: str) -> bytes:
    """
    Convert HTML to a non-empty .docx. Uses HTML2Docx class API.
    If conversion yields no visible content, add a debug paragraph so file isn't blank.
    """
    from io import BytesIO
    from docx import Document
    from html2docx import HTML2Docx

    doc = Document()
    # Try the class-based converter (more reliable across versions)
    try:
        HTML2Docx().add_html_to_document(html or "", doc)
    except Exception as e:
        # Absolute fallback: dump raw text (no extra deps)
        import re
        doc.add_heading("HTML conversion failed, showing raw text fallback", level=2)
        raw = re.sub(r"\s+", " ", (html or ""))[:2000]
        doc.add_paragraph(raw)

    # Safety: if converter produced nothing, inject a tiny marker
    if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
        doc.add_paragraph("⚠️ HTML conversion produced no visible content.")

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


@app.get("/local/{path:path}")
def get_local_file(path: str):
    full = os.path.join(LOCAL_SAVE_DIR, path)
    if not os.path.isfile(full):
        raise HTTPException(404, "Not found")
    return FileResponse(full, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ------------------------------------------------------------
# Models
# ------------------------------------------------------------
class BuildDocx(BaseModel):
    auditee_name: str
    ein: str
    audit_year: int
    body_html: str
    dest_path: Optional[str] = None
    filename: Optional[str] = None

class BuildFromFAC(BaseModel):
    auditee_name: str
    ein: str
    audit_year: int
    fac_general: List[Dict[str, Any]] = []
    fac_findings: List[Dict[str, Any]] = []
    fac_findings_text: List[Dict[str, Any]] = []
    fac_caps: List[Dict[str, Any]] = []
    federal_awards: List[Dict[str, Any]] = []
    dest_path: Optional[str] = None
    filename: Optional[str] = None

# ------------------------------------------------------------
# Routes
# ------------------------------------------------------------
@app.get("/healthz")
def healthz():
    return {"ok": True, "time": datetime.utcnow().isoformat()}

@app.post("/echo")
def echo(payload: Dict[str, Any]):
    return {"received": payload, "ts": datetime.utcnow().isoformat()}

@app.post("/build-docx")
def build_docx(req: BuildDocx):
    # Convert HTML -> DOCX
    docx_io = BytesIO()
    document = Document()
    html2docx(req.body_html, document)
    document.save(docx_io)
    data = docx_io.getvalue()

    container = os.getenv("AZURE_BLOB_CONTAINER", "mdl-output")
    folder = (req.dest_path or "").lstrip("/")
    base = req.filename or f"MDL-{sanitize(req.auditee_name)}-{sanitize(req.ein)}-{req.audit_year}.docx"
    blob_name = f"{folder}{base}" if folder else base

    url = (upload_and_sas(container, blob_name, data)
           if os.getenv("AZURE_STORAGE_CONNECTION_STRING")
           else save_local_and_url(blob_name, data))
    return {"url": url, "blob_path": f"{container}/{blob_name}", "size_bytes": len(data)}

def _short(s: Optional[str], limit: int) -> str:
    if not s:
        return ""
    s = s.strip()
    return (s[:limit] + "…") if len(s) > limit else s

def _compose_html_from_fac(req: BuildFromFAC) -> str:
    g = (req.fac_general[0] if req.fac_general else {})
    rid = g.get("report_id") or "N/A"
    fac_date = g.get("fac_accepted_date") or "N/A"

    text_map = { (t.get("finding_ref_number") or t.get("reference_number")): (t.get("finding_text") or "")
                 for t in (req.fac_findings_text or []) }
    cap_map  = { (c.get("finding_ref_number") or c.get("reference_number")): (c.get("planned_action") or "")
                 for c in (req.fac_caps or []) }
    prog_map = { a.get("award_reference"): a.get("federal_program_name")
                 for a in (req.federal_awards or []) if a.get("award_reference") }

    def relevant(row: Dict[str, Any]) -> bool:
        flags = [
            row.get("is_material_weakness"),
            row.get("is_significant_deficiency"),
            row.get("is_questioned_costs"),
            row.get("is_modified_opinion"),
            row.get("is_other_findings"),
            row.get("is_other_matters"),
            row.get("is_repeat_finding"),
        ]
        return any(bool(x) for x in flags)

    findings = [f for f in (req.fac_findings or []) if relevant(f)]
    findings.sort(key=lambda x: str(x.get("reference_number") or ""))
    findings = findings[:50]

    rows = []
    for f in findings:
        ref = str(f.get("reference_number") or "")
        program = prog_map.get(f.get("award_reference")) or "N/A"
        req_type = f.get("type_requirement") or ""
        sev = []
        if f.get("is_material_weakness"): sev.append("material_weakness")
        if f.get("is_significant_deficiency"): sev.append("significant_deficiency")
        if f.get("is_questioned_costs"): sev.append("questioned_costs")
        if f.get("is_modified_opinion"): sev.append("modified_opinion")
        if f.get("is_repeat_finding"): sev.append("repeat_finding")
        if f.get("is_other_matters"): sev.append("other_matters")
        if f.get("is_other_findings"): sev.append("other_findings")
        summary = _short(text_map.get(ref, ""), 900)
        cap_txt = _short(cap_map.get(ref, ""), 400)
        sev_str = ", ".join(sev) if sev else "—"
        rows.append(f"<tr><td>{ref}</td><td>{program}</td><td>{req_type}</td><td>{sev_str}</td><td>{summary}</td><td>{cap_txt}</td></tr>")

    total = len(findings)
    exec_summary = f"{total} MDL-relevant finding(s) identified." if total else "No MDL-relevant findings identified per FAC records."

    return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>MDL Preview - {req.auditee_name} - {req.ein} - {req.audit_year}</title></head>
<body>
<h1>Master Decision Letter – Preview (No-LLM)</h1>
<p><b>Auditee:</b> {req.auditee_name}<br/>
<b>EIN:</b> {req.ein} &nbsp; <b>Audit Year:</b> {req.audit_year}<br/>
<b>FAC Report ID:</b> {rid} &nbsp; <b>FAC Accepted:</b> {fac_date}<br/>
<b>Date:</b> {datetime.utcnow().date().isoformat()}</p>

<h2>Executive Summary</h2>
<p>{exec_summary}</p>

<h2>Findings (first {total} shown)</h2>
<table border="1" cellspacing="0" cellpadding="6">
<thead><tr><th>Ref #</th><th>Program</th><th>Requirement Type</th><th>Severity</th><th>Summary</th><th>CAP</th></tr></thead>
<tbody>
{''.join(rows) if rows else '<tr><td colspan="6">None</td></tr>'}
</tbody></table>

<div style="page-break-after: always;"></div>
<h2>Appendix A — Raw counts</h2>
<ul>
  <li>general rows: {len(req.fac_general or [])}</li>
  <li>findings rows (MDL-relevant): {total}</li>
  <li>findings_text rows: {len(req.fac_findings_text or [])}</li>
  <li>corrective_action_plans rows: {len(req.fac_caps or [])}</li>
  <li>federal_awards rows: {len(req.federal_awards or [])}</li>
</ul>
</body></html>"""

@app.post("/build-docx-from-fac")
def build_docx_from_fac(req: BuildFromFAC):
    html = _compose_html_from_fac(req)
    docx_io = BytesIO()
    document = Document()
    html2docx(html, document)
    document.save(docx_io)
    data = docx_io.getvalue()

    container = os.getenv("AZURE_BLOB_CONTAINER", "mdl-output")
    folder = (req.dest_path or "").lstrip("/")
    base = req.filename or f"MDL-Preview-{sanitize(req.auditee_name)}-{sanitize(req.ein)}-{req.audit_year}.docx"
    blob_name = f"{folder}{base}" if folder else base

    url = (upload_and_sas(container, blob_name, data)
           if os.getenv("AZURE_STORAGE_CONNECTION_STRING")
           else save_local_and_url(blob_name, data))
    return {"url": url, "blob_path": f"{container}/{blob_name}", "size_bytes": len(data)}


@app.post("/build-docx-demo")
def build_docx_demo():
    from io import BytesIO
    from docx import Document
    document = Document()
    document.add_heading("Hello from the DOCX demo ✅", level=1)
    document.add_paragraph("If you can read this in Word/Pages, your write/upload pipeline is good.")
    bio = BytesIO()
    document.save(bio)
    data = bio.getvalue()

    # Reuse your existing storage logic
    container = os.getenv("AZURE_BLOB_CONTAINER", "mdl-output")
    blob_name = "demo/hello.docx"

    if os.getenv("AZURE_STORAGE_CONNECTION_STRING"):
        url = upload_and_sas(container, blob_name, data)
    elif 'save_local_and_url' in globals():
        url = save_local_and_url(blob_name, data)
    else:
        raise HTTPException(500, "No storage configured. Set AZURE_STORAGE_CONNECTION_STRING or LOCAL_SAVE_DIR.")

    return {"url": url, "size_bytes": len(data)}


# Optional convenience launcher
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=int(os.getenv("PORT", "8000")), reload=True)

