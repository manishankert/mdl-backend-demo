# main.py
from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional, Tuple
from io import BytesIO
from urllib.parse import quote
import os, re, base64, html as htmlmod, requests

# DOCX / HTML
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
from html2docx import HTML2Docx
from copy import deepcopy
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

# Azure
from azure.storage.blob import BlobServiceClient, BlobSasPermissions, generate_blob_sas

# ------------------------------------------------------------------------------
# FastAPI app
# ------------------------------------------------------------------------------
app = FastAPI(title="MDL DOCX Builder")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ------------------------------------------------------------------------------
# Environment
# ------------------------------------------------------------------------------
FAC_BASE = os.getenv("FAC_API_BASE", "https://api.fac.gov")
FAC_KEY  = os.getenv("FAC_API_KEY")

AZURE_CONTAINER = os.getenv("AZURE_BLOB_CONTAINER", "mdl-output")
AZURE_CONN_STR  = os.getenv("AZURE_STORAGE_CONNECTION_STRING")  # optional
AZURITE_SAS_VERSION = os.getenv("AZURITE_SAS_VERSION", "2021-08-06")

LOCAL_SAVE_DIR = os.getenv("LOCAL_SAVE_DIR", "./_out")
PUBLIC_BASE_URL = os.getenv("PUBLIC_BASE_URL", "http://localhost:8000")

# Path to the client’s Treasury template (must exist on disk)
MDL_TEMPLATE_PATH = os.getenv(
    "MDL_TEMPLATE_PATH",
    "./Template 1 - Findings Sustained, CAPs Accepted, No QC.docx"
)

# Default Treasury programs ordering (if you later want to filter)
TREASURY_DEFAULT_LISTINGS = ["21.023", "21.026", "21.027"]  # ERA, HAF, SLFRF


# ---- Optional: program acronym map (fill from your workbook) ----
PROGRAM_ACRONYMS = {
    "21.029": "CPF",   # Capital Projects Fund
    "21.019": "CRF",   # Coronavirus Relief Fund
    "21.023": "ERA",   # Emergency Rental Assistance Program
    "21.026": "HAF",   # Homeowner Assistance Fund
    "21.032": "LATCF", # Local Assistance and Tribal Consistency Fund
    "21.031": "SSBCI", # State Small Business Credit Initiative
    "21.027": "SLFRF"  # State and Local Fiscal Recovery Funds
    }
# ------------------------------------------------------------------------------
# Utilities
# ------------------------------------------------------------------------------
def sanitize(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9._-]+", "_", name or "").strip("_")

def _short(s: Optional[str], limit: int = 900) -> str:
    if not s:
        return ""
    s = re.sub(r"\s+", " ", s.strip())
    return (s[: limit - 1] + "…") if len(s) > limit else s

def _norm_ref(s: Optional[str]) -> str:
    return re.sub(r"\s+", "", (s or "")).upper()

def _shade_cell(cell, hex_fill="E7E6E6"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)

def _set_col_widths(table, widths):
    for col_idx, w in enumerate(widths):
        for cell in table.columns[col_idx].cells:
            cell.width = w

def _tight_paragraph(p):
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)

# Extract numbers like 21.027, 10.123, etc.
_AL_RE = re.compile(r"\b\d{2}\.\d{3}\b")

def _derive_assistance_listing(name: str) -> str:
    if not name:
        return "Unknown"
    m = _AL_RE.search(name)
    return m.group(0) if m else "Unknown"

def _aln_sort_key(aln: Optional[str]) -> tuple:
    """Turn '21.027' -> (21, 27) for numeric sorting; otherwise put at end."""
    if not aln:
        return (999, 999999)
    m = re.search(r"(\d+)\.(\d+)", aln)
    if not m:
        return (999, 999999)
    return (int(m.group(1)), int(m.group(2)))

def _extract_repeat_of(text: str) -> Optional[str]:
    """
    Try to pull 'Repeat of 2023-003' style ref from narrative like:
    'Repeat Finding: Yes; 2023-003'
    """
    if not text:
        return None
    m = re.search(r"Repeat\s+Finding:\s*Yes;?\s*([0-9]{4}-[0-9]{3,})", text, re.I)
    return m.group(1) if m else None

# ------------------------------------------------------------------------------
# Storage helpers
# ------------------------------------------------------------------------------
def _parse_conn_str(conn: str) -> Dict[str, Optional[str]]:
    """
    Return dict with AccountName, AccountKey, BlobEndpoint.
    Supports UseDevelopmentStorage=true and Azurite-like "AccountName=...;AccountKey=...;BlobEndpoint=..."
    """
    if not conn:
        return {"AccountName": None, "AccountKey": None, "BlobEndpoint": None}

    if "UseDevelopmentStorage=true" in conn:
        return {
            "AccountName": "devstoreaccount1",
            "AccountKey": (
                "Eby8vdM02xNOcqFlqUwJPLlmEtlCDXJ1OUzFT50uSRZ6IFsu"
                "Fq2UVErCz4I6tq/K1SZFPTOtr/KBHBeksoGMGw=="
            ),
            "BlobEndpoint": "http://127.0.0.1:10000/devstoreaccount1",
        }

    parts = dict(p.split("=", 1) for p in conn.split(";") if "=" in p)
    return {
        "AccountName": parts.get("AccountName"),
        "AccountKey": parts.get("AccountKey"),
        "BlobEndpoint": parts.get("BlobEndpoint"),
    }

def _blob_service_client():
    if not AZURE_CONN_STR:
        raise RuntimeError("AZURE_STORAGE_CONNECTION_STRING not set")

    info = _parse_conn_str(AZURE_CONN_STR)
    if info.get("BlobEndpoint") and info.get("AccountKey"):
        return BlobServiceClient(account_url=info["BlobEndpoint"], credential=info["AccountKey"])
    return BlobServiceClient.from_connection_string(AZURE_CONN_STR)

def upload_and_sas(container: str, blob_name: str, data: bytes, ttl_minutes: int = 120) -> str:
    if not AZURE_CONN_STR:
        raise RuntimeError("AZURE_STORAGE_CONNECTION_STRING not set")

    info = _parse_conn_str(AZURE_CONN_STR)
    account_name = info["AccountName"]
    account_key  = info["AccountKey"]
    blob_endpoint = info.get("BlobEndpoint")

    bsc = _blob_service_client()
    cc = bsc.get_container_client(container)
    try:
        cc.create_container()
    except Exception:
        pass
    cc.upload_blob(name=blob_name, data=data, overwrite=True)

    proto = "http" if (blob_endpoint and ("127.0.0.1" in blob_endpoint or "localhost" in blob_endpoint)) else None
    sas = generate_blob_sas(
        account_name=account_name,
        account_key=account_key,
        container_name=container,
        blob_name=blob_name,
        permission=BlobSasPermissions(read=True),
        expiry=datetime.utcnow() + timedelta(minutes=ttl_minutes),
        version=AZURITE_SAS_VERSION,
        protocol=proto,
    )
    sas_q = quote(sas, safe="=&")

    base = blob_endpoint.rstrip("/") if blob_endpoint else f"https://{account_name}.blob.core.windows.net"
    return f"{base}/{container}/{blob_name}?{sas_q}"

def save_local_and_url(blob_name: str, data: bytes) -> str:
    full_path = os.path.join(LOCAL_SAVE_DIR, blob_name)
    os.makedirs(os.path.dirname(full_path), exist_ok=True)
    with open(full_path, "wb") as f:
        f.write(data)
    return f"{PUBLIC_BASE_URL}/local/{blob_name}"

@app.get("/local/{path:path}")
def get_local_file(path: str):
    full = os.path.join(LOCAL_SAVE_DIR, path)
    if not os.path.isfile(full):
        raise HTTPException(404, "Not found")
    return FileResponse(
        full,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# ------------------------------------------------------------------------------
# FAC helpers
# ------------------------------------------------------------------------------
def _fac_headers():
    key = os.getenv("FAC_API_KEY")
    if not key:
        raise HTTPException(500, "FAC_API_KEY not configured on the docx service")
    return {"X-Api-Key": key}

def _fac_get(path: str, params: Dict[str, Any]) -> Any:
    try:
        r = requests.get(f"{FAC_BASE.rstrip('/')}/{path.lstrip('/')}",
                         headers=_fac_headers(), params=params, timeout=20)
        r.raise_for_status()
        return r.json()
    except requests.HTTPError as e:
        raise HTTPException(r.status_code if 'r' in locals() else 500,
                            f"FAC GET {path} failed: {getattr(r,'text','')}") from e
    except Exception as e:
        raise HTTPException(500, f"FAC GET {path} failed: {e}") from e

def _or_param(field: str, values: List[str]) -> str:
    inner = ",".join([f"{field}.eq.{v}" for v in values])
    return f"({inner})"

# ------------------------------------------------------------------------------
# HTML → DOCX (fallback/preview renderer)
# ------------------------------------------------------------------------------
def _apply_inline_formatting(paragraph, node):
    def add_text(text, bold=False, italic=False, underline=False):
        if text is None:
            return
        r = paragraph.add_run(text)
        r.bold = bool(bold)
        r.italic = bool(italic)
        r.underline = bool(underline)

    for child in getattr(node, "children", []):
        name = getattr(child, "name", None)
        if name is None:
            add_text(str(child))
        elif name in ("b", "strong"):
            add_text(child.get_text(), bold=True)
        elif name in ("i", "em"):
            add_text(child.get_text(), italic=True)
        elif name == "u":
            add_text(child.get_text(), underline=True)
        elif name == "br":
            paragraph.add_run().add_break()
        else:
            if hasattr(child, "children"):
                _apply_inline_formatting(paragraph, child)

def _basic_html_to_docx(doc: Document, html_str: str):
    soup = BeautifulSoup(html_str, "html.parser")
    body = soup.body or soup

    for element in body.children:
        if getattr(element, "name", None) is None:
            txt = str(element).strip()
            if txt:
                doc.add_paragraph(txt)
            continue

        tag = element.name.lower()

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(tag[1])
            text = element.get_text(strip=False)
            doc.add_heading(text, level=min(max(level, 1), 6))
            continue

        if tag == "p":
            p = doc.add_paragraph()
            _apply_inline_formatting(p, element)
            continue

        if tag in ("ul", "ol"):
            style = "List Bullet" if tag == "ul" else "List Number"
            for li in element.find_all("li", recursive=False):
                p = doc.add_paragraph(style=style)
                _apply_inline_formatting(p, li)
            continue

        if tag == "table":
            rows = element.find_all("tr", recursive=False)
            if not rows:
                continue
            first_cells = rows[0].find_all(["th", "td"], recursive=False)
            cols = max(1, len(first_cells))
            first_is_header = any(c.name == "th" for c in first_cells)

            tbl = doc.add_table(rows=len(rows), cols=cols)
            tbl.style = "Table Grid"

            sect = doc.sections[0]
            content_width = sect.page_width - sect.left_margin - sect.right_margin
            col_w = int(content_width / cols)
            _set_col_widths(tbl, [col_w] * cols)

            for r_idx, tr in enumerate(rows):
                cells = tr.find_all(["th", "td"], recursive=False)
                for c_idx in range(cols):
                    cell = tbl.cell(r_idx, c_idx)
                    if not cell.paragraphs:
                        cell.add_paragraph()
                    p = cell.paragraphs[0]
                    if hasattr(p, "clear"):
                        p.clear()
                    if c_idx < len(cells):
                        _apply_inline_formatting(p, cells[c_idx])
                    else:
                        p.text = ""
                    _tight_paragraph(p)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            if first_is_header:
                for c in tbl.rows[0].cells:
                    _shade_cell(c, "E7E6E6")
                    for r in c.paragraphs[0].runs:
                        r.bold = True
                    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if tag in ("div", "section", "article"):
            p = doc.add_paragraph()
            _apply_inline_formatting(p, element)
            continue

        txt = element.get_text(strip=True)
        if txt:
            doc.add_paragraph(txt)

def html_to_docx_bytes(html_str: str, *, force_basic: bool = False) -> bytes:
    doc = Document()
    tried_html2docx = False
    if not force_basic:
        try:
            tried_html2docx = True
            HTML2Docx().add_html_to_document(html_str or "", doc)
        except Exception:
            _basic_html_to_docx(doc, html_str or "")
    else:
        _basic_html_to_docx(doc, html_str or "")

    if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
        doc.add_paragraph("⚠️ HTML was empty." if force_basic else "⚠️ html2docx failed; fallback also produced no content.")
    bio = BytesIO(); doc.save(bio)
    return bio.getvalue()

def _docx_global_replace(doc, mapping: Dict[str, str]):
    """
    Replace placeholders anywhere in the document: body, tables, headers/footers,
    textboxes/shapes. Works by walking every w:t node and doing string replace.
    NOTE: If a single placeholder is split across multiple runs, Word sometimes
    breaks it; most templates keep bracketed tokens intact so this is fine.
    """
    # body + shapes/textboxes
    for t in doc.element.xpath('.//w:t'):
        if t.text:
            txt = t.text
            for k, v in mapping.items():
                if k in txt:
                    txt = txt.replace(k, v)
            t.text = txt

    # headers/footers
    for part in doc.part.related_parts.values():
        el = getattr(part, 'element', None)
        if el is not None:
            for t in el.xpath('.//w:t'):
                if t.text:
                    txt = t.text
                    for k, v in mapping.items():
                        if k in txt:
                            txt = txt.replace(k, v)
                    t.text = txt

def _iter_block_items(parent):
    """
    Yield paragraphs and tables in document order.
    """
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    parent_elm = parent._element
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def _find_program_section(doc: Document):
    """
    Locate the first 'program section' = (para: ALN header, para: ALN line, table)
    Returns (aln_header_para, aln_line_para, table, start_index, end_index)
    """
    blocks = list(_iter_block_items(doc))
    for i, blk in enumerate(blocks):
        if isinstance(blk, Paragraph) and "Assistance Listing Number/Program Name" in (blk.text or ""):
            # expect next is the ALN line para
            if i + 1 < len(blocks) and isinstance(blocks[i+1], Paragraph):
                aln_line = blocks[i+1]
                # then expect a table after that
                if i + 2 < len(blocks) and isinstance(blocks[i+2], Table):
                    tbl = blocks[i+2]
                    # sanity: header cells look like the template’s table
                    hdr_texts = " ".join([c.text.strip() for c in tbl.rows[0].cells])
                    if ("Audit Finding" in hdr_texts and "Compliance Type" in hdr_texts
                        and "Questioned Cost" in hdr_texts and "CAP" in hdr_texts):
                        return (blk, aln_line, tbl, i, i+2)
    return (None, None, None, -1, -1)

def _insert_block_after(doc: Document, ref, new_elm):
    """
    Insert a raw XML block (paragraph or table) after `ref` block.
    """
    ref_elm = ref._element
    ref_elm.addnext(new_elm)

def _set_cell_text(cell: _Cell, text: str):
    # clear paragraph runs and write single run for stability
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    # if no paragraphs, add one
    if not cell.paragraphs:
        cell.add_paragraph()
    cell.paragraphs[0].runs[0].text = text if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(text).text

def _clone_table(tbl: Table) -> Table:
    # deep-copy underlying XML and re-wrap as python-docx Table
    new_tbl_elm = deepcopy(tbl._element)
    return Table(new_tbl_elm, tbl._parent)

def _fill_program_block(doc: Document, aln_header_para: Paragraph, aln_line_para: Paragraph, tbl: Table,
                        program: Dict[str, Any]):
    """
    Mutates the given block in-place for one program.
    program = {
      "assistance_listing": "21.027",
      "program_name": "Coronavirus State and Local Fiscal Recovery Funds",
      "program_acronym": "SLFRF",
      "findings": [{ finding_id, compliance_type, summary, audit_determination, questioned_cost_determination, cap_determination }, ...]
    }
    """
    # 1) ALN line exactly like template's placeholder line
    al = program.get("assistance_listing") or "Unknown"
    pname = program.get("program_name") or "Unknown Program"
    pacr = program.get("program_acronym") or PROGRAM_ACRONYMS.get(al, "")
    aln_line_para.text = f"{al}/ {pname}" + (f" ({pacr})" if pacr else "")

    # 2) Identify the prototype data row:
    # We keep header row(s) untouched; remove existing data rows; re-add for each finding.
    # Assuming the first *non-header* row is the prototype.
    # Heuristics: table has one header row; rest are data.
    header_row = tbl.rows[0]
    data_rows = list(tbl.rows)[1:]
    proto = None
    if data_rows:
        proto = deepcopy(data_rows[0]._tr)
    # wipe all data rows
    while len(tbl.rows) > 1:
        tbl._element.remove(tbl.rows[1]._tr)

    findings = program.get("findings") or []
    if not findings:
        # one-row “no findings” message (matches template tone)
        row = deepcopy(proto) if proto is not None else deepcopy(header_row._tr)
        tbl._element.append(row)
        rwrap = tbl.rows[-1]
        # set columns to: Finding#, Compliance/summary, Audit det, QC det, CAP det
        _set_cell_text(rwrap.cells[0], "—")
        _set_cell_text(rwrap.cells[1], "No MDL-relevant findings identified for this program.")
        _set_cell_text(rwrap.cells[2], "—")
        _set_cell_text(rwrap.cells[3], "—")
        _set_cell_text(rwrap.cells[4], "—")
        return

    for f in findings:
        row = deepcopy(proto) if proto is not None else deepcopy(header_row._tr)
        tbl._element.append(row)
        rwrap = tbl.rows[-1]
        _set_cell_text(rwrap.cells[0], f.get("finding_id",""))
        # The combined “Compliance Type – Audit Finding Summary” column:
        ctype = f.get("compliance_type","")
        csum  = f.get("summary","")
        if ctype and csum:
            combo = f"{ctype} — {csum}"
        else:
            combo = ctype or csum or ""
        _set_cell_text(rwrap.cells[1], combo)
        _set_cell_text(rwrap.cells[2], f.get("audit_determination","Sustained"))
        _set_cell_text(rwrap.cells[3], f.get("questioned_cost_determination","None"))
        _set_cell_text(rwrap.cells[4], f.get("cap_determination","Accepted"))

# ------------------------------------------------------------------------------
# Treasury-style MDL model + HTML (for preview/alt)
# ------------------------------------------------------------------------------
def summarize_finding_text(raw: str, max_chars: int = 1000) -> str:
    if not raw:
        return ""
    text = re.sub(r"\s+", " ", raw).strip()
    parts = re.split(r"(?<=[.?!])\s+", text)
    picked = []
    for p in parts:
        if len(picked) >= 3:
            break
        if re.search(r"\b(Assistance Listing|Award Period|Federal Program|Identification Number|CFDA)\b", p, re.I):
            continue
        picked.append(p)
    out = " ".join(picked) or text
    return _short(out, max_chars)

def format_letter_date(date_iso: Optional[str] = None) -> Tuple[str, str]:
    dt = datetime.fromisoformat(date_iso) if date_iso else datetime.utcnow()
    return dt.strftime("%Y-%m-%d"), dt.strftime("%B %d, %Y")

def render_mdl_html(model: Dict[str, Any]) -> str:
    letter_date_iso = model.get("letter_date_iso")
    _, letter_date_long = format_letter_date(letter_date_iso)

    auditee_name = model.get("auditee_name", "Recipient")
    ein = model.get("ein", "")
    address_lines = model.get("address_lines", [])
    attention_line = model.get("attention_line")
    period_end_text = model.get("period_end_text", str(model.get("audit_year", "")))
    include_no_qc_line = model.get("include_no_qc_line", True)

    address_block = "<br>".join(htmlmod.escape(x) for x in address_lines) if address_lines else ""
    attention_block = f"<p><strong>{htmlmod.escape(attention_line)}</strong></p>" if attention_line else ""

    def _render_program_table(p: Dict[str, Any]) -> str:
        rows_html = []
        for f in p.get("findings", []):
            rows_html.append(f"""
              <tr>
                <td>{htmlmod.escape(f.get('finding_id',''))}</td>
                <td>{htmlmod.escape(f.get('compliance_type',''))}</td>
                <td>{htmlmod.escape(f.get('summary',''))}</td>
                <td>{htmlmod.escape(f.get('audit_determination',''))}</td>
                <td>{htmlmod.escape(f.get('questioned_cost_determination',''))}</td>
                <td>{htmlmod.escape(f.get('cap_determination',''))}</td>
              </tr>
            """)
        if not rows_html:
            rows_html.append("<tr><td colspan=\"6\"><em>No MDL-relevant findings identified for this program.</em></td></tr>")
        table = f"""
          <h3>Assistance Listing Number/Program Name: {htmlmod.escape(p.get('assistance_listing','Unknown'))} / {htmlmod.escape(p.get('program_name','Unknown'))}</h3>
          <table border="1" cellspacing="0" cellpadding="6" style="border-collapse:collapse; width:100%; font-size:10.5pt;">
            <tr>
              <th>Audit<br>Finding #</th>
              <th>Compliance Type -<br>Audit Finding</th>
              <th>Summary</th>
              <th>Audit Finding<br>Determination</th>
              <th>Questioned Cost<br>Determination</th>
              <th>CAP<br>Determination</th>
            </tr>
            {''.join(rows_html)}
          </table>
        """
        cap_blocks = []
        for f in p.get("findings", []):
            cap_text = f.get("cap_text")
            if cap_text:
                cap_blocks.append(f"""
                  <h4>Corrective Action Plan – {htmlmod.escape(f.get('finding_id',''))}</h4>
                  <p>{htmlmod.escape(cap_text)}</p>
                """)
        return table + ("\n".join(cap_blocks) if cap_blocks else "")

    programs = model.get("programs", [])
    programs_html = "\n".join(_render_program_table(p) for p in programs) if programs else "<p><em>No MDL-relevant findings identified per FAC records.</em></p>"

    not_sustained_notes = model.get("not_sustained_notes", [])
    not_sustained_html = ""
    if not_sustained_notes:
        notes_paras = "\n".join(f"<p>{htmlmod.escape(n)}</p>" for n in not_sustained_notes if n)
        not_sustained_html = f"<h3>FINDINGS NOT SUSTAINED</h3>\n{notes_paras}"

    chunks = []
    chunks.append(f'<p style="text-align:right; margin:0 0 12pt 0;">{htmlmod.escape(letter_date_long)}</p>')
    chunks.append("""
      <p style="margin:0 0 12pt 0;">
        <strong>DEPARTMENT OF THE TREASURY</strong><br>
        WASHINGTON, D.C.
      </p>
    """)
    chunks.append(f"""
      <p style="margin:0 0 12pt 0;">
        <strong>{htmlmod.escape(auditee_name)}</strong><br>
        EIN: {htmlmod.escape(ein)}<br>
      </p>
    """)
    if attention_block:
        chunks.append(attention_block)
    chunks.append(f"""
      <p style="margin:12pt 0 12pt 0;">
        <strong>Subject:</strong> U.S. Department of the Treasury’s Management Decision Letter (MDL) for Single Audit Report for the period ending on {htmlmod.escape(period_end_text)}
      </p>
    """)
    chunks.append("""
      <p>
        In accordance with 2 C.F.R. § 200.521(b), the U.S. Department of the Treasury (Treasury)
        is required to issue a management decision for single audit findings pertaining to awards under
        Treasury’s programs. Treasury’s review as part of its responsibilities under 2 C.F.R § 200.513(c)
        includes an assessment of Treasury’s award recipients’ single audit findings, corrective action plans (CAPs),
        and questioned costs, if any.
      </p>
    """)
    chunks.append(f"""
      <p>
        Treasury has reviewed the single audit report for {htmlmod.escape(auditee_name)}.
        Treasury has made the following determinations regarding the audit finding(s) and CAP(s) listed below.
      </p>
    """)
    if model.get("include_no_qc_line", True):
        chunks.append("<p>No questioned costs are included in this single audit report.</p>")

    chunks.append(programs_html)
    if not_sustained_html:
        chunks.append(not_sustained_html)

    chunks.append("""
      <p>
        Please note, the corrective action(s) are subject to review during the recipient’s next annual single audit
        or program-specific audit, as applicable, to determine adequacy. If the same audit finding(s) appear in a future single
        audit report for this recipient, its current or future award funding under Treasury’s programs may be adversely impacted.
      </p>
      <p>
        The recipient may appeal Treasury’s decision for the audit finding(s) listed above. A written appeal must be submitted within
        30 calendar days of the date of this management decision letter to Treasury via email at
        <a href="mailto:ORP_SingleAudits@treasury.gov">ORP_SingleAudits@treasury.gov</a>.
        The appeal must include: (1) the specific reasons for disputing Treasury’s determination; (2) relevant documentation
        to support the recipient’s position; (3) an alternative course of action with an anticipated completion date of the action; and
        (4) the contact information of the managing official responsible for implementing the proposed alternative course of action.
      </p>
      <p>For questions regarding the audit finding(s), please email us at <a href="mailto:ORP_SingleAudits@treasury.gov">ORP_SingleAudits@treasury.gov</a>. Thank you.</p>
      <p style="margin-top:18pt;">Sincerely,<br><br>
      Audit and Compliance Resolution Team<br>
      Office of Capital Access<br>
      U.S. Department of the Treasury</p>
    """)

    return f'<div style="font-family: Calibri, Arial, sans-serif; font-size:11pt; line-height:1.4;">{"".join(chunks)}</div>'

def build_mdl_model_from_fac(
    *,
    auditee_name: str,
    ein: str,
    audit_year: int,
    fac_general: List[Dict[str, Any]],
    fac_findings: List[Dict[str, Any]],
    fac_findings_text: List[Dict[str, Any]],
    fac_caps: List[Dict[str, Any]],
    federal_awards: List[Dict[str, Any]],
    # Optional enrichments / knobs
    period_end_text: Optional[str] = None,
    address_lines: Optional[List[str]] = None,
    attention_line: Optional[str] = None,
    only_flagged: bool = False,
    max_refs: int = 10,
    auto_cap_determination: bool = True,
    # NEW knobs (AnythingLLM or your routes may pass these)
    include_no_qc_line: bool = False,
    include_no_cap_line: bool = False,
    treasury_listings: Optional[List[str]] = None,
    **_  # swallow unknown kwargs
) -> Dict[str, Any]:
    # Map award_reference -> program meta (derive ALN from name)
    award2meta: Dict[str, Dict[str, str]] = {}
    for a in federal_awards or []:
        ref = a.get("award_reference")
        pname = (a.get("federal_program_name") or "").strip()
        if ref:
            award2meta[ref] = {
                "program_name": pname or "Unknown Program",
                "assistance_listing": _derive_assistance_listing(pname),
                "program_acronym": "",  # placeholder if you later map acronyms
            }

    # Normalize narrative & CAP by finding ref
    norm = lambda s: re.sub(r"\s+", "", (s or "")).upper()
    text_by_ref = { norm(t.get("finding_ref_number")): (t.get("finding_text") or "").strip()
                    for t in (fac_findings_text or []) }
    cap_by_ref  = { norm(c.get("finding_ref_number")): (c.get("planned_action") or "").strip()
                    for c in (fac_caps or []) }

    def _is_flagged(f: dict) -> bool:
        return any([
            f.get("is_material_weakness") is True,
            f.get("is_significant_deficiency") is True,
            f.get("is_questioned_costs") is True,
            f.get("is_modified_opinion") is True,
            f.get("is_other_findings") is True,
            f.get("is_other_matters") is True,
            f.get("is_repeat_finding") is True,
        ])

    # Seed refs from /findings (respect only_flagged)
    base_refs = []
    for f in fac_findings or []:
        if only_flagged and not _is_flagged(f):
            continue
        r = f.get("reference_number")
        if r:
            base_refs.append(r)

    # Fallback: from narrative rows
    if not base_refs and fac_findings_text:
        base_refs = [t.get("finding_ref_number") for t in fac_findings_text if t.get("finding_ref_number")]

    # De-dupe + cap
    seen = set()
    norm_refs: List[Tuple[str, str]] = []
    for r in base_refs:
        if not r:
            continue
        k = norm(r)
        if k not in seen:
            seen.add(k)
            norm_refs.append((r, k))
    norm_refs = norm_refs[: max_refs or 10]
    keep_keys = {k for _, k in norm_refs}

    # Group by award_reference
    programs_map: Dict[str, Dict[str, Any]] = {}
    for f in fac_findings or []:
        r = f.get("reference_number")
        if not r:
            continue
        k = norm(r)
        if k not in keep_keys:
            continue

        award_ref = f.get("award_reference") or "UNKNOWN"
        meta = award2meta.get(award_ref, {})
        group = programs_map.setdefault(award_ref, {
            "assistance_listing": meta.get("assistance_listing", "Unknown"),
            "program_name": meta.get("program_name", "Unknown Program"),
            "program_acronym": meta.get("program_acronym", ""),
            "findings": []
        })

        narrative = text_by_ref.get(k, "")
        repeat_of = _extract_repeat_of(narrative)

        # allow multiple compliance types (split on common separators)
        ctype_raw = f.get("type_requirement") or ""
        compliance_types = [s.strip() for s in re.split(r"(?:\band\b|;|,|/)", ctype_raw, flags=re.I) if s.strip()]

        summary  = summarize_finding_text(narrative)
        cap_text = cap_by_ref.get(k)

        qcost_det = "No questioned costs identified" if include_no_qc_line else "None"
        cap_det   = ("Accepted" if (auto_cap_determination and cap_text)
                    else ("No CAP required" if include_no_cap_line else "Not Applicable"))

        group["findings"].append({
            "finding_id": f.get("reference_number") or "",
            "is_repeat_finding": bool(f.get("is_repeat_finding")),
            "repeat_of": repeat_of,
            "compliance_type": ctype_raw,
            "compliance_types": compliance_types,
            "summary": summary,
            "audit_determination": "Sustained",
            "questioned_cost_determination": qcost_det,
            "disallowed_cost_determination": "None",
            "cap_determination": cap_det,
            "cap_text": cap_text,
        })

    # Catch-all if we only had narrative rows
    if not programs_map and norm_refs:
        catchall = {
            "assistance_listing": "Unknown",
            "program_name": "Unknown Program",
            "program_acronym": "",
            "findings": []
        }
        for orig, key in norm_refs:
            narrative = text_by_ref.get(key, "")
            repeat_of = _extract_repeat_of(narrative)
            summary = summarize_finding_text(narrative)
            cap_text = cap_by_ref.get(key)
            qcost_det = "No questioned costs identified" if include_no_qc_line else "None"
            cap_det   = ("Accepted" if (auto_cap_determination and cap_text)
                        else ("No CAP required" if include_no_cap_line else "Not Applicable"))
            catchall["findings"].append({
                "finding_id": orig,
                "is_repeat_finding": False,
                "repeat_of": repeat_of,
                "compliance_type": "",
                "compliance_types": [],
                "summary": summary,
                "audit_determination": "Sustained",
                "questioned_cost_determination": qcost_det,
                "disallowed_cost_determination": "None",
                "cap_determination": cap_det,
                "cap_text": cap_text,
            })
        programs_map["UNKNOWN"] = catchall

    # Optional filter by Treasury ALNs
    programs = list(programs_map.values())
    if treasury_listings:
        wanted = set(treasury_listings)
        programs = [p for p in programs if _derive_assistance_listing(p.get("program_name","")) in wanted
                    or p.get("assistance_listing") in wanted]

    # Sort by ALN
    programs_sorted = sorted(programs, key=lambda p: _aln_sort_key(p.get("assistance_listing")))

    letter_date_iso, _ = format_letter_date(None)
    model = {
        "letter_date_iso": letter_date_iso,
        "auditee_name": auditee_name,
        "ein": f"{ein[:2]}-{ein[2:]}" if ein and ein.isdigit() and len(ein) == 9 else ein,
        "address_lines": address_lines or [],
        "attention_line": attention_line or "",
        "period_end_text": period_end_text or f"June 30, {audit_year}",
        "audit_year": audit_year,
        "programs": programs_sorted,
        "not_sustained_notes": [],
        "include_no_qc_line": include_no_qc_line,
    }
    return model

# ------------------------------------------------------------------------------
# DOCX Template filling (Treasury template EXACT layout)
# ------------------------------------------------------------------------------
def _replace_text_in_paragraph_runsafe(paragraph, mapping: Dict[str, str]):
    """
    Robust placeholder replacement for {{TOKENS}} even when split across runs.
    We rebuild the paragraph text once, then set as a single run (formatting loss only for tokens).
    """
    if not paragraph.runs:
        text = paragraph.text
        for k, v in mapping.items():
            text = text.replace(k, v)
        paragraph.text = text
        return

    full_text = "".join(run.text for run in paragraph.runs)
    changed = False
    for k, v in mapping.items():
        if k in full_text:
            full_text = full_text.replace(k, v)
            changed = True
    if changed:
        for _ in range(len(paragraph.runs)-1, -1, -1):
            paragraph.runs[_].clear() if hasattr(paragraph.runs[_], "clear") else None
            paragraph.runs[_].text = ""
        paragraph.runs[0].text = full_text

def _replace_placeholders_everywhere(doc: Document, mapping: Dict[str, str]):
    # body paragraphs
    for p in doc.paragraphs:
        _replace_text_in_paragraph_runsafe(p, mapping)
    # tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_text_in_paragraph_runsafe(p, mapping)

def _find_program_anchor_and_table(doc: Document) -> Tuple[Optional[Any], Optional[Any]]:
    """
    Find the anchor paragraph that begins the program block and the table that follows.
    We look for the literal label the client template uses.
    """
    anchor = None
    table = None
    label = "Assistance Listing Number/Program Name"
    for idx, p in enumerate(doc.paragraphs):
        if label.lower() in p.text.lower():
            anchor = p
            # find the next table in the document (python-docx keeps tables also in .tables order,
            # but we need the first table after this paragraph; scan the body)
            # A portable way: walk the document element order:
            # However, python-docx doesn't give us a straightforward block iterator with ordering.
            # We'll approximate: return the first table in doc.tables that has at least 1 header row (6 columns typical).
            # If multiple, the first is usually correct in the provided template.
            table = doc.tables[0] if doc.tables else None
            break
    # If we couldn't find by label, but there is at least one table, treat that as prototype
    if not anchor and doc.paragraphs:
        anchor = doc.paragraphs[-1]  # append copies to end if needed
    if not table and doc.tables:
        table = doc.tables[0]
    return anchor, table

def _delete_all_data_rows_keep_header(table):
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

def _add_finding_row(table, cells_text: List[str]):
    row = table.add_row()
    for i, txt in enumerate(cells_text):
        cell = row.cells[i] if i < len(row.cells) else None
        if not cell:
            continue
        # clear any existing para(s)
        for p in cell.paragraphs:
            if hasattr(p, "clear"):
                p.clear()
            else:
                for r in p.runs:
                    r.text = ""
        # add with line breaks
        lines = (txt or "").split("\n")
        if not cell.paragraphs:
            para = cell.add_paragraph()
        else:
            para = cell.paragraphs[0]
        for j, line in enumerate(lines):
            if j > 0:
                para.add_run().add_break()
            para.add_run(line)
        _tight_paragraph(para)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

def _insert_paragraph_after(ref_paragraph, text: str = ""):
    new_p = ref_paragraph._element.addnext(OxmlElement("w:p"))
    p = ref_paragraph._parent.add_paragraph()
    # python-docx API does not support inserting at arbitrary position simply,
    # but the above ensures a following paragraph; fallback: append at end
    p.add_run(text)
    return p

def _clone_table_after(doc: Document, ref_table):
    # Append a new table with same column count; copy header row cells' text
    cols = len(ref_table.columns)
    new_tbl = doc.add_table(rows=1, cols=cols)
    new_tbl.style = ref_table.style
    # copy header text (row 0)
    for ci in range(cols):
        src = ref_table.cell(0, ci)
        dst = new_tbl.cell(0, ci)
        # copy text
        for p in dst.paragraphs:
            if hasattr(p, "clear"):
                p.clear()
            else:
                for r in p.runs:
                    r.text = ""
        txt = "\n".join([r.text for r in src.paragraphs[0].runs]) if src.paragraphs else src.text
        dst.paragraphs[0].add_run(src.paragraphs[0].text if src.paragraphs else txt)
        _shade_cell(dst, "E7E6E6")
        for r in dst.paragraphs[0].runs:
            r.bold = True
        dst.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return new_tbl

def _fill_one_program_block(anchor_par, table, *, assistance_listing: str, program_name: str, program_acronym: str, findings: List[Dict[str, Any]]):
    # Anchor text with optional extra ALNs (from findings[*].also_programs, if present)
    anchor_prefix = "Assistance Listing Number/Program Name: "
    header_line = f"{assistance_listing}/ {program_name}".strip()
    if program_acronym:
        header_line += f" ({program_acronym})"

    extra_lines: List[str] = []
    for f in findings or []:
        for ap in f.get("also_programs", []) or []:
            line = f"{ap.get('assistance_listing','')}/ {ap.get('program_name','')}".strip()
            if ap.get("program_acronym"):
                line += f" ({ap['program_acronym']})"
            if line and line not in extra_lines:
                extra_lines.append(line)

    # Set anchor paragraph text; put extra ALNs on subsequent lines
    try:
        # clear paragraph content
        for r in anchor_par.runs:
            r.text = ""
    except Exception:
        pass
    anchor_par.text = f"{anchor_prefix}{header_line}" + (("\n" + "\n".join(extra_lines)) if extra_lines else "")

    # Reset table to header only
    _delete_all_data_rows_keep_header(table)

    for f in findings or []:
        finding_no = f.get("finding_id", "") or ""
        repeat_of = f.get("repeat_of")
        if repeat_of:
            first_cell = f"{finding_no}\nRepeat of {repeat_of}"
        elif f.get("is_repeat_finding"):
            first_cell = f"{finding_no}\nRepeat finding"
        else:
            first_cell = finding_no

        comp_lines = f.get("compliance_types") or []
        if not comp_lines and f.get("compliance_type"):
            comp_lines = [f.get("compliance_type")]
        comp_block = "\n".join([c for c in comp_lines if c]) if comp_lines else ""
        summary = f.get("summary", "") or ""
        if comp_block and summary:
            comp_plus_summary = comp_block + "\n" + summary
        else:
            comp_plus_summary = comp_block or summary

        det_cell = "Sustained"
        qc_cell  = "Questioned Cost:\nNone\nDisallowed Cost:\nNone"
        # If you want the phrasing “No questioned costs identified” instead, adjust above build_mdl_model_from_fac to pass this wording here.
        cap_cell = "Accepted"

        _add_finding_row(table, [first_cell, comp_plus_summary, det_cell, qc_cell, cap_cell])

def _render_docx_from_template(model: Dict[str, Any], template_path: str) -> bytes:
    if not os.path.isfile(template_path):
        raise HTTPException(400, f"Treasury template not found at: {template_path}")

    doc = Document(template_path)

    # 1) Replace simple placeholders if the template contains them
    #    (kept generic to avoid formatting drift in other parts)
    _, letter_date_long = format_letter_date(model.get("letter_date_iso"))
    placeholders = {
        "{{DATE}}": letter_date_long,
        "{{AUDITEE_NAME}}": model.get("auditee_name", ""),
        "{{EIN}}": model.get("ein", ""),
        "{{FY_END}}": model.get("period_end_text", ""),
        "{{RECIPIENT_NAME}}": model.get("auditee_name", ""),
        "{{ATTENTION_LINE}}": model.get("attention_line", ""),
        # add more if template contains them; unknown tokens are harmless
    }
    _replace_placeholders_everywhere(doc, placeholders)

    # 2) Find the anchor paragraph + prototype table (first “Assistance Listing Number/Program Name” section)
    anchor, proto_table = _find_program_anchor_and_table(doc)
    if not proto_table:
        # Fall back: generate a basic HTML letter (rare edge) to avoid total failure
        html_str = render_mdl_html(model)
        return html_to_docx_bytes(html_str, force_basic=True)

    programs = model.get("programs", [])
    if not programs:
        # No findings/programs – leave template table as-is but ensure at least one "None" row.
        _delete_all_data_rows_keep_header(proto_table)
        _add_finding_row(proto_table, [
            "—", "No MDL-relevant findings identified for this program.", "—", "—", "—"
        ])
        bio = BytesIO(); doc.save(bio); return bio.getvalue()

    # 3) Fill the first program into the prototype
    first = programs[0]
    _fill_one_program_block(
        anchor, proto_table,
        assistance_listing = first.get("assistance_listing", "Unknown"),
        program_name = first.get("program_name", "Unknown Program"),
        program_acronym = first.get("program_acronym", ""),
        findings = first.get("findings", []),
    )

    # 4) For remaining programs, clone a new header+table block each time and fill
    for p in programs[1:]:
        # Create a spacer paragraph
        doc.add_paragraph()  # visual separation

        # Add a fresh anchor header paragraph
        anchor2 = doc.add_paragraph("Assistance Listing Number/Program Name: ")
        # Clone a new table with header row like prototype
        new_tbl = _clone_table_after(doc, proto_table)

        _fill_one_program_block(
            anchor2, new_tbl,
            assistance_listing = p.get("assistance_listing", "Unknown"),
            program_name = p.get("program_name", "Unknown Program"),
            program_acronym = p.get("program_acronym", ""),
            findings = p.get("findings", []),
        )

    bio = BytesIO(); doc.save(bio); return bio.getvalue()


from typing import Optional, Any

# Accept both explicit token and common phrasings
PROGRAM_TABLES_ANCHORS = [
    "[[PROGRAM_TABLES]]",
    "Assistance Listing Number/Program Name",
    "Assistance Listing Number / Program Name",
    "Assistance Listing Number – Program Name",
    "Assistance Listing Number — Program Name",
]

def _iter_all_paragraphs(doc):
    """Yield ALL paragraphs: top-level and inside tables (recursively)."""
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        yield from _iter_table_paragraphs(t)

def _iter_table_paragraphs(tbl):
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for inner_tbl in cell.tables:
                yield from _iter_table_paragraphs(inner_tbl)

def _find_program_anchor_paragraph(doc) -> Optional[Any]:
    """Return the first paragraph whose full text contains any anchor (case-insensitive)."""
    anchors = [a.lower() for a in PROGRAM_TABLES_ANCHORS]
    for p in _iter_all_paragraphs(doc):
        txt = (p.text or "").strip().lower()
        if any(a in txt for a in anchors):
            return p
    return None

def _clear_paragraph_text(p):
    for r in p.runs:
        r.text = ""

def _containing_tbl_element(paragraph):
    """If paragraph is inside a table, return the <w:tbl> element; else None."""
    elm = paragraph._p
    while elm is not None:
        if elm.tag.endswith('}tbl'):
            return elm
        elm = elm.getparent()
    return None

def _addnext(anchor_elm, new_obj):
    """Move new_obj immediately after anchor_elm (works for paragraphs or tables)."""
    new_elm = getattr(new_obj, "_p", None) or getattr(new_obj, "_tbl", None)
    anchor_elm.addnext(new_elm)

def _insert_program_sections(doc, anchor_p, programs: list[dict]):
    """
    Insert per-program headings & tables at anchor (if present);
    otherwise append them to the end of the document.
    """
    # Choose an insertion point element: either the paragraph itself or its containing table
    insert_after = anchor_p._p if anchor_p is not None else None
    tbl_container = _containing_tbl_element(anchor_p) if anchor_p is not None else None
    if tbl_container is not None:
        insert_after = tbl_container

    if anchor_p is not None:
        _clear_paragraph_text(anchor_p)  # hide the token/anchor text

    # Helper to insert a block after `insert_after`, then update `insert_after`
    def _append_block(new_block):
        nonlocal insert_after
        if insert_after is None:
            # no anchor: just let python-docx append to end
            return
        _addnext(insert_after, new_block)
        # set new last-inserted element as the next insertion point
        insert_after = getattr(new_block, "_p", None) or getattr(new_block, "_tbl", None)

    # Sort programs by ALN (e.g., 21.023, 21.026, 21.027)
    def _aln_key(p):
        a = p.get("assistance_listing") or ""
        try:
            parts = a.split(".")
            return (int(parts[0]), int(parts[1])) if len(parts) == 2 else (999, 999)
        except Exception:
            return (999, 999)

    for prog in sorted(programs, key=_aln_key):
        aln = prog.get("assistance_listing") or "Unknown"
        pname = prog.get("program_name") or "Unknown Program"

        # Heading
        heading_txt = f"Assistance Listing Number/Program Name: {aln} / {pname}"
        ph = doc.add_paragraph(heading_txt)
        _append_block(ph)

        # Table (6 columns, header row + N findings)
        f_list = prog.get("findings", []) or []
        rows = 1 + max(1, len(f_list))
        tbl = doc.add_table(rows=rows, cols=6)
        # keep template’s default table style; otherwise uncomment:
        # tbl.style = "Table Grid"

        hdr = tbl.rows[0].cells
        hdr[0].text = "Audit\nFinding #"
        hdr[1].text = "Compliance Type -\nAudit Finding"
        hdr[2].text = "Summary"
        hdr[3].text = "Audit Finding\nDetermination"
        hdr[4].text = "Questioned Cost\nDetermination"
        hdr[5].text = "CAP\nDetermination"

        if not f_list:
            r = tbl.rows[1].cells
            r[0].text = "—"
            r[1].text = "—"
            r[2].text = "No MDL-relevant findings identified for this program."
            r[3].text = "—"
            r[4].text = "—"
            r[5].text = "—"
        else:
            for i, f in enumerate(f_list, start=1):
                c = tbl.rows[i].cells
                c[0].text = f.get("finding_id", "")
                c[1].text = f.get("compliance_type", "")
                c[2].text = f.get("summary", "")
                c[3].text = f.get("audit_determination", "")
                c[4].text = f.get("questioned_cost_determination", "")
                c[5].text = f.get("cap_determination", "")

        _append_block(tbl)
        
def _address_lines_from_req(req) -> List[str]:
    if getattr(req, "address_lines", None):
        return [x for x in req.address_lines if x]
    parts = []
    if req.street_address: parts.append(req.street_address)
    city_state_zip = " ".join([p for p in [req.city and req.city.strip(", "), req.state, req.zip_code] if p])
    if city_state_zip: parts.append(city_state_zip)
    return parts
# ------------------------------------------------------------------------------
# Schemas
# ------------------------------------------------------------------------------
class BuildDocx(BaseModel):
    auditee_name: str
    ein: str
    audit_year: int
    body_html: Optional[str] = None
    body_html_b64: Optional[str] = None
    dest_path: Optional[str] = None
    filename: Optional[str] = None

class BuildFromFAC(BaseModel):
    auditee_name: str
    ein: str
    audit_year: int
    fac_general: List[Dict[str, Any]] = []
    fac_findings: List[Dict[str, Any]] = []
    fac_findings_text: List[Dict[str, Any]] = []
    fac_caps: List[Dict[str, Any]] = []
    federal_awards: List[Dict[str, Any]] = []
    dest_path: Optional[str] = None
    filename: Optional[str] = None

class BuildByReport(BaseModel):
    auditee_name: str
    ein: str
    audit_year: int
    report_id: str
    dest_path: Optional[str] = None
    only_flagged: bool = False
    max_refs: int = 15
    include_awards: bool = True     # turn ON to get program names
    treasury_listings: Optional[List[str]] = None  # e.g. ["21.027","21.023"]

class BuildByReportTemplated(BaseModel):
    auditee_name: str
    ein: str
    audit_year: int
    report_id: str
    dest_path: Optional[str] = None
    only_flagged: bool = False
    max_refs: int = 15
    include_awards: bool = True
    treasury_listings: Optional[List[str]] = None
    # Header enrichments (optional – pass from SF-SAC if you have it)
    fy_end_text: Optional[str] = None
    auditor_name: Optional[str] = None
    recipient_name: Optional[str] = None
    address_lines: Optional[List[str]] = None        # [street, city, state, zip]
    attention_line: Optional[str] = None             # "Jane Doe, Title"
    # Template path
    template_path: Optional[str] = None

class BuildByReportTemplated(BuildByReport):
    auditor_name: Optional[str] = None
    fy_end_text: Optional[str] = None
    recipient_name: Optional[str] = None

    # Either pass address_lines directly OR use the pieces below:
    address_lines: Optional[List[str]] = None
    street_address: Optional[str] = None
    city: Optional[str] = None
    state: Optional[str] = None
    zip_code: Optional[str] = None

    poc_name: Optional[str] = None
    poc_title: Optional[str] = None

    template_path: Optional[str] = None
# ------------------------------------------------------------------------------
# Routes
# ------------------------------------------------------------------------------
@app.get("/healthz")
def healthz():
    return {"ok": True, "time": datetime.utcnow().isoformat()}

@app.post("/echo")
def echo(payload: Dict[str, Any]):
    return {"received": payload, "ts": datetime.utcnow().isoformat()}

@app.get("/debug/env")
def debug_env():
    key = os.getenv("FAC_API_KEY") or ""
    masked = (key[:4] + "…" + key[-2:]) if key else None
    return {"fac_api_key_present": bool(key), "fac_api_key_masked": masked}

@app.get("/debug/storage")
def debug_storage():
    info = _parse_conn_str(AZURE_CONN_STR) if AZURE_CONN_STR else {}
    return {
        "using_storage": bool(AZURE_CONN_STR),
        "account": info.get("AccountName"),
        "blob_endpoint": info.get("BlobEndpoint")
    }

@app.get("/debug/sas")
def debug_sas():
    if not AZURE_CONN_STR:
        raise HTTPException(400, "Set AZURE_STORAGE_CONNECTION_STRING to test SAS.")
    url = upload_and_sas(AZURE_CONTAINER, "debug/hello.txt", b"hi", ttl_minutes=5)
    return {"url": url}

@app.post("/build-docx-demo")
def build_docx_demo():
    document = Document()
    document.add_heading("Hello from the DOCX demo ✅", level=1)
    document.add_paragraph("If you can read this, your write/upload pipeline is good.")
    bio = BytesIO(); document.save(bio); data = bio.getvalue()
    blob_name = "demo/hello.docx"
    url = upload_and_sas(AZURE_CONTAINER, blob_name, data) if AZURE_CONN_STR else save_local_and_url(blob_name, data)
    return {"url": url, "size_bytes": len(data)}

@app.post("/build-docx")
def build_docx(req: BuildDocx):
    html_str = (req.body_html or "").strip()
    if (not html_str) and req.body_html_b64:
        try:
            html_str = base64.b64decode(req.body_html_b64).decode("utf-8", errors="ignore")
        except Exception:
            raise HTTPException(400, "Invalid base64 in body_html_b64")
    if not html_str:
        raise HTTPException(400, "body_html (or body_html_b64) is required")

    data = html_to_docx_bytes(html_str, force_basic=False)
    folder = (req.dest_path or "").lstrip("/")
    base = req.filename or f"MDL-{sanitize(req.auditee_name)}-{sanitize(req.ein)}-{req.audit_year}.docx"
    blob_name = f"{folder}{base}" if folder else base
    url = upload_and_sas(AZURE_CONTAINER, blob_name, data) if AZURE_CONN_STR else save_local_and_url(blob_name, data)
    return {"url": url, "blob_path": f"{AZURE_CONTAINER}/{blob_name}", "size_bytes": len(data)}

# ---------- Treasury-style: arrays -> DOCX via TEMPLATE ----------
@app.post("/build-docx-from-fac")
def build_docx_from_fac(req: BuildFromFAC):
    model = build_mdl_model_from_fac(
        auditee_name=req.auditee_name,
        ein=req.ein,
        audit_year=req.audit_year,
        fac_general=req.fac_general,
        fac_findings=req.fac_findings,
        fac_findings_text=req.fac_findings_text,
        fac_caps=req.fac_caps,
        federal_awards=req.federal_awards,
        only_flagged=False,
        max_refs=25,
        include_no_qc_line=True,
    )
    # Use the exact Treasury template if available; fallback to HTML
    if os.path.isfile(MDL_TEMPLATE_PATH):
        data = _render_docx_from_template(model, MDL_TEMPLATE_PATH)
    else:
        html_str = render_mdl_html(model)
        data = html_to_docx_bytes(html_str, force_basic=True)

    folder = (req.dest_path or "").lstrip("/")
    base = req.filename or f"MDL-Preview-{sanitize(req.auditee_name)}-{sanitize(req.ein)}-{req.audit_year}.docx"
    blob_name = f"{folder}{base}" if folder else base
    url = upload_and_sas(AZURE_CONTAINER, blob_name, data) if AZURE_CONN_STR else save_local_and_url(blob_name, data)
    return {"url": url, "blob_path": f"{AZURE_CONTAINER}/{blob_name}", "size_bytes": len(data)}

# ---------- Treasury-style: report_id -> FAC -> arrays -> DOCX via TEMPLATE ----------
@app.post("/build-docx-by-report")
def build_docx_by_report(req: BuildByReport):
    # 1) Minimal general (accepted date, etc.)
    fac_general = _fac_get("general", {
        "report_id": f"eq.{req.report_id}",
        "select": "report_id,fac_accepted_date",
        "limit": 1
    })

    # 2) Findings (optionally only flagged)
    findings_params = {
        "report_id": f"eq.{req.report_id}",
        "select": "reference_number,award_reference,type_requirement,"
                  "is_material_weakness,is_significant_deficiency,is_questioned_costs,"
                  "is_modified_opinion,is_other_findings,is_other_matters,is_repeat_finding",
        "order": "reference_number.asc",
        "limit": str(req.max_refs)
    }
    if req.only_flagged:
        flagged = [
            "is_material_weakness","is_significant_deficiency","is_questioned_costs",
            "is_modified_opinion","is_other_findings","is_other_matters","is_repeat_finding"
        ]
        findings_params["or"] = "(" + ",".join([f"{f}.eq.true" for f in flagged]) + ")"
    fac_findings = _fac_get("findings", findings_params)

    # 3) selected refs
    refs = [row.get("reference_number") for row in fac_findings if row.get("reference_number")]
    refs = refs[: req.max_refs]

    # 4) narrative & CAP for those refs only
    if refs:
        fac_findings_text = _fac_get("findings_text", {
            "report_id": f"eq.{req.report_id}",
            "select": "finding_ref_number,finding_text",
            "order": "finding_ref_number.asc",
            "limit": str(len(refs)),
            "or": _or_param("finding_ref_number", refs)
        })
        fac_caps = _fac_get("corrective_action_plans", {
            "report_id": f"eq.{req.report_id}",
            "select": "finding_ref_number,planned_action",
            "order": "finding_ref_number.asc",
            "limit": str(len(refs)),
            "or": _or_param("finding_ref_number", refs)
        })
    else:
        fac_findings_text, fac_caps = [], []

    # 5) Awards (program names) – DO NOT request non-existent columns
    federal_awards = []
    if req.include_awards:
        federal_awards = _fac_get("federal_awards", {
            "report_id": f"eq.{req.report_id}",
            "select": "award_reference,federal_program_name",
            "order": "award_reference.asc",
            "limit": "200"
        })

    # 6) Build model and render using template
    model = build_mdl_model_from_fac(
        auditee_name=req.auditee_name,
        ein=req.ein,
        audit_year=req.audit_year,
        fac_general=fac_general,
        fac_findings=fac_findings,
        fac_findings_text=fac_findings_text,
        fac_caps=fac_caps,
        federal_awards=federal_awards,
        only_flagged=req.only_flagged,
        max_refs=req.max_refs,
        include_no_qc_line=True,
        treasury_listings=req.treasury_listings or TREASURY_DEFAULT_LISTINGS,
    )

    if os.path.isfile(MDL_TEMPLATE_PATH):
        data = _render_docx_from_template(model, MDL_TEMPLATE_PATH)
    else:
        html_str = render_mdl_html(model)
        data = html_to_docx_bytes(html_str, force_basic=True)

    folder = (req.dest_path or "").lstrip("/")
    base = f"MDL-Preview-{sanitize(req.auditee_name)}-{sanitize(req.ein)}-{req.audit_year}.docx"
    blob_name = f"{folder}{base}" if folder else base
    url = upload_and_sas(AZURE_CONTAINER, blob_name, data) if AZURE_CONN_STR else save_local_and_url(blob_name, data)
    return {"url": url, "blob_path": f"{AZURE_CONTAINER}/{blob_name}", "size_bytes": len(data)}

def build_docx_from_template(mdl_model: Dict[str, Any], *, template_path: str) -> bytes:
    """
    Fill the client's DOCX template exactly:
      - Replace header placeholders everywhere ([Recipient Name], [EIN], [Fiscal Year End Date], etc.)
      - For each program in mdl_model['programs'], duplicate the ALN section + table and fill rows
      - Preserve all formatting in the template
    """
    if not os.path.isfile(template_path):
        raise HTTPException(400, f"Template not found at {template_path}")

    doc = Document(template_path)

    # 1) Header replacements
    recip = mdl_model.get("auditee_name") or ""
    ein   = mdl_model.get("ein") or ""
    fyend = mdl_model.get("period_end_text") or ""
    attention = mdl_model.get("attention_line") or ""
    addr_lines = mdl_model.get("address_lines") or []
    street = addr_lines[0] if len(addr_lines) > 0 else ""
    city   = addr_lines[1] if len(addr_lines) > 1 else ""
    state  = addr_lines[2] if len(addr_lines) > 2 else ""
    zipc   = addr_lines[3] if len(addr_lines) > 3 else ""
    auditor= mdl_model.get("auditor_name") or ""  # you can inject this from SF-SAC if/when available

    header_map = {
        "[Recipient Name]": recip,
        "[EIN]": ein,
        "[Street Address]": street,
        "[City]": city,
        "[State]": state,
        "[Zip Code]": zipc,
        "[Point of Contact]": attention.split(",")[0] if attention else "",
        "[Point of Contact Title]": (attention.split(",")[1].strip() if ("," in attention) else ""),
        "[Fiscal Year End Date]": fyend,
        "[Auditor Name]": auditor,
        # Date line: if your template keeps "Date XX, 2025", you can replace "XX, 2025" too:
        "Date XX, 2025": datetime.utcnow().strftime("Date %B %d, %Y"),
    }
    _docx_global_replace(doc, header_map)

    # 2) Find the first ALN/Program + table block (prototype)
    aln_hdr, aln_line, table, start_i, end_i = _find_program_section(doc)
    if aln_hdr is None:
        raise HTTPException(500, "Template does not contain the expected 'Assistance Listing Number/Program Name' section.")

    programs = mdl_model.get("programs") or []

    if not programs:
        # Fill prototype with a single 'no relevant findings' section
        _fill_program_block(doc, aln_hdr, aln_line, table, {
            "assistance_listing": "Unknown",
            "program_name": "Unknown Program",
            "program_acronym": "",
            "findings": []
        })
    else:
        # Fill the prototype with the first program
        first = programs[0]
        first = {
            **first,
            "program_acronym": first.get("program_acronym") or PROGRAM_ACRONYMS.get(first.get("assistance_listing"), "")
        }
        _fill_program_block(doc, aln_hdr, aln_line, table, first)

        # For each remaining program, clone and insert a new section
        for p in programs[1:]:
            p = {**p, "program_acronym": p.get("program_acronym") or PROGRAM_ACRONYMS.get(p.get("assistance_listing"), "")}
            # clone paragraph + paragraph + table
            new_hdr = deepcopy(aln_hdr._element)
            new_line = deepcopy(aln_line._element)
            new_tbl = deepcopy(table._element)

            # insert after current table
            _insert_block_after(doc, table, new_hdr)
            hdr_para = Paragraph(new_hdr, doc)
            _insert_block_after(doc, hdr_para, new_line)
            line_para = Paragraph(new_line, doc)
            _insert_block_after(doc, line_para, new_tbl)
            tbl_obj = Table(new_tbl, doc)

            # fill that block
            _fill_program_block(doc, hdr_para, line_para, tbl_obj, p)

            # move the reference table forward for the next insertion point
            table = tbl_obj

    # 3) Save to bytes
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

@app.post("/build-mdl-docx-by-report-templated")
def build_mdl_docx_by_report_templated(req: BuildByReportTemplated):
    # 1) General (you can add more selects if you want SF-SAC dates, etc.)
    fac_general = _fac_get("general", {
        "report_id": f"eq.{req.report_id}",
        "select": "report_id,fac_accepted_date",
        "limit": 1
    })

    # 2) Findings (respect only_flagged)
    findings_params = {
        "report_id": f"eq.{req.report_id}",
        "select": "reference_number,award_reference,type_requirement,"
                  "is_material_weakness,is_significant_deficiency,is_questioned_costs,"
                  "is_modified_opinion,is_other_findings,is_other_matters,is_repeat_finding",
        "order": "reference_number.asc",
        "limit": str(req.max_refs)
    }
    if req.only_flagged:
        flagged = ["is_material_weakness","is_significant_deficiency","is_questioned_costs",
                   "is_modified_opinion","is_other_findings","is_other_matters","is_repeat_finding"]
        findings_params["or"] = "(" + ",".join([f"{f}.eq.true" for f in flagged]) + ")"
    fac_findings = _fac_get("findings", findings_params)

    # 3) refs we actually fetched
    refs = [row.get("reference_number") for row in fac_findings if row.get("reference_number")]
    refs = refs[: req.max_refs]

    # 4) Finding text & CAP for those refs
    if refs:
        fac_findings_text = _fac_get("findings_text", {
            "report_id": f"eq.{req.report_id}",
            "select": "finding_ref_number,finding_text",
            "order": "finding_ref_number.asc",
            "limit": str(len(refs)),
            "or": _or_param("finding_ref_number", refs)
        })
        fac_caps = _fac_get("corrective_action_plans", {
            "report_id": f"eq.{req.report_id}",
            "select": "finding_ref_number,planned_action",
            "order": "finding_ref_number.asc",
            "limit": str(len(refs)),
            "or": _or_param("finding_ref_number", refs)
        })
    else:
        fac_findings_text, fac_caps = [], []

    # 5) Awards for program names
    federal_awards = []
    if req.include_awards:
        federal_awards = _fac_get("federal_awards", {
            "report_id": f"eq.{req.report_id}",
            "select": "award_reference,federal_program_name",
            "order": "award_reference.asc",
            "limit": "200"
        })

    # 6) Build MDL model (your existing function)
    mdl_model = build_mdl_model_from_fac(
        auditee_name=req.auditee_name,
        ein=req.ein,
        audit_year=req.audit_year,
        fac_general=fac_general,
        fac_findings=fac_findings,
        fac_findings_text=fac_findings_text,
        fac_caps=fac_caps,
        federal_awards=federal_awards,
        only_flagged=req.only_flagged,
        max_refs=req.max_refs,
    )

    # Apply client header enrichments if provided
    if req.fy_end_text:      mdl_model["period_end_text"] = req.fy_end_text
    if req.auditor_name:     mdl_model["auditor_name"] = req.auditor_name
    if req.recipient_name:   mdl_model["auditee_name"] = req.recipient_name
    if req.address_lines:    mdl_model["address_lines"] = req.address_lines
    if req.attention_line:   mdl_model["attention_line"] = req.attention_line

    # 7) Render via template
    template_path = req.template_path or os.getenv("MDL_TEMPLATE_PATH")
    if not template_path:
        raise HTTPException(400, "Template path not provided (set template_path or MDL_TEMPLATE_PATH).")
    data = build_docx_from_template(mdl_model, template_path=template_path)

    # 8) Upload
    folder = (req.dest_path or "").lstrip("/")
    base = f"MDL-{sanitize(req.auditee_name)}-{sanitize(req.ein)}-{req.audit_year}.docx"
    blob_name = f"{folder}{base}" if folder else base
    url = upload_and_sas(AZURE_CONTAINER, blob_name, data) if AZURE_CONN_STR else save_local_and_url(blob_name, data)
    return {"url": url, "blob_path": f"{AZURE_CONTAINER}/{blob_name}", "size_bytes": len(data)}

from docx import Document

def render_with_template_and_programs(template_path: str, mdl_model: dict) -> bytes:
    doc = Document(template_path)

    # Replace simple placeholders (run-wise)
    letter_date_iso = mdl_model.get("letter_date_iso")
    _, letter_date_long = format_letter_date(letter_date_iso)

    mapping = {
        "{{DATE_LONG}}": letter_date_long,
        "{{RECIPIENT_NAME}}": mdl_model.get("auditee_name",""),
        "{{EIN}}": mdl_model.get("ein",""),
        "{{FY_END_TEXT}}": mdl_model.get("period_end_text",""),
        "{{POC_NAME}}": mdl_model.get("attention_line",""),
        # add any other placeholders you have in the template
    }
    _replace_placeholders_everywhere(doc, mapping)

    # Insert program tables at anchor (or append if none)
    anchor_p = _find_program_anchor_paragraph(doc)
    _insert_program_sections(doc, anchor_p, mdl_model.get("programs", []))

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

@app.post("/debug/scan-template")
def debug_scan_template(payload: Dict[str,str]):
    path = payload.get("template_path")
    if not path or not os.path.isfile(path):
        raise HTTPException(400, f"Template not found at: {path!r}")

    from docx import Document
    doc = Document(path)
    found = []
    for p in _iter_all_paragraphs(doc):
        txt = (p.text or "").strip()
        if not txt:
            continue
        for a in PROGRAM_TABLES_ANCHORS:
            if a.lower() in txt.lower():
                found.append(txt)

    return {
        "template_path": path,
        "anchors_scanned": PROGRAM_TABLES_ANCHORS,
        "matches_found": found,
        "match_count": len(found),
    }