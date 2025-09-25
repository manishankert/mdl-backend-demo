# main.py
from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional, Tuple
from io import BytesIO
from urllib.parse import quote
import os, re, base64, html, requests

# DOCX / HTML
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
from html2docx import HTML2Docx

# Azure
from azure.storage.blob import BlobServiceClient, BlobSasPermissions, generate_blob_sas

import copy
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ------------------------------------------------------------------------------
# FastAPI app
# ------------------------------------------------------------------------------
app = FastAPI(title="MDL DOCX Builder")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ------------------------------------------------------------------------------
# Environment
# ------------------------------------------------------------------------------
FAC_BASE = os.getenv("FAC_API_BASE", "https://api.fac.gov")
FAC_KEY  = os.getenv("FAC_API_KEY")

AZURE_CONTAINER = os.getenv("AZURE_BLOB_CONTAINER", "mdl-output")
AZURE_CONN_STR  = os.getenv("AZURE_STORAGE_CONNECTION_STRING")  # optional
AZURITE_SAS_VERSION = os.getenv("AZURITE_SAS_VERSION", "2021-08-06")

LOCAL_SAVE_DIR = os.getenv("LOCAL_SAVE_DIR", "./_out")
PUBLIC_BASE_URL = os.getenv("PUBLIC_BASE_URL", "http://localhost:8000")

# ------------------------------------------------------------------------------
# Utilities
# ------------------------------------------------------------------------------
def sanitize(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9._-]+", "_", name or "").strip("_")

def _short(s: Optional[str], limit: int = 900) -> str:
    if not s:
        return ""
    s = re.sub(r"\s+", " ", s.strip())
    return (s[: limit - 1] + "…") if len(s) > limit else s

def _norm_ref(s: Optional[str]) -> str:
    return re.sub(r"\s+", "", (s or "")).upper()

def _norm_award_ref(val: Optional[str]) -> str:
    v = (val or "").strip()
    return v if v else "UNKNOWN"

def _pretty_compliance_type(raw: Optional[str], finding_text: str) -> str:
    t = (raw or "").strip()
    if not t or len(t) <= 2:  # fix junk like "I"
        s = (finding_text or "").lower()
        if "suspension" in s and "debar" in s:
            return "Procurement and Suspension and Debarment"
        if "subrecipient" in s:
            return "Subrecipient Monitoring"
        if "eligibility" in s:
            return "Eligibility"
        if "report" in s:
            return "Reporting – Inaccurate Treasury Reporting"
        if "allowed" in s or "unallowed" in s:
            return "Activities Allowed or Unallowed"
    return t

def _shade_cell(cell, hex_fill="E7E6E6"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)

def _set_col_widths(table, widths):
    # widths are in "EMU-like" integers; for our use a simple int split is fine
    for col_idx, w in enumerate(widths):
        for cell in table.columns[col_idx].cells:
            cell.width = w

def _tight_paragraph(p):
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)

TREASURY_DEFAULT_LISTINGS = ["21.027", "21.023"]  # SLFRF, ERA

def _looks_like_treasury_program(program_name: Optional[str]) -> bool:
    if not program_name:
        return False
    n = program_name.lower()
    return ("coronavirus state and local fiscal recovery funds" in n) or ("emergency rental assistance" in n)

def _is_treasury_listing(al: Optional[str], allow_list: List[str]) -> bool:
    if not al:
        return False
    return any(al.strip().startswith(prefix) for prefix in allow_list)

# Extract numbers like 21.027, 10.123, etc.
LISTING_RE = re.compile(r"\b(\d{2}\.\d{3})\b")

def extract_assistance_listings(text: Optional[str]) -> list[str]:
    if not text:
        return []
    return LISTING_RE.findall(text)

import re

_AL_RE = re.compile(r"\b\d{2}\.\d{3}\b")

def _derive_assistance_listing(name: str) -> str:
    """
    Best-effort: extract 'NN.NNN' from a program name if present,
    else return 'Unknown'.
    """
    if not name:
        return "Unknown"
    m = _AL_RE.search(name)
    return m.group(0) if m else "Unknown"

# ------------------------------------------------------------------------------
# Storage helpers
# ------------------------------------------------------------------------------
def _parse_conn_str(conn: str) -> Dict[str, Optional[str]]:
    """
    Return dict with AccountName, AccountKey, BlobEndpoint.
    Supports UseDevelopmentStorage=true and Azurite-like "AccountName=...;AccountKey=...;BlobEndpoint=..."
    """
    if not conn:
        return {"AccountName": None, "AccountKey": None, "BlobEndpoint": None}

    if "UseDevelopmentStorage=true" in conn:
        return {
            "AccountName": "devstoreaccount1",
            "AccountKey": (
                "Eby8vdM02xNOcqFlqUwJPLlmEtlCDXJ1OUzFT50uSRZ6IFsu"
                "Fq2UVErCz4I6tq/K1SZFPTOtr/KBHBeksoGMGw=="
            ),
            "BlobEndpoint": "http://127.0.0.1:10000/devstoreaccount1",
        }

    parts = dict(p.split("=", 1) for p in conn.split(";") if "=" in p)
    return {
        "AccountName": parts.get("AccountName"),
        "AccountKey": parts.get("AccountKey"),
        "BlobEndpoint": parts.get("BlobEndpoint"),  # presence implies Azurite/manual endpoint
    }

def _blob_service_client():
    """
    Create a BlobServiceClient that works for both Azure and Azurite.
    Avoids Azure's strict connection-string parser when we have only BlobEndpoint.
    """
    if not AZURE_CONN_STR:
        raise RuntimeError("AZURE_STORAGE_CONNECTION_STRING not set")

    info = _parse_conn_str(AZURE_CONN_STR)
    if info.get("BlobEndpoint") and info.get("AccountKey"):
        # Azurite or custom endpoint path
        return BlobServiceClient(account_url=info["BlobEndpoint"], credential=info["AccountKey"])
    # Otherwise let SDK parse the full Azure connection string
    return BlobServiceClient.from_connection_string(AZURE_CONN_STR)

def upload_and_sas(container: str, blob_name: str, data: bytes, ttl_minutes: int = 120) -> str:
    if not AZURE_CONN_STR:
        raise RuntimeError("AZURE_STORAGE_CONNECTION_STRING not set")

    info = _parse_conn_str(AZURE_CONN_STR)
    account_name = info["AccountName"]
    account_key  = info["AccountKey"]
    blob_endpoint = info.get("BlobEndpoint")

    bsc = _blob_service_client()
    cc = bsc.get_container_client(container)
    try:
        cc.create_container()
    except Exception:
        pass
    cc.upload_blob(name=blob_name, data=data, overwrite=True)

    # SAS tuned for Azurite and Azure
    proto = "http" if (blob_endpoint and ("127.0.0.1" in blob_endpoint or "localhost" in blob_endpoint)) else None
    sas = generate_blob_sas(
        account_name=account_name,
        account_key=account_key,
        container_name=container,
        blob_name=blob_name,
        permission=BlobSasPermissions(read=True),
        expiry=datetime.utcnow() + timedelta(minutes=ttl_minutes),
        version=AZURITE_SAS_VERSION,  # works on Azurite; Azure accepts older versions too
        protocol=proto,
    )
    sas_q = quote(sas, safe="=&")

    base = blob_endpoint.rstrip("/") if blob_endpoint else f"https://{account_name}.blob.core.windows.net"
    return f"{base}/{container}/{blob_name}?{sas_q}"

def save_local_and_url(blob_name: str, data: bytes) -> str:
    full_path = os.path.join(LOCAL_SAVE_DIR, blob_name)
    os.makedirs(os.path.dirname(full_path), exist_ok=True)
    with open(full_path, "wb") as f:
        f.write(data)
    return f"{PUBLIC_BASE_URL}/local/{blob_name}"

@app.get("/local/{path:path}")
def get_local_file(path: str):
    full = os.path.join(LOCAL_SAVE_DIR, path)
    if not os.path.isfile(full):
        raise HTTPException(404, "Not found")
    return FileResponse(
        full,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# ------------------------------------------------------------------------------
# FAC helpers
# ------------------------------------------------------------------------------
def _fac_headers():
    key = os.getenv("FAC_API_KEY")
    if not key:
        raise HTTPException(500, "FAC_API_KEY not configured on the docx service")
    return {"X-Api-Key": key}

def _fac_get(path: str, params: Dict[str, Any]) -> Any:
    try:
        r = requests.get(f"{FAC_BASE.rstrip('/')}/{path.lstrip('/')}",
                         headers=_fac_headers(), params=params, timeout=15)
        r.raise_for_status()
        return r.json()
    except requests.HTTPError as e:
        raise HTTPException(r.status_code if 'r' in locals() else 500,
                            f"FAC GET {path} failed: {getattr(r,'text','')}") from e
    except Exception as e:
        raise HTTPException(500, f"FAC GET {path} failed: {e}") from e

def _or_param(field: str, values: List[str]) -> str:
    inner = ",".join([f"{field}.eq.{v}" for v in values])
    return f"({inner})"

# ------------------------------------------------------------------------------
# HTML → DOCX (sturdy renderer)
# ------------------------------------------------------------------------------
def _apply_inline_formatting(paragraph, node):
    def add_text(text, bold=False, italic=False, underline=False):
        if text is None:
            return
        r = paragraph.add_run(text)
        r.bold = bool(bold)
        r.italic = bool(italic)
        r.underline = bool(underline)

    for child in getattr(node, "children", []):
        name = getattr(child, "name", None)
        if name is None:
            add_text(str(child))
        elif name in ("b", "strong"):
            add_text(child.get_text(), bold=True)
        elif name in ("i", "em"):
            add_text(child.get_text(), italic=True)
        elif name == "u":
            add_text(child.get_text(), underline=True)
        elif name == "br":
            paragraph.add_run().add_break()
        else:
            if hasattr(child, "children"):
                _apply_inline_formatting(paragraph, child)

def _basic_html_to_docx(doc: Document, html_str: str):
    soup = BeautifulSoup(html_str, "html.parser")
    body = soup.body or soup

    for element in body.children:
        if getattr(element, "name", None) is None:
            txt = str(element).strip()
            if txt:
                doc.add_paragraph(txt)
            continue

        tag = element.name.lower()

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(tag[1])
            text = element.get_text(strip=False)
            doc.add_heading(text, level=min(max(level, 1), 6))
            continue

        if tag == "p":
            p = doc.add_paragraph()
            _apply_inline_formatting(p, element)
            continue

        if tag in ("ul", "ol"):
            style = "List Bullet" if tag == "ul" else "List Number"
            for li in element.find_all("li", recursive=False):
                p = doc.add_paragraph(style=style)
                _apply_inline_formatting(p, li)
            continue

        if tag == "table":
            rows = element.find_all("tr", recursive=False)
            if not rows:
                continue
            first_cells = rows[0].find_all(["th", "td"], recursive=False)
            cols = max(1, len(first_cells))
            first_is_header = any(c.name == "th" for c in first_cells)

            tbl = doc.add_table(rows=len(rows), cols=cols)
            tbl.style = "Table Grid"

            sect = doc.sections[0]
            content_width = sect.page_width - sect.left_margin - sect.right_margin
            col_w = int(content_width / cols)
            _set_col_widths(tbl, [col_w] * cols)

            for r_idx, tr in enumerate(rows):
                cells = tr.find_all(["th", "td"], recursive=False)
                for c_idx in range(cols):
                    cell = tbl.cell(r_idx, c_idx)
                    if not cell.paragraphs:
                        cell.add_paragraph()
                    p = cell.paragraphs[0]
                    if hasattr(p, "clear"):
                        p.clear()   # not always present in python-docx
                    if c_idx < len(cells):
                        _apply_inline_formatting(p, cells[c_idx])
                    else:
                        p.text = ""
                    _tight_paragraph(p)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            if first_is_header:
                for c in tbl.rows[0].cells:
                    _shade_cell(c, "E7E6E6")
                    for r in c.paragraphs[0].runs:
                        r.bold = True
                    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if tag in ("div", "section", "article"):
            p = doc.add_paragraph()
            _apply_inline_formatting(p, element)
            continue

        txt = element.get_text(strip=True)
        if txt:
            doc.add_paragraph(txt)

def _replace_text_in_paragraph(paragraph, mapping: Dict[str, str]):
    # simple run-wise replacement; works when placeholders aren't split across runs
    for run in paragraph.runs:
        text = run.text
        for k, v in mapping.items():
            if k in text:
                run.text = text.replace(k, v)
                text = run.text

def _replace_placeholders_everywhere(doc, mapping: Dict[str, str]):
    # body paragraphs
    for p in doc.paragraphs:
        _replace_text_in_paragraph(p, mapping)
    # tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_text_in_paragraph(p, mapping)


def html_to_docx_bytes(html_str: str, *, force_basic: bool = False) -> bytes:
    doc = Document()
    tried_html2docx = False
    if not force_basic:
        try:
            tried_html2docx = True
            HTML2Docx().add_html_to_document(html_str or "", doc)
        except Exception:
            _basic_html_to_docx(doc, html_str or "")
    else:
        _basic_html_to_docx(doc, html_str or "")

    if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
        doc.add_paragraph("⚠️ HTML was empty." if force_basic else "⚠️ html2docx failed; fallback also produced no content.")
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ------------------------------------------------------------------------------
# Treasury-style MDL rendering
# ------------------------------------------------------------------------------
def summarize_finding_text(raw: str, max_chars: int = 1000) -> str:
    if not raw:
        return ""
    text = re.sub(r"\s+", " ", raw).strip()
    parts = re.split(r"(?<=[.?!])\s+", text)
    picked = []
    for p in parts:
        if len(picked) >= 3:
            break
        if re.search(r"\b(Assistance Listing|Award Period|Federal Program|Identification Number|CFDA)\b", p, re.I):
            continue
        picked.append(p)
    out = " ".join(picked) or text
    return _short(out, max_chars)

def format_letter_date(date_iso: Optional[str] = None) -> Tuple[str, str]:
    dt = datetime.fromisoformat(date_iso) if date_iso else datetime.utcnow()
    return dt.strftime("%Y-%m-%d"), dt.strftime("%B %d, %Y")

def render_mdl_html(model: Dict[str, Any]) -> str:
    letter_date_iso = model.get("letter_date_iso")
    _, letter_date_long = format_letter_date(letter_date_iso)

    auditee_name = model.get("auditee_name", "Recipient")
    ein = model.get("ein", "")
    address_lines = model.get("address_lines", [])
    attention_line = model.get("attention_line")
    period_end_text = model.get("period_end_text", str(model.get("audit_year", "")))
    include_no_qc_line = model.get("include_no_qc_line", True)

    address_block = "<br>".join(html.escape(x) for x in address_lines) if address_lines else ""
    attention_block = f"<p><strong>{html.escape(attention_line)}</strong></p>" if attention_line else ""

    def _render_program_table(p: Dict[str, Any]) -> str:
        rows_html = []
        for f in p.get("findings", []):
            rows_html.append(f"""
              <tr>
                <td>{html.escape(f.get('finding_id',''))}</td>
                <td>{html.escape(f.get('compliance_type',''))}</td>
                <td>{html.escape(f.get('summary',''))}</td>
                <td>{html.escape(f.get('audit_determination',''))}</td>
                <td>{html.escape(f.get('questioned_cost_determination',''))}</td>
                <td>{html.escape(f.get('cap_determination',''))}</td>
              </tr>
            """)
        if not rows_html:
            rows_html.append("""
              <tr><td colspan="6"><em>No MDL-relevant findings identified for this program.</em></td></tr>
            """)
        table = f"""
          <h3>Assistance Listing Number/Program Name: {html.escape(p.get('assistance_listing','Unknown'))} / {html.escape(p.get('program_name','Unknown'))}</h3>
          <table border="1" cellspacing="0" cellpadding="6" style="border-collapse:collapse; width:100%; font-size:10.5pt;">
            <tr>
              <th>Audit<br>Finding #</th>
              <th>Compliance Type -<br>Audit Finding</th>
              <th>Summary</th>
              <th>Audit Finding<br>Determination</th>
              <th>Questioned Cost<br>Determination</th>
              <th>CAP<br>Determination</th>
            </tr>
            {''.join(rows_html)}
          </table>
        """
        cap_blocks = []
        for f in p.get("findings", []):
            cap_text = f.get("cap_text")
            if cap_text:
                cap_blocks.append(f"""
                  <h4>Corrective Action Plan – {html.escape(f.get('finding_id',''))}</h4>
                  <p>{html.escape(cap_text)}</p>
                """)
        return table + ("\n".join(cap_blocks) if cap_blocks else "")

    programs = model.get("programs", [])
    programs_html = "\n".join(_render_program_table(p) for p in programs) if programs else "<p><em>No MDL-relevant findings identified per FAC records.</em></p>"

    not_sustained_notes = model.get("not_sustained_notes", [])
    not_sustained_html = ""
    if not_sustained_notes:
        notes_paras = "\n".join(f"<p>{html.escape(n)}</p>" for n in not_sustained_notes if n)
        not_sustained_html = f"<h3>FINDINGS NOT SUSTAINED</h3>\n{notes_paras}"

    # Build the letter (small chunks to keep it readable)
    chunks = []
    chunks.append(f'<p style="text-align:right; margin:0 0 12pt 0;">{html.escape(letter_date_long)}</p>')
    chunks.append("""
      <p style="margin:0 0 12pt 0;">
        <strong>DEPARTMENT OF THE TREASURY</strong><br>
        WASHINGTON, D.C.
      </p>
    """)
    chunks.append(f"""
      <p style="margin:0 0 12pt 0;">
        <strong>{html.escape(auditee_name)}</strong><br>
        EIN: {html.escape(ein)}<br>
        {address_block}
      </p>
    """)
    if attention_block:
        chunks.append(attention_block)
    chunks.append(f"""
      <p style="margin:12pt 0 12pt 0;">
        <strong>Subject:</strong> U.S. Department of the Treasury’s Management Decision Letter (MDL) for Single Audit Report for the period ending on {html.escape(period_end_text)}
      </p>
    """)
    chunks.append("""
      <p>
        In accordance with 2 C.F.R. § 200.521(b), the U.S. Department of the Treasury (Treasury)
        is required to issue a management decision for single audit findings pertaining to awards under
        Treasury’s programs. Treasury’s review as part of its responsibilities under 2 C.F.R § 200.513(c)
        includes an assessment of Treasury’s award recipients’ single audit findings, corrective action plans (CAPs),
        and questioned costs, if any.
      </p>
    """)
    chunks.append(f"""
      <p>
        Treasury has reviewed the single audit report for {html.escape(auditee_name)}.
        Treasury has made the following determinations regarding the audit finding(s) and CAP(s) listed below.
      </p>
    """)
    if include_no_qc_line:
        chunks.append("<p>No questioned costs are included in this single audit report.</p>")

    chunks.append(programs_html)
    if not_sustained_html:
        chunks.append(not_sustained_html)

    chunks.append("""
      <p>
        Please note, the corrective action(s) are subject to review during the recipient’s next annual single audit
        or program-specific audit, as applicable, to determine adequacy. If the same audit finding(s) appear in a future single
        audit report for this recipient, its current or future award funding under Treasury’s programs may be adversely impacted.
      </p>
      <p>
        The recipient may appeal Treasury’s decision for the audit finding(s) listed above. A written appeal must be submitted within
        30 calendar days of the date of this management decision letter to Treasury via email at
        <a href="mailto:ORP_SingleAudits@treasury.gov">ORP_SingleAudits@treasury.gov</a>.
        The appeal must include: (1) the specific reasons for disputing Treasury’s determination; (2) relevant documentation
        to support the recipient’s position; (3) an alternative course of action with an anticipated completion date of the action; and
        (4) the contact information of the managing official responsible for implementing the proposed alternative course of action.
      </p>
      <p>For questions regarding the audit finding(s), please email us at <a href="mailto:ORP_SingleAudits@treasury.gov">ORP_SingleAudits@treasury.gov</a>. Thank you.</p>
      <p style="margin-top:18pt;">Sincerely,<br><br>
      Audit and Compliance Resolution Team<br>
      Office of Capital Access<br>
      U.S. Department of the Treasury</p>
    """)

    return f'<div style="font-family: Calibri, Arial, sans-serif; font-size:11pt; line-height:1.4;">{"".join(chunks)}</div>'

from typing import Dict, Any, List, Optional, Tuple
import re
from datetime import datetime

def build_mdl_model_from_fac(
    *,
    auditee_name: str,
    ein: str,
    audit_year: int,
    fac_general: List[Dict[str, Any]],
    fac_findings: List[Dict[str, Any]],
    fac_findings_text: List[Dict[str, Any]],
    fac_caps: List[Dict[str, Any]],
    federal_awards: List[Dict[str, Any]],
    # Optional enrichments / knobs
    period_end_text: Optional[str] = None,
    address_lines: Optional[List[str]] = None,
    attention_line: Optional[str] = None,
    only_flagged: bool = False,
    max_refs: int = 10,
    auto_cap_determination: bool = True,
    # NEW knobs (AnythingLLM may pass these)
    include_no_qc_line: bool = False,
    include_no_cap_line: bool = False,
    **_  # swallow any future/unknown keywords safely
) -> Dict[str, Any]:
    """
    Normalize FAC payloads into the MDL model used by render_mdl_html().

    - Joins findings with narrative text & CAP by normalized ref #
    - Groups findings by award_reference -> program name (from federal_awards)
    - If /findings is empty but /findings_text exists, creates a catch-all section
    - include_no_qc_line/include_no_cap_line toggle the wording in the output table
    """

    # --- tiny helper: derive assistance listing from program text if not present elsewhere
    def _derive_assistance_listing(program_name: str) -> str:
        # Look for patterns like "21.027" etc.
        m = re.search(r"\b\d{2}\.\d{3}\b", program_name or "")
        return m.group(0) if m else "Unknown"

    # --- map award_reference -> program meta
    award2meta: Dict[str, Dict[str, str]] = {}
    for a in federal_awards or []:
        ref = a.get("award_reference")
        pname = (a.get("federal_program_name") or "").strip()
        if ref:
            award2meta[ref] = {
                "program_name": pname or "Unknown Program",
                "assistance_listing": _derive_assistance_listing(pname),
            }

    # --- normalize text/CAP by finding ref (uppercase, strip spaces)
    def _norm_ref(x: Optional[str]) -> str:
        return re.sub(r"\s+", "", (x or "")).upper()

    text_by_ref = { _norm_ref(t.get("finding_ref_number")): (t.get("finding_text") or "").strip()
                    for t in (fac_findings_text or []) }
    cap_by_ref  = { _norm_ref(c.get("finding_ref_number")): (c.get("planned_action") or "").strip()
                    for c in (fac_caps or []) }

    def _is_flagged(f: dict) -> bool:
        return any([
            f.get("is_material_weakness") is True,
            f.get("is_significant_deficiency") is True,
            f.get("is_questioned_costs") is True,
            f.get("is_modified_opinion") is True,
            f.get("is_other_findings") is True,
            f.get("is_other_matters") is True,
            f.get("is_repeat_finding") is True,
        ])

    # --- seed refs from /findings (respect only_flagged)
    base_refs: List[str] = []
    for f in fac_findings or []:
        if only_flagged and not _is_flagged(f):
            continue
        r = f.get("reference_number")
        if r:
            base_refs.append(r)

    # fallback: if /findings empty, derive refs from /findings_text
    if not base_refs and fac_findings_text:
        base_refs = [t.get("finding_ref_number") for t in fac_findings_text if t.get("finding_ref_number")]

    # de-dupe + cap
    seen = set()
    norm_refs: List[Tuple[str, str]] = []
    for r in base_refs:
        if not r:
            continue
        k = _norm_ref(r)
        if k not in seen:
            seen.add(k)
            norm_refs.append((r, k))
    norm_refs = norm_refs[: max_refs or 10]

    # --- build program groups from /findings
    programs_map: Dict[str, Dict[str, Any]] = {}
    for f in fac_findings or []:
        r = f.get("reference_number")
        if not r:
            continue
        k = _norm_ref(r)
        if k not in {kn for _, kn in norm_refs}:
            continue

        award_ref = f.get("award_reference") or "UNKNOWN"
        meta = award2meta.get(award_ref, {})
        group = programs_map.setdefault(award_ref, {
            "assistance_listing": meta.get("assistance_listing", "Unknown"),
            "program_name": meta.get("program_name", "Unknown Program"),
            "findings": []
        })

        # summarize the narrative (your existing helper)
        summary  = summarize_finding_text(text_by_ref.get(k, ""))
        cap_text = cap_by_ref.get(k)

        qcost_det = "No questioned costs identified" if include_no_qc_line else "None"
        cap_det   = ("Accepted" if (auto_cap_determination and cap_text)
                     else ("No CAP required" if include_no_cap_line else "Not Applicable"))

        group["findings"].append({
            "finding_id": f.get("reference_number") or "",
            "compliance_type": f.get("type_requirement") or "",
            "summary": summary,
            "audit_determination": "Sustained",
            "questioned_cost_determination": qcost_det,
            "disallowed_cost_determination": "None",
            "cap_determination": cap_det,
            "cap_text": cap_text,
        })

    # --- catch-all group if we have narrative but no /findings rows
    if not programs_map and norm_refs:
        catchall = {
            "assistance_listing": "Unknown",
            "program_name": "Unknown Program",
            "findings": []
        }
        for orig, key in norm_refs:
            cap_text = cap_by_ref.get(key)
            qcost_det = "No questioned costs identified" if include_no_qc_line else "None"
            cap_det   = ("Accepted" if (auto_cap_determination and cap_text)
                         else ("No CAP required" if include_no_cap_line else "Not Applicable"))
            catchall["findings"].append({
                "finding_id": orig,
                "compliance_type": "",
                "summary": summarize_finding_text(text_by_ref.get(key, "")),
                "audit_determination": "Sustained",
                "questioned_cost_determination": qcost_det,
                "disallowed_cost_determination": "None",
                "cap_determination": cap_det,
                "cap_text": cap_text,
            })
        programs_map["UNKNOWN"] = catchall

    # --- letter header basics
    model = {
        "letter_date_iso": datetime.utcnow().strftime("%Y-%m-%d"),
        "auditee_name": auditee_name,
        "ein": f"{ein[:2]}-{ein[2:]}" if ein and ein.isdigit() and len(ein) == 9 else ein,
        "address_lines": address_lines or [],
        "attention_line": attention_line or "",
        # placeholder unless you fetch actual period end elsewhere
        "period_end_text": period_end_text or f"June 30, {audit_year}",
        "audit_year": audit_year,
        "programs": list(programs_map.values()),
        "not_sustained_notes": [],
    }
    return model



# ------------------------------------------------------------------------------
# Schemas
# ------------------------------------------------------------------------------
class BuildDocx(BaseModel):
    auditee_name: str
    ein: str
    audit_year: int
    body_html: Optional[str] = None
    body_html_b64: Optional[str] = None
    dest_path: Optional[str] = None
    filename: Optional[str] = None

class BuildFromFAC(BaseModel):
    auditee_name: str
    ein: str
    audit_year: int
    fac_general: List[Dict[str, Any]] = []
    fac_findings: List[Dict[str, Any]] = []
    fac_findings_text: List[Dict[str, Any]] = []
    fac_caps: List[Dict[str, Any]] = []
    federal_awards: List[Dict[str, Any]] = []
    dest_path: Optional[str] = None
    filename: Optional[str] = None

class BuildByReport(BaseModel):
    auditee_name: str
    ein: str
    audit_year: int
    report_id: str
    dest_path: Optional[str] = None
    only_flagged: bool = False
    max_refs: int = 15
    include_awards: bool = False  # ignored now; we fetch minimal awards anyway

class BuildByReport(BaseModel):
    auditee_name: str
    ein: str
    audit_year: int
    report_id: str
    dest_path: Optional[str] = None
    only_flagged: bool = False
    max_refs: int = 15
    include_awards: bool = True     # turn ON to get program names
    treasury_listings: Optional[List[str]] = None  # e.g. ["21.027","21.023"]


class BuildByReportTemplated(BuildByReport):
    # letter header & addressing (all optional; template will be filled with blanks if missing)
    auditor_name: Optional[str] = None
    fy_end_text: Optional[str] = None             # e.g., "June 30, 2024"
    recipient_name: Optional[str] = None          # if different from auditee_name
    street_address: Optional[str] = None
    city: Optional[str] = None
    state: Optional[str] = None
    zip_code: Optional[str] = None
    poc_name: Optional[str] = None
    poc_title: Optional[str] = None

    # template path (optional; env fallback below)
    template_path: Optional[str] = None

# ------------------------------------------------------------------------------
# Routes
# ------------------------------------------------------------------------------
@app.get("/healthz")
def healthz():
    return {"ok": True, "time": datetime.utcnow().isoformat()}

@app.post("/echo")
def echo(payload: Dict[str, Any]):
    return {"received": payload, "ts": datetime.utcnow().isoformat()}

@app.get("/debug/env")
def debug_env():
    key = os.getenv("FAC_API_KEY") or ""
    masked = (key[:4] + "…" + key[-2:]) if key else None
    return {"fac_api_key_present": bool(key), "fac_api_key_masked": masked}

@app.get("/debug/storage")
def debug_storage():
    info = _parse_conn_str(AZURE_CONN_STR) if AZURE_CONN_STR else {}
    return {"using_storage": bool(AZURE_CONN_STR), "account": info.get("AccountName"), "blob_endpoint": info.get("BlobEndpoint")}

@app.get("/debug/sas")
def debug_sas():
    if not AZURE_CONN_STR:
        raise HTTPException(400, "Set AZURE_STORAGE_CONNECTION_STRING to test SAS.")
    url = upload_and_sas(AZURE_CONTAINER, "debug/hello.txt", b"hi", ttl_minutes=5)
    return {"url": url}

@app.post("/build-docx-demo")
def build_docx_demo():
    document = Document()
    document.add_heading("Hello from the DOCX demo ✅", level=1)
    document.add_paragraph("If you can read this, your write/upload pipeline is good.")
    bio = BytesIO(); document.save(bio); data = bio.getvalue()
    blob_name = "demo/hello.docx"
    url = upload_and_sas(AZURE_CONTAINER, blob_name, data) if AZURE_CONN_STR else save_local_and_url(blob_name, data)
    return {"url": url, "size_bytes": len(data)}

@app.post("/build-docx")
def build_docx(req: BuildDocx):
    html_str = (req.body_html or "").strip()
    if (not html_str) and req.body_html_b64:
        try:
            html_str = base64.b64decode(req.body_html_b64).decode("utf-8", errors="ignore")
        except Exception:
            raise HTTPException(400, "Invalid base64 in body_html_b64")
    if not html_str:
        raise HTTPException(400, "body_html (or body_html_b64) is required")

    data = html_to_docx_bytes(html_str, force_basic=False)  # general HTML can use html2docx first
    folder = (req.dest_path or "").lstrip("/")
    base = req.filename or f"MDL-{sanitize(req.auditee_name)}-{sanitize(req.ein)}-{req.audit_year}.docx"
    blob_name = f"{folder}{base}" if folder else base
    url = upload_and_sas(AZURE_CONTAINER, blob_name, data) if AZURE_CONN_STR else save_local_and_url(blob_name, data)
    return {"url": url, "blob_path": f"{AZURE_CONTAINER}/{blob_name}", "size_bytes": len(data)}

# ---------- Treasury-style: arrays -> DOCX ----------
@app.post("/build-docx-from-fac")
def build_docx_from_fac(req: BuildFromFAC):
    model = build_mdl_model_from_fac(
        auditee_name=req.auditee_name,
        ein=req.ein,
        audit_year=req.audit_year,
        fac_general=req.fac_general,
        fac_findings=req.fac_findings,
        fac_findings_text=req.fac_findings_text,
        fac_caps=req.fac_caps,
        federal_awards=req.federal_awards,
        only_flagged=False,
        max_refs=25,
        include_no_qc_line=True,
    )
    html_str = render_mdl_html(model)
    data = html_to_docx_bytes(html_str, force_basic=True)  # FORCE basic for faithful table layout

    folder = (req.dest_path or "").lstrip("/")
    base = req.filename or f"MDL-Preview-{sanitize(req.auditee_name)}-{sanitize(req.ein)}-{req.audit_year}.docx"
    blob_name = f"{folder}{base}" if folder else base
    url = upload_and_sas(AZURE_CONTAINER, blob_name, data) if AZURE_CONN_STR else save_local_and_url(blob_name, data)
    return {"url": url, "blob_path": f"{AZURE_CONTAINER}/{blob_name}", "size_bytes": len(data)}

# ---------- Treasury-style: report_id -> FAC -> arrays -> DOCX ----------
@app.post("/build-docx-by-report")
def build_docx_by_report(req: BuildByReport):
    # 1) minimal general
    fac_general = _fac_get("general", {
        "report_id": f"eq.{req.report_id}",
        "select": "report_id,fac_accepted_date",
        "limit": 1
    })

    # 2) findings (optionally only flagged)
    findings_params = {
        "report_id": f"eq.{req.report_id}",
        "select": "reference_number,award_reference,type_requirement,"
                  "is_material_weakness,is_significant_deficiency,is_questioned_costs,"
                  "is_modified_opinion,is_other_findings,is_other_matters,is_repeat_finding",
        "order": "reference_number.asc",
        "limit": str(req.max_refs)
    }
    if req.only_flagged:
        flagged = [
            "is_material_weakness","is_significant_deficiency","is_questioned_costs",
            "is_modified_opinion","is_other_findings","is_other_matters","is_repeat_finding"
        ]
        findings_params["or"] = "(" + ",".join([f"{f}.eq.true" for f in flagged]) + ")"
    fac_findings = _fac_get("findings", findings_params)

    # selected refs
    refs = [row.get("reference_number") for row in fac_findings if row.get("reference_number")]
    refs = refs[: req.max_refs]

    # 3) text & CAP for those refs only
    if refs:
        fac_findings_text = _fac_get("findings_text", {
            "report_id": f"eq.{req.report_id}",
            "select": "finding_ref_number,finding_text",
            "order": "finding_ref_number.asc",
            "limit": str(len(refs)),
            "or": _or_param("finding_ref_number", refs)
        })
        fac_caps = _fac_get("corrective_action_plans", {
            "report_id": f"eq.{req.report_id}",
            "select": "finding_ref_number,planned_action",
            "order": "finding_ref_number.asc",
            "limit": str(len(refs)),
            "or": _or_param("finding_ref_number", refs)
        })
    else:
        fac_findings_text, fac_caps = [], []

    # 4) (Optional) awards kept small/off by default
    federal_awards = []
    if req.include_awards:
        # NOTE: assistance_listing is NOT a column on federal_awards
        federal_awards = _fac_get("federal_awards", {
            "report_id": f"eq.{req.report_id}",
            "select": "award_reference,federal_program_name",
            "order": "award_reference.asc",
            "limit": "50"
        })


    # 5) build model + render + upload
    model = build_mdl_model_from_fac(
        auditee_name=req.auditee_name,
        ein=req.ein,
        audit_year=req.audit_year,
        fac_general=fac_general,
        fac_findings=fac_findings,
        fac_findings_text=fac_findings_text,
        fac_caps=fac_caps,
        federal_awards=federal_awards,
        only_flagged=req.only_flagged,
        max_refs=req.max_refs,
        include_no_qc_line=True,
    )
    html_str = render_mdl_html(model)
    data = html_to_docx_bytes(html_str, force_basic=True)

    folder = (req.dest_path or "").lstrip("/")
    base = f"MDL-Preview-{sanitize(req.auditee_name)}-{sanitize(req.ein)}-{req.audit_year}.docx"
    blob_name = f"{folder}{base}" if folder else base
    url = upload_and_sas(AZURE_CONTAINER, blob_name, data) if AZURE_CONN_STR else save_local_and_url(blob_name, data)
    return {"url": url, "blob_path": f"{AZURE_CONTAINER}/{blob_name}", "size_bytes": len(data)}

@app.post("/build-mdl-docx-by-report")
def build_mdl_docx_by_report(req: BuildByReport):
    # 1) General
    fac_general = _fac_get("general", {
        "report_id": f"eq.{req.report_id}",
        "select": "report_id,fac_accepted_date",
        "limit": 1
    })

    # 2) Findings (respect only_flagged)
    findings_params = {
        "report_id": f"eq.{req.report_id}",
        "select": "reference_number,award_reference,type_requirement,"
                  "is_material_weakness,is_significant_deficiency,is_questioned_costs,"
                  "is_modified_opinion,is_other_findings,is_other_matters,is_repeat_finding",
        "order": "reference_number.asc",
        "limit": str(req.max_refs)
    }
    if req.only_flagged:
        flagged = [
            "is_material_weakness","is_significant_deficiency","is_questioned_costs",
            "is_modified_opinion","is_other_findings","is_other_matters","is_repeat_finding"
        ]
        findings_params["or"] = "(" + ",".join([f"{f}.eq.true" for f in flagged]) + ")"
    fac_findings = _fac_get("findings", findings_params)

    # 3) get refs we actually fetched
    refs = [row.get("reference_number") for row in fac_findings if row.get("reference_number")]
    refs = refs[: req.max_refs]

    # 4) Finding text & CAP for those refs
    if refs:
        fac_findings_text = _fac_get("findings_text", {
            "report_id": f"eq.{req.report_id}",
            "select": "finding_ref_number,finding_text",
            "order": "finding_ref_number.asc",
            "limit": str(len(refs)),
            "or": _or_param("finding_ref_number", refs)
        })
        fac_caps = _fac_get("corrective_action_plans", {
            "report_id": f"eq.{req.report_id}",
            "select": "finding_ref_number,planned_action",
            "order": "finding_ref_number.asc",
            "limit": str(len(refs)),
            "or": _or_param("finding_ref_number", refs)
        })
    else:
        fac_findings_text, fac_caps = [], []

    # 5) Awards (needed for program names)
    federal_awards = []
    if req.include_awards:
        federal_awards = _fac_get("federal_awards", {
            "report_id": f"eq.{req.report_id}",
            "select": "award_reference,federal_program_name",  # ← only columns that exist
            "order": "award_reference.asc",
            "limit": "200"
        })

    # 6) Build MDL model (Treasury filtered) and render
    mdl_model = build_mdl_model_from_fac(
        auditee_name=req.auditee_name,
        ein=req.ein,
        audit_year=req.audit_year,
        fac_general=fac_general,
        fac_findings=fac_findings,
        fac_findings_text=fac_findings_text,
        fac_caps=fac_caps,
        federal_awards=federal_awards,
        only_flagged=req.only_flagged,
        max_refs=req.max_refs,
        treasury_listings=req.treasury_listings or TREASURY_DEFAULT_LISTINGS,
    )
    html = render_mdl_html(mdl_model)

    # 7) DOCX & upload
    data = html_to_docx_bytes(html)
    container = os.getenv("AZURE_BLOB_CONTAINER", "mdl-output")
    folder = (req.dest_path or "").lstrip("/")
    base = f"MDL-{sanitize(req.auditee_name)}-{sanitize(req.ein)}-{req.audit_year}.docx"
    blob_name = f"{folder}{base}" if folder else base

    url = upload_and_sas(container, blob_name, data) if os.getenv("AZURE_STORAGE_CONNECTION_STRING") else save_local_and_url(blob_name, data)
    return {"url": url, "blob_path": f"{container}/{blob_name}", "size_bytes": len(data)}
