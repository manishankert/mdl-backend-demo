# tests/test_html_converter.py
"""Tests for HTML to DOCX converter service."""

import pytest
from io import BytesIO
from docx import Document
from docx.shared import Pt

from services.html_converter import (
    apply_inline_formatting,
    basic_html_to_docx,
    set_font_size_to_12,
    html_to_docx_bytes,
)


@pytest.fixture
def empty_document():
    """Create an empty DOCX document."""
    return Document()


class TestApplyInlineFormatting:
    """Tests for apply_inline_formatting function."""

    def test_apply_inline_formatting_bold(self, empty_document):
        """Test apply_inline_formatting handles bold tags."""
        from bs4 import BeautifulSoup

        p = empty_document.add_paragraph()
        soup = BeautifulSoup("<b>Bold text</b>", "html.parser")
        apply_inline_formatting(p, soup)

        assert len(p.runs) > 0
        assert "Bold text" in p.text

    def test_apply_inline_formatting_italic(self, empty_document):
        """Test apply_inline_formatting handles italic tags."""
        from bs4 import BeautifulSoup

        p = empty_document.add_paragraph()
        soup = BeautifulSoup("<i>Italic text</i>", "html.parser")
        apply_inline_formatting(p, soup)

        assert "Italic text" in p.text

    def test_apply_inline_formatting_underline(self, empty_document):
        """Test apply_inline_formatting handles underline tags."""
        from bs4 import BeautifulSoup

        p = empty_document.add_paragraph()
        soup = BeautifulSoup("<u>Underlined text</u>", "html.parser")
        apply_inline_formatting(p, soup)

        assert "Underlined text" in p.text

    def test_apply_inline_formatting_plain_text(self, empty_document):
        """Test apply_inline_formatting handles plain text."""
        from bs4 import BeautifulSoup

        p = empty_document.add_paragraph()
        soup = BeautifulSoup("Plain text content", "html.parser")
        apply_inline_formatting(p, soup)

        assert "Plain text content" in p.text


class TestBasicHtmlToDocx:
    """Tests for basic_html_to_docx function."""

    def test_basic_html_to_docx_paragraph(self, empty_document):
        """Test basic_html_to_docx converts paragraphs."""
        html = "<p>Test paragraph</p>"
        basic_html_to_docx(empty_document, html)

        assert len(empty_document.paragraphs) > 0
        assert "Test paragraph" in empty_document.paragraphs[-1].text

    def test_basic_html_to_docx_heading(self, empty_document):
        """Test basic_html_to_docx converts headings."""
        html = "<h1>Main Heading</h1>"
        basic_html_to_docx(empty_document, html)

        # Document should have content
        assert len(empty_document.paragraphs) > 0

    def test_basic_html_to_docx_unordered_list(self, empty_document):
        """Test basic_html_to_docx converts unordered lists."""
        html = "<ul><li>Item 1</li><li>Item 2</li></ul>"
        basic_html_to_docx(empty_document, html)

        text = " ".join(p.text for p in empty_document.paragraphs)
        assert "Item 1" in text
        assert "Item 2" in text

    def test_basic_html_to_docx_ordered_list(self, empty_document):
        """Test basic_html_to_docx converts ordered lists."""
        html = "<ol><li>First</li><li>Second</li></ol>"
        basic_html_to_docx(empty_document, html)

        text = " ".join(p.text for p in empty_document.paragraphs)
        assert "First" in text
        assert "Second" in text

    def test_basic_html_to_docx_table(self, empty_document):
        """Test basic_html_to_docx converts tables."""
        html = """
        <table>
            <tr><th>Header 1</th><th>Header 2</th></tr>
            <tr><td>Cell 1</td><td>Cell 2</td></tr>
        </table>
        """
        basic_html_to_docx(empty_document, html)

        assert len(empty_document.tables) > 0
        table = empty_document.tables[0]
        assert table.cell(0, 0).text == "Header 1"
        assert table.cell(1, 0).text == "Cell 1"

    def test_basic_html_to_docx_empty_html(self, empty_document):
        """Test basic_html_to_docx handles empty HTML."""
        basic_html_to_docx(empty_document, "")
        # Should not raise an exception


class TestSetFontSizeTo12:
    """Tests for set_font_size_to_12 function."""

    def test_set_font_size_to_12_paragraphs(self, empty_document):
        """Test set_font_size_to_12 sets font size on paragraphs."""
        p = empty_document.add_paragraph("Test text")
        set_font_size_to_12(empty_document)

        for run in p.runs:
            assert run.font.size == Pt(12)

    def test_set_font_size_to_12_tables(self, empty_document):
        """Test set_font_size_to_12 sets font size on table cells."""
        table = empty_document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell text"

        set_font_size_to_12(empty_document)

        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        assert run.font.size == Pt(12)


class TestHtmlToDocxBytes:
    """Tests for html_to_docx_bytes function."""

    def test_html_to_docx_bytes_returns_bytes(self):
        """Test html_to_docx_bytes returns bytes."""
        html = "<p>Test content</p>"
        result = html_to_docx_bytes(html)

        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_html_to_docx_bytes_creates_valid_docx(self):
        """Test html_to_docx_bytes creates a valid DOCX file."""
        html = "<p>Test paragraph</p>"
        result = html_to_docx_bytes(html)

        # Should be able to open as a Document
        doc = Document(BytesIO(result))
        assert len(doc.paragraphs) > 0

    def test_html_to_docx_bytes_empty_html(self):
        """Test html_to_docx_bytes handles empty HTML."""
        result = html_to_docx_bytes("")

        doc = Document(BytesIO(result))
        # Should add a placeholder paragraph
        assert len(doc.paragraphs) > 0
        assert "empty" in doc.paragraphs[0].text.lower()

    def test_html_to_docx_bytes_force_basic(self):
        """Test html_to_docx_bytes with force_basic flag."""
        html = "<p>Test content</p>"
        result = html_to_docx_bytes(html, force_basic=True)

        doc = Document(BytesIO(result))
        assert len(doc.paragraphs) > 0

    def test_html_to_docx_bytes_complex_html(self):
        """Test html_to_docx_bytes with complex HTML."""
        html = """
        <h1>Main Title</h1>
        <p>This is a <b>bold</b> and <i>italic</i> paragraph.</p>
        <ul>
            <li>Item one</li>
            <li>Item two</li>
        </ul>
        <table>
            <tr><th>Col 1</th><th>Col 2</th></tr>
            <tr><td>Data 1</td><td>Data 2</td></tr>
        </table>
        """
        result = html_to_docx_bytes(html)

        doc = Document(BytesIO(result))
        # Should have content
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "bold" in full_text or "Main Title" in full_text
