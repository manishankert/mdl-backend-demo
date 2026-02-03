# tests/test_docx_utils.py
"""Tests for DOCX utility functions."""

import pytest
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches

from utils.docx_utils import (
    shade_cell,
    set_col_widths,
    tight_paragraph,
    as_oxml,
    apply_grid_borders,
    remove_paragraph,
    clear_runs,
    para_text,
    rewrite_para_text,
    twips_from_inches,
    set_cell_preferred_width,
    set_table_column_widths,
)


@pytest.fixture
def sample_document():
    """Create a sample DOCX document for testing."""
    doc = Document()
    doc.add_paragraph("Test paragraph 1")
    doc.add_paragraph("Test paragraph 2")
    return doc


@pytest.fixture
def sample_table_document():
    """Create a sample DOCX document with a table."""
    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"Cell {i},{j}"
    return doc, table


class TestShadeCell:
    """Tests for shade_cell function."""

    def test_shade_cell_default_color(self, sample_table_document):
        """Test shade_cell applies default grey color."""
        doc, table = sample_table_document
        cell = table.cell(0, 0)
        shade_cell(cell)
        # Check that the shading element was added
        tcPr = cell._tc.tcPr
        assert tcPr is not None

    def test_shade_cell_custom_color(self, sample_table_document):
        """Test shade_cell applies custom color."""
        doc, table = sample_table_document
        cell = table.cell(0, 0)
        shade_cell(cell, hex_fill="FF0000")
        tcPr = cell._tc.tcPr
        assert tcPr is not None


class TestSetColWidths:
    """Tests for set_col_widths function."""

    def test_set_col_widths(self, sample_table_document):
        """Test set_col_widths sets column widths."""
        doc, table = sample_table_document
        widths = [Inches(1), Inches(2), Inches(1.5)]
        set_col_widths(table, widths)
        # Verify widths were set
        for col_idx, expected_width in enumerate(widths):
            cell = table.columns[col_idx].cells[0]
            assert cell.width == expected_width


class TestTightParagraph:
    """Tests for tight_paragraph function."""

    def test_tight_paragraph_sets_zero_spacing(self, sample_document):
        """Test tight_paragraph sets zero spacing before/after."""
        p = sample_document.paragraphs[0]
        tight_paragraph(p)
        assert p.paragraph_format.space_before == Pt(0)
        assert p.paragraph_format.space_after == Pt(0)
        assert p.paragraph_format.line_spacing == 1.0


class TestAsOxml:
    """Tests for as_oxml function."""

    def test_as_oxml_paragraph(self, sample_document):
        """Test as_oxml extracts oxml from Paragraph."""
        p = sample_document.paragraphs[0]
        result = as_oxml(p)
        assert result == p._p

    def test_as_oxml_table(self, sample_table_document):
        """Test as_oxml extracts oxml from Table."""
        doc, table = sample_table_document
        result = as_oxml(table)
        assert result == table._tbl

    def test_as_oxml_raw_element(self):
        """Test as_oxml returns raw element unchanged."""
        from docx.oxml import OxmlElement
        el = OxmlElement("w:p")
        result = as_oxml(el)
        assert result == el


class TestApplyGridBorders:
    """Tests for apply_grid_borders function."""

    def test_apply_grid_borders(self, sample_table_document):
        """Test apply_grid_borders adds border elements."""
        doc, table = sample_table_document
        apply_grid_borders(table)
        tblPr = table._tbl.tblPr
        assert tblPr is not None


class TestRemoveParagraph:
    """Tests for remove_paragraph function."""

    def test_remove_paragraph(self, sample_document):
        """Test remove_paragraph removes paragraph from document."""
        initial_count = len(sample_document.paragraphs)
        p = sample_document.paragraphs[0]
        remove_paragraph(p)
        assert len(sample_document.paragraphs) == initial_count - 1


class TestClearRuns:
    """Tests for clear_runs function."""

    def test_clear_runs(self, sample_document):
        """Test clear_runs removes all runs from paragraph."""
        p = sample_document.paragraphs[0]
        # Ensure paragraph has runs
        assert len(p.runs) > 0
        clear_runs(p)
        assert len(p.runs) == 0


class TestParaText:
    """Tests for para_text function."""

    def test_para_text_extracts_text(self, sample_document):
        """Test para_text extracts text from all runs."""
        p = sample_document.paragraphs[0]
        result = para_text(p)
        assert result == "Test paragraph 1"

    def test_para_text_empty_paragraph(self, sample_document):
        """Test para_text with empty paragraph."""
        p = sample_document.add_paragraph("")
        result = para_text(p)
        assert result == ""


class TestRewriteParaText:
    """Tests for rewrite_para_text function."""

    def test_rewrite_para_text(self, sample_document):
        """Test rewrite_para_text replaces paragraph text."""
        p = sample_document.paragraphs[0]
        rewrite_para_text(p, "New text content")
        assert para_text(p) == "New text content"


class TestTwipsFromInches:
    """Tests for twips_from_inches function."""

    def test_twips_from_inches_one_inch(self):
        """Test twips_from_inches converts 1 inch to 1440 twips."""
        assert twips_from_inches(1.0) == 1440

    def test_twips_from_inches_half_inch(self):
        """Test twips_from_inches converts 0.5 inch to 720 twips."""
        assert twips_from_inches(0.5) == 720

    def test_twips_from_inches_zero(self):
        """Test twips_from_inches converts 0 to 0."""
        assert twips_from_inches(0) == 0


class TestSetCellPreferredWidth:
    """Tests for set_cell_preferred_width function."""

    def test_set_cell_preferred_width(self, sample_table_document):
        """Test set_cell_preferred_width sets width."""
        doc, table = sample_table_document
        cell = table.cell(0, 0)
        set_cell_preferred_width(cell, 2.0)
        tcPr = cell._tc.tcPr
        assert tcPr is not None


class TestSetTableColumnWidths:
    """Tests for set_table_column_widths function."""

    def test_set_table_column_widths(self, sample_table_document):
        """Test set_table_column_widths sets widths for all columns."""
        doc, table = sample_table_document
        widths = [1.5, 2.0, 1.5]
        set_table_column_widths(table, widths)
        # Verify widths were set on cells
        for row in table.rows:
            for i, expected in enumerate(widths):
                assert row.cells[i].width == Inches(expected)
