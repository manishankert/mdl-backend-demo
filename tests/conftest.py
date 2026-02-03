# tests/conftest.py
"""
Pytest configuration and fixtures for the MDL DOCX Builder test suite.
"""

import os
import sys
import pytest
from unittest.mock import MagicMock, patch
from io import BytesIO

# Add the project root to the path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from fastapi.testclient import TestClient


@pytest.fixture
def test_client():
    """Create a FastAPI test client."""
    from main import app
    return TestClient(app)


@pytest.fixture
def sample_fac_general():
    """Sample FAC general data."""
    return [{
        "report_id": "2024-06-GSAFAC-0000123456",
        "fac_accepted_date": "2024-06-15",
        "auditee_address_line_1": "123 Main Street",
        "auditee_city": "Ann Arbor",
        "auditee_state": "MI",
        "auditee_zip": "48104",
        "auditor_firm_name": "REHMANN ROBSON LLC",
        "fy_end_date": "2023-06-30",
        "auditee_contact_name": "John Smith",
        "auditee_contact_title": "Finance Director",
        "auditee_name": "CITY OF ANN ARBOR",
    }]


@pytest.fixture
def sample_fac_findings():
    """Sample FAC findings data."""
    return [
        {
            "reference_number": "2024-001",
            "award_reference": "AWARD-001",
            "type_requirement": "I",
            "is_material_weakness": False,
            "is_significant_deficiency": True,
            "is_questioned_costs": False,
            "is_modified_opinion": False,
            "is_other_findings": False,
            "is_other_matters": False,
            "is_repeat_finding": False,
        },
        {
            "reference_number": "2024-002",
            "award_reference": "AWARD-001",
            "type_requirement": "M",
            "is_material_weakness": True,
            "is_significant_deficiency": False,
            "is_questioned_costs": False,
            "is_modified_opinion": False,
            "is_other_findings": False,
            "is_other_matters": False,
            "is_repeat_finding": False,
        },
    ]


@pytest.fixture
def sample_fac_findings_text():
    """Sample FAC findings text data."""
    return [
        {
            "finding_ref_number": "2024-001",
            "finding_text": "The City did not verify suspension and debarment status for contractors prior to award."
        },
        {
            "finding_ref_number": "2024-002",
            "finding_text": "The City did not adequately monitor subrecipients receiving federal funds."
        },
    ]


@pytest.fixture
def sample_fac_caps():
    """Sample FAC corrective action plans data."""
    return [
        {
            "finding_ref_number": "2024-001",
            "planned_action": "The City will implement a verification process for all contractors."
        },
        {
            "finding_ref_number": "2024-002",
            "planned_action": "The City will establish a subrecipient monitoring program."
        },
    ]


@pytest.fixture
def sample_federal_awards():
    """Sample federal awards data."""
    return [
        {
            "award_reference": "AWARD-001",
            "federal_program_name": "Coronavirus State and Local Fiscal Recovery Funds",
            "assistance_listing": "21.027",
        },
    ]


@pytest.fixture
def sample_mdl_model():
    """Sample MDL model for template testing."""
    return {
        "letter_date_iso": "2024-06-15",
        "auditee_name": "City of Ann Arbor",
        "ein": "38-6004534",
        "address_lines": ["123 Main Street", "Ann Arbor, MI 48104"],
        "attention_line": "Finance Director",
        "period_end_text": "June 30, 2023",
        "audit_year": 2023,
        "auditor_name": "Rehmann Robson LLC",
        "street_address": "123 Main Street",
        "city": "Ann Arbor",
        "state": "MI",
        "zip_code": "48104",
        "poc_name": "John Smith",
        "poc_title": "Finance Director",
        "treasury_contact_email": "ORP_SingleAudits@treasury.gov",
        "programs": [
            {
                "assistance_listing": "21.027",
                "program_name": "Coronavirus State and Local Fiscal Recovery Funds (SLFRF)",
                "findings": [
                    {
                        "finding_id": "2024-001",
                        "compliance_type": "Procurement and suspension and debarment",
                        "summary": "Lack of evidence of suspension and debarment verification",
                        "compliance_and_summary": "Procurement and suspension and debarment - Lack of evidence of suspension and debarment verification",
                        "audit_determination": "Sustained",
                        "questioned_cost_determination": "None",
                        "disallowed_cost_determination": "None",
                        "cap_determination": "Accepted",
                        "cap_text": "The City will implement a verification process.",
                    }
                ],
            }
        ],
        "not_sustained_notes": [],
    }


@pytest.fixture
def mock_azure_storage():
    """Mock Azure storage client."""
    with patch('services.storage.BlobServiceClient') as mock:
        mock_client = MagicMock()
        mock.return_value = mock_client
        mock.from_connection_string.return_value = mock_client

        mock_container = MagicMock()
        mock_client.get_container_client.return_value = mock_container

        yield mock


@pytest.fixture
def mock_fac_api():
    """Mock FAC API responses."""
    with patch('services.fac_api.requests.get') as mock:
        yield mock


@pytest.fixture
def mock_openai_api():
    """Mock OpenAI API responses."""
    with patch('services.mdl_builder.requests.post') as mock:
        mock.return_value.json.return_value = {
            "choices": [{"message": {"content": "Lack of evidence of suspension and debarment verification"}}]
        }
        yield mock


@pytest.fixture
def temp_docx_file(tmp_path):
    """Create a temporary DOCX template file for testing."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Date XX, 2025")
    doc.add_paragraph("[Recipient Name]")
    doc.add_paragraph("EIN: [EIN]")
    doc.add_paragraph("[Street Address]")
    doc.add_paragraph("[City], [State] [Zip Code]")
    doc.add_paragraph("Assistance Listing Number/Program Name:")
    doc.add_paragraph("[[PROGRAM_TABLES]]")
    doc.add_paragraph("Treasury has reviewed the single audit report for [Recipient Name], prepared by [Auditor Name] for the fiscal year ending on [Fiscal Year End Date].")

    template_path = tmp_path / "test_template.docx"
    doc.save(str(template_path))

    return str(template_path)


@pytest.fixture
def set_test_env_vars():
    """Set environment variables for testing."""
    original_env = os.environ.copy()

    os.environ["FAC_API_KEY"] = "test-api-key"
    os.environ["FAC_API_BASE"] = "https://api.fac.gov"
    os.environ["AZURE_STORAGE_CONNECTION_STRING"] = ""
    os.environ["LOCAL_SAVE_DIR"] = "./_test_out"
    os.environ["PUBLIC_BASE_URL"] = "http://localhost:8000"

    yield

    os.environ.clear()
    os.environ.update(original_env)
