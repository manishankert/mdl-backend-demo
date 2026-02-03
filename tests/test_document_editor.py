# tests/test_document_editor.py
"""Tests for document editor service functions."""

import pytest
from io import BytesIO
from docx import Document
from docx.shared import Pt
from unittest.mock import patch, MagicMock

from services.document_editor import (
    para_text,
    force_paragraph_font_size,
    set_font_size_to_12,
    looks_like_optional_plural_text,
    fix_mdl_grammar_text,
    apply_mdl_grammar,
    replace_email_with_mailto_link,
    postprocess_docx,
)


@pytest.fixture
def sample_document():
    """Create a sample DOCX document."""
    doc = Document()
    doc.add_paragraph("Test paragraph content")
    doc.add_paragraph("Another paragraph")
    return doc


@pytest.fixture
def document_with_table():
    """Create a document with a table."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header 1"
    table.cell(0, 1).text = "Header 2"
    table.cell(1, 0).text = "Cell 1"
    table.cell(1, 1).text = "Cell 2"
    return doc


class TestParaText:
    """Tests for para_text function."""

    def test_para_text_extracts_text(self, sample_document):
        """Test para_text extracts text from paragraph."""
        p = sample_document.paragraphs[0]
        result = para_text(p)

        assert result == "Test paragraph content"

    def test_para_text_empty_paragraph(self, sample_document):
        """Test para_text with empty paragraph."""
        p = sample_document.add_paragraph("")
        result = para_text(p)

        assert result == ""


class TestForceParagraphFontSize:
    """Tests for force_paragraph_font_size function."""

    def test_force_paragraph_font_size_default(self, sample_document):
        """Test force_paragraph_font_size with default size."""
        p = sample_document.paragraphs[0]
        force_paragraph_font_size(p)

        for run in p.runs:
            assert run.font.size == Pt(12)

    def test_force_paragraph_font_size_custom(self, sample_document):
        """Test force_paragraph_font_size with custom size."""
        p = sample_document.paragraphs[0]
        force_paragraph_font_size(p, size_pt=14)

        for run in p.runs:
            assert run.font.size == Pt(14)


class TestSetFontSizeTo12:
    """Tests for set_font_size_to_12 function."""

    def test_set_font_size_to_12_paragraphs(self, sample_document):
        """Test set_font_size_to_12 on paragraphs."""
        set_font_size_to_12(sample_document)

        for p in sample_document.paragraphs:
            for run in p.runs:
                assert run.font.size == Pt(12)

    def test_set_font_size_to_12_tables(self, document_with_table):
        """Test set_font_size_to_12 on table cells."""
        set_font_size_to_12(document_with_table)

        for table in document_with_table.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            assert run.font.size == Pt(12)


class TestLooksLikeOptionalPluralText:
    """Tests for looks_like_optional_plural_text function."""

    def test_looks_like_optional_plural_text_with_s(self):
        """Test detection of (s) pattern."""
        assert looks_like_optional_plural_text("The finding(s) are listed below")

    def test_looks_like_optional_plural_text_with_es(self):
        """Test detection of (es) pattern."""
        assert looks_like_optional_plural_text("This address(es) the issue")

    def test_looks_like_optional_plural_text_violate(self):
        """Test detection of violate(s) pattern."""
        assert looks_like_optional_plural_text("This issue violate(s) the rule")

    def test_looks_like_optional_plural_text_audit_finding(self):
        """Test detection of 'audit finding' text."""
        assert looks_like_optional_plural_text("The audit finding is sustained")

    def test_looks_like_optional_plural_text_corrective_action(self):
        """Test detection of 'corrective action' text."""
        assert looks_like_optional_plural_text("The corrective action is accepted")

    def test_looks_like_optional_plural_text_normal(self):
        """Test normal text without pluralization issues."""
        assert not looks_like_optional_plural_text("This is normal text")

    def test_looks_like_optional_plural_text_empty(self):
        """Test empty string."""
        assert not looks_like_optional_plural_text("")

    def test_looks_like_optional_plural_text_none(self):
        """Test None input."""
        assert not looks_like_optional_plural_text(None)


class TestFixMdlGrammarText:
    """Tests for fix_mdl_grammar_text function."""

    def test_fix_mdl_grammar_text_singular_s(self):
        """Test fixing (s) for singular."""
        text = "The finding(s) is listed"
        result = fix_mdl_grammar_text(text, n_findings=1)

        assert "(s)" not in result
        assert "finding" in result.lower()

    def test_fix_mdl_grammar_text_plural_s(self):
        """Test fixing (s) for plural."""
        text = "The finding(s) are listed"
        result = fix_mdl_grammar_text(text, n_findings=2)

        assert "(s)" not in result
        assert "findings" in result.lower()

    def test_fix_mdl_grammar_text_is_are_singular(self):
        """Test fixing [is/are] for singular."""
        text = "The finding [is/are] sustained"
        result = fix_mdl_grammar_text(text, n_findings=1)

        assert "[is/are]" not in result
        assert " is " in result

    def test_fix_mdl_grammar_text_is_are_plural(self):
        """Test fixing [is/are] for plural."""
        text = "The findings [is/are] sustained"
        result = fix_mdl_grammar_text(text, n_findings=2)

        assert "[is/are]" not in result
        assert " are " in result

    def test_fix_mdl_grammar_text_nbsp(self):
        """Test NBSP conversion to space."""
        text = "Some\u00A0text\u00A0here"
        result = fix_mdl_grammar_text(text, n_findings=1)

        assert "\u00A0" not in result

    def test_fix_mdl_grammar_text_multiple_spaces(self):
        """Test cleaning up multiple spaces."""
        text = "Too   many   spaces"
        result = fix_mdl_grammar_text(text, n_findings=1)

        assert "   " not in result


class TestApplyMdlGrammar:
    """Tests for apply_mdl_grammar function."""

    def test_apply_mdl_grammar_updates_paragraphs(self):
        """Test apply_mdl_grammar updates paragraph text."""
        doc = Document()
        doc.add_paragraph("The finding(s) is listed below")

        apply_mdl_grammar(doc, n_findings=1)

        text = doc.paragraphs[0].text
        assert "(s)" not in text

    def test_apply_mdl_grammar_updates_tables(self):
        """Test apply_mdl_grammar updates table cells."""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "The finding(s) is here"

        apply_mdl_grammar(doc, n_findings=1)

        text = table.cell(0, 0).text
        assert "(s)" not in text


class TestReplaceEmailWithMailtoLink:
    """Tests for replace_email_with_mailto_link function."""

    def test_replace_email_with_mailto_link_found(self):
        """Test replacing email with mailto link."""
        doc = Document()
        p = doc.add_paragraph("Contact us at test@example.com for questions.")

        result = replace_email_with_mailto_link(p, "test@example.com")

        assert result is True

    def test_replace_email_with_mailto_link_not_found(self):
        """Test when email is not in paragraph."""
        doc = Document()
        p = doc.add_paragraph("No email in this paragraph.")

        result = replace_email_with_mailto_link(p, "test@example.com")

        assert result is False


class TestPostprocessDocx:
    """Tests for postprocess_docx function."""

    def test_postprocess_docx_returns_bytes(self):
        """Test postprocess_docx returns bytes."""
        doc = Document()
        doc.add_paragraph("Test content")

        bio = BytesIO()
        doc.save(bio)
        doc_bytes = bio.getvalue()

        model = {
            "auditee_name": "Test City",
            "treasury_contact_email": "test@treasury.gov"
        }

        result = postprocess_docx(doc_bytes, model)

        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_postprocess_docx_preserves_content(self):
        """Test postprocess_docx preserves document content."""
        doc = Document()
        doc.add_paragraph("Test paragraph that should be preserved")

        bio = BytesIO()
        doc.save(bio)
        doc_bytes = bio.getvalue()

        model = {"auditee_name": "Test City"}

        result = postprocess_docx(doc_bytes, model)

        # Open resulting document and check content
        result_doc = Document(BytesIO(result))
        all_text = " ".join(p.text for p in result_doc.paragraphs)
        assert "preserved" in all_text

    def test_postprocess_docx_creates_valid_docx(self):
        """Test postprocess_docx creates a valid DOCX."""
        doc = Document()
        doc.add_paragraph("Test")

        bio = BytesIO()
        doc.save(bio)
        doc_bytes = bio.getvalue()

        model = {"auditee_name": "Test City"}

        result = postprocess_docx(doc_bytes, model)

        # Should be able to open as a Document without error
        result_doc = Document(BytesIO(result))
        assert result_doc is not None
