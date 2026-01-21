"""
MDL Generator - Complete Implementation Based on Template Comments

Requirements extracted from 40 comments in the template document:

FIELD REQUIREMENTS:
1. Current Date - Format: "Month DD, YYYY" (e.g., December 12, 2025)
2. Recipient Name - Standard case (not ALL CAPS), from [auditee_name]
   - Address block: WITHOUT "The" prefix
   - Narrative: WITH "The" prefix
3. EIN - Format: XX-XXXXXXX (with dash), from [auditee_ein]
4. Street Address - Standard case, from [auditee_address_line_1]
5. City, State, Zip - Standard case, from [auditee_city], [auditee_state], [auditee_zip]
6. Point of Contact - Standard case, from [auditee_contact_name], [auditee_contact_title]
7. Fiscal Year End Date - Format: "Month Day, Year", from [fy_end_date]
8. Auditor Name - Add "the" prefix in narrative, standard case

PLURALIZATION (based on total finding count):
- is/are, finding/findings, issue/issues, violates/violate, CAP/CAPs, corrective action/actions

TABLE REQUIREMENTS:
- One table per program (ALN)
- Tables in ALN order (21.023, then 21.026, then 21.027, etc.)
- Format: [ALN]/ [Program Name] ([Acronym])

FINDING REQUIREMENTS:
- Finding Number: Reference number from SF-SAC
- Repeat Finding: If [is_repeat_finding]=Y, add "Repeat of [prior_finding_ref_numbers]"
- Compliance Type: Mapped from [type_requirement] letter code
- Finding Summary: Matched from standardized list
- Audit Finding Determination: "Sustained"
- Questioned Cost: "Questioned Cost:\nNone\nDisallowed Cost:\nNone"
- CAP Determination: "Accepted" if CAP exists, else "Not Applicable"
"""

import os
import re
import logging
from io import BytesIO
from datetime import datetime, timedelta
from typing import Optional, List, Dict, Any, Tuple
from dataclasses import dataclass, field

import requests
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)


def add_hyperlink(paragraph, url: str, text: str, color: str = "0000FF", underline: bool = True):
    """
    Add a hyperlink to a paragraph.
    
    Args:
        paragraph: The paragraph to add the hyperlink to
        url: The URL (e.g., "mailto:email@example.com" or "https://...")
        text: The display text for the hyperlink
        color: Hex color code (default blue)
        underline: Whether to underline the link
    
    Returns:
        The hyperlink element
    
    # Get the document part
      - type_map: {'I': 'Procurement and suspension and debarment', ...}
      - summary_labels: ['Lack of evidence of suspension and debarment verification', ...]
    Tolerant to header naming; no-op if workbook missing.
    """
    type_map, summary_labels = {}, []
    if not xlsx_path:
        return type_map, summary_labels
    try:
        import openpyxl
    except Exception:
        return type_map, summary_labels

    try:
        wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    except Exception:
        return type_map, summary_labels

    def _find_sheet(*names):
        for ws in wb.worksheets:
            nm = (ws.title or "").strip().lower().replace(" ", "_")
            for want in names:
                if want in nm:
                    return ws
        return None

    # 1) Finding Types sheet
    ws_types = _find_sheet("finding_types", "findingtype", "finding_types_sheet", "types")
    if ws_types and ws_types.max_row >= 2:
        hdrs = [ (c.value or "") for c in ws_types[1] ]
        hl = [str(h).strip().lower() for h in hdrs]

        def colidx(cands):
            for i,h in enumerate(hl):
                for c in cands:
                    c = c.lower()
                    if h == c or c in h:
                        return i
            return None

        i_code = colidx(["code","compliance type","compliance_type","ctype"])
        i_name = colidx(["name","description","label","type name","type"])
        if i_code is not None and i_name is not None:
            for row in ws_types.iter_rows(min_row=2, values_only=True):
                code = (row[i_code] or "")
                name = (row[i_name] or "")
                code = str(code).strip().upper()
                name = str(name).strip()
                if code and name:
                    type_map[code] = name

    # 2) Finding_summaries sheet
    ws_summ = _find_sheet("finding_summaries", "finding_summ", "summaries", "summary")
    if ws_summ and ws_summ.max_row >= 2:
        # Use first non-empty cell in each row (or a column named 'summary' / 'label')
        hdrs = [ (c.value or "") for c in ws_summ[1] ]
        hl = [str(h).strip().lower() for h in hdrs]
        def cidx(cands):
            for i,h in enumerate(hl):
                for c in cands:
                    c = c.lower()
                    if h == c or c in h:
                        return i
            return None
        i_lbl = cidx(["summary","label","finding summary","finding_label"]) or 0
        for row in ws_summ.iter_rows(min_row=2, values_only=True):
            cell = row[i_lbl] if i_lbl < len(row) else None
            txt = str(cell or "").strip()
            if txt:
                summary_labels.append(txt)

    return type_map, summary_labels


def _best_summary_label(summary: str, labels: List[str]) -> Optional[str]:
    """
    Offline fuzzy match: pick the label with highest similarity to the summary.
    """
    if not summary or not labels:
        return None
    import difflib
    cand = difflib.get_close_matches(summary, labels, n=1, cutoff=0.0)
    if cand:
        return cand[0]
    return None

# ------------------------------------------------------------------------------
# Storage helpers
# ------------------------------------------------------------------------------
def _parse_conn_str(conn: str) -> Dict[str, Optional[str]]:
    if not conn:
        return {"AccountName": None, "AccountKey": None, "BlobEndpoint": None}

    if "UseDevelopmentStorage=true" in conn:
        return {
            "AccountName": "devstoreaccount1",
            "AccountKey": (
                "Eby8vdM02xNOcqFlqUwJPLlmEtlCDXJ1OUzFT50uSRZ6IFsu"
                "Fq2UVErCz4I6tq/K1SZFPTOtr/KBHBeksoGMGw=="
            ),
            "BlobEndpoint": "http://127.0.0.1:10000/devstoreaccount1",
        }

    parts = dict(p.split("=", 1) for p in conn.split(";") if "=" in p)
    return {
        "AccountName": parts.get("AccountName"),
        "AccountKey": parts.get("AccountKey"),
        "BlobEndpoint": parts.get("BlobEndpoint"),
    }

def _blob_service_client():
    if not AZURE_CONN_STR:
        raise RuntimeError("AZURE_STORAGE_CONNECTION_STRING not set")
    info = _parse_conn_str(AZURE_CONN_STR)
    if info.get("BlobEndpoint") and info.get("AccountKey"):
        return BlobServiceClient(account_url=info["BlobEndpoint"], credential=info["AccountKey"])
    return BlobServiceClient.from_connection_string(AZURE_CONN_STR)

# def _best_summary_label_openai(summary: str, labels: List[str]) -> Optional[str]:
#     import os, json, requests
#     api_key = os.getenv("OPENAI_API_KEY")
#     if not api_key or not labels:
#         return None
#     prompt = {
#         "summary": summary,
#         "labels": labels,
#         "task": "Pick exactly one label from 'labels' that best matches 'summary'. Respond with just the label text."
#     }
#     try:
#         r = requests.post(
#             "https://api.openai.com/v1/chat/completions",
#             headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
#             data=json.dumps({
#                 "model": "gpt-4o-mini",
#                 "messages": [{"role":"user","content": json.dumps(prompt)}],
#                 "temperature": 0
#             }),
#             timeout=12,
#         )
#         out = r.json()
#         txt = (out.get("choices",[{}])[0].get("message",{}).get("content") or "").strip()
#         if txt in labels:
#             return txt
#     except Exception:
#         pass
#     return None

def _best_summary_label_openai(summary: str, labels: List[str]) -> Optional[str]:
    import os, json, requests
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key or not labels:
        return None
    
    # Create a clear prompt for matching
    prompt = (
        f"Given the following audit finding text, select the SINGLE best matching category from the list below.\n\n"
        f"Finding text:\n{summary}\n\n"
        f"Categories:\n" + "\n".join(f"- {label}" for label in labels) + "\n\n"
        f"Respond with ONLY the exact category text from the list above that best matches this finding."
    )
    
    try:
        r = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            data=json.dumps({
                "model": "gpt-4o-mini",
                "messages": [{"role":"user","content": prompt}],
                "temperature": 0
            }),
            timeout=12,
        )
        out = r.json()
        txt = (out.get("choices",[{}])[0].get("message",{}).get("content") or "").strip()
        if txt in labels:
            return txt
    except Exception:
        pass
    return None

# def upload_and_sas(container: str, blob_name: str, data: bytes, ttl_minutes: int = 120) -> str:
#     if not AZURE_CONN_STR:
#         raise RuntimeError("AZURE_STORAGE_CONNECTION_STRING not set")

#     info = _parse_conn_str(AZURE_CONN_STR)
#     account_name = info["AccountName"]
#     account_key  = info["AccountKey"]
#     blob_endpoint = info.get("BlobEndpoint")

#     bsc = _blob_service_client()
#     cc = bsc.get_container_client(container)
#     try:
#         cc.create_container()
#     except Exception:
#         pass
#     cc.upload_blob(name=blob_name, data=data, overwrite=True)

#     proto = "http" if (blob_endpoint and ("127.0.0.1" in blob_endpoint or "localhost" in blob_endpoint)) else None
#     sas = generate_blob_sas(
#         account_name=account_name,
#         account_key=account_key,
#         container_name=container,
#         blob_name=blob_name,
#         permission=BlobSasPermissions(read=True),
#         expiry=datetime.utcnow() + timedelta(minutes=ttl_minutes),
#         version=AZURITE_SAS_VERSION,
#         protocol=proto,
#     )
#     sas_q = quote(sas, safe="=&")

#     base = blob_endpoint.rstrip("/") if blob_endpoint else f"https://{account_name}.blob.core.windows.net"
#     return f"{base}/{container}/{blob_name}?{sas_q}"


from urllib.parse import quote
from datetime import datetime, timedelta
from azure.storage.blob import BlobSasPermissions, generate_blob_sas

def upload_and_sas(container: str, blob_name: str, data: bytes, ttl_minutes: int = 120) -> str:
    if not AZURE_CONN_STR:
        raise RuntimeError("AZURE_STORAGE_CONNECTION_STRING not set")

    info = _parse_conn_str(AZURE_CONN_STR)
    account_name  = info["AccountName"]
    account_key   = info["AccountKey"]
    blob_endpoint = info.get("BlobEndpoint")  # e.g. https://<acct>.blob.core.windows.net

    bsc = _blob_service_client()
    cc = bsc.get_container_client(container)
    try:
        cc.create_container()
    except Exception:
        pass

    cc.upload_blob(name=blob_name, data=data, overwrite=True)

    # Build SAS (no extra encoding)
    sas_kwargs = dict(
        account_name=account_name,
        account_key=account_key,
        container_name=container,
        blob_name=blob_name,
        permission=BlobSasPermissions(read=True),
        # allow 5 min clock skew
        start=datetime.utcnow() - timedelta(minutes=5),
        expiry=datetime.utcnow() + timedelta(minutes=ttl_minutes),
    )

    # Only force http/version when running against Azurite
    if blob_endpoint and ("127.0.0.1" in blob_endpoint or "localhost" in blob_endpoint):
        sas_kwargs["protocol"] = "http"               # ok for Azurite
        if AZURITE_SAS_VERSION:
            sas_kwargs["version"] = AZURITE_SAS_VERSION

    sas = generate_blob_sas(**sas_kwargs)

    base = blob_endpoint.rstrip("/") if blob_endpoint else f"https://{account_name}.blob.core.windows.net"
    # Important: DO NOT quote/encode the SAS. It is already correctly encoded.
    # Optionally quote the blob path in case of spaces or special chars.
    return f"{base}/{container}/{quote(blob_name, safe='/')}?{sas}"

def save_local_and_url(blob_name: str, data: bytes) -> str:
    full_path = os.path.join(LOCAL_SAVE_DIR, blob_name)
    os.makedirs(os.path.dirname(full_path), exist_ok=True)
    with open(full_path, "wb") as f:
        f.write(data)
    return f"{PUBLIC_BASE_URL}/local/{blob_name}"

def _title_with_acronyms(s: str, keep_all_caps=True) -> str:
    # simple title-caser with stop-words; preserves ALL-CAPS tokens and acronyms in ( )
    lowers = {"and","or","the","of","for","to","in","on","by","with","a","an"}
    parts = []
    for word in s.split():
        base = word.strip()
        if keep_all_caps and base.isupper() and len(base) > 1:
            parts.append(base)
        else:
            w = base.lower()
            if w in lowers:
                parts.append(w)
            else:
                parts.append(w.capitalize())
    return " ".join(parts)

# ------------------------------------------------------------------------------
# FAC helpers
# ------------------------------------------------------------------------------
def _fac_headers():
    key = os.getenv("FAC_API_KEY")
    if not key:
        raise HTTPException(500, "FAC_API_KEY not configured on the docx service")
    return {"X-Api-Key": key}

def _fac_get(path: str, params: Dict[str, Any]) -> Any:
    try:
        r = requests.get(f"{FAC_BASE.rstrip('/')}/{path.lstrip('/')}",
                         headers=_fac_headers(), params=params, timeout=20)
        r.raise_for_status()
        return r.json()
    except requests.HTTPError as e:
        raise HTTPException(r.status_code if 'r' in locals() else 500,
                            f"FAC GET {path} failed: {getattr(r,'text','')}") from e
    except Exception as e:
        raise HTTPException(500, f"FAC GET {path} failed: {e}") from e

def _or_param(field: str, values: List[str]) -> str:
    inner = ",".join([f"{field}.eq.{v}" for v in values])
    return f"({inner})"

# ------------------------------------------------------------------------------
# HTML → DOCX (preview)
# ------------------------------------------------------------------------------
def _apply_inline_formatting(paragraph, node):
    def add_text(text, bold=False, italic=False, underline=False):
        if text is None:
            return
        r = paragraph.add_run(text)
        r.bold = bool(bold)
        r.italic = bool(italic)
        r.underline = bool(underline)

    for child in getattr(node, "children", []):
        name = getattr(child, "name", None)
        if name is None:
            add_text(str(child))
        elif name in ("b", "strong"):
            add_text(child.get_text(), bold=True)
        elif name in ("i", "em"):
            add_text(child.get_text(), italic=True)
        elif name == "u":
            add_text(child.get_text(), underline=True)
        elif name == "br":
            paragraph.add_run().add_break()
        else:
            if hasattr(child, "children"):
                _apply_inline_formatting(paragraph, child)

def _basic_html_to_docx(doc: Document, html_str: str):
    soup = BeautifulSoup(html_str, "html.parser")
    body = soup.body or soup

    for element in body.children:
        if getattr(element, "name", None) is None:
            txt = str(element).strip()
            if txt:
                doc.add_paragraph(txt)
            continue

        tag = element.name.lower()
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(tag[1])
            text = element.get_text(strip=False)
            doc.add_heading(text, level=min(max(level, 1), 6))
            continue

        if tag == "p":
            p = doc.add_paragraph()
            _apply_inline_formatting(p, element)
            continue

        if tag in ("ul", "ol"):
            style = "List Bullet" if tag == "ul" else "List Number"
            for li in element.find_all("li", recursive=False):
                p = doc.add_paragraph(style=style)
                _apply_inline_formatting(p, li)
            continue

        if tag == "table":
            rows = element.find_all("tr", recursive=False)
            if not rows: continue
            first_cells = rows[0].find_all(["th","td"], recursive=False)
            cols = max(1, len(first_cells))
            first_is_header = any(c.name == "th" for c in first_cells)

            tbl = doc.add_table(rows=len(rows), cols=cols)

            try:
                tbl.style = "Table Grid"
            except Exception:
                pass
            _apply_grid_borders(tbl)

            sect = doc.sections[0]
            content_width = sect.page_width - sect.left_margin - sect.right_margin
            col_w = int(content_width / cols)
            if cols != 5:
                _set_col_widths(tbl, [col_w]*cols)

            for r_idx, tr in enumerate(rows):
                cells = tr.find_all(["th","td"], recursive=False)
                for c_idx in range(cols):
                    cell = tbl.cell(r_idx, c_idx)
                    if not cell.paragraphs:
                        cell.add_paragraph()
                    p = cell.paragraphs[0]
                    if hasattr(p, "clear"):
                        p.clear()
                    if c_idx < len(cells):
                        _apply_inline_formatting(p, cells[c_idx])
                    else:
                        p.text = ""
                    #_tight_paragraph(p)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            if first_is_header:
                for c in tbl.rows[0].cells:
                    #_shade_cell(c, "E7E6E6")
                    for r in c.paragraphs[0].runs:
                        r.bold = True
                    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
            
            

        if tag in ("div","section","article"):
            p = doc.add_paragraph()
            _apply_inline_formatting(p, element)
            continue

        txt = element.get_text(strip=True)
        if txt:
            doc.add_paragraph(txt)

def html_to_docx_bytes(html_str: str, *, force_basic: bool = False) -> bytes:
    doc = Document()
    if not force_basic:
        try:
            HTML2Docx().add_html_to_document(html_str or "", doc)
        except Exception:
            _basic_html_to_docx(doc, html_str or "")
    else:
        _basic_html_to_docx(doc, html_str or "")

    if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
        doc.add_paragraph("⚠️ HTML result is empty.")
    # ADD THIS:
    _set_font_size_to_12(doc)
    bio = BytesIO(); doc.save(bio)
    return bio.getvalue()

# ------------------------------------------------------------------------------
# MDL model & renderer (preview)
# ------------------------------------------------------------------------------
def summarize_finding_text(raw: str, max_chars: int = 1000) -> str:
    if not raw: return ""
    text = re.sub(r"\s+", " ", raw).strip()
    parts = re.split(r"(?<=[.?!])\s+", text)
    picked = []
    for p in parts:
        if len(picked) >= 3: break
        if re.search(r"\b(Assistance Listing|Award Period|Federal Program|Identification Number|CFDA)\b", p, re.I):
            continue
        picked.append(p)
    out = " ".join(picked) or text
    return _short(out, max_chars)

def format_letter_date(date_iso: Optional[str] = None) -> Tuple[str, str]:
    dt = datetime.fromisoformat(date_iso) if date_iso else datetime.utcnow()
    return dt.strftime("%Y-%m-%d"), dt.strftime("%B %d, %Y")

def _allcaps(s: str) -> str:
    return (s or "").strip().upper()

def _with_The_allcaps(name: str) -> str:
    # "The CITY OF ..." — ensure capital T + no double “The”
    raw = (name or "").strip()
    core = raw[4:].strip() if raw.lower().startswith("the ") else raw
    return f"The {_allcaps(core)}"

def _with_the_allcaps(name: str) -> str:
    # "the REHMANN ROBSON LLC" — ensure lowercase “the” + ALL CAPS entity
    raw = (name or "").strip()
    core = raw[4:].strip() if raw.lower().startswith("the ") else raw
    return f"the {_allcaps(core)}"

# def _add_hyperlink(paragraph, url, text):
#     """
#     Add a hyperlink to a paragraph.
#     """
#     # This gets access to the document.xml.rels file and gets a new relation id value
#     part = paragraph.part
#     r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

#     # Create the w:hyperlink tag and add needed values
#     hyperlink = OxmlElement('w:hyperlink')
#     hyperlink.set(qn('r:id'), r_id)

#     # Create a new run object (a wrapper over a w:r element)
#     new_run = OxmlElement('w:r')

#     # Set the run's style to Hyperlink style
#     rPr = OxmlElement('w:rPr')
#     rStyle = OxmlElement('w:rStyle')
#     rStyle.set(qn('w:val'), 'Hyperlink')
#     rPr.append(rStyle)
#     new_run.append(rPr)

#     # Add the text
#     new_run.text = text
#     hyperlink.append(new_run)

#     # Add the hyperlink to the paragraph
#     paragraph._p.append(hyperlink)

#     return hyperlink

def _add_hyperlink(paragraph, url, text):
    """
    Add a hyperlink element to a paragraph.
    Returns the hyperlink OxmlElement (does NOT append it).
    """
    # Get relationship ID
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create a new run for the hyperlink text
    new_run = OxmlElement('w:r')
    
    # Set run properties (color, underline)
    rPr = OxmlElement('w:rPr')
    
    # Color
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)
    
    # Underline
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    
    new_run.append(rPr)
    
    # Add the text
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink


# ============================================================
# CONFIGURATION
# ============================================================

class Config:
    FAC_API_BASE = os.getenv("FAC_API_BASE", "https://api.fac.gov")
    FAC_API_KEY = os.getenv("FAC_API_KEY", "")
    OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
    AZURE_CONN_STR = os.getenv("AZURE_STORAGE_CONNECTION_STRING", "")
    AZURE_CONTAINER = os.getenv("AZURE_BLOB_CONTAINER", "mdl-output")
    LOCAL_SAVE_DIR = os.getenv("LOCAL_SAVE_DIR", "./_out")
    PUBLIC_BASE_URL = os.getenv("PUBLIC_BASE_URL", "http://localhost:8000")
    MDL_TEMPLATE_PATH = os.getenv("MDL_TEMPLATE_PATH", "templates/MDL_Template.docx")
    TREASURY_EMAIL = os.getenv("TREASURY_CONTACT_EMAIL", "ORP_SingleAudits@treasury.gov")


# ============================================================
# STATIC DATA - From Comments & mdl_helpers.py
# ============================================================

# Treasury Programs (Comment #26, #27): ALN -> (Full Name, Acronym)
TREASURY_PROGRAMS = {
    "21.019": ("Coronavirus Relief Fund", "CRF"),
    "21.023": ("Emergency Rental Assistance Program", "ERA"),
    "21.026": ("Homeowner Assistance Fund", "HAF"),
    "21.027": ("Coronavirus State and Local Fiscal Recovery Funds", "SLFRF"),
    "21.029": ("Capital Projects Fund", "CPF"),
    "21.031": ("State Small Business Credit Initiative", "SSBCI"),
    "21.032": ("Local Assistance and Tribal Consistency Fund", "LATCF"),
}

# Compliance Types (Comment #32, #34): Letter -> Full Description
COMPLIANCE_TYPES = {
    "A": "Activities allowed or unallowed",
    "B": "Allowable costs/cost principles",
    "C": "Cash management",
    "E": "Eligibility",
    "F": "Equipment and real property management",
    "G": "Matching, level of effort, earmarking",
    "H": "Period of performance (or availability) of Federal funds",
    "I": "Procurement and suspension and debarment",
    "J": "Program income",
    "L": "Reporting",
    "M": "Subrecipient monitoring",
    "N": "Special tests and provisions",
    "P": "Other",
}

# Standard Finding Summaries (Comment #37, #38)
FINDING_SUMMARIES = [
    "Deficient Subrecipient Monitoring or Deficient Subaward",
    "Failure to file FFATA report for subawards",
    "Lack of evidence of competitive procurement",
    "Lack of evidence of suspension and debarment verification",
    "Lack of time and effort documentation",
    "Failure to retain adequate supporting documentation",
    "Inaccurate Treasury Reporting",
    "Lack of Eligibility Support",
    "Unallowable expenditures due to being incurred outside of period of performance",
    "Lack of Written Policies and/or Procedures - Management of Federal Funds",
    "Lack of Segregation of Duties",
    "Lack of Internal Controls - Grants Management",
]

# Keywords for classification fallback
FINDING_KEYWORDS = {
    "subrecipient": "Deficient Subrecipient Monitoring or Deficient Subaward",
    "subaward": "Deficient Subrecipient Monitoring or Deficient Subaward",
    "sub-recipient": "Deficient Subrecipient Monitoring or Deficient Subaward",
    "ffata": "Failure to file FFATA report for subawards",
    "procurement": "Lack of evidence of competitive procurement",
    "competitive bid": "Lack of evidence of competitive procurement",
    "sole source": "Lack of evidence of competitive procurement",
    "bid": "Lack of evidence of competitive procurement",
    "suspension and debarment": "Lack of evidence of suspension and debarment verification",
    "sam.gov": "Lack of evidence of suspension and debarment verification",
    "debarment": "Lack of evidence of suspension and debarment verification",
    "suspended": "Lack of evidence of suspension and debarment verification",
    "debarred": "Lack of evidence of suspension and debarment verification",
    "time and effort": "Lack of time and effort documentation",
    "timesheet": "Lack of time and effort documentation",
    "labor cost": "Lack of time and effort documentation",
    "personnel": "Lack of time and effort documentation",
    "documentation": "Failure to retain adequate supporting documentation",
    "supporting documentation": "Failure to retain adequate supporting documentation",
    "records": "Failure to retain adequate supporting documentation",
    "treasury report": "Inaccurate Treasury Reporting",
    "quarterly report": "Inaccurate Treasury Reporting",
    "project and expenditure": "Inaccurate Treasury Reporting",
    "p&e report": "Inaccurate Treasury Reporting",
    "period of performance": "Unallowable expenditures due to being incurred outside of period of performance",
    "outside the period": "Unallowable expenditures due to being incurred outside of period of performance",
    "eligibility": "Lack of Eligibility Support",
    "eligible": "Lack of Eligibility Support",
    "internal control": "Lack of Internal Controls - Grants Management",
    "policies and procedures": "Lack of Written Policies and/or Procedures - Management of Federal Funds",
    "written policies": "Lack of Written Policies and/or Procedures - Management of Federal Funds",
    "segregation of duties": "Lack of Segregation of Duties",
    "segregation": "Lack of Segregation of Duties",
}


# ============================================================
# FORMATTING UTILITIES (Based on Comments #1, #5, #14, #16)
# ============================================================

def to_standard_case(name: str) -> str:
    """
    Convert ALL CAPS to Standard Case (Comment #1: "Everything must be in standard case")
    Preserves known acronyms like LLC, LLP, etc.
    """
    if not name:
        return ""
    name = name.strip()
    
    if not name.isupper():
        return name  # Already mixed case, preserve
    
    # Known acronyms to preserve
    acronyms = {"LLC", "LLP", "PC", "PA", "CPA", "USA", "US", "II", "III", "IV"}
    # Words to keep lowercase (except at start)
    lowercase_words = {"and", "or", "the", "of", "for", "to", "in", "on", "by", "with", "a", "an"}
    
    words = []
    for i, word in enumerate(name.split()):
        word_upper = word.upper()
        word_lower = word.lower()
        
        if word_upper in acronyms:
            words.append(word_upper)
        elif word_lower in lowercase_words and i > 0:
            words.append(word_lower)
        else:
            words.append(word.capitalize())
    
    return " ".join(words)


def format_ein(ein: str) -> str:
    """
    Format EIN as XX-XXXXXXX (Comment #5: "Must be XX-XXXXXXX format")
    """
    ein = (ein or "").replace("-", "").replace(" ", "").strip()
    if len(ein) == 9 and ein.isdigit():
        return f"{ein[:2]}-{ein[2:]}"
    return ein


#def format_date(date_str: Optional[str]) -> str:
def render_mdl_html(model: Dict[str, Any]) -> str:
    letter_date_iso = model.get("letter_date_iso")
    _, letter_date_long = format_letter_date(letter_date_iso)

    auditee_name = model.get("auditee_name", "Recipient")
    ein = model.get("ein", "")
    address_lines = model.get("address_lines", [])
    attention_line = model.get("attention_line")
    period_end_text = model.get("period_end_text", str(model.get("audit_year", "")))
    include_no_qc_line = model.get("include_no_qc_line", True)
    treasury_contact_email = model.get("treasury_contact_email", "ORP_SingleAudits@treasury.gov ")
    address_block = "<br>".join(html.escape(x) for x in address_lines) if address_lines else ""
    attention_block = f"<p><strong>{html.escape(attention_line)}</strong></p>" if attention_line else ""

    def _render_program_table(p: Dict[str, Any]) -> str:
        rows_html = []
        for f in p.get("findings", []):
            rows_html.append(f"""
              <tr>
                <td>{html.escape(f.get('finding_id',''))}</td>
                <td>{html.escape(_combine_comp_summary(f))}</td>
                <td>{html.escape(f.get('audit_determination',''))}</td>
                <td>{html.escape(f.get('questioned_cost_determination',''))}</td>
                <td>{html.escape(f.get('cap_determination',''))}</td>
              </tr>
            """)
        if not rows_html:
            rows_html.append("""
              <tr><td colspan="5"><em>No MDL-relevant findings identified for this program.</em></td></tr>
            """)
        table = f"""
          <table border="1" cellspacing="0" cellpadding="6" style="border-collapse:collapse; width:100%; font-size:10.5pt;">
            <tr>
              <th>Audit<br>Finding #</th>
              <th>Compliance Type -<br>Audit Finding Summary</th>
              <th>Audit Finding<br>Determination</th>
              <th>Questioned Cost<br>Determination</th>
              <th>CAP<br>Determination</th>
            </tr>
            {''.join(rows_html)}
          </table>
        """
        cap_blocks = []
        for f in p.get("findings", []):
            cap_text = f.get("cap_text")
            if cap_text:
                cap_blocks.append(f"""
                  <h4>Corrective Action Plan – {html.escape(f.get('finding_id',''))}</h4>
                  <p>{html.escape(cap_text)}</p>
                """)
        return table + ("\n".join(cap_blocks) if cap_blocks else "")

    programs = model.get("programs", [])
    programs_html = "\n".join(_render_program_table(p) for p in programs) if programs else "<p><em>No MDL-relevant findings identified per FAC records.</em></p>"

    not_sustained_notes = model.get("not_sustained_notes", [])
    not_sustained_html = ""
    if not_sustained_notes:
        notes_paras = "\n".join(f"<p>{html.escape(n)}</p>" for n in not_sustained_notes if n)
        not_sustained_html = f"<h3>FINDINGS NOT SUSTAINED</h3>\n{notes_paras}"

    chunks = []
    chunks.append(f'<p style="text-align:right; margin:0 0 12pt 0;">{html.escape(letter_date_long)}</p>')
    chunks.append("""
      <p style="margin:0 0 12pt 0;">
        <strong>DEPARTMENT OF THE TREASURY</strong><br>
        WASHINGTON, D.C.
      </p>
    """)
    chunks.append(f"""
      <p style="margin:0 0 12pt 0;">
        <strong>{html.escape(_allcaps(auditee_name))}</strong><br>
        EIN: {html.escape(ein)}<br>
        {address_block}
      </p>
    """)
    if attention_block:
        chunks.append(attention_block)
    chunks.append(f"""
      <p style="margin:12pt 0 12pt 0;">
        <strong>Subject:</strong> U.S. Department of the Treasury’s Management Decision Letter (MDL) for Single Audit Report for the period ending on {html.escape(period_end_text)}
      </p>
    """)
    chunks.append("""
      <p>
        In accordance with 2 C.F.R. § 200.521(b), the U.S. Department of the Treasury (Treasury)
        is required to issue a management decision for single audit findings pertaining to awards under
        Treasury’s programs. Treasury’s review as part of its responsibilities under 2 C.F.R § 200.513(c)
        includes an assessment of Treasury’s award recipients’ single audit findings, corrective action plans (CAPs),
        and questioned costs, if any.
      </p>
    """)
    # chunks.append(f"""
    #   <p>
    #     Treasury has reviewed the single audit report for {html.escape(auditee_name)}.
    #     Treasury has made the following determinations regarding the audit finding(s) and CAP(s) listed below.
    #   </p>
    # """)
    chunks.append(f"""
    <p>
        Treasury has reviewed the single audit report for {html.escape(_with_The_allcaps(auditee_name))},
        prepared by {html.escape(_with_the_allcaps(model.get("auditor_name","")))} for the fiscal year ending on
        {html.escape(period_end_text)}.
        Treasury has made the following determinations regarding the audit finding(s) and CAP(s) listed below.
    </p>
    """)
    if include_no_qc_line:
        chunks.append("<p>No questioned costs are included in this single audit report.</p>")

    chunks.append(programs_html)
    if not_sustained_html:
        chunks.append(not_sustained_html)

    # Email sentence removed per feedback.
    chunks.append("""
      <p>
        Please note, the corrective action(s) are subject to review during the recipient’s next annual single audit
        or program-specific audit, as applicable, to determine adequacy. If the same audit finding(s) appear in a future single
        audit report for this recipient, its current or future award funding under Treasury’s programs may be adversely impacted.
      </p>
      <p>
        For questions regarding the audit finding(s), please email us at
        <a href="mailto:{html.escape(treasury_contact_email)}">{html.escape(treasury_contact_email)}</a>.
        Thank you.
    </p>
      <p style="margin-top:18pt;">Sincerely,<br><br>
      Audit and Compliance Resolution Team<br>
      Office of Capital Access<br>
      U.S. Department of the Treasury</p>
    """)

    return f'<div style="font-family: Calibri, Arial, sans-serif; font-size:11pt; line-height:1.4;">{"".join(chunks)}</div>'

def build_mdl_model_from_fac(
    *,
    auditee_name: str,
    ein: str,
    audit_year: int,
    fac_general: List[Dict[str, Any]],
    fac_findings: List[Dict[str, Any]],
    fac_findings_text: List[Dict[str, Any]],
    fac_caps: List[Dict[str, Any]],
    federal_awards: List[Dict[str, Any]],
    period_end_text: Optional[str] = None,
    address_lines: Optional[List[str]] = None,
    attention_line: Optional[str] = None,
    only_flagged: bool = False,
    max_refs: int = 10,
    auto_cap_determination: bool = True,
    include_no_qc_line: bool = False,
    include_no_cap_line: bool = False,
    treasury_listings: Optional[List[str]] = None,
    aln_reference_xlsx: Optional[str] = None,
    aln_overrides_by_finding: Optional[Dict[str, str]] = None,
    **_
) -> Dict[str, Any]:
    """
    Builds the MDL model.

    Point C addressed:
      - Ensures ALN is not "Unknown" by:
        1) Loading Excel mapping (ALN -> canonical name; name -> ALN).
        2) Canonicalizing each grouped program after grouping:
           - If ALN known, use canonical label.
           - Else map by program name -> ALN.
           - Else apply Treasury heuristics (SLFRF/ERA/HAF/CPF/SSBCI/LATCF).
        3) THEN apply treasury_listings filter.
    """
        # Temporary hardcoded mapping for testing
    # type_map = {
    #     "A": "Activities allowed or unallowed",
    #     "B": "Allowable costs/cost principles",
    #     "C": "Cash management",
    #     "E": "Eligibility",
    #     "F": "Equipment and real property management",
    #     "G": "Matching, level of effort, earmarking",
    #     "H": "Period of performance (or availability) of Federal funds",
    #     "I": "Procurement and suspension and debarment",
    #     "J": "Program income",
    #     "L": "Reporting",
    #     "M": "Subrecipient monitoring",
    #     "N": "Special tests and provisions",
    #     "P": "Other"
    # }
    # # Still try to load from file (will override if successful)
    # if aln_reference_xlsx:
    #     #loaded_type_map, summary_labels = _load_finding_mappings(aln_reference_xlsx)
    #     if loaded_type_map:
    #         type_map = loaded_type_map
    #     logging.info(f"Loaded {len(summary_labels)} summary labels from Excel")
    # else:
    #     summary_labels = finding_summaries_list  # Use the one from mdl_helpers.py
    #     logging.warning("No aln_reference_xlsx provided, using hardcoded type_map")
    
    # logging.info(f"Final type_map: {type_map}")
    # logging.info(f"Looking up 'I': {type_map.get('I')}")

    # ========== LOAD MAPPINGS ONCE AT THE TOP ==========
    # Default hardcoded mappings as fallback
    type_map = {
        "A": "Activities allowed or unallowed",
        "B": "Allowable costs/cost principles",
        "C": "Cash management",
        "E": "Eligibility",
        "F": "Equipment and real property management",
        "G": "Matching, level of effort, earmarking",
        "H": "Period of performance (or availability) of Federal funds",
        "I": "Procurement and suspension and debarment",
        "J": "Program income",
        "L": "Reporting",
        "M": "Subrecipient monitoring",
        "N": "Special tests and provisions",
        "P": "Other"
    }
    summary_labels = finding_summaries_list  # from mdl_helpers.py
    
    # Try to load from Excel (only if path provided and file exists)
    if aln_reference_xlsx:
        loaded_type_map, loaded_summary_labels = _load_finding_mappings(aln_reference_xlsx)
        if loaded_type_map:  # Only override if we got data
            type_map = loaded_type_map
            logging.info(f"✅ Loaded type_map from Excel with {len(loaded_type_map)} entries")
        if loaded_summary_labels:
            summary_labels = loaded_summary_labels
            logging.info(f"✅ Loaded {len(loaded_summary_labels)} summary labels from Excel")
    
    logging.info(f"Final type_map: {type_map}")
    logging.info(f"Looking up 'I': {type_map.get('I')}")
    # --------- helpers ----------
    def _derive_assistance_listing(program_name: str, fallback: str = "Unknown") -> str:
        m = re.search(r"\b\d{2}\.\d{3}\b", program_name or "")
        return m.group(0) if m else fallback

    def _title_with_acronyms(s: str) -> str:
        """Title-case but preserve ALL-CAPS tokens (e.g., SLFRF) and common stop words."""
        if not s:
            return ""
        lowers = {"and","or","the","of","for","to","in","on","by","with","a","an"}
        out = []
        for tok in str(s).split():
            if tok.isupper() and len(tok) > 1:
                out.append(tok)  # keep acronym
            else:
                w = tok.lower()
                out.append(w if w in lowers else w.capitalize())
        return " ".join(out)

    def _load_aln_mapping(xlsx_path: Optional[str]):
        """
        Returns:
          aln_to_label: {'21.027': 'Coronavirus State and Local Fiscal Recovery Funds (SLFRF)', ...}
          name_to_aln:  {'coronavirus state and local fiscal recovery funds': '21.027', ...}
        Tolerant to header naming; safe no-op if workbook missing or openpyxl not installed.
        """
        aln_to_label, name_to_aln = {}, {}
        if not xlsx_path:
            return aln_to_label, name_to_aln
        try:
            import openpyxl
        except Exception:
            return aln_to_label, name_to_aln
        try:
            wb = openpyxl.load_workbook(xlsx_path, data_only=True)
        except Exception:
            return aln_to_label, name_to_aln

        # Heuristic: find the first sheet that looks like a mapping
        for ws in wb.worksheets:
            hdrs = [(c.value.strip() if isinstance(c.value, str) else (c.value or "")) for c in ws[1]]
            if not any(hdrs):
                continue
            hl = [str(h).strip().lower() for h in hdrs]

            def _find_col(candidates: List[str]) -> Optional[int]:
                for i, h in enumerate(hl):
                    for cand in candidates:
                        cand = cand.lower()
                        if h == cand or cand in h:
                            return i
                return None

            i_aln = _find_col(["aln", "assistance listing", "assistance listing number", "cfda", "cfda number"])
            i_prog = _find_col(["program", "program name", "assistance listing title", "program title"])
            i_acr = _find_col(["acronym", "short", "short name", "abbrev", "abbreviation"])

            if i_prog is None or (i_aln is None and i_acr is None):
                continue  # not a mapping sheet

            for row in ws.iter_rows(min_row=2, values_only=True):
                raw_aln = (row[i_aln] if i_aln is not None else "") or ""
                raw_prog = (row[i_prog] or "")
                raw_acr = (row[i_acr] if i_acr is not None else "") or ""

                aln = str(raw_aln).strip()
                prog_name = str(raw_prog).strip()
                acr = str(raw_acr).strip()
                if not prog_name:
                    continue

                canonical_name = _title_with_acronyms(prog_name)
                if acr:
                    canonical_name = f"{canonical_name} ({acr})"

                if aln:
                    aln_to_label[aln] = canonical_name
                name_to_aln[prog_name.lower()] = aln  # may be "" if not provided
            break  # first matching sheet is enough

        return aln_to_label, name_to_aln

    def _apply_canonicalization_after_grouping(
        group: Dict[str, Any],
        aln_to_label: Dict[str, str],
        name_to_aln: Dict[str, str],
    ):
        """
        Normalize 'assistance_listing' and 'program_name' in-place for a group:
          - Use Excel ALN→label if ALN known.
          - Else map by name→ALN.
          - Else apply Treasury heuristics.
          - Always tidy program name casing.
        """
        cur_aln  = (group.get("assistance_listing") or "").strip()
        cur_name = (group.get("program_name") or "").strip()

        # 1) If we already have a valid ALN, use canonical label
        if cur_aln and cur_aln in aln_to_label:
            group["assistance_listing"] = cur_aln
            group["program_name"] = aln_to_label[cur_aln]
            return

        # 2) If ALN missing/Unknown, try via name → ALN
        guess_aln = name_to_aln.get(cur_name.lower())
        if (not cur_aln or cur_aln == "Unknown") and guess_aln:
            group["assistance_listing"] = guess_aln
            group["program_name"] = aln_to_label.get(guess_aln, _title_with_acronyms(cur_name or "Unknown Program"))
            return

        # 3) Treasury heuristics (common programs) — last resort
        nm = cur_name.lower()
        heuristics = [
            ("slfrf", ("21.027", "Coronavirus State and Local Fiscal Recovery Funds (SLFRF)")),  # noqa: E501,
            ("fiscal recovery", ("21.027", "Coronavirus State and Local Fiscal Recovery Funds (SLFRF)")),
            ("emergency rental assistance", ("21.023", "Emergency Rental Assistance Program (ERA)")),
            ("homeowner assistance", ("21.026", "Homeowner Assistance Fund (HAF)")),
            ("capital projects fund", ("21.029", "Capital Projects Fund (CPF)")),
            ("state small business credit", ("21.031", "State Small Business Credit Initiative (SSBCI)")),
            ("local assistance and tribal consistency", ("21.032", "Local Assistance and Tribal Consistency Fund (LATCF)")),
        ]
        for key, (aln_guess, label) in heuristics:
            if key in nm:
                group["assistance_listing"] = aln_guess
                group["program_name"] = label
                break

        # 4) Final tidy for program name casing if still raw/all-caps
        final_name = (group.get("program_name") or "").strip()
        if final_name.isupper() or not final_name:
            group["program_name"] = _title_with_acronyms(final_name or cur_name or "Unknown Program")
        if not group.get("assistance_listing"):
            group["assistance_listing"] = "Unknown"

    # --------- load mapping once ----------
    aln_to_label, name_to_aln = _load_aln_mapping(aln_reference_xlsx)
    # ADD THIS:
    #type_map, summary_labels = _load_finding_mappings(aln_reference_xlsx)
    logging.info(f"Loaded type_map: {type_map}")
    logging.info(f"Loaded {len(summary_labels)} summary labels")
    # --------- award lookups from FAC ----------
    award2meta: Dict[str, Dict[str, str]] = {}
    for a in (federal_awards or []):
        ref = a.get("award_reference")
        pname = (a.get("federal_program_name") or "").strip()
        explicit_aln = (a.get("assistance_listing") or a.get("assistance_listing_number") or "").strip()
        derived_aln = _derive_assistance_listing(pname, fallback="")
        aln = explicit_aln or derived_aln or "Unknown"

        # Prefer Excel canonical if we have the ALN
        if aln != "Unknown" and aln in aln_to_label:
            canonical_name = aln_to_label[aln]
        else:
            # Try map by name -> ALN
            mapped_aln = name_to_aln.get(pname.lower())
            if mapped_aln:
                aln = mapped_aln
                canonical_name = aln_to_label.get(mapped_aln, _title_with_acronyms(pname or "Unknown Program"))
            else:
                canonical_name = _title_with_acronyms(pname or "Unknown Program")

        if ref:
            award2meta[ref] = {
                "program_name": canonical_name or "Unknown Program",
                "assistance_listing": aln or "Unknown",
            }
    # ADD THIS DEBUG:
    logging.info(f"📋 Built award2meta with {len(award2meta)} entries:")
    for k, v in award2meta.items():
        logging.info(f" AWARDSSSSSSSSSSS {k}: {v}")
    # --------- text / CAP lookups ----------
    text_by_ref = {
        _norm_ref(t.get("finding_ref_number")): (t.get("finding_text") or "").strip()
        for t in (fac_findings_text or [])
    }
    cap_by_ref = {
        _norm_ref(c.get("finding_ref_number")): (c.get("planned_action") or "").strip()
        for c in (fac_caps or [])
    }

    def _is_flagged(f: dict) -> bool:
        return any([
            f.get("is_material_weakness") is True,
            f.get("is_significant_deficiency") is True,
            f.get("is_questioned_costs") is True,
            f.get("is_modified_opinion") is True,
            f.get("is_other_findings") is True,
            f.get("is_other_matters") is True,
            f.get("is_repeat_finding") is True,
        ])

    # --------- choose finding refs ----------
    base_refs: List[str] = []
    for f in (fac_findings or []):
        if only_flagged and not _is_flagged(f):
            continue
        r = f.get("reference_number")
        if r:
            base_refs.append(r)

    if not base_refs and fac_findings_text:
        base_refs = [t.get("finding_ref_number") for t in fac_findings_text if t.get("finding_ref_number")]

    seen = set()
    norm_refs: List[Tuple[str, str]] = []
    for r in base_refs:
        if not r:
            continue
        k = _norm_ref(r)
        if k not in seen:
            seen.add(k)
            norm_refs.append((r, k))
    norm_refs = norm_refs[: max_refs or 10]
    chosen_keys = {kn for _, kn in norm_refs}

    # --------- group findings under award_reference ----------
    programs_map: Dict[str, Dict[str, Any]] = {}
    for f in (fac_findings or []):
        r = f.get("reference_number")
        if not r:
            continue
        k = _norm_ref(r)
        if k not in chosen_keys:
            continue

        award_ref = f.get("award_reference") or "UNKNOWN"
        logging.info(f"🔍 Finding {r} → award_ref: {award_ref}")
        
        # Try to get metadata from award lookup
        meta = award2meta.get(award_ref, {})
        
        # If not found in award2meta, try aln_overrides_by_finding
        if not meta.get("assistance_listing") or meta.get("assistance_listing") == "Unknown":
            if aln_overrides_by_finding and r in aln_overrides_by_finding:
                override_aln = aln_overrides_by_finding[r]
                logging.info(f"   ✅ Using finding-level ALN override: {override_aln}")
                meta["assistance_listing"] = override_aln
                # Update program name if we have ALN mapping
                if override_aln in aln_to_label:
                    meta["program_name"] = aln_to_label[override_aln]
        
        logging.info(f"   Final meta: {meta}")

        group = programs_map.setdefault(award_ref, {
            "assistance_listing": meta.get("assistance_listing", "Unknown"),
            "program_name": meta.get("program_name", "Unknown Program"),
            "findings": []
        })
        
        # If ALN is Unknown, try to fill from finding-level override (XLSX)
        if group.get("assistance_listing") in (None, "", "Unknown"):
            # use the original finding reference if available
            orig_ref = f.get("reference_number") or ""
            # try both raw and normalized keys
            cand_aln = None
            if aln_overrides_by_finding:
                cand_aln = (aln_overrides_by_finding.get(orig_ref)
                            or aln_overrides_by_finding.get(_norm_ref(orig_ref)))
            if cand_aln:
                group["assistance_listing"] = cand_aln
                # if we have an ALN→label map, upgrade program_name too
                if cand_aln in aln_to_label:
                    group["program_name"] = aln_to_label[cand_aln]

        summary  = summarize_finding_text(text_by_ref.get(k, ""))
        cap_text = cap_by_ref.get(k)

        #qcost_det = "No questioned costs identified" if include_no_qc_line else "None"
        qcost_det = "Questioned Cost:\nNone\nDisallowed Cost:\nNone" if include_no_qc_line else "None"
        cap_det   = (
            "Accepted" if (auto_cap_determination and cap_text)
            else ("No CAP required" if include_no_cap_line else "Not Applicable")
        )

        # ctype_code = (f.get("type_requirement") or "").strip().upper()[:1]
        # ctype_label = type_map.get(ctype_code) or ctype_code or ""
        # matched_label = (_best_summary_label_openai(summary, summary_labels)
        #          or _best_summary_label(summary, summary_labels)
        #          or summary)
        # Get the full compliance type label (e.g., "Procurement and suspension and debarment")
        ctype_code = (f.get("type_requirement") or "").strip().upper()[:1]
        ctype_label = type_map.get(ctype_code) or ctype_code or ""
        
        # Get the complete finding summary text
        complete_summary = text_by_ref.get(k, "")
        
        # Match against standardized summaries - try OpenAI first with COMPLETE text
        matched_label = None
        if complete_summary:
            matched_label = (_best_summary_label_openai(complete_summary, summary_labels)
                           or _best_summary_label(complete_summary, summary_labels))
        
        # Fallback to shortened summary if no match
        if not matched_label:
            matched_label = summary
        
        logging.info(f"Finding {f.get('reference_number')}: {ctype_label} - {matched_label}")
        logging.info(f"Matched label: {matched_label}")
        #print("\n")
        logging.info(f"Compliance type: {ctype_label}")
        logging.info(f"Summary: {summary}")
        #logging.info(" for finding {f.get('reference_number')}")
        logging.info(f" {ctype_label}, {summary}, {cap_text}, {qcost_det}, {cap_det}")
        # group["findings"].append({
        #     "finding_id": f.get("reference_number") or "",
        #     "compliance_type": ctype_label,  # use the full label, not just 'I'
        #     "summary": summary,
        #     "compliance_and_summary": f"{ctype_label} - {matched_label}".strip(" -"),
        #     "audit_determination": "Sustained",
        #     "questioned_cost_determination": qcost_det,
        #     "disallowed_cost_determination": "None",
        #     "cap_determination": cap_det,
        #     "cap_text": cap_text,
        # })
        group["findings"].append({
            "finding_id": f.get("reference_number") or "",
            "compliance_type": ctype_label,  # Full label: "Procurement and suspension and debarment"
            "summary": matched_label,  # Matched standardized summary
            "compliance_and_summary": f"{ctype_label} - {matched_label}".strip(" -"),  # Combined for display
            "audit_determination": "Sustained",
            "questioned_cost_determination": qcost_det,
            "disallowed_cost_determination": "None",
            "cap_determination": cap_det,
            "cap_text": cap_text,
        })
    # If nothing grouped but we have refs, emit a catch-all
    # if not programs_map and norm_refs:
    #     catchall = {"assistance_listing": "Unknown", "program_name": "Unknown Program", "findings": []}
    #     ctype_code = (f.get("type_requirement") or "").strip().upper()[:1]
    #     ctype_label = type_map.get(ctype_code) or ctype_code or ""
    #     matched_label = _best_summary_label(summary, summary_labels) or summary
    #     for orig, key in norm_refs:
    #         cap_text = cap_by_ref.get(key)
    #         qcost_det = "No questioned costs identified" if include_no_qc_line else "None"
    #         cap_det   = (
    #             "Accepted" if (auto_cap_determination and cap_text)
    #             else ("No CAP required" if include_no_cap_line else "Not Applicable")
    #         )
    #         catchall["findings"].append({
    #             "finding_id": orig,
    #             "compliance_type": ctype_label, # use the full label not just 'I'
    #             "summary": summarize_finding_text(text_by_ref.get(key, "")),
    #             "compliance_and_summary": f"{ctype_label} - {matched_label}".strip(" -"),
    #             "audit_determination": "Sustained",
    #             "questioned_cost_determination": qcost_det,
    #             "disallowed_cost_determination": "None",
    #             "cap_determination": cap_det,
    #             "cap_text": cap_text,
    #         })
    #     programs_map["UNKNOWN"] = catchall
    # If nothing grouped but we have refs, emit a catch-all
    if not programs_map and norm_refs:
        catchall = {"assistance_listing": "Unknown", "program_name": "Unknown Program", "findings": []}
        
        for orig, key in norm_refs:
            # Get compliance type for THIS finding
            # Find the finding in fac_findings to get its type_requirement
            finding_data = None
            for f in (fac_findings or []):
                if _norm_ref(f.get("reference_number")) == key:
                    finding_data = f
                    break
            
            if finding_data:
                ctype_code = (finding_data.get("type_requirement") or "").strip().upper()[:1]
            else:
                ctype_code = ""
            
            ctype_label = type_map.get(ctype_code) or ctype_code or ""
            
            # Get complete finding text and match
            complete_summary = text_by_ref.get(key, "")
            matched_label = None
            if complete_summary:
                matched_label = (_best_summary_label_openai(complete_summary, summary_labels)
                            or _best_summary_label(complete_summary, summary_labels))
            if not matched_label:
                matched_label = summarize_finding_text(complete_summary)
            
            cap_text = cap_by_ref.get(key)
            qcost_det = "No questioned costs identified" if include_no_qc_line else "None"
            cap_det   = (
                "Accepted" if (auto_cap_determination and cap_text)
                else ("No CAP required" if include_no_cap_line else "Not Applicable")
            )
            
            catchall["findings"].append({
                "finding_id": orig,
                "compliance_type": ctype_label,  # Full label like "Procurement and suspension and debarment"
                "summary": matched_label,  # Matched standardized summary
                "compliance_and_summary": f"{ctype_label} - {matched_label}".strip(" -"),
                "audit_determination": "Sustained",
                "questioned_cost_determination": qcost_det,
                "disallowed_cost_determination": "None",
                "cap_determination": cap_det,
                "cap_text": cap_text,
            })
        programs_map["UNKNOWN"] = catchall

    # ----- Canonicalize program fields using the Excel mapping (now that groups exist)
    for grp in programs_map.values():
        _apply_canonicalization_after_grouping(grp, aln_to_label, name_to_aln)

    # ----- Apply Treasury ALN filter AFTER canonicalization
    if treasury_listings:
        allowed = {(aln or "").strip() for aln in treasury_listings if aln}
        logging.info(f" Treasury listings filter: {allowed}")
        logging.info(f"  Programs before filter: {list(programs_map.keys())}")
        programs_map = {k: v for k, v in programs_map.items() if v.get("assistance_listing") in allowed}
        logging.info(f" Programs after filter: {list(programs_map.keys())}")
    # ----- Build final model
    model = {
        "letter_date_iso": datetime.utcnow().strftime("%Y-%m-%d"),
        "auditee_name": auditee_name,
        "ein": f"{ein[:2]}-{ein[2:]}" if ein and ein.isdigit() and len(ein) == 9 else ein,
        "address_lines": address_lines or [],
        "attention_line": attention_line or "",
        "period_end_text": period_end_text or f"June 30, {audit_year}",
        "audit_year": audit_year,
        "programs": list(programs_map.values()),
        "not_sustained_notes": [],
    }
    return model

# ------------------------------------------------------------------------------
# Template helpers
# ------------------------------------------------------------------------------
def _para_text(p: Paragraph) -> str:
    return "".join(run.text for run in p.runs)

def _clear_runs(p: Paragraph):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)

def _replace_in_paragraph_run_aware(p: Paragraph, mapping: Dict[str, str]) -> bool:
    original = _para_text(p)
    if not original:
        return False
    new_text = original
    for k, v in mapping.items():
        if k in new_text:
            new_text = new_text.replace(k, v)
    if new_text != original:
        _clear_runs(p)
        p.add_run(new_text)
        return True
    return False

def _iter_cells_in_table(tbl: Table):
    for row in tbl.rows:
        for cell in row.cells:
            yield cell

def _iter_all_paragraphs_in_container(container) -> list[Paragraph]:
    items = []
    if hasattr(container, "paragraphs"):
        items.extend(container.paragraphs)
    if hasattr(container, "tables"):
        for t in container.tables:
            for c in _iter_cells_in_table(t):
                items.extend(c.paragraphs)
                for nt in c.tables:
                    for nc in _iter_cells_in_table(nt):
                        items.extend(nc.paragraphs)
    return items

def _replace_placeholders_docwide(doc: Document, mapping: Dict[str, str]):
    for p in _iter_all_paragraphs_in_container(doc):
        _replace_in_paragraph_run_aware(p, mapping)
    for sec in doc.sections:
        for p in _iter_all_paragraphs_in_container(sec.header):
            _replace_in_paragraph_run_aware(p, mapping)
        for p in _iter_all_paragraphs_in_container(sec.footer):
            _replace_in_paragraph_run_aware(p, mapping)

def _find_anchor_paragraph(doc: Document, anchor: str) -> Optional[Paragraph]:
    for p in _iter_all_paragraphs_in_container(doc):
        if anchor in _para_text(p):
            return p
    for sec in doc.sections:
        for p in _iter_all_paragraphs_in_container(sec.header):
            if anchor in _para_text(p):
                return p
        for p in _iter_all_paragraphs_in_container(sec.footer):
            if anchor in _para_text(p):
                return p
    return None

def _delete_immediate_next_table(anchor_para: Paragraph):
    """If the template has a placeholder table immediately after the anchor paragraph, delete it."""
    p_el = anchor_para._p
    nxt = p_el.getnext()
    if nxt is not None and nxt.tag.endswith("tbl"):
        parent = nxt.getparent()
        parent.remove(nxt)

def _pick_table_style(doc: Document) -> Optional[str]:
    if getattr(doc, "tables", None):
        for t in doc.tables:
            try:
                if t.style and t.style.name:
                    return t.style.name
            except Exception:
                pass
    try:
        _ = doc.styles["Table Grid"]
        return "Table Grid"
    except KeyError:
        pass
    for st in doc.styles:
        try:
            if st.type == WD_STYLE_TYPE.TABLE:
                return st.name
        except Exception:
            continue
    return None

# def _build_program_table(doc: Document, program: Dict[str, Any]) -> Table:
#     findings = program.get("findings", []) or []
#     rows = max(1, len(findings)) + 1

#     tbl = doc.add_table(rows=rows, cols=5)
#     _style = _pick_table_style(doc)
#     if _style:
#         try:
#             tbl.style = _style
#         except Exception:
#             pass
#     _apply_grid_borders(tbl)  # ensure borders even without style

#     headers = [
#         "Audit\nFinding #",
#         "Compliance Type -\nAudit Finding",
#         "Audit Finding\nDetermination",
#         "Questioned Cost\nDetermination",
#         "CAP\nDetermination",
#     ]
#     for i, h in enumerate(headers):
#         cell = tbl.cell(0, i)
#         _clear_runs(cell.paragraphs[0])
#         cell.paragraphs[0].add_run(h)
#         _shade_cell(cell, "E7E6E6")
#         cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
#         cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
#         _tight_paragraph(cell.paragraphs[0])

#     if findings:
#         for r, f in enumerate(findings, start=1):
#             vals = [
#                 f.get("finding_id", ""),
#                 f.get("compliance_and_summary", ""),  # ← Use the combined field
#                 f.get("audit_determination", "Sustained"),
#                 f.get("questioned_cost_determination", "None"),
#                 f.get("cap_determination", "Not Applicable"),
#             ]
#             for c, val in enumerate(vals):
#                 cell = tbl.cell(r, c)
#                 _clear_runs(cell.paragraphs[0])
#                 cell.paragraphs[0].add_run(str(val))
#                 cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
#                 _tight_paragraph(cell.paragraphs[0])
#     else:
#         cell = tbl.cell(1, 0)
#         _clear_runs(cell.paragraphs[0])
#         cell.paragraphs[0].add_run("—")

#     return tbl

def _build_program_table(doc: Document, program: Dict[str, Any]) -> Table:
    findings = program.get("findings", []) or []
    rows = max(1, len(findings)) + 1

    tbl = doc.add_table(rows=rows, cols=5)  # Fixed to 5 columns
    _style = _pick_table_style(doc)
    if _style:
        try:
            tbl.style = _style
        except Exception:
            pass
    _apply_grid_borders(tbl)

    headers = [
        "Audit\nFinding #",
        "Compliance Type -\nAudit Finding Summary",
        "Audit Finding\nDetermination",
        "Questioned Cost\nDetermination",
        "CAP\nDetermination",
    ]
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        _clear_runs(cell.paragraphs[0])
        run = cell.paragraphs[0].add_run(h)
        run.bold = True  # ✅ Make header text bold
        #cell.paragraphs[0].add_run(h)
        #_shade_cell(cell, "E7E6E6")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _tight_paragraph(cell.paragraphs[0])

    # if findings:
    #     for r, f in enumerate(findings, start=1):
    #         vals = [
    #             f.get("finding_id", ""),
    #             f.get("compliance_and_summary", ""),  # ← Use the combined field that has proper mapping
    #             f.get("audit_determination", "Sustained"),
    #             f.get("questioned_cost_determination", "None"),
    #             f.get("cap_determination", "Not Applicable"),
    #         ]
    #         for c, val in enumerate(vals):
    #             cell = tbl.cell(r, c)
    #             _clear_runs(cell.paragraphs[0])
    #             cell.paragraphs[0].add_run(str(val))
    #             cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    #             # ✅ Center align all columns EXCEPT column 1 (compliance_and_summary)
    #             if c != 1:
    #                 cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    #             _tight_paragraph(cell.paragraphs[0])
    if findings:
        for r, f in enumerate(findings, start=1):
            for c in range(5):
                cell = tbl.cell(r, c)
                _clear_runs(cell.paragraphs[0])
                
                # Column-specific formatting
                if c == 0:  # Finding ID
                    cell.paragraphs[0].add_run(f.get("finding_id", ""))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                elif c == 1:  # Compliance Type - Audit Finding (SPECIAL FORMATTING)
                    compliance_type = f.get("compliance_type", "")
                    summary = f.get("summary", "")
                    
                    # Add compliance type in BOLD
                    if compliance_type:
                        bold_run = cell.paragraphs[0].add_run(compliance_type)
                        bold_run.bold = True
                    
                    # Add hyphen with spaces
                    if compliance_type and summary:
                        cell.paragraphs[0].add_run(" - ")
                        cell.paragraphs[0].add_run("\n")
                    
                    # Add summary (not bold)
                    if summary:
                        cell.paragraphs[0].add_run("\n")
                        cell.paragraphs[0].add_run(summary)
                    
                    # Left align this column
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                elif c == 2:  # Audit Finding Determination
                    cell.paragraphs[0].add_run(f.get("audit_determination", "Sustained"))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                elif c == 3:  # Questioned Cost Determination
                    cell.paragraphs[0].add_run(f.get("questioned_cost_determination", "None"))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                elif c == 4:  # CAP Determination
                    cell.paragraphs[0].add_run(f.get("cap_determination", "Not Applicable"))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                #_tight_paragraph(cell.paragraphs[0])
        
    else:
        cell = tbl.cell(1, 0)
        _clear_runs(cell.paragraphs[0])
        cell.paragraphs[0].add_run("—")
    
    # FORMAT START
    
    set_table_cell_margins(tbl, top_in=0.00, bottom_in=0.00, left_in=0.06, right_in=0.06)

    # ---- Program table formatting (ONLY if 5 columns) ----
    set_table_preferred_width_and_indent(tbl, width_in=6.25, indent_in=0.05)

    for r in tbl.rows:
        set_row_height_and_allow_break(r, height_in=0.49, allow_break_across_pages=True)

    set_table_column_widths(tbl, [0.83, 1.59, 1.2, 1.44, 1.19])
    # ---- end program table formatting ----

    # SPACING MUST BE LAST so nothing overwrites it
    apply_program_table_spacing(tbl)

    set_table_bold_borders(tbl, size=12)

    # END FORMAT
            
    return tbl

# def _insert_program_tables_at_anchor(doc: Document, anchor_para: Paragraph, programs: List[Dict[str, Any]]):
#     # Clean anchor text and delete any placeholder table immediately following it
#     text = _para_text(anchor_para).replace("[[PROGRAM_TABLES]]", "")
#     _clear_runs(anchor_para)
#     if text.strip():
#         anchor_para.add_run(text)

#     _delete_immediate_next_table(anchor_para)

#     # Order programs by ALN
#     def _al_key(p):
#         return (p.get("assistance_listing") or "99.999")
#     programs_sorted = sorted(programs or [], key=_al_key)

#     last = anchor_para
#     for p in programs_sorted:
#         al = p.get("assistance_listing", "Unknown")
#         name = p.get("program_name", "Unknown Program")
#         heading = f"Assistance Listing Number/Program Name: {al} / {name}"
#         heading_para = doc.add_paragraph()
#         _clear_runs(heading_para); heading_para.add_run(heading)

#         # splice heading after 'last'
#         heading_el = heading_para._p
#         heading_el.getparent().remove(heading_el)
#         _insert_after(last, heading_el)
#         last = heading_el

#         # table
#         tbl = _build_program_table(doc, p)
#         tbl_el = tbl._tbl
#         tbl_el.getparent().remove(tbl_el)
#         _insert_after(last, tbl_el)
#         last = tbl_el

#         # CAPs
#         for f in p.get("findings", []):
#             cap_text = (f or {}).get("cap_text")
#             if cap_text:
#                 cap_title = doc.add_paragraph()
#                 _clear_runs(cap_title); cap_title.add_run(f"Corrective Action Plan – {f.get('finding_id','')}")
#                 cap_text_para = doc.add_paragraph()
#                 _clear_runs(cap_text_para); cap_text_para.add_run(cap_text)

#                 cap_title_el = cap_title._p; cap_text_el = cap_text_para._p
#                 cap_title_el.getparent().remove(cap_title_el)
#                 cap_text_el.getparent().remove(cap_text_el)
#                 _insert_after(last, cap_title_el)
#                 _insert_after(cap_title_el, cap_text_el)
#                 last = cap_text_el

#         # spacer
#         spacer = doc.add_paragraph()
#         spacer_el = spacer._p
#         spacer_el.getparent().remove(spacer_el)
#         _insert_after(last, spacer_el)
#         last = spacer_el

# def _insert_program_tables_at_anchor_no_headers(doc: Document, anchor_para: Paragraph, programs: List[Dict[str, Any]]):
#     """
#     Insert program tables without creating duplicate headers.
#     The template already has the header paragraph, we just insert tables.
#     """
#     # Clean anchor text
#     text = _para_text(anchor_para).replace("[[PROGRAM_TABLES]]", "")
#     _clear_runs(anchor_para)
#     if text.strip():
#         anchor_para.add_run(text)

#     # Delete any placeholder table immediately following the anchor
#     _delete_immediate_next_table(anchor_para)

#     # Order programs by ALN
#     def _al_key(p):
#         return (p.get("assistance_listing") or "99.999")
#     programs_sorted = sorted(programs or [], key=_al_key)

#     last = anchor_para
    
#     # For SINGLE program: just insert table (header already exists in template)
#     # For MULTIPLE programs: insert header + table for 2nd, 3rd, etc.
#     for idx, p in enumerate(programs_sorted):
#         al = p.get("assistance_listing", "Unknown")
#         name = p.get("program_name", "Unknown Program")
        
#         # Only add header for 2nd+ programs (first uses the template header)
#         if idx > 0:
#             heading = f"Assistance Listing Number/Program Name: {al} / {name}"
#             heading_para = doc.add_paragraph()
#             _clear_runs(heading_para)
#             heading_para.add_run(heading)
            
#             # Add spacing before the header
#             heading_para.paragraph_format.space_before = Pt(12)
            
#             # Splice heading after 'last'
#             heading_el = heading_para._p
#             heading_el.getparent().remove(heading_el)
#             _insert_after(last, heading_el)
#             last = heading_el

#         # Insert table
#         tbl = _build_program_table(doc, p)
#         tbl_el = tbl._tbl
#         tbl_el.getparent().remove(tbl_el)
#         _insert_after(last, tbl_el)
#         last = tbl_el

#         # Insert CAPs after the table
#         for f in p.get("findings", []):
#             cap_text = (f or {}).get("cap_text")
#             if cap_text:
#                 cap_title = doc.add_paragraph()
#                 _clear_runs(cap_title)
#                 cap_title.add_run(f"Corrective Action Plan – {f.get('finding_id','')}")
                
#                 cap_text_para = doc.add_paragraph()
#                 _clear_runs(cap_text_para)
#                 cap_text_para.add_run(cap_text)

#                 cap_title_el = cap_title._p
#                 cap_text_el = cap_text_para._p
#                 cap_title_el.getparent().remove(cap_title_el)
#                 cap_text_el.getparent().remove(cap_text_el)
                
#                 _insert_after(last, cap_title_el)
#                 _insert_after(cap_title_el, cap_text_el)
#                 last = cap_text_el

#         # Spacer between programs (if multiple)
#         if idx < len(programs_sorted) - 1:
#             spacer = doc.add_paragraph()
#             spacer_el = spacer._p
#             spacer_el.getparent().remove(spacer_el)
#             _insert_after(last, spacer_el)
#             last = spacer_el


def _insert_program_tables_at_anchor_no_headers(doc: Document, anchor_para: Paragraph, programs: List[Dict[str, Any]]):
    """
    Format date as "Month Day, Year" (Comment #14: "Must be in [Month] [Day], [Year] format")
    Example: "June 30, 2024"
    """
    if not date_str:
        return ""
    try:
        # Try ISO format (2024-06-30)
        if "-" in date_str and len(date_str) >= 10:
            dt = datetime.fromisoformat(date_str.replace("Z", "+00:00").split("T")[0])
            return dt.strftime("%B %d, %Y")
        return date_str
    except:
        return date_str


def add_the_prefix(name: str) -> str:
    """
    Add "The" prefix if not present (Comment #16: "Must add 'The' before recipient name")
    """
    if not name:
        return ""
    name = name.strip()
    if name.lower().startswith("the "):
        return name
    return f"The {name}"


def add_the_lowercase_prefix(name: str) -> str:
    """
    Add "the" prefix (lowercase) for auditor name in narrative.
    Removes WordArt/VML watermark shapes (e.g., 'PowerPlusWaterMarkObject', 'DRAFT')
    from headers/footers/body. Works with typical Word 'DRAFT' watermarks.
    """
    for part in doc.part.package.parts:
        if not hasattr(part, "element"):
            continue
        root = part.element
        ns = root.nsmap
        # Remove shapes named like watermark or with textpath 'DRAFT'
        for pict in list(root.xpath(".//*[local-name()='pict']")):
            kill_pict = False
            for shp in pict.xpath(".//*[local-name()='shape']"):
                sid = (shp.get("id") or "") + " " + (shp.get("{urn:schemas-microsoft-com:office:office}spid") or "")
                if "PowerPlusWaterMarkObject" in sid or "WaterMark" in sid:
                    kill_pict = True
                    break
                for tx in shp.xpath(".//*[local-name()='textpath'][@string]"):
                    if "DRAFT" in (tx.get("string") or "").upper():
                        kill_pict = True
                        break
            if kill_pict:
                parent = pict.getparent()
                parent.remove(pict)

def _insert_paragraph_after(p: Paragraph, text: str = "") -> Paragraph:
    new_p = p._element.addnext(p._element.__class__())
    # Create a real python-docx Paragraph around that element
    from docx.text.paragraph import Paragraph as _Para
    new_par = _Para(new_p, p._parent)
    if text:
        new_par.add_run(text)
    return new_par

def _find_para_by_contains(doc: Document, needle: str) -> Optional[Paragraph]:
    def _norm(s: str) -> str:
        s = (s or "").replace("\u00A0"," ").replace("\xa0"," ")
        s = s.replace("\u200b","").replace("\u200c","").replace("\u200d","")
        return " ".join(s.split())
    N = _norm(needle)
    for p in _iter_all_paragraphs_in_container(doc):
        if N in _norm(_para_text(p)):
            return p
    for sec in doc.sections:
        for p in _iter_all_paragraphs_in_container(sec.header):
            if N in _norm(_para_text(p)):
                return p
        for p in _iter_all_paragraphs_in_container(sec.footer):
            if N in _norm(_para_text(p)):
                return p
    return None

LOWERCASE_WORDS = {"and", "of", "the", "for", "to", "in", "on", "at", "by", "with", "from"}

def _format_name_standard_case(name: str) -> str:
    """
    Format name in standard title case, removing 'The' article if present.
    Example: "CITY OF ANN ARBOR, MICHIGAN" -> "City Of Ann Arbor, Michigan"
    """
    if not name:
        return ""
    name = name.strip()
    if name.lower().startswith("the "):
        return name
    return f"the {name}"


def get_current_date() -> str:
    """
    Get current date in required format (Comment #0: "Current Date")
    """
    return datetime.now().strftime("%B %d, %Y")


# ============================================================
# PLURALIZATION (Comments #19, #20, #40)
# ============================================================

def get_pluralization(count: int) -> Dict[str, str]:
    """
    Get singular/plural forms based on finding count.
    Comment #19: "Add S if plural, remove if singular finding"
    """
    singular = (count == 1)
    
    return {
        # For "[is/are]" placeholder
        "is/are": "is" if singular else "are",
        
        # For "[The]" placeholder before recipient name in body
        # Comment #16: Must add "The" before recipient name
        # But for singular, we may not want "The" (based on template context)
        "The": "" if singular else "The",
        
        # For "[the]" placeholder before auditor name  
        "the": "the",  # Always lowercase "the" before auditor
        
        # Additional pluralization helpers (if template uses these)
        "finding_s": "finding" if singular else "findings",
        "issue_s": "issue" if singular else "issues",
        "violate_s": "violates" if singular else "violate",
        "CAP_s": "CAP" if singular else "CAPs",
    }


# ============================================================
# FAC CLIENT
# ============================================================

class FACClient:
    """Federal Audit Clearinghouse API client."""
    
    def __init__(self, api_key: str = ""):
        self.base_url = Config.FAC_API_BASE
        self.session = requests.Session()
        self.session.headers["Accept"] = "application/json"
        if api_key:
            self.session.headers["X-Api-Key"] = api_key
    
    def _get(self, endpoint: str, params: dict) -> list:
        url = f"{self.base_url}/{endpoint}"
        try:
            resp = self.session.get(url, params=params, timeout=30)
            resp.raise_for_status()
            return resp.json()
        except requests.HTTPError as e:
            logger.error(f"FAC API error: {e}")
            return []
        except Exception as e:
            logger.error(f"FAC API error: {e}")
            return []
    
    def _or_param(self, field: str, values: List[str]) -> str:
        inner = ",".join([f"{field}.eq.{v}" for v in values])
        return f"({inner})"
    
    def find_report(self, ein: str, year: int) -> Optional[dict]:
        """Find most recent report for EIN/year."""
        logger.info(f"Searching FAC: EIN={ein}, year={year}")
        data = self._get("general", {
            "audit_year": f"eq.{year}",
            "auditee_ein": f"eq.{ein}",
            "select": "report_id,fac_accepted_date,auditee_name,auditee_address_line_1,"
                     "auditee_city,auditee_state,auditee_zip,auditor_firm_name,"
                     "fy_end_date,auditee_contact_name,auditee_contact_title",
            "order": "fac_accepted_date.desc",
            "limit": "1",
        })
        if data:
            logger.info(f"Found report: {data[0].get('report_id')}")
            return data[0]
        logger.warning("No report found")
        return None
    
    def get_findings(self, report_id: str, max_refs: int = 15, only_flagged: bool = False) -> List[dict]:
        """Get findings for a report."""
        params = {
            "report_id": f"eq.{report_id}",
            "select": "reference_number,award_reference,type_requirement,"
                     "is_material_weakness,is_significant_deficiency,is_questioned_costs,"
                     "is_modified_opinion,is_other_findings,is_other_matters,"
                     "is_repeat_finding,prior_finding_ref_numbers",
            "order": "reference_number.asc",
            "limit": str(max_refs),
        }
        if only_flagged:
            flagged = ["is_material_weakness", "is_significant_deficiency", "is_questioned_costs",
                      "is_modified_opinion", "is_other_findings", "is_other_matters", "is_repeat_finding"]
            params["or"] = "(" + ",".join(f"{f}.eq.true" for f in flagged) + ")"
        return self._get("findings", params) or []
    
    def get_findings_text(self, report_id: str, refs: List[str]) -> Dict[str, str]:
        """Get finding text by reference."""
        if not refs:
            return {}
        data = self._get("findings_text", {
            "report_id": f"eq.{report_id}",
            "select": "finding_ref_number,finding_text",
            "order": "finding_ref_number.asc",
            "limit": str(len(refs) + 5),
            "or": self._or_param("finding_ref_number", refs),
        }) or []
        return {d.get("finding_ref_number", ""): d.get("finding_text", "") for d in data}
    
    def get_caps(self, report_id: str, refs: List[str]) -> Dict[str, str]:
        """Get corrective action plans by reference."""
        if not refs:
            return {}
        data = self._get("corrective_action_plans", {
            "report_id": f"eq.{report_id}",
            "select": "finding_ref_number,planned_action",
            "order": "finding_ref_number.asc",
            "limit": str(len(refs) + 5),
            "or": self._or_param("finding_ref_number", refs),
        }) or []
        return {d.get("finding_ref_number", ""): d.get("planned_action", "") for d in data}
    
    def get_awards(self, report_id: str) -> List[dict]:
        """Get federal awards for a report."""
        return self._get("federal_awards", {
            "report_id": f"eq.{report_id}",
            "select": "award_reference,federal_program_name,assistance_listing",
            "order": "award_reference.asc",
            "limit": "200",
        }) or []
    
    def get_aln_from_summary_excel(self, report_id: str) -> Tuple[Dict[str, str], Dict[str, str]]:
        """
        Download FAC summary Excel and extract ALN mappings.
        
        Returns:
            (aln_by_award, aln_by_finding) where:
            - aln_by_award: {award_reference: aln}
            - aln_by_finding: {reference_number: aln}
        
        This is needed because the FAC API's federal_awards table doesn't always 
        have the assistance_listing populated, but the Excel summary has complete data.
        """
        url = f"https://app.fac.gov/dissemination/summary-report/xlsx/{report_id}"
        logger.info(f"📥 Downloading FAC summary from: {url}")
        
        aln_by_award = {}
        aln_by_finding = {}
        
        try:
            resp = requests.get(url, timeout=30)
            resp.raise_for_status()
            
            import openpyxl
            wb = openpyxl.load_workbook(BytesIO(resp.content), data_only=True)
            
            logger.info(f"📑 Excel sheets available: {wb.sheetnames}")
            
            # Process FEDERALAWARD sheet
            if 'federalaward' in wb.sheetnames:
                ws_fed = wb['federalaward']
                headers = [str(cell.value or "").strip().lower() for cell in ws_fed[1]]
                logger.info(f"   federalaward headers: {headers}")
                
                try:
                    i_award_ref = headers.index('award_reference')
                    i_aln = headers.index('aln')
                    
                    for row in ws_fed.iter_rows(min_row=2, values_only=True):
                        if not row or all(c is None for c in row):
                            continue
                        
                        award_ref = str(row[i_award_ref] or "").strip()
                        aln = str(row[i_aln] or "").strip()
                        
                        # Validate ALN format (should be like 21.027)
                        if award_ref and aln and re.match(r'^\d{2}\.\d{3}', aln):
                            aln_by_award[award_ref] = aln
                    
                    logger.info(f"   ✅ Loaded {len(aln_by_award)} award→ALN mappings")
                    
                except ValueError as e:
                    logger.warning(f"   ⚠️ Could not find columns in federalaward: {e}")
            
            # Process FINDING sheet
            if 'finding' in wb.sheetnames:
                ws_find = wb['finding']
                headers = [str(cell.value or "").strip().lower() for cell in ws_find[1]]
                logger.info(f"   finding headers: {headers}")
                
                try:
                    i_ref_num = headers.index('reference_number')
                    i_aln = headers.index('aln')
                    
                    for row in ws_find.iter_rows(min_row=2, values_only=True):
                        if not row or all(c is None for c in row):
                            continue
                        
                        ref_num = str(row[i_ref_num] or "").strip()
                        aln = str(row[i_aln] or "").strip()
                        
                        if ref_num and aln and re.match(r'^\d{2}\.\d{3}', aln):
                            aln_by_finding[ref_num] = aln
                    
                    logger.info(f"   ✅ Loaded {len(aln_by_finding)} finding→ALN mappings")
                    
                except ValueError as e:
                    logger.warning(f"   ⚠️ Could not find columns in finding: {e}")
            
            # Log sample mappings
            if aln_by_award:
                samples = list(aln_by_award.items())[:3]
                logger.info(f"   Sample award mappings: {samples}")
            if aln_by_finding:
                samples = list(aln_by_finding.items())[:3]
                logger.info(f"   Sample finding mappings: {samples}")
                
        except Exception as e:
            logger.warning(f"⚠️ Failed to load FAC summary Excel: {e}")
        
        return aln_by_award, aln_by_finding
    
    clean = name.strip()
    
    # Remove "The" or "the" if it exists at the beginning
    if clean.lower().startswith("the "):
        clean = clean[4:].strip()
    
    
    # If input is ALL CAPS (or mostly caps), normalize first
    letters = [ch for ch in clean if ch.isalpha()]
    if letters and sum(ch.isupper() for ch in letters) / len(letters) > 0.8:
        clean = clean.lower()

    titled = _title_case(clean)

    # Lowercase connector words unless first word
    parts = titled.split(" ")
    for i, w in enumerate(parts):
        if i > 0 and w.lower() in LOWERCASE_WORDS:
            parts[i] = w.lower()

    #return " ".join(parts)

# def _remove_duplicate_program_headers(doc: Document, first_label: Paragraph):
#     """
#     Remove any duplicate 'Assistance Listing Number/Program Name' paragraphs 
#     that appear after the first one (the template's original).
#     """
#     # Get all paragraphs
#     all_paras = list(doc.paragraphs)
    
#     # Find the index of the first label
#     try:
#         first_idx = all_paras.index(first_label)
#     except ValueError:
#         return  # Can't find it, give up
    
#     # Look for duplicates after the first one (within the next 5 paragraphs)
#     for i in range(first_idx + 1, min(first_idx + 6, len(all_paras))):
#         p = all_paras[i]
#         text = _para_text(p)
        
#         # If this paragraph also starts with "Assistance Listing Number/Program Name"
#         if "Assistance Listing Number/Program Name" in text:
#             logging.info(f"🗑️  Removing duplicate header: {text[:80]}")
#             _remove_paragraph(p)
#             break  # Only remove one duplicate

def _remove_duplicate_program_headers(doc: Document, first_label: Paragraph):
    """
    Remove any duplicate 'Assistance Listing Number/Program Name' paragraphs 
    that appear after the first one (the template's original).
    """
    # Get all paragraphs
    all_paras = list(doc.paragraphs)
    
    # Find the index of the first label
    try:
        first_idx = all_paras.index(first_label)
    except ValueError:
        return  # Can't find it, give up
    
    # Look for duplicates after the first one (within the next 5 paragraphs)
    for i in range(first_idx + 1, min(first_idx + 6, len(all_paras))):
        p = all_paras[i]
        text = _para_text(p)
        
        # If this paragraph also contains "Assistance Listing Number/Program Name"
        if "Assistance Listing Number/Program Name" in text:
            logging.info(f"🗑️  Removing duplicate header: {text[:80]}")
            _remove_paragraph(p)
            break  # Only remove one duplicate

def _fix_questioned_costs_grammar(doc):
    """Fix 'No questioned cost is' to 'No questioned costs are'."""
    for p in _iter_all_paragraphs_full(doc):
        text = _para_text(p)
        if "No questioned cost is included" in text:
            new_text = text.replace(
                "No questioned cost is included in this single audit report",
                "No questioned costs are included in this single audit report"
            )
            if new_text != text:
                _clear_runs(p)
                p.add_run(new_text)
                logging.info("✅ Fixed questioned costs grammar")
                break


def build_docx_from_template(model: Dict[str, Any], *, template_path: str) -> bytes:
    """
    Open a .docx template and:
      1) Replace placeholders across the whole document (headers/footers too)
      2) Insert program tables at the [[PROGRAM_TABLES]] anchor
    """
    if not os.path.isfile(template_path):
        raise HTTPException(400, f"Template not found: {template_path}")

    doc = Document(template_path)
    #_remove_watermarks(doc)  # remove DRAFT/Watermark shapes immediately

    # Dates
    _, letter_date_long = format_letter_date(model.get("letter_date_iso"))

    # Header fields (defaults -> empty so placeholders never leak through)
    auditee = (model.get("auditee_name")
               or model.get("recipient_name")
               or "")
    # if not auditee.lower().startswith("the "):
    #     auditee = "The " + auditee
    ein     = model.get("ein", "") or ""
    street  = model.get("street_address", "") or ""
    city    = model.get("city", "") or ""
    state   = model.get("state", "") or ""
    zipc    = model.get("zip_code", "") or ""
    poc     = model.get("poc_name", "") or ""
    poc_t   = model.get("poc_title", "") or ""
    auditor = model.get("auditor_name", "") or ""
    # if auditor and not auditor.lower().startswith("the "):
    #     auditor = "the " + auditor
    logging.info(f"Auditor: {auditor}")
    logging.info(f"Auditee: {auditee}")
    logging.info(f"POC: {poc} ({poc_t})")
    fy_end  = (model.get("period_end_text")
               or str(model.get("audit_year", ""))) or ""
    # Treasury contact email
    treasury_contact_email = "ORP_SingleAudits@treasury.gov "

    # Map BOTH styles of placeholders used by the template
    mapping = {
        # date stub used in some templates
        "Date XX, 2025": letter_date_long,

        # [bracket] style
        "[Recipient Name]": auditee,
        "[EIN]": ein,
        "[Street Address]": street,
        "[City]": city,
        "[State]": state,
        "[Zip Code]": zipc,
        "[Point of Contact]": poc,
        "[Point of Contact Title]": poc_t,
        "[Auditor Name]": auditor,
        "[Fiscal Year End Date]": fy_end,
        "[The]": "The", 
        "[the]": "the",

        # ${curly} style
        "${recipient_name}": auditee,
        "${ein}": ein,
        "${street_address}": street,
        "${city}": city,
        "${state}": state,
        "${zip_code}": zipc,
        "${poc_name}": poc,
        "${poc_title}": poc_t,
        "${auditor_name}": auditor,
        "${fy_end_text}": fy_end,
    }

    # Ensure no None values sneak in
    mapping = {k: (v if v is not None else "") for k, v in mapping.items()}
    email = (model.get("treasury_contact_email") or "ORP_SingleAudits@treasury.gov ").strip()

    mapping.update({
        # bracket style used by template
        "[treasury_contact_email]": f" {email} ",
        # curly style just in case
        "${treasury_contact_email}": f" {email} "
    })
    # 1) Replace placeholders everywhere (body + headers/footers + nested tables)
    _replace_placeholders_docwide(doc, mapping)
    # 2) Fix questioned costs grammar
    _fix_questioned_costs_grammar(doc)
    # 3) Run email cleanup BEFORE hyperlink creation
    _email_postfix_cleanup(doc, email)
    _strip_leading_token_artifacts(doc)
    # 4) NOW convert email to hyperlink (after all text manipulation)
    #_replace_email_with_hyperlink(doc, email)

    #_email_postfix_cleanup(doc, email)
    #_fix_treasury_email(doc, model.get("treasury_contact_email") or "ORP_SingleAudits@treasury.gov")
    # 5) Final cleanups that don't touch text
    _unset_all_caps_everywhere(doc)
    #_fix_narrative_article(doc, auditee, auditor)

    # 2) Insert program tables at the anchor (do this BEFORE stripping bracketed tokens,
    # because cleanup would otherwise delete the [[PROGRAM_TABLES]] marker)
    anchor = _find_anchor_paragraph(doc, "[[PROGRAM_TABLES]]")
    if not anchor:
        raise HTTPException(400, "Template does not contain the [[PROGRAM_TABLES]] anchor paragraph.")
    programs = model.get("programs", []) or []
    # Find the visible label paragraph and fill it with ALN/Program from the first program
    # try:
    #     label_p = _find_para_by_contains(doc, "Assistance Listing Number/Program Name")
    #     progs = model.get("programs") or []
    #     if label_p is not None and progs:
    #         first = progs[0]
    #         aln = (first.get("assistance_listing") or "").strip()
    #         pname = (first.get("program_name") or "").strip()
    #         # Title-case the program if it somehow stayed all-caps
    #         # def _fix_case(s: str) -> str:
    #         #     if s.isupper():
    #         #         lowers = {"and","or","the","of","for","to","in","on","by","with","a","an"}
    #         #         parts = []
    #         #         for w in s.split():
    #         #             lw = w.lower()
    #         #             parts.append(lw if lw in lowers else lw.capitalize())
    #         #         return " ".join(parts)
    #         #     return s
    #         # pname = _fix_case(pname)
    #         _clear_runs(label_p)
    #         label_p.add_run(f"Assistance Listing Number/Program Name: {aln} / {pname}")
    #         # ✅ ADD THIS: Remove any duplicate headers that follow
    #         _remove_duplicate_program_headers(doc, label_p)
    # except Exception:
    #     pass
    try:
        label_p = _find_para_by_contains(doc, "Assistance Listing Number/Program Name")
        progs = model.get("programs") or []
        if label_p is not None and progs:
            first = progs[0]
            aln = (first.get("assistance_listing") or "").strip()
            pname = (first.get("program_name") or "").strip()
            
            # Clear the paragraph and add formatted text
            _clear_runs(label_p)
            
            # Add bold header text
            header_run = label_p.add_run("Assistance Listing Number/Program Name:")
            header_run.bold = True
            # Add a line break (not new paragraph)
            label_p.add_run("\n")
            # Add the ALN and program name (not bold)
            label_p.add_run(f"{aln} / {pname}")
            # ✅ KEY FIX: Set tight spacing - use _tight_paragraph for consistent removal
            _tight_paragraph(label_p)
            
            # Add the ALN and program name (not bold)
            #label_p.add_run(f"{aln} / {pname}")
            # After creating label_p:
            pf = label_p.paragraph_format
            logging.info(f"Label para - space_before: {pf.space_before}, space_after: {pf.space_after}")

            # After creating table:
            logging.info(f"Table spacing check")
            # Remove any duplicate headers that follow
            _remove_duplicate_program_headers(doc, label_p)
    except Exception as e:
        logging.warning(f"Error handling program headers: {e}")
        pass
    #_insert_program_tables_at_anchor(doc, anchor, programs)
    _insert_program_tables_at_anchor_no_headers(doc, anchor, programs)

    # Remove duplicate narrative blocks under the table
    # Remove duplicate narrative that appears below the table
    try:
        _cleanup_post_table_narrative(doc, model)
    except Exception:
        pass
    # Compute total findings and pluralize boilerplate
    # try:
    #     total_findings = sum(len(p.get("findings") or []) for p in (model.get("programs") or []))
    #     _pluralize_text_everywhere(doc, total_findings)
    # except Exception:
    #     pass
    # Grammar-fix optional plurals via OpenAI (if key set)
    try:
        total_findings = sum(len(prog.get("findings") or []) for prog in (model.get("programs") or []))
        _ai_fix_pluralization_in_doc(doc, total_findings)
    except Exception:
        pass

    # Final tidy: strip any *remaining* token patterns like ${...} or [...]
    def _strip_leftovers_in_container(container):
        for p in _iter_all_paragraphs_in_container(container):
            t = _para_text(p)
            if not t:
                continue
            new_t = t
            if "${" in new_t:
                new_t = re.sub(r"\$\{[^}]+\}", "", new_t)
            if "[" in new_t and "]" in new_t:
                new_t = re.sub(r"\[[^\]]+\]", "", new_t)
            if new_t != t:
                _clear_runs(p)
                p.add_run(new_t)

    _strip_leftovers_in_container(doc)
    for sec in doc.sections:
        _strip_leftovers_in_container(sec.header)
        _strip_leftovers_in_container(sec.footer)
    # ADD THIS LINE HERE:
    _set_font_size_to_12(doc)
    # ========== FORCE FIX NARRATIVE PARAGRAPH ==========
    correct_auditee = model.get("auditee_name") or model.get("recipient_name") or ""
    correct_auditor = model.get("auditor_name") or ""
    # ✅ Ensure no "The" in auditee name
    if correct_auditee.lower().startswith("the "):
        correct_auditee = correct_auditee[4:].strip()
    # Ensure auditee name is bold
    # for p in doc.paragraphs:
    #     text = _para_text(p)
    #     if "Treasury has reviewed the single audit report for" in text:
    #         # Brute force replace the entire sentence
    #         pattern = r'(Treasury has reviewed the single audit report for )([^,]+)(, prepared by )(.+?)( for the fiscal year)'
    #         new_text = re.sub(pattern, f"\\1{correct_auditee}\\3{correct_auditor}\\5", text)
    #         if new_text != text:
    #             _clear_runs(p)
    #             p.add_run(new_text)
    #         break
    for p in doc.paragraphs:
        text = _para_text(p)
        if "Treasury has reviewed the single audit report for" in text:
            # Pattern: "for [NAME], prepared by [AUDITOR] for the fiscal year"
            # This will match "for The City..." or "for City..." and replace with correct name
            pattern = r'(Treasury has reviewed the single audit report for )(The |the )?([^,]+)(, prepared by )(.+?)( for the fiscal year)'
            
            def replacer(match):
                return f"{match.group(1)}{correct_auditee}{match.group(4)}{correct_auditor}{match.group(6)}"
            
            new_text = re.sub(pattern, replacer, text)
            
            if new_text != text:
                _clear_runs(p)
                p.add_run(new_text)
                logging.info(f"✅ Fixed narrative: {correct_auditee}")
            break
    # ========== END FIX ==========
    # ========== FIX APPEALS PARAGRAPH ==========
    for p in doc.paragraphs:
        text = _para_text(p)
        if "may appeal Treasury's decision" in text:
            # Remove "The" from beginning
            new_text = text
            
            # Pattern: "The CITY..." or "The City..." at start
            new_text = re.sub(
                r'^(The |THE )',
                '',
                new_text
            )
            
            # Replace with correct formatted name
            # Pattern: [NAME] may appeal
            pattern = r'^([^,]+)(may appeal)'
            new_text = re.sub(pattern, f'{correct_auditee} \\2', new_text)
            
            if new_text != text:
                _clear_runs(p)
                p.add_run(new_text)
                logging.info(f"✅ Fixed appeals paragraph - removed 'The'")
            break
    # ========== END FIX ==========
    bio = BytesIO()

    # ========== FORCE FIX NARRATIVE PARAGRAPH (FINAL, BOLD-SAFE) ==========
    correct_auditee = model.get("auditee_name") or model.get("recipient_name") or ""
    correct_auditor = model.get("auditor_name") or ""

    # Strip leading "The "
    if correct_auditee.lower().startswith("the "):
        correct_auditee = correct_auditee[4:].strip()

    for p in _iter_all_paragraphs_in_container(doc):
        text = _para_text(p)

        if "Treasury has reviewed the single audit report for" not in text:
            continue

        pattern = (
            r'(Treasury has reviewed the single audit report for )'
            r'(The |the )?(.+?)'
            r'(, prepared by )(.+?)'
            r'( for the fiscal year)'
        )

        m = re.search(pattern, text)
        if not m:
            continue

        _clear_runs(p)

        # Build paragraph with run-level formatting
        p.add_run(m.group(1))                 # fixed intro text

        r = p.add_run(correct_auditee)        # auditee
        r.bold = True                         # ✅ GUARANTEED bold

        p.add_run(m.group(4))                 # ", prepared by "
        p.add_run(correct_auditor)            # auditor
        p.add_run(m.group(6))                 # trailing text

        logging.info(f"✅ Narrative fixed + bolded auditee: {correct_auditee}")
        break

    
    # ========== END FIX ==========

    # Adjust grammar depending on number of findings
    apply_mdl_grammar(doc, total_findings)
    
    doc.save(bio)
    return bio.getvalue()


# --- placeholder cleanup helpers ---
PLACEHOLDER_RE = re.compile(r"^\s*\$\{[^}]+\}\s*$")

def _none_if_placeholder(v):
    """Return None if value looks like an unresolved ${var} placeholder."""
    return None if isinstance(v, str) and PLACEHOLDER_RE.match(v.strip()) else v

def _str_or_default(v, default=""):
    """If v is placeholder/blank/None return default, else v."""
    v = _none_if_placeholder(v)
    if isinstance(v, str) and v.strip():
        return v
    return default

import requests
from io import BytesIO

def _read_headers(ws):
    return [ (c.value or "").strip() if isinstance(c.value, str) else (c.value or "") for c in ws[1] ]

def _find_col(headers, candidates):
    hl = [str(h).strip().lower() for h in headers]
    for i, h in enumerate(hl):
        for cand in candidates:
            cl = cand.strip().lower()
            if h == cl or cl in h:
                return i
    return None

# def _aln_overrides_from_summary(report_id: str):
#     """
#     Returns (aln_by_award, aln_by_finding) by parsing the public FAC summary XLSX.
#     """
#     url = f"https://app.fac.gov/dissemination/summary-report/xlsx/{report_id}"
#     r = requests.get(url, timeout=20)
#     r.raise_for_status()

#     import openpyxl
#     wb = openpyxl.load_workbook(BytesIO(r.content), data_only=True)

#     aln_by_award, aln_by_finding = {}, {}
#     # Look for a sheet with findings
#     for ws in wb.worksheets:
#         headers = _read_headers(ws)
#         if not any(headers):
#             continue
#         i_findref = _find_col(headers, ["finding_ref_number", "finding reference number", "reference_number"])
#         i_award   = _find_col(headers, ["award_reference", "award reference"])
#         i_aln     = _find_col(headers, ["assistance listing", "assistance listing number", "aln", "cfda", "cfda number"])
#         if i_aln is None:
#             continue

#         for row in ws.iter_rows(min_row=2, values_only=True):
#             findref = (row[i_findref] if i_findref is not None else "") or ""
#             award   = (row[i_award]   if i_award   is not None else "") or ""
#             aln     = (row[i_aln]     if i_aln     is not None else "") or ""
#             aln = str(aln).strip()
#             if not aln:
#                 continue
#             if award:
#                 aln_by_award[str(award).strip()] = aln
#             if findref:
#                 aln_by_finding[str(findref).strip()] = aln
#         break  # first matching sheet is enough

#     return aln_by_award, aln_by_finding

def _aln_overrides_from_summary(report_id: str):
    """
    Returns (aln_by_award, aln_by_finding) by parsing the public FAC summary XLSX.
    Updated to handle the actual FAC Excel structure correctly.
    """
    url = f"https://app.fac.gov/dissemination/summary-report/xlsx/{report_id}"
    logging.info(f"📥 Downloading FAC summary from: {url}")
    
    r = requests.get(url, timeout=20)
    r.raise_for_status()

    import openpyxl
    wb = openpyxl.load_workbook(BytesIO(r.content), data_only=True)

    aln_by_award = {}
    aln_by_finding = {}
    
    logging.info(f"📑 Excel sheets available: {wb.sheetnames}")
    
    # ============================================================
    # PART 1: Process FEDERALAWARD sheet
    # ============================================================
    if 'federalaward' in wb.sheetnames:
        ws_fed = wb['federalaward']
        logging.info(f"\n🔍 Processing 'federalaward' sheet (range: {ws_fed.dimensions})")
        
        # Read headers from row 1
        headers = [str(cell.value or "").strip().lower() for cell in ws_fed[1]]
        logging.info(f"   Headers: {headers}")
        
        # Find column indices
        try:
            i_award_ref = headers.index('award_reference')
            i_aln = headers.index('aln')
            i_program = headers.index('federal_program_name')
            
            logging.info(f"   Column indices - award_ref:{i_award_ref}, aln:{i_aln}, program:{i_program}")
            
            # Process data rows (starting from row 2)
            award_count = 0
            for row in ws_fed.iter_rows(min_row=2, values_only=True):
                if not row or all(c is None for c in row):
                    continue
                
                award_ref = str(row[i_award_ref] or "").strip()
                aln = str(row[i_aln] or "").strip()
                program_name = str(row[i_program] or "").strip()
                
                # Validate ALN format (should be like 21.027)
                if award_ref and aln and re.match(r'^\d{2}\.\d{3}', aln):
                    aln_by_award[award_ref] = aln
                    award_count += 1
                    logging.info(f"   ✅ Award: {award_ref} → {aln} ({program_name[:50]})")
            
            logging.info(f"   📊 Processed {award_count} federal awards")
            
        except ValueError as e:
            logging.warning(f"   ⚠️  Could not find required columns in federalaward sheet: {e}")
    
    # ============================================================
    # PART 2: Process FINDING sheet
    # ============================================================
    if 'finding' in wb.sheetnames:
        ws_find = wb['finding']
        logging.info(f"\n🔍 Processing 'finding' sheet (range: {ws_find.dimensions})")
        
        # Read headers from row 1
        headers = [str(cell.value or "").strip().lower() for cell in ws_find[1]]
        logging.info(f"   Headers: {headers}")
        
        # Find column indices
        try:
            i_ref_num = headers.index('reference_number')
            i_aln = headers.index('aln')
            i_award_ref = headers.index('award_reference')
            
            logging.info(f"   Column indices - ref_num:{i_ref_num}, aln:{i_aln}, award_ref:{i_award_ref}")
            
            # Process data rows
            finding_count = 0
            for row in ws_find.iter_rows(min_row=2, values_only=True):
                if not row or all(c is None for c in row):
                    continue
                
                ref_num = str(row[i_ref_num] or "").strip()
                aln = str(row[i_aln] or "").strip()
                award_ref = str(row[i_award_ref] or "").strip()
                
                # Validate ALN format
                if ref_num and aln and re.match(r'^\d{2}\.\d{3}', aln):
                    aln_by_finding[ref_num] = aln
                    finding_count += 1
                    logging.info(f"   ✅ Finding: {ref_num} → {aln} (Award: {award_ref})")
            
            logging.info(f"   📊 Processed {finding_count} findings")
            
        except ValueError as e:
            logging.warning(f"   ⚠️  Could not find required columns in finding sheet: {e}")
    
    logging.info(f"\n✅ FINAL RESULTS:")
    logging.info(f"   Award mappings: {len(aln_by_award)}")
    logging.info(f"   Finding mappings: {len(aln_by_finding)}")
    
    if aln_by_award:
        logging.info(f"\n   Sample award mappings:")
        for k, v in list(aln_by_award.items())[:3]:
            logging.info(f"     {k} → {v}")
    
    if aln_by_finding:
        logging.info(f"\n   Sample finding mappings:")
        for k, v in list(aln_by_finding.items())[:3]:
            logging.info(f"     {k} → {v}")
    
    return aln_by_award, aln_by_finding


# ============================================================
# CLASSIFIER (Comments #37, #38)
# ============================================================

def classify_finding_openai(text: str, labels: List[str] = FINDING_SUMMARIES) -> Optional[str]:
    """Classify finding using OpenAI."""
    if not Config.OPENAI_API_KEY or not text:
        return None
    
    prompt = f"""Classify this audit finding into exactly one category.

Categories:
{chr(10).join(f"- {label}" for label in labels)}

Finding text (first 2000 chars):
{text[:2000]}

Reply with ONLY the category name from the list above, nothing else."""

    try:
        resp = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {Config.OPENAI_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "model": "gpt-4o-mini",
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0,
            },
            timeout=15,
        )
        result = resp.json()
        answer = (result.get("choices", [{}])[0].get("message", {}).get("content") or "").strip()
        if answer in labels:
            return answer
        # Try partial match
        for label in labels:
            if label.lower() in answer.lower() or answer.lower() in label.lower():
                return label
    except Exception as e:
        logger.warning(f"OpenAI classification failed: {e}")
    return None


def classify_finding_keywords(text: str) -> Optional[str]:
    """Classify finding using keyword matching."""
    if not text:
        return None
    text_lower = text.lower()
    for keyword, summary in FINDING_KEYWORDS.items():
        if keyword in text_lower:
            return summary
    return None


def classify_finding(text: str) -> str:
    """
    Classify finding text (Comment #37: "Finding Summary to be selected from list")
    Tries OpenAI first, then keywords, then extracts first sentence.
    """
    if not text:
        return "Other"
    
    # Try OpenAI first
    result = classify_finding_openai(text)
    if result:
        logger.info(f"OpenAI classified: {result}")
        return result
    
    # Fallback to keywords
    result = classify_finding_keywords(text)
    if result:
        logger.info(f"Keyword classified: {result}")
        return result
    
    # Last resort: extract first meaningful sentence
    clean = re.sub(r"\s+", " ", text.strip())
    # Try to get first sentence
    match = re.match(r'^([^.!?]+[.!?])', clean)
    if match and len(match.group(1)) < 150:
        return match.group(1)
    
    return clean[:100] + "..." if len(clean) > 100 else clean


# ============================================================
# DOCUMENT HELPERS
# ============================================================

def add_table_borders(table):
    """Add borders to table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def add_shading(cell, hex_color: str = "D9D9D9"):
    """Add background shading to cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def normalize_aln(aln: str) -> str:
    """Normalize ALN to XX.XXX format."""
    aln = (aln or "").strip()
    if aln and "." not in aln and len(aln) >= 5:
        return f"{aln[:2]}.{aln[2:]}"
    return aln


# ============================================================
# MDL GENERATOR
# ============================================================

@dataclass
class Finding:
    """Single finding data."""
    finding_id: str
    compliance_type: str  # Full text like "Procurement and suspension and debarment"
    summary: str  # Matched summary like "Lack of evidence of suspension and debarment verification"
    audit_determination: str = "Sustained"
    questioned_cost: str = "Questioned Cost:\nNone\nDisallowed Cost:\nNone"
    cap_determination: str = "Accepted"
    is_repeat: bool = False
    prior_finding_ref: str = ""


@dataclass  
class Program:
    """Program with findings."""
    aln: str  # e.g., "21.027"
    name: str  # e.g., "Coronavirus State and Local Fiscal Recovery Funds"
    acronym: str  # e.g., "SLFRF"
    findings: List[Finding] = field(default_factory=list)
    
    @property
    def header(self) -> str:
        """Format: ALN/ Program Name (Acronym)"""
        return f"{self.aln}/ {self.name} ({self.acronym})"


class MDLGenerator:
    """
    MDL Generator using simple template replacement.
    Implements all requirements from template comments.
    """
    
    def __init__(self, template_path: str = ""):
        self.template_path = template_path or Config.MDL_TEMPLATE_PATH
        self.fac = FACClient(Config.FAC_API_KEY)
    
    def generate_from_fac(
        self,
        auditee_name: str,
        ein: str,
        audit_year: int,
        treasury_listings: Optional[List[str]] = None,
        max_refs: int = 15,
        only_flagged: bool = False,
        # Optional overrides
        recipient_name: Optional[str] = None,
        street_address: Optional[str] = None,
        city: Optional[str] = None,
        state: Optional[str] = None,
        zip_code: Optional[str] = None,
        poc_name: Optional[str] = None,
        poc_title: Optional[str] = None,
        auditor_name: Optional[str] = None,
        fiscal_year_end: Optional[str] = None,
        treasury_email: Optional[str] = None,
    ) -> Dict[str, Any]:
        """Generate MDL from FAC data."""
        
        if not treasury_listings:
            treasury_listings = ["21.027", "21.023", "21.026"]
        
        try:
            # 1. Find report
            report = self.fac.find_report(ein, audit_year)
            if not report:
                return {"ok": False, "message": f"No FAC report found for EIN {ein} in {audit_year}"}
            
            report_id = report["report_id"]
            
            # 2. Fetch all data
            findings_raw = self.fac.get_findings(report_id, max_refs, only_flagged)
            refs = [f["reference_number"] for f in findings_raw if f.get("reference_number")]
            findings_text = self.fac.get_findings_text(report_id, refs)
            caps = self.fac.get_caps(report_id, refs)
            awards = self.fac.get_awards(report_id)
            
            # 2b. Get ALN mappings from FAC summary Excel (more reliable than API)
            aln_by_award, aln_by_finding = self.fac.get_aln_from_summary_excel(report_id)
            
            # 3. Build award lookup (award_reference -> ALN)
            # Priority: Excel mapping > API data
            award_to_aln = {}
            for a in awards:
                ref = a.get("award_reference")
                if not ref:
                    continue
                # First try Excel mapping
                if ref in aln_by_award:
                    aln = aln_by_award[ref]
                else:
                    # Fallback to API data
                    aln = normalize_aln(a.get("assistance_listing") or "")
                if aln:
                    award_to_aln[ref] = aln
            
            logger.info(f"📋 Award→ALN mappings: {len(award_to_aln)}")
            if award_to_aln:
                logger.info(f"   Samples: {list(award_to_aln.items())[:5]}")
            
            # 4. Group findings by program (Comment #21, #22: separate table per program)
            programs_map: Dict[str, Program] = {}
            
            for f in findings_raw:
                ref = f.get("reference_number")
                if not ref:
                    continue
                
                # Get ALN - try multiple sources
                # Priority: 1) Direct finding→ALN from Excel, 2) award_reference→ALN, 3) API data
                aln = ""
                
                # First: direct finding→ALN mapping from Excel
                if ref in aln_by_finding:
                    aln = aln_by_finding[ref]
                    logger.info(f"   Finding {ref}: ALN from Excel finding sheet → {aln}")
                
                # Second: via award_reference
                if not aln:
                    award_ref = f.get("award_reference") or ""
                    if award_ref in award_to_aln:
                        aln = award_to_aln[award_ref]
                        logger.info(f"   Finding {ref}: ALN from award {award_ref} → {aln}")
                
                if not aln:
                    logger.warning(f"   Finding {ref}: No ALN found, skipping")
                    continue
                
                # Skip non-Treasury programs
                if aln not in treasury_listings:
                    logger.info(f"   Finding {ref}: ALN {aln} not in Treasury listings, skipping")
                    continue
                
                # Get program info (Comment #26, #27)
                if aln in TREASURY_PROGRAMS:
                    prog_name, prog_acronym = TREASURY_PROGRAMS[aln]
                else:
                    prog_name = "Unknown Program"
                    prog_acronym = "UNK"
                
                # Get or create program
                if aln not in programs_map:
                    programs_map[aln] = Program(
                        aln=aln,
                        name=prog_name,
                        acronym=prog_acronym,
                    )
                
                # Classify finding (Comment #37)
                text = findings_text.get(ref, "")
                summary = classify_finding(text)
                
                # Get compliance type (Comment #32, #34)
                ctype_code = (f.get("type_requirement") or "")[:1].upper()
                ctype_label = COMPLIANCE_TYPES.get(ctype_code, "Other")
                
                # If type is "P" (Other), use heading from finding text (Comment #33)
                if ctype_code == "P" and text:
                    # Try to extract a heading from the finding
                    first_line = text.split('\n')[0].strip()
                    if first_line and len(first_line) < 100:
                        ctype_label = first_line
                
                # CAP determination
                cap_text = caps.get(ref)
                cap_det = "Accepted" if cap_text else "Not Applicable"
                
                # Check for repeat finding (Comment #29, #30)
                is_repeat = f.get("is_repeat_finding") in [True, "Y", "Yes", "true", "yes"]
                prior_ref = f.get("prior_finding_ref_numbers") or ""
                
                programs_map[aln].findings.append(Finding(
                    finding_id=ref,
                    compliance_type=ctype_label,
                    summary=summary,
                    cap_determination=cap_det,
                    is_repeat=is_repeat,
                    prior_finding_ref=prior_ref,
                ))
            
            if not programs_map:
                return {"ok": False, "message": f"No Treasury findings found for {treasury_listings}"}
            
            # Sort programs by ALN (Comment #21: "Put tables in ALN order")
            programs = [programs_map[aln] for aln in sorted(programs_map.keys())]
            
            # 5. Build template replacement data
            # Count total findings for pluralization (Comment #19, #20, #40)
            total_findings = sum(len(p.findings) for p in programs)
            plural = get_pluralization(total_findings)
            
            # Get recipient name in standard case (Comment #1, #3)
            raw_recipient = recipient_name or report.get("auditee_name") or auditee_name
            recipient_standard = to_standard_case(raw_recipient)
            
            # Get auditor name in standard case
            raw_auditor = auditor_name or report.get("auditor_firm_name") or ""
            auditor_standard = to_standard_case(raw_auditor)
            
            # Build data dict
            data = {
                # Date (Comment #0)
                "Date XX, 2025": get_current_date(),
                
                # Recipient (Comments #1, #3) - standard case, NO "The" in address
                "Recipient Name": recipient_standard,
                
                # EIN (Comments #4, #5, #6) - XX-XXXXXXX format
                "EIN": format_ein(ein),
                
                # Address (Comments #7, #8, #9, #10)
                "Street Address": to_standard_case(street_address or report.get("auditee_address_line_1") or ""),
                "City": to_standard_case(city or report.get("auditee_city") or ""),
                "State": (state or report.get("auditee_state") or "").upper(),
                "Zip Code": zip_code or report.get("auditee_zip") or "",
                
                # Point of Contact (Comments #11, #12)
                "Point of Contact": to_standard_case(poc_name or report.get("auditee_contact_name") or ""),
                "Point of Contact Title": to_standard_case(poc_title or report.get("auditee_contact_title") or ""),
                
                # Auditor (with "the" prefix for narrative)
                "Auditor Name": auditor_standard,
                
                # Fiscal Year End (Comments #13, #14, #15)
                "Fiscal Year End Date": fiscal_year_end or format_date(report.get("fy_end_date")),
                
                # Email
                "treasury_contact_email": treasury_email or Config.TREASURY_EMAIL,
                
                # Pluralization (Comments #19, #20, #40)
                **plural,
            }
            
            # 6. Generate document
            docx_bytes = self._generate_document(data, programs, auditor_standard)
            
            # 7. Save
            filename = f"MDL-{self._sanitize(auditee_name)}-{ein.replace('-', '')}-{audit_year}.docx"
            url = self._save_document(docx_bytes, filename)
            
            logger.info(f"Generated: {filename} ({total_findings} findings in {len(programs)} programs)")
            
            return {
                "ok": True,
                "url": url,
                "blob_path": f"{Config.AZURE_CONTAINER}/{filename}" if Config.AZURE_CONN_STR else filename,
                "report_id": report_id,
                "findings_count": total_findings,
                "programs_count": len(programs),
            }
            
        except Exception as e:
            logger.exception("Generation failed")
            return {"ok": False, "message": str(e)}
    
    def _generate_document(
        self,
        data: Dict[str, str],
        programs: List[Program],
        auditor_name: str,
    ) -> bytes:
        """Generate document from template."""
        
        # Load template
        if os.path.exists(self.template_path):
            doc = Document(self.template_path)
            logger.info(f"Loaded template: {self.template_path}")
        else:
            raise FileNotFoundError(f"Template not found: {self.template_path}")
        
        # Build replacement map
        replacements = {}
        for key, value in data.items():
            # Special handling for "Date XX, 2025" (literal text in template)
            if key.startswith("Date "):
                replacements[key] = value or ""
            else:
                # All other keys are bracketed placeholders
                replacements[f"[{key}]"] = value or ""
        
        # Special: "[The] [Recipient Name]" -> "The City of Ann Arbor" (Comment #16)
        # But only in narrative, not in address block
        # The address block just has "[Recipient Name]" without "[The]"
        
        # Special: "[the] [Auditor Name]" -> "the Rehmann Robson LLC"
        replacements["[the]"] = "the"
        
        # Replace in all paragraphs
        for para in doc.paragraphs:
            self._replace_in_paragraph(para, replacements)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, replacements)
        
        # Replace in headers/footers
        for section in doc.sections:
            for para in section.header.paragraphs:
                self._replace_in_paragraph(para, replacements)
            for para in section.footer.paragraphs:
                self._replace_in_paragraph(para, replacements)
        
        # Find "[ALN]/ [Program Name] [(Program Acronym)]" and replace with actual program headers
        # Then insert findings tables
        self._insert_program_sections(doc, programs)
        
        # Convert email addresses to hyperlinks
        self._convert_emails_to_hyperlinks(doc)
        
        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    
    def _convert_emails_to_hyperlinks(self, doc: Document):
        """
        Find email addresses in the document and convert them to mailto: hyperlinks.
        """
        import re
        email_pattern = re.compile(r'[\w.+-]+@[\w-]+\.[\w.-]+')
        
        for para in doc.paragraphs:
            text = para.text
            match = email_pattern.search(text)
            
            if match:
                email = match.group()
                logger.info(f"Converting email to hyperlink: {email}")
                
                # Split text around the email
                before = text[:match.start()]
                after = text[match.end():]
                
                # Clear the paragraph
                for run in para.runs:
                    run.text = ""
                
                # Rebuild: text before + hyperlink + text after
                if para.runs:
                    para.runs[0].text = before
                else:
                    para.add_run(before)
                
                # Add the hyperlink
                add_hyperlink(para, f"mailto:{email}", email)
                
                # Add text after
                para.add_run(after)
    
    def _replace_in_paragraph(self, para, replacements: Dict[str, str]):
        """
        Replace placeholders in paragraph - handles placeholders split across runs.
        Key insight: Always check the full paragraph text, not just individual runs.
        """
        # Get full text first
        full_text = para.text
        
        # Check if any replacement needed
        needs_replacement = False
        for old in replacements:
            if old in full_text:
                needs_replacement = True
                break
        
        if not needs_replacement:
            return
        
        # Apply replacements to full text
        new_text = full_text
        for old, new in replacements.items():
            new_text = new_text.replace(old, new)
        
        if new_text == full_text:
            return  # Nothing changed
        
        # Rebuild paragraph: put all text in first run, clear others
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""
    
    def _insert_program_sections(self, doc: Document, programs: List[Program]):
        """
        Find the [[PROGRAM_TABLES]] placeholder and insert program sections there.
        
        Template structure:
        - "Assistance Listing Number/Program Name:" (already exists for first program)
        - [[PROGRAM_TABLES]] <- we insert here
        
        For first program: just add name + table
        For additional programs: add header + name + table
        """
        # Find the [[PROGRAM_TABLES]] placeholder paragraph
        anchor_para = None
        
        for para in doc.paragraphs:
            if "[[PROGRAM_TABLES]]" in para.text:
                anchor_para = para
                logger.info(f"Found [[PROGRAM_TABLES]] placeholder")
                break
        
        if not anchor_para:
            logger.warning("Could not find [[PROGRAM_TABLES]] placeholder")
            return
        
        # Get the anchor element for inserting
        anchor_element = anchor_para._element
        
        # Process programs in order
        is_first = True
        insert_after = anchor_element
        
        for program in programs:
            logger.info(f"Inserting table for {program.aln} with {len(program.findings)} findings")
            
            # For programs after the first, add the header
            if not is_first:
                # Add spacing
                blank_p = doc.add_paragraph()
                insert_after.addnext(blank_p._element)
                insert_after = blank_p._element
                
                # Add header: "Assistance Listing Number/Program Name:"
                header_para = doc.add_paragraph()
                run = header_para.add_run("Assistance Listing Number/Program Name:")
                run.bold = True
                insert_after.addnext(header_para._element)
                insert_after = header_para._element
            
            # Add program name: "21.027/ Coronavirus State and Local Fiscal Recovery Funds (SLFRF)"
            name_para = doc.add_paragraph()
            name_para.add_run(f"{program.aln}/ {program.name} ({program.acronym})")
            insert_after.addnext(name_para._element)
            insert_after = name_para._element
            
            # Add blank line before table
            blank_p2 = doc.add_paragraph()
            insert_after.addnext(blank_p2._element)
            insert_after = blank_p2._element
            
            # Create and insert the findings table
            table = self._create_findings_table(doc, program.findings)
            insert_after.addnext(table._element)
            insert_after = table._element
            
            is_first = False
        
        # Remove the anchor paragraph (the [[PROGRAM_TABLES]] placeholder)
        parent = anchor_para._element.getparent()
        if parent is not None:
            parent.remove(anchor_para._element)
        
        logger.info(f"Inserted {len(programs)} program tables")
    
    def _fill_template_row(self, row, finding: Finding):
        """Fill the template row by replacing placeholders in all paragraphs."""
        replacements = {
            "[Finding Number]": finding.finding_id,
            "[Compliance Type]": finding.compliance_type,
            "[Audit Finding Summary]": finding.summary,
        }
        
        for cell in row.cells:
            # Process ALL paragraphs in the cell (template has 2 paragraphs in column 1)
            for para in cell.paragraphs:
                full_text = para.text
                
                # Check if any replacement needed
                needs_replacement = False
                for old in replacements:
                    if old in full_text:
                        needs_replacement = True
                        break
                
                if needs_replacement:
                    new_text = full_text
                    for old, new in replacements.items():
                        new_text = new_text.replace(old, new)
                    
                    # Rebuild paragraph
                    if para.runs:
                        para.runs[0].text = new_text
                        for r in para.runs[1:]:
                            r.text = ""
    
    def _fill_findings_table(self, table, findings: List[Finding]):
        """Fill existing template table with findings."""
        # The template has a header row and one data row template
        # We need to fill the first finding, then add rows for rest
        
        if len(table.rows) < 2:
            return
        
        # Fill first finding in existing row
        if findings:
            self._fill_finding_row(table.rows[1], findings[0])
        
        # Add additional rows
        for finding in findings[1:]:
            # Add a new row by copying structure
            new_row = table.add_row()
            self._fill_finding_row(new_row, finding)
    
    def _fill_finding_row(self, row, finding: Finding):
        """Fill a table row with finding data."""
        cells = row.cells
        
        # Column 0: Finding Number (with repeat indicator if applicable)
        # Comment #29, #30: If repeat, add "Repeat of [prior_ref]"
        finding_text = finding.finding_id
        if finding.is_repeat and finding.prior_finding_ref:
            finding_text += f"\nRepeat of {finding.prior_finding_ref}"
        
        cells[0].text = ""
        p = cells[0].paragraphs[0]
        run = p.add_run(finding_text)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Column 1: Compliance Type - Summary
        cells[1].text = ""
        p = cells[1].paragraphs[0]
        run = p.add_run(finding.compliance_type)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        p.add_run("  ̶\n")  # Em dash
        run = p.add_run(finding.summary)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        
        # Column 2: Audit Finding Determination
        cells[2].text = ""
        p = cells[2].paragraphs[0]
        run = p.add_run(finding.audit_determination)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Column 3: Questioned Cost Determination
        cells[3].text = ""
        p = cells[3].paragraphs[0]
        run = p.add_run(finding.questioned_cost)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Column 4: CAP Determination
        cells[4].text = ""
        p = cells[4].paragraphs[0]
        run = p.add_run(finding.cap_determination)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _create_findings_table(self, doc: Document, findings: List[Finding]):
        """Create a new findings table."""
        table = doc.add_table(rows=len(findings) + 1, cols=5)
        add_table_borders(table)
        
        # Column widths
        widths = [Inches(0.8), Inches(2.8), Inches(1.0), Inches(1.2), Inches(0.8)]
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                cell.width = widths[idx]
        
        # Header row
        headers = ["Audit\nFinding #", "Compliance Type - Audit Finding Summary",
                   "Audit Finding\nDetermination", "Questioned Cost\nDetermination", "CAP\nDetermination"]
        
        for i, h in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            add_shading(cell)
        
        # Data rows
        for row_idx, finding in enumerate(findings, start=1):
            self._fill_finding_row(table.rows[row_idx], finding)
        
        return table
    
    def _save_document(self, data: bytes, filename: str) -> str:
        """Save document and return URL."""
        if Config.AZURE_CONN_STR:
            return self._save_to_azure(data, filename)
        return self._save_local(data, filename)
    
    def _save_to_azure(self, data: bytes, filename: str) -> str:
        """Save to Azure Blob Storage."""
        from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions
        
        parts = dict(p.split("=", 1) for p in Config.AZURE_CONN_STR.split(";") if "=" in p)
        account_name = parts.get("AccountName", "")
        account_key = parts.get("AccountKey", "")
        
        client = BlobServiceClient.from_connection_string(Config.AZURE_CONN_STR)
        container = client.get_container_client(Config.AZURE_CONTAINER)
        try:
            container.create_container()
        except:
            pass
        
        blob = container.get_blob_client(filename)
        blob.upload_blob(data, overwrite=True)
        
        sas = generate_blob_sas(
            account_name=account_name,
            container_name=Config.AZURE_CONTAINER,
            blob_name=filename,
            account_key=account_key,
            permission=BlobSasPermissions(read=True),
            expiry=datetime.utcnow() + timedelta(hours=2),
        )
        
        return f"https://{account_name}.blob.core.windows.net/{Config.AZURE_CONTAINER}/{filename}?{sas}"
    
    def _save_local(self, data: bytes, filename: str) -> str:
        """Save locally."""
        os.makedirs(Config.LOCAL_SAVE_DIR, exist_ok=True)
        path = os.path.join(Config.LOCAL_SAVE_DIR, filename)
        with open(path, "wb") as f:
            f.write(data)
        return f"{Config.PUBLIC_BASE_URL}/local/{filename}"
    
    def _sanitize(self, s: str) -> str:
        """Sanitize for filename."""
        return re.sub(r"[^A-Za-z0-9._-]+", "_", s or "").strip("_")[:50]


# ============================================================
# FASTAPI APP
# ============================================================

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel, Field

app = FastAPI(title="MDL DOCX Builder", version="2.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


class BuildRequest(BaseModel):
    """Request model."""
    auditee_name: str
    ein: str
    audit_year: int
    treasury_listings: Optional[List[str]] = None
    max_refs: int = Field(default=15, ge=1, le=100)
    only_flagged: bool = False
    include_awards: bool = True
    recipient_name: Optional[str] = None
    street_address: Optional[str] = None
    city: Optional[str] = None
    state: Optional[str] = None
    zip_code: Optional[str] = None
    poc_name: Optional[str] = None
    poc_title: Optional[str] = None
    auditor_name: Optional[str] = None
    fy_end_text: Optional[str] = None
    template_path: Optional[str] = None
    aln_reference_xlsx: Optional[str] = None
    dest_path: Optional[str] = None


@app.get("/healthz")
def healthz():
    return {"ok": True, "service": "mdl-generator", "version": "2.0.0"}

@app.get("/local/{path:path}")
def get_local_file(path: str):
    full = os.path.join(LOCAL_SAVE_DIR, path)
    if not os.path.isfile(full):
        raise HTTPException(404, "Not found")
    return FileResponse(
        full,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.post("/build-mdl-docx-auto")
def build_mdl_docx_auto(req: BuildRequest):
    """Main endpoint."""
    template = req.template_path or Config.MDL_TEMPLATE_PATH
    gen = MDLGenerator(template_path=template)
    
    result = gen.generate_from_fac(
        auditee_name=req.auditee_name,
        ein=req.ein,
        audit_year=req.audit_year,
        treasury_listings=req.treasury_listings,
        max_refs=req.max_refs,
        only_flagged=req.only_flagged,
        recipient_name=req.recipient_name,
        street_address=req.street_address,
        city=req.city,
        state=req.state,
        zip_code=req.zip_code,
        poc_name=req.poc_name,
        poc_title=req.poc_title,
        auditor_name=req.auditor_name,
        fiscal_year_end=req.fy_end_text,
    )
    
    return result
        
# ------------------------------------------------------------------------------
# Schemas
# ------------------------------------------------------------------------------
class BuildDocx(BaseModel):
    auditee_name: str
    ein: str
    audit_year: int
    body_html: Optional[str] = None
    body_html_b64: Optional[str] = None
    dest_path: Optional[str] = None
    filename: Optional[str] = None

class BuildFromFAC(BaseModel):
    auditee_name: str
    ein: str
    audit_year: int
    fac_general: List[Dict[str, Any]] = []
    fac_findings: List[Dict[str, Any]] = []
    fac_findings_text: List[Dict[str, Any]] = []
    fac_caps: List[Dict[str, Any]] = []
    federal_awards: List[Dict[str, Any]] = []
    dest_path: Optional[str] = None
    filename: Optional[str] = None

class BuildByReport(BaseModel):
    auditee_name: str
    ein: str
    audit_year: int
    report_id: str
    dest_path: Optional[str] = None
    only_flagged: bool = False
    max_refs: int = 15
    include_awards: bool = True

class BuildByReportTemplated(BuildByReport):
    auditor_name: Optional[str] = None
    fy_end_text: Optional[str] = None
    recipient_name: Optional[str] = None
    street_address: Optional[str] = None
    city: Optional[str] = None
    state: Optional[str] = None
    zip_code: Optional[str] = None
    poc_name: Optional[str] = None
    poc_title: Optional[str] = None
    template_path: Optional[str] = None
    treasury_listings: Optional[List[str]] = None

# ------------------------------------------------------------------------------
# Routes
# ------------------------------------------------------------------------------
@app.get("/healthz")
def healthz():
    logging.info(f"Incoming payload to healthsz endpoint")
    return {"ok": True, "time": datetime.utcnow().isoformat()}

@app.post("/echo")
def echo(payload: Dict[str, Any]):
    logging.info(f"Incoming payload to echo endpoint: {payload}")
    #logging.info(f"Incoming payload: {payload.dict()}")
    return {"received": payload, "ts": datetime.utcnow().isoformat()}

@app.get("/debug/env")
def debug_env():
    key = os.getenv("FAC_API_KEY") or ""
    masked = (key[:4] + "…" + key[-2:]) if key else None
    return {"fac_api_key_present": bool(key), "fac_api_key_masked": masked}

@app.get("/debug/storage")
def debug_storage():
    info = _parse_conn_str(AZURE_CONN_STR) if AZURE_CONN_STR else {}
    return {"using_storage": bool(AZURE_CONN_STR), "account": info.get("AccountName"), "blob_endpoint": info.get("BlobEndpoint")}

@app.get("/debug/sas")
def debug_sas():
    if not AZURE_CONN_STR:
        raise HTTPException(400, "Set AZURE_STORAGE_CONNECTION_STRING to test SAS.")
    url = upload_and_sas(AZURE_CONTAINER, "debug/hello.txt", b"hi", ttl_minutes=5)
    return {"url": url}

@app.post("/build-docx-demo")
def build_docx_demo():
    document = Document()
    document.add_heading("Hello from the DOCX demo ✅", level=1)
    document.add_paragraph("If you can read this, your write/upload pipeline is good.")
    bio = BytesIO(); document.save(bio); data = bio.getvalue()
    blob_name = "demo/hello.docx"
    url = upload_and_sas(AZURE_CONTAINER, blob_name, data) if AZURE_CONN_STR else save_local_and_url(blob_name, data)
    return {"url": url, "size_bytes": len(data)}

@app.post("/build-docx")
def build_docx(req: BuildDocx):
    html_str = (req.body_html or "").strip()
    if (not html_str) and req.body_html_b64:
        try:
            html_str = base64.b64decode(req.body_html_b64).decode("utf-8", errors="ignore")
        except Exception:
            raise HTTPException(400, "Invalid base64 in body_html_b64")
    if not html_str:
        raise HTTPException(400, "body_html (or body_html_b64) is required")

    data = html_to_docx_bytes(html_str, force_basic=False)
    folder = (req.dest_path or "").lstrip("/")
    base = req.filename or f"MDL-{sanitize(req.auditee_name)}-{sanitize(req.ein)}-{req.audit_year}.docx"
    blob_name = f"{folder}{base}" if folder else base
    url = upload_and_sas(AZURE_CONTAINER, blob_name, data) if AZURE_CONN_STR else save_local_and_url(blob_name, data)
    return {"url": url, "blob_path": f"{AZURE_CONTAINER}/{blob_name}", "size_bytes": len(data)}

@app.post("/build-docx-from-fac")
def build_docx_from_fac(req: BuildFromFAC):
    model = build_mdl_model_from_fac(
        auditee_name=req.auditee_name,
        ein=req.ein,
        audit_year=req.audit_year,
        fac_general=req.fac_general,
        fac_findings=req.fac_findings,
        fac_findings_text=req.fac_findings_text,
        fac_caps=req.fac_caps,
        federal_awards=req.federal_awards,
        only_flagged=False,
        max_refs=25,
        include_no_qc_line=True,
    )
    html_str = render_mdl_html(model)
    data = html_to_docx_bytes(html_str, force_basic=True)

    folder = (req.dest_path or "").lstrip("/")
    base = req.filename or f"MDL-Preview-{sanitize(req.auditee_name)}-{sanitize(req.ein)}-{req.audit_year}.docx"
    blob_name = f"{folder}{base}" if folder else base
    url = upload_and_sas(AZURE_CONTAINER, blob_name, data) if AZURE_CONN_STR else save_local_and_url(blob_name, data)
    return {"url": url, "blob_path": f"{AZURE_CONTAINER}/{blob_name}", "size_bytes": len(data)}



@app.post("/build-docx-by-report")
def build_docx_by_report(req: BuildByReport):
    fac_general = _fac_get("general", {
        "report_id": f"eq.{req.report_id}",
        "select": "report_id,fac_accepted_date",
        "limit": 1
    })

    findings_params = {
        "report_id": f"eq.{req.report_id}",
        "select": "reference_number,award_reference,type_requirement,"
                  "is_material_weakness,is_significant_deficiency,is_questioned_costs,"
                  "is_modified_opinion,is_other_findings,is_other_matters,is_repeat_finding",
        "order": "reference_number.asc",
        "limit": str(req.max_refs)
    }
    if req.only_flagged:
        flagged = [
            "is_material_weakness","is_significant_deficiency","is_questioned_costs",
            "is_modified_opinion","is_other_findings","is_other_matters","is_repeat_finding"
        ]
        findings_params["or"] = "(" + ",".join([f"{f}.eq.true" for f in flagged]) + ")"
    fac_findings = _fac_get("findings", findings_params)

    refs = [row.get("reference_number") for row in fac_findings if row.get("reference_number")]
    refs = refs[: req.max_refs]

    if refs:
        fac_findings_text = _fac_get("findings_text", {
            "report_id": f"eq.{req.report_id}",
            "select": "finding_ref_number,finding_text",
            "order": "finding_ref_number.asc",
            "limit": str(len(refs)),
            "or": _or_param("finding_ref_number", refs)
        })
        fac_caps = _fac_get("corrective_action_plans", {
            "report_id": f"eq.{req.report_id}",
            "select": "finding_ref_number,planned_action",
            "order": "finding_ref_number.asc",
            "limit": str(len(refs)),
            "or": _or_param("finding_ref_number", refs)
        })
    else:
        fac_findings_text, fac_caps = [], []

    federal_awards = []
    if req.include_awards:
        federal_awards = _fac_get("federal_awards", {
            "report_id": f"eq.{req.report_id}",
            "select": "award_reference,federal_program_name",
            "order": "award_reference.asc",
            "limit": "200"
        })

    model = build_mdl_model_from_fac(
        auditee_name=req.auditee_name,
        ein=req.ein,
        audit_year=req.audit_year,
        fac_general=fac_general,
        fac_findings=fac_findings,
        fac_findings_text=fac_findings_text,
        fac_caps=fac_caps,
        federal_awards=federal_awards,
        only_flagged=req.only_flagged,
        max_refs=req.max_refs,
        include_no_qc_line=True,
    )
    html_str = render_mdl_html(model)
    data = html_to_docx_bytes(html_str, force_basic=True)

    folder = (req.dest_path or "").lstrip("/")
    base = f"MDL-Preview-{sanitize(req.auditee_name)}-{sanitize(req.ein)}-{req.audit_year}.docx"
    blob_name = f"{folder}{base}" if folder else base
    url = upload_and_sas(AZURE_CONTAINER, blob_name, data) if AZURE_CONN_STR else save_local_and_url(blob_name, data)
    return {"url": url, "blob_path": f"{AZURE_CONTAINER}/{blob_name}", "size_bytes": len(data)}

from fastapi.responses import JSONResponse

@app.post("/build-mdl-docx-by-report-templated")
def build_mdl_docx_by_report_templated(req: BuildByReportTemplated):
    try:
        # 1) General
        fac_general = _fac_get("general", {
            "report_id": f"eq.{req.report_id}",
            "select": "report_id,fac_accepted_date",
            "limit": 1
        })

        # 2) Findings
        findings_params = {
            "report_id": f"eq.{req.report_id}",
            "select": "reference_number,award_reference,type_requirement,"
                      "is_material_weakness,is_significant_deficiency,is_questioned_costs,"
                      "is_modified_opinion,is_other_findings,is_other_matters,is_repeat_finding",
            "order": "reference_number.asc",
            "limit": str(req.max_refs)
        }
        if req.only_flagged:
            flagged = [
                "is_material_weakness","is_significant_deficiency","is_questioned_costs",
                "is_modified_opinion","is_other_findings","is_other_matters","is_repeat_finding"
            ]
            findings_params["or"] = "(" + ",".join([f"{f}.eq.true" for f in flagged]) + ")"
        fac_findings = _fac_get("findings", findings_params)

        # 3) refs
        refs = [row.get("reference_number") for row in fac_findings if row.get("reference_number")]
        refs = refs[: req.max_refs]

        # 4) texts & CAPs
        if refs:
            fac_findings_text = _fac_get("findings_text", {
                "report_id": f"eq.{req.report_id}",
                "select": "finding_ref_number,finding_text",
                "order": "finding_ref_number.asc",
                "limit": str(len(refs)),
                "or": _or_param("finding_ref_number", refs)
            })
            fac_caps = _fac_get("corrective_action_plans", {
                "report_id": f"eq.{req.report_id}",
                "select": "finding_ref_number,planned_action",
                "order": "finding_ref_number.asc",
                "limit": str(len(refs)),
                "or": _or_param("finding_ref_number", refs)
            })
        else:
            fac_findings_text, fac_caps = [], []

        # 5) awards
        federal_awards = []
        if req.include_awards:
            federal_awards = _fac_get("federal_awards", {
                "report_id": f"eq.{req.report_id}",
                "select": "award_reference,federal_program_name",
                "order": "award_reference.asc",
                "limit": "200"
            })

        # 6) model
        mdl_model = build_mdl_model_from_fac(
            auditee_name=req.auditee_name,
            ein=req.ein,
            audit_year=req.audit_year,
            fac_general=fac_general,
            fac_findings=fac_findings,
            fac_findings_text=fac_findings_text,
            fac_caps=fac_caps,
            federal_awards=federal_awards,
            only_flagged=req.only_flagged,
            max_refs=req.max_refs,
            include_no_qc_line=True,
        )

        # 7) header autofill (if provided)
        if req.fy_end_text:     mdl_model["period_end_text"] = req.fy_end_text
        if req.auditor_name:    mdl_model["auditor_name"]    = req.auditor_name
        if req.recipient_name:  mdl_model["auditee_name"]    = req.recipient_name
        if req.street_address:  mdl_model["street_address"]  = req.street_address
        if req.city:            mdl_model["city"]            = req.city
        if req.state:           mdl_model["state"]           = req.state
        if req.zip_code:        mdl_model["zip_code"]        = req.zip_code
        if req.poc_name:        mdl_model["poc_name"]        = req.poc_name
        if req.poc_title:       mdl_model["poc_title"]       = req.poc_title

        # 8) template path
        template_path = req.template_path or MDL_TEMPLATE_PATH
        if not template_path:
            return JSONResponse(status_code=200, content={"ok": False, "message": "Template path not provided (template_path or MDL_TEMPLATE_PATH)."})

        # 9) build docx
        try:
            data = build_docx_from_template(mdl_model, template_path=template_path)
        except HTTPException as e:
            return JSONResponse(status_code=200, content={"ok": False, "message": f"Template error: {e.detail}"})
        except Exception as e:
            return JSONResponse(status_code=200, content={"ok": False, "message": f"Unexpected template error: {e}"})

        # 10) upload
        folder = (req.dest_path or "").lstrip("/")
        base = f"MDL-{sanitize(req.auditee_name)}-{sanitize(req.ein)}-{req.audit_year}.docx"
        blob_name = f"{folder}{base}" if folder else base
        url = upload_and_sas(AZURE_CONTAINER, blob_name, data) if AZURE_CONN_STR else save_local_and_url(blob_name, data)
        return {"ok": True, "url": url, "blob_path": f"{AZURE_CONTAINER}/{blob_name}", "size_bytes": len(data)}

    except HTTPException as e:
        # Convert any other HTTPException into a soft error so AnythingLLM can surface the message
        return JSONResponse(status_code=200, content={"ok": False, "message": f"{e.status_code}: {e.detail}"})
    except Exception as e:
        return JSONResponse(status_code=200, content={"ok": False, "message": f"Unhandled error: {e}"})

# --- NEW: single-call builder that looks up report_id first -------------------
from pydantic import BaseModel

# --- Pydantic input model (replace your BuildAuto) ---
class BuildAuto(BaseModel):
    # required
    auditee_name: str
    ein: str
    audit_year: int

    # options (all optional)
    dest_path: Optional[str] = None
    max_refs: int = 15
    only_flagged: bool = False
    include_awards: bool = True
    treasury_listings: Optional[List[str]] = None

    # header overrides (all optional)
    recipient_name: Optional[str] = None
    fy_end_text: Optional[str] = None
    auditor_name: Optional[str] = None
    street_address: Optional[str] = None
    city: Optional[str] = None
    state: Optional[str] = None
    zip_code: Optional[str] = None
    poc_name: Optional[str] = None
    poc_title: Optional[str] = None

    # template & mappings (optional)
    template_path: Optional[str] = None
    aln_reference_xlsx: Optional[str] = None
    treasury_contact_email: Optional[str] = None

    # optional flags
    include_no_qc_line: bool = True
    include_no_cap_line: bool = False


from fastapi.responses import JSONResponse

def _from_fac_general(gen_rows):
    """
    Pulls reasonable defaults from FAC /general for headers.
    Expects select to include:
      auditee_address_line_1,auditee_city,auditee_state,auditee_zip,
      auditor_firm_name,fy_end_date
    """
    if not gen_rows:
        return {}
    g = gen_rows[0]
    # fy_end_date can be 'YYYY-MM-DD' -> convert to 'Month DD, YYYY' if present
    per_end = None
    try:
        from datetime import datetime
        if g.get("fy_end_date"):
            dt = datetime.fromisoformat(g["fy_end_date"])
            per_end = dt.strftime("%B %-d, %Y") if hasattr(dt, "strftime") else None
    except Exception:
        pass

    return {
        "street_address": g.get("auditee_address_line_1") or "",
        "city": g.get("auditee_city") or "",
        "state": g.get("auditee_state") or "",
        "zip_code": g.get("auditee_zip") or "",
        "auditor_name": g.get("auditor_firm_name") or "",
        "period_end_text": per_end,
        "poc_name":  g.get("auditee_contact_name")  or "",
        "poc_title": g.get("auditee_contact_title") or "",
    }

def _unset_all_caps_everywhere(doc):
    # body paragraphs
    for p in doc.paragraphs:
        for r in p.runs:
            r.font.all_caps = False
            r.font.small_caps = False
    # tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.all_caps = False
                        r.font.small_caps = False
    # headers/footers
    for sec in doc.sections:
        for container in (sec.header, sec.footer):
            for p in container.paragraphs:
                for r in p.runs:
                    r.font.all_caps = False
                    r.font.small_caps = False

# def _rewrite_paragraph(p, text):
#     _clear_runs(p); p.add_run(text)

def _rewrite_paragraph(p, text):
    """Rewrite paragraph text safely."""
    _clear_runs(p)
    p.add_run(text)

def _iter_all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
        
def _iter_all_paragraphs_full(doc):
    # body
    for p in doc.paragraphs: yield p
    # tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    # header/footer
    for sec in doc.sections:
        for container in (sec.header, sec.footer):
            for p in container.paragraphs:
                yield p

# def _email_postfix_cleanup(doc, email):
#     # strip leading bracket/curly tokens at paragraph start; fix ".The" joins
#     pat_leading = re.compile(r"^\s*(\[\s*treasury_contact_email\s*\]|\$\{treasury_contact_email\})\.?\s*")
#     for p in _iter_all_paragraphs_full(doc):
#         t = _para_text(p)
#         if not t: continue
#         new = pat_leading.sub("", t)
#         if email and f"{email}.The" in new:
#             new = new.replace(f"{email}.The", f"{email}. The")
#         if new != t:
#             _rewrite_paragraph(p, new)


def _email_postfix_cleanup(doc, email):
    """
    Strip leading bracket/curly tokens at paragraph start; fix ".The" joins.
    Safe to run before hyperlink insertion.
    """
    pat_leading = re.compile(r"^\s*(\[\s*treasury_contact_email\s*\]|\$\{treasury_contact_email\})\.?\s*")
    for p in _iter_all_paragraphs_full(doc):
        t = _para_text(p)
        if not t:
            continue
        
        new = pat_leading.sub("", t)
        if email and f"{email}.The" in new:
            new = new.replace(f"{email}.The", f"{email}. The")
        
        if new != t:
            _rewrite_paragraph(p, new)

def _fix_treasury_email(doc, email: str):
    if not email:
        return
    email = email.strip()

    found_token = False
    # 1) Replace tokens anywhere (body, tables, hdr/ftr)
    for p in _iter_all_paragraphs(doc):
        t = _para_text(p)
        if "${treasury_contact_email}" in t:
            found_token = True
            new_t = t.replace("${treasury_contact_email}", email)
            # tidy double spaces and missing space before next sentence
            new_t = new_t.replace(f"{email}.The", f"{email}. The")
            new_t = re.sub(r"\s{2,}", " ", new_t)
            _rewrite_paragraph(p, new_t)

    if found_token:
        return

    # 2) No token in template → inject into the “For questions…” line only
    target = _find_para_by_contains(doc, "For questions regarding the audit finding")
    if target:
        t = _para_text(target)
        if email not in t:
            new_t = re.sub(r"(?i)(please email us at)(\s*)", rf"\1 {email}. ", t, count=1)
            _rewrite_paragraph(target, new_t)

# def _replace_email_with_hyperlink(doc, email):
#     """Replace email text with clickable hyperlink."""
#     if not email:
#         return
    
#     for p in _iter_all_paragraphs_full(doc):
#         text = _para_text(p)
        
#         # Find paragraphs containing the email
#         if email in text and "email us at" in text.lower():
#             # Split the text around the email
#             parts = text.split(email)
            
#             if len(parts) == 2:
#                 # Clear and rebuild with hyperlink
#                 _clear_runs(p)
#                 p.add_run(parts[0])  # Text before email
#                 _add_hyperlink(p, f"mailto:{email}", email)  # Email as hyperlink
#                 p.add_run(parts[1])  # Text after email
#                 break

# def _replace_email_with_hyperlink(doc, email):
#     """Replace email text with clickable hyperlink - DEBUG VERSION."""
#     if not email:
#         return
    
#     email = email.strip()
#     logging.info(f"🔍 Looking for email: {email}")
    
#     for p in _iter_all_paragraphs_full(doc):
#         text = _para_text(p)
        
#         if email in text:
#             logging.info(f"📧 Found email in: {text[:80]}...")
            
#             # Find position
#             email_pos = text.find(email)
#             text_before = text[:email_pos]
#             text_after = text[email_pos + len(email):]
            
#             logging.info(f"   Before: '{text_before[-20:]}'")
#             logging.info(f"   Email: '{email}'")
#             logging.info(f"   After: '{text_after[:20]}'")
            
#             # Rebuild
#             _clear_runs(p)
#             if text_before:
#                 p.add_run(text_before)
            
#             hyperlink = _add_hyperlink(p, f"mailto:{email}", email)
#             p._p.append(hyperlink)
            
#             if text_after:
#                 p.add_run(text_after)
            
#             logging.info(f"   ✅ Hyperlink created")

# def _replace_email_with_hyperlink(doc, email):
#     """
#     Replace email text with clickable hyperlink.
#     Simpler approach: build new runs in order.
#     """
#     if not email:
#         return
    
#     email = email.strip()
    
#     for p in _iter_all_paragraphs_full(doc):
#         text = _para_text(p)
        
#         if email not in text:
#             continue
        
#         # Find email position
#         email_start = text.find(email)
#         if email_start == -1:
#             continue
        
#         text_before = text[:email_start]
#         text_after = text[email_start + len(email):]
        
#         # Clear paragraph
#         p_element = p._p
        
#         # Remove all runs
#         for child in list(p_element):
#             if child.tag.endswith('}r') or child.tag.endswith('}hyperlink'):
#                 p_element.remove(child)
        
#         # Add text before email
#         if text_before:
#             run1 = OxmlElement('w:r')
#             t1 = OxmlElement('w:t')
#             t1.set(qn('xml:space'), 'preserve')
#             t1.text = text_before
#             run1.append(t1)
#             p_element.append(run1)
        
#         # Add hyperlink
#         hyperlink = _add_hyperlink(p, f"mailto:{email}", email)
#         p_element.append(hyperlink)
        
#         # Add text after email
#         if text_after:
#             run2 = OxmlElement('w:r')
#             t2 = OxmlElement('w:t')
#             t2.set(qn('xml:space'), 'preserve')
#             t2.text = text_after
#             run2.append(t2)
#             p_element.append(run2)
        
#         logging.info(f"✅ Email hyperlink inserted correctly")
#         logging.info(f"   Before: '{text_before[-30:]}'")
#         logging.info(f"   Email: '{email}'")
#         logging.info(f"   After: '{text_after[:30]}'")

# def _replace_email_with_hyperlink(doc, email):
#     """
#     Replace email text with clickable hyperlink in proper position.
#     Ensures proper XML ordering so hyperlink appears in correct location.
#     """
#     if not email:
#         return
    
#     email = email.strip()
    
#     for p in _iter_all_paragraphs_full(doc):
#         text = _para_text(p)
        
#         # Check if email exists in this paragraph
#         if email not in text:
#             continue
        
#         # Find the position of the email
#         email_pos = text.find(email)
#         if email_pos == -1:
#             continue
        
#         text_before = text[:email_pos]
#         text_after = text[email_pos + len(email):]
        
#         logging.info(f"📧 Found email in paragraph")
#         logging.info(f"   Before: ...{text_before[-40:]}")
#         logging.info(f"   Email: {email}")
#         logging.info(f"   After: {text_after[:40]}...")
        
#         # Get paragraph element
#         p_elem = p._p
        
#         # Clear all existing content (runs and hyperlinks)
#         for child in list(p_elem):
#             if child.tag.endswith('}r') or child.tag.endswith('}hyperlink') or child.tag.endswith('}bookmarkStart') or child.tag.endswith('}bookmarkEnd'):
#                 p_elem.remove(child)
        
#         # Rebuild in correct order
#         # 1. Add text before email
#         if text_before:
#             run_before = OxmlElement('w:r')
#             t_before = OxmlElement('w:t')
#             t_before.set(qn('xml:space'), 'preserve')  # Preserve spaces
#             t_before.text = text_before
#             run_before.append(t_before)
#             p_elem.append(run_before)
        
#         # 2. Add hyperlink with email
#         hyperlink = _add_hyperlink(p, f"mailto:{email}", email)
#         p_elem.append(hyperlink)
        
#         # 3. Add text after email
#         if text_after:
#             run_after = OxmlElement('w:r')
#             t_after = OxmlElement('w:t')
#             t_after.set(qn('xml:space'), 'preserve')  # Preserve spaces
#             t_after.text = text_after
#             run_after.append(t_after)
#             p_elem.append(run_after)
        
#         logging.info(f"   ✅ Rebuilt paragraph with hyperlink in correct position")


# def _replace_email_with_hyperlink(doc, email: str):
#     """
#     Replace a plain-text email with a clickable hyperlink in-place,
#     keeping the email at its original position within the sentence.
#     """
#     if not email:
#         return
#     email = email.strip()
#     if not email:
#         return

#     for p in _iter_all_paragraphs_full(doc):
#         text = _para_text(p)
#         if not text or email not in text:
#             continue

#         # Find email position within this paragraph's full text
#         pos = text.find(email)
#         if pos < 0:
#             continue

#         before = text[:pos]
#         after  = text[pos + len(email):]

#         # Clear all existing content (runs/hyperlinks/bookmarks)
#         p_elem = p._p
#         for child in list(p_elem):
#             # remove runs, hyperlinks, and bookmarks cleanly
#             tag = child.tag.rsplit('}', 1)[-1]
#             if tag in ('r', 'hyperlink', 'bookmarkStart', 'bookmarkEnd'):
#                 p_elem.remove(child)

#         # Helper to append a plain-text run with whitespace preserved
#         def _append_text_run(parent, s: str):
#             if not s:
#                 return
#             r = OxmlElement('w:r')
#             t = OxmlElement('w:t')
#             t.set(qn('xml:space'), 'preserve')    # keep spaces exactly
#             t.text = s
#             r.append(t)
#             parent.append(r)

#         # 1) text BEFORE the email
#         _append_text_run(p_elem, before)

#         # 2) the EMAIL as a hyperlink (mailto:)
#         hyperlink = _add_hyperlink(p, f"mailto:{email}", email)
#         p_elem.append(hyperlink)

#         # 3) text AFTER the email
#         _append_text_run(p_elem, after)

def _replace_email_with_hyperlink(doc, email):
    """
    Replace email text with clickable hyperlink in correct position.
    """
    if not email:
        return
    
    email = email.strip()
    replaced_count = 0
    
    for p in _iter_all_paragraphs_full(doc):
        text = _para_text(p)
        
        if email not in text:
            continue
        
        parts = text.split(email)
        
        if len(parts) < 2:
            continue
        
        logging.info(f"📧 Found email in: {text[:80]}...")
        
        p_elem = p._p
        
        # Remove only runs and hyperlinks, keep pPr and other elements
        for child in list(p_elem):
            if child.tag.endswith('}r') or child.tag.endswith('}hyperlink'):
                p_elem.remove(child)
        
        # Append new content (pPr stays at the beginning automatically)
        for i, part in enumerate(parts):
            if part:
                run = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.set(qn('xml:space'), 'preserve')
                t.text = part
                run.append(t)
                p_elem.append(run)
            
            if i < len(parts) - 1:
                hyperlink = _add_hyperlink(p, f"mailto:{email}", email)
                p_elem.append(hyperlink)
                replaced_count += 1
        
        logging.info(f"   ✅ Replaced {len(parts) - 1} email occurrence(s) with hyperlink")
    
    if replaced_count > 0:
        logging.info(f"✅ Total emails replaced with hyperlinks: {replaced_count}")


def _strip_leading_token_artifacts(doc):
    pat = re.compile(r"^\s*\$\{[^}]+\}\.?\s*")
    for p in _iter_all_paragraphs(doc):
        t = _para_text(p)
        if not t:
            continue
        new = pat.sub("", t)
        if new != t:
            _clear_runs(p); p.add_run(new)
# def _fix_narrative_article(doc, auditee_with_article: str, auditor_with_article: str):
#     """Fix the narrative paragraph to include 'The' before recipient and auditor names."""
#     target = _find_para_by_contains(doc, "Treasury has reviewed the single audit report for")
#     if target:
#         t = _para_text(target)
#         # Replace patterns without "The" with proper article
#         # Pattern: "for City of Ann Arbor" -> "for The City of Ann Arbor"
#         # Pattern: "prepared by REHMANN" -> "prepared by the Rehmann"
        
#         # Extract just the name without "The" for matching
#         auditee_no_article = auditee_with_article.replace("The ", "", 1).replace("the ", "", 1)
#         auditor_no_article = auditor_with_article.replace("the ", "", 1).replace("The ", "", 1)
        
#         new_t = t
#         # Fix auditee (case-insensitive)
#         if auditee_no_article:
#             # Match "for [auditee]" and replace with "for The [auditee]"
#             new_t = re.sub(
#                 rf"\bfor\s+{re.escape(auditee_no_article)}\b",
#                 f"for {auditee_with_article}",
#                 new_t,
#                 flags=re.IGNORECASE
#             )
        
#         # Fix auditor (case-insensitive)
#         if auditor_no_article:
#             # Match "prepared by [auditor]" and replace with "prepared by the [auditor]"
#             new_t = re.sub(
#                 rf"\bprepared by\s+{re.escape(auditor_no_article)}\b",
#                 f"prepared by {auditor_with_article}",
#                 new_t,
#                 flags=re.IGNORECASE
#             )
        
#         if new_t != t:
#             _rewrite_paragraph(target, new_t)

def _fix_narrative_article(doc, auditee_exact: str, auditor_exact: str):
    """Replace narrative text with exact names from model (including articles and casing)."""
    target = _find_para_by_contains(doc, "Treasury has reviewed the single audit report for")
    if not target:
        return
    
    text = _para_text(target)
    
    # Pattern: "for [NAME], prepared by [AUDITOR] for the fiscal"
    import re
    pattern = r'(for\s+)([^,]+)(,\s+prepared by\s+)(.+?)(\s+for the fiscal year)'
    
    def replacer(match):
        return f"{match.group(1)}{auditee_exact}{match.group(3)}{auditor_exact}{match.group(5)}"
    
    new_text = re.sub(pattern, replacer, text, count=1)
    
    if new_text != text:
        _rewrite_paragraph(target, new_text)

def _set_font_size_to_12(doc):
    """Set all text in document to 12pt font."""
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.size = Pt(12)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(12)
    
    for section in doc.sections:
        for container in (section.header, section.footer):
            for p in container.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(12)

# ===================================== BEGIN DOC EDITING =====================================   

import re
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt
from docx.opc.constants import RELATIONSHIP_TYPE as RT

def add_hyperlink(paragraph, text, url, bold=False, font_pt=12):
    """
    Add a clickable hyperlink to a paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Style like a normal hyperlink (blue + underline)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)

    # FORCE FONT SIZE for hyperlink run
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_pt * 2)))
    rPr.append(sz)

    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(font_pt * 2)))
    rPr.append(szCs)

    run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    run.append(t)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)

def replace_email_with_mailto_link(p, email: str):
    """Replace occurrences of the email in paragraph text with a clickable mailto hyperlink."""
    if email not in p.text:
        return False

    full = p.text
    parts = full.split(email)

    _clear_runs(p)

    # rebuild: text + hyperlink + text (+ possible repeats)
    for i, chunk in enumerate(parts):
        if chunk:
            p.add_run(chunk)
        if i < len(parts) - 1:
            add_hyperlink(p, email, f"mailto:{email}", font_pt=12)

    return True



def set_table_bold_borders(tbl, size=12, color="000000"):
    tblPr = tbl._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)

    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = tblBorders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            tblBorders.append(border)

        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))   # “bold”
        border.set(qn("w:color"), color)

def set_table_cell_margins(tbl, top_in=0.06, bottom_in=0.06, left_in=0.06, right_in=0.06):
    """Adds internal cell padding for the whole table (Word: tblCellMar)."""
    def twips(inches: float) -> str:
        return str(int(round(inches * 1440)))

    tblPr = tbl._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)

    cellMar = tblPr.find(qn("w:tblCellMar"))
    if cellMar is None:
        cellMar = OxmlElement("w:tblCellMar")
        tblPr.append(cellMar)

    for side, val in (("top", top_in), ("bottom", bottom_in), ("left", left_in), ("right", right_in)):
        node = cellMar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            cellMar.append(node)
        node.set(qn("w:w"), twips(val))
        node.set(qn("w:type"), "dxa")

def _set_cell_paragraph_spacing_before(cell, before_pt: float):
    
    #Apply spacing-before to ALL paragraphs in a cell.
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(before_pt)

def apply_program_table_spacing(tbl):
    # Header row spacing-before = 3.8pt for all header cells
    header_row = tbl.rows[0]
    for cell in header_row.cells:
        _set_cell_paragraph_spacing_before(cell, 3.8)

    # Subsequent rows: per-column spacing-before
    col_before_pts = [10.0, 0.0, 10.0, 3.8, 10.0]  # cols 1..5

    for r_i in range(1, len(tbl.rows)):
        row = tbl.rows[r_i]
        for c_i, before_pt in enumerate(col_before_pts):
            _set_cell_paragraph_spacing_before(row.cells[c_i], before_pt)


def _twips_from_inches(inches: float) -> int:
    return int(round(inches * 1440))

def set_table_preferred_width_and_indent(table, width_in=6.25, indent_in=0.05):
    # Disable autofit so Word respects widths
    table.autofit = False

    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Preferred table width
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(_twips_from_inches(width_in)))

    # Table indent from left
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:type"), "dxa")
    tblInd.set(qn("w:w"), str(_twips_from_inches(indent_in)))

def set_row_height_and_allow_break(row, height_in=0.48, allow_break_across_pages=True):
    tr = row._tr
    trPr = tr.get_or_add_trPr()

    trHeight = trPr.find(qn("w:trHeight"))
    if trHeight is None:
        trHeight = OxmlElement("w:trHeight")
        trPr.append(trHeight)

    trHeight.set(qn("w:val"), str(_twips_from_inches(height_in)))
    trHeight.set(qn("w:hRule"), "atLeast")  # allows taller rows when needed

    cantSplit = trPr.find(qn("w:cantSplit"))
    if allow_break_across_pages:
        if cantSplit is not None:
            trPr.remove(cantSplit)
    else:
        if cantSplit is None:
            trPr.append(OxmlElement("w:cantSplit"))

    # Row height (exact)
    trHeight = trPr.find(qn("w:trHeight"))
    if trHeight is None:
        trHeight = OxmlElement("w:trHeight")
        trPr.append(trHeight)
    trHeight.set(qn("w:val"), str(_twips_from_inches(height_in)))
    trHeight.set(qn("w:hRule"), "atLeast")

    # Allow row to break across pages:
    # Word uses <w:cantSplit/> to PREVENT breaking. So remove it if present.
    cantSplit = trPr.find(qn("w:cantSplit"))
    if allow_break_across_pages:
        if cantSplit is not None:
            trPr.remove(cantSplit)
    else:
        if cantSplit is None:
            cantSplit = OxmlElement("w:cantSplit")
            trPr.append(cantSplit)

def set_cell_preferred_width(cell, width_in: float):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:type"), "dxa")
    tcW.set(qn("w:w"), str(_twips_from_inches(width_in)))

def set_table_column_widths(table, col_widths_in):
    # col_widths_in: list[float] length == number of cols
    for row in table.rows:
        for i, w in enumerate(col_widths_in):
            # python-docx visible width
            row.cells[i].width = Inches(w)
            # Word preferred width
            set_cell_preferred_width(row.cells[i], w)


def fix_mdl_grammar_text(text: str, n_findings: int) -> str:
    singular = (n_findings == 1)
    be = "is" if singular else "are"

    # NBSP -> space
    out = text.replace("\u00A0", " ")  

    # If tokens are still present, resolve them
    out = re.sub(r"\[\s*is\s*/\s*are\s*\]", be, out, flags=re.IGNORECASE)
    out = re.sub(r"\[\s*The\s*\]", "The" if singular else "", out, flags=re.IGNORECASE)
    out = re.sub(r"\(s\)", "" if singular else "s", out)
    out = re.sub(r"\bviolate\s*\(s\)\b", "violates" if singular else "violate", out, flags=re.IGNORECASE)
    out = re.sub(r"\bappear\s*\(s\)\b",  "appears"  if singular else "appear",  out, flags=re.IGNORECASE)
    out = re.sub(r"\baddress\s*\(es\)\b", "addresses" if singular else "address", out, flags=re.IGNORECASE)
    out = re.sub(r"\baddresses\s*\(es\)\b", "addresses", out, flags=re.IGNORECASE)
    out = re.sub(r"\(es\)", "", out)


    # FIX the remaining grammar
    # Insert missing "is/are" after these subjects if it's missing
    out = re.sub(r"\b(The audit finding(?:s)?)\s+(?=sustained\b)", rf"\1 {be} ", out, flags=re.IGNORECASE)
    out = re.sub(r"\b(The CAP(?:s)?)\s*,?\s*if implemented,\s+(?!is\b|are\b)", rf"\1, if implemented, {be} ", out, flags=re.IGNORECASE)
    out = re.sub(r"\b(the corrective action(?:s)?)\s+(?=subject\b)", rf"\1 {be} ", out, flags=re.IGNORECASE)

    # Fix singular verb agreement when subject is singular
    if singular:
        out = re.sub(r"\bissue\s+violate\b", "issue violates", out, flags=re.IGNORECASE)
        out = re.sub(r"\bfinding\s+appear\b", "finding appears", out, flags=re.IGNORECASE)
        out = re.sub(r"\bCAP,\s*if implemented,\s*is responsive to the audit finding,\s*address\b",
                     "CAP, if implemented, is responsive to the audit finding, addresses",
                     out, flags=re.IGNORECASE)

    # Cleanup spacing/punctuation
    out = re.sub(r"[ \t]{2,}", " ", out)
    out = re.sub(r"\s+([,.;:])", r"\1", out)
    return out

def _force_paragraph_font_size(p, size_pt=12):
    for r in p.runs:
        r.font.size = Pt(size_pt)

def _clear_runs(p):
    for r in p.runs[::-1]:
        p._p.remove(r._r)

def postprocess_docx(doc_bytes: bytes, model: dict) -> bytes:
    bio = BytesIO(doc_bytes)
    doc = Document(bio)

    email = model.get("treasury_contact_email", "ORP_SingleAudits@treasury.gov")

    correct_auditee = (model.get("auditee_name") or model.get("recipient_name") or "").strip()
    if correct_auditee.lower().startswith("the "):
        correct_auditee = correct_auditee[4:].strip()

    date_text = (model.get("fy_end_text") or model.get("fy_end_date") or model.get("fiscal_year_end") or "").strip()

    for p in doc.paragraphs:
        text = p.text
        if "Treasury has reviewed the single audit report for" not in text:
            continue

        # Remove leading "the" before the auditee in the sentence
        text = re.sub(
            r'(Treasury has reviewed the single audit report for )the\s+',
            r'\1',
            text,
            flags=re.IGNORECASE
)
        # Find the date in the paragraph
        date_in_doc = None
        if date_text and date_text in text:
            date_in_doc = date_text
        else:
            m = re.search(r'([A-Za-z]+ \d{1,2}, \d{4})', text)
            if m:
                date_in_doc = m.group(1)

        _clear_runs(p)

        # rebuild auditee and date with two bold runs
        if correct_auditee and date_in_doc and (correct_auditee in text) and (date_in_doc in text):
            # Split around auditee first
            pre_a, rest = text.split(correct_auditee, 1)
            # Then split the remaining text around date
            pre_d, post_d = rest.split(date_in_doc, 1)

            p.add_run(pre_a)
            r1 = p.add_run(correct_auditee)
            r1.bold = True

            p.add_run(pre_d)
            r2 = p.add_run(date_in_doc)
            r2.bold = True

            p.add_run(post_d)
        elif correct_auditee and (correct_auditee in text):
            # Fallback: bold only auditee
            pre, post = text.split(correct_auditee, 1)
            p.add_run(pre)
            r = p.add_run(correct_auditee)
            r.bold = True
            p.add_run(post)
        elif date_in_doc and (date_in_doc in text):
            # Fallback: bold only date
            pre, post = text.split(date_in_doc, 1)
            p.add_run(pre)
            r = p.add_run(date_in_doc)
            r.bold = True
            p.add_run(post)
        else:
            # Fallback: just keep text
            p.add_run(text)

        _force_paragraph_font_size(p, 12)

        break

    #hyperlink treasury email everywhere in the doc
    for p in doc.paragraphs:
        changed = replace_email_with_mailto_link(p, email)
        if changed:
            _force_paragraph_font_size(p, 12)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()

# Grammar map for findings
def fix_mdl_template_grammar(text: str, n_findings: int) -> str:
    singular = (n_findings == 1)

    out = text

    out = out.replace("[is/are]", "is" if singular else "are")
    out = out.replace("[The]", "The" if singular else "")

    out = out.replace("(s)", "" if singular else "s")

    out = out.replace("violate(s)", "violates" if singular else "violate")
    out = out.replace("appear(s)", "appears" if singular else "appear")

    out = out.replace("address(es)", "addresses" if singular else "address")
    out = re.sub(r"\baddresses\s*\(es\)\b", "addresses", out, flags=re.IGNORECASE)

    # General leftover cleanup for (es) if it survived
    out = re.sub(r"\(es\)", "", out)

    # cleanup
    out = re.sub(r"[ \t]{2,}", " ", out)
    out = re.sub(r"\s+([,.;:])", r"\1", out)

    return out

def apply_mdl_grammar(doc, n_findings: int):
    def rewrite_paragraph(p):
        old = p.text
        new = fix_mdl_grammar_text(old, n_findings)
        if new != old:
            for r in p.runs[::-1]:
                p._p.remove(r._r)
            p.add_run(new)

            # FORCE font size to 12pt for entire paragraph
            _force_paragraph_font_size(p, 12)

    # body
    for p in doc.paragraphs:
        rewrite_paragraph(p)

    # tables
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    rewrite_paragraph(p)

# ===================================== END DOC EDITING =====================================                  

@app.post("/build-mdl-docx-auto")
def build_mdl_docx_auto(req: BuildAuto):
    try:
        # 1) Find newest report_id for EIN/year
        gen = _fac_get("general", {
            "audit_year": f"eq.{req.audit_year}",
            "auditee_ein": f"eq.{req.ein}",
            "select": "report_id, fac_accepted_date, auditee_address_line_1, auditee_city, auditee_state, auditee_zip, auditor_firm_name, fy_end_date, auditee_contact_name,auditee_contact_title",
            "order": "fac_accepted_date.desc",
            "limit": 1
        })
        if not gen:
            return {"ok": False, "message": f"No FAC report found for EIN {req.ein} in {req.audit_year}."}

        report_id = gen[0]["report_id"]
        logging.info(f"Found report_id {report_id} for EIN {req.ein} in {req.audit_year}")
        try:
            aln_by_award, aln_by_finding = _aln_overrides_from_summary(report_id)
        except Exception:
            aln_by_award, aln_by_finding = {}, {}
        logging.info(f"ALN overrides loaded: {len(aln_by_award)} awards, {len(aln_by_finding)} findings")
        logging.info("aln_by_award")
        logging.info(aln_by_award)
        logging.info("aln_by_finding")
        logging.info(aln_by_finding)
        # 2) (unchanged) fetch findings / texts / caps / awards ...
        #    ... your existing code here ...

        # 2) Fetch findings / finding text / CAPs (ALWAYS initialize lists)
        findings_params = {
            "report_id": f"eq.{report_id}",
            "select": ("reference_number,award_reference,type_requirement,"
                    "is_material_weakness,is_significant_deficiency,is_questioned_costs,"
                    "is_modified_opinion,is_other_findings,is_other_matters,is_repeat_finding"),
            "order": "reference_number.asc",
            "limit": str(req.max_refs or 15),
        }
        if req.only_flagged:
            flagged = [
                "is_material_weakness", "is_significant_deficiency", "is_questioned_costs",
                "is_modified_opinion", "is_other_findings", "is_other_matters", "is_repeat_finding",
            ]
            findings_params["or"] = "(" + ",".join(f"{f}.eq.true" for f in flagged) + ")"

        try:
            fac_findings = _fac_get("findings", findings_params) or []
        except Exception:
            fac_findings = []

        # refs are safe even if no findings
        refs = [r.get("reference_number") for r in fac_findings if r.get("reference_number")]
        refs = refs[: (req.max_refs or 15)]

        # Always define these
        if refs:
            try:
                fac_findings_text = _fac_get("findings_text", {
                    "report_id": f"eq.{report_id}",
                    "select": "finding_ref_number,finding_text",
                    "order": "finding_ref_number.asc",
                    "limit": str(len(refs)),
                    "or": _or_param("finding_ref_number", refs),
                }) or []
            except Exception:
                fac_findings_text = []
            try:
                fac_caps = _fac_get("corrective_action_plans", {
                    "report_id": f"eq.{report_id}",
                    "select": "finding_ref_number,planned_action",
                    "order": "finding_ref_number.asc",
                    "limit": str(len(refs)),
                    "or": _or_param("finding_ref_number", refs),
                }) or []
            except Exception:
                fac_caps = []
        else:
            fac_findings_text, fac_caps = [], []

        # Awards (optional, still safe)
        # federal_awards = []
        # if req.include_awards:
        #     try:
        #         federal_awards = _fac_get("federal_awards", {
        #             "report_id": f"eq.{report_id}",
        #             "select": "award_reference,federal_program_name,assistance_listing",
        #             "order": "award_reference.asc",
        #             "limit": "200",
        #         }) or []
        #     except Exception:
        #         federal_awards = []
        # if federal_awards:
        #     for a in federal_awards:
        #         if not (a.get("assistance_listing") or "").strip():
        #             ar = (a.get("award_reference") or "").strip()
        #             if ar and ar in aln_by_award:
        #                 a["assistance_listing"] = aln_by_award[ar]

        # Awards (optional, still safe)
        federal_awards = []
        if req.include_awards:
            try:
                federal_awards = _fac_get("federal_awards", {
                    "report_id": f"eq.{report_id}",
                    "select": "award_reference,federal_program_name,assistance_listing",
                    "order": "award_reference.asc",
                    "limit": "200",
                }) or []
            except Exception:
                federal_awards = []

        # APPLY ALN OVERRIDES BEFORE BUILDING MODEL
        # if federal_awards and aln_by_award:
        #     for a in federal_awards:
        #         ar = (a.get("award_reference") or "").strip()
        #         # If no ALN, try to get from override
        #         if not (a.get("assistance_listing") or "").strip() and ar in aln_by_award:
        #             a["assistance_listing"] = aln_by_award[ar]
        #             logging.info(f"🔧 Applied ALN override for {ar}: {aln_by_award[ar]}")
        if federal_awards and aln_by_award:
            logging.info(f"🔧 Applying ALN overrides from {len(aln_by_award)} award mappings")
            for a in federal_awards:
                ar = (a.get("award_reference") or "").strip()
                current_aln = (a.get("assistance_listing") or "").strip()
                
                logging.info(f"   Award {ar}: current ALN = '{current_aln}'")
                
                # If no ALN or it's 'Unknown', try to get from override
                if (not current_aln or current_aln == "Unknown") and ar in aln_by_award:
                    a["assistance_listing"] = aln_by_award[ar]
                    logging.info(f"   ✅ Applied override: {ar} → {aln_by_award[ar]}")
                elif ar in aln_by_award:
                    # Even if there's an ALN, if override exists and differs, consider using it
                    override_aln = aln_by_award[ar]
                    if override_aln != current_aln:
                        logging.info(f"   ⚠️  ALN mismatch for {ar}: '{current_aln}' vs override '{override_aln}'")
                        # Uncomment to prefer override:
                        # a["assistance_listing"] = override_aln
        # ========== HARDCODED ALN FIX ==========
        # TREASURY_PROGRAMS = {
        #     "21.027": "Coronavirus State and Local Fiscal Recovery Funds (SLFRF)",
        #     "21.023": "Emergency Rental Assistance Program (ERA)",
        #     "21.026": "Homeowner Assistance Fund (HAF)",
        # }
        TREASURY_PROGRAMS = {
            "21.027": "Coronavirus State and Local Fiscal Recovery Funds (SLFRF)",  # ✅ Added "Coronavirus"
            "21.023": "Emergency Rental Assistance Program (ERA)",
            "21.026": "Homeowner Assistance Fund (HAF)",
            "21.029": "Capital Projects Fund (CPF)",
            "21.031": "State Small Business Credit Initiative (SSBCI)",
            "21.032": "Local Assistance and Tribal Consistency Fund (LATCF)",
        }
        for a in federal_awards:
            aln = (a.get("assistance_listing") or "").strip()
            if aln in TREASURY_PROGRAMS:
                a["federal_program_name"] = TREASURY_PROGRAMS[aln]
                logging.info(f"✅ Set program name for {aln}")
        # ========== END FIX ==========
        template_path = _none_if_placeholder(req.template_path) or "templates/MDL_Template_Data_Mapping_Comments.docx"
        aln_xlsx = _none_if_placeholder(req.aln_reference_xlsx) or "templates/Additional_Reference_Documentation_MDLs.xlsx"
        # ---------- NEW: build the model -------------
        mdl_model = build_mdl_model_from_fac(
            auditee_name=req.auditee_name,
            ein=req.ein,
            audit_year=req.audit_year,
            fac_general=gen,
            fac_findings=fac_findings,
            fac_findings_text=fac_findings_text,
            fac_caps=fac_caps,
            federal_awards=federal_awards,
            only_flagged=req.only_flagged,
            max_refs=req.max_refs,
            include_no_qc_line=True,
            treasury_listings=req.treasury_listings,
            aln_reference_xlsx=aln_xlsx,
            aln_overrides_by_finding=aln_by_finding,
        )

        # ---------- NEW: enrich headers from FAC + defaults ----------
        fac_defaults = _from_fac_general(gen)

        # def _normalize_auditor_name(name: str) -> str:
        #     if not name:
        #         return ""
        #     clean = name.strip()
        #     return clean if clean.lower().startswith("the ") else f"the {clean}"

        # def _normalize_auditor_name(name: str) -> str:
        #     if not name:
        #         return ""
        #     clean = name.strip()
            
        #     # Apply title casing if the name is all caps
        #     if clean.isupper():
        #         # Use the existing _title_case or _title_with_acronyms function
        #         clean = _title_with_acronyms(clean, keep_all_caps=False)
            
        #     # Add "the" if not present
        #     return clean if clean.lower().startswith("the ") else f"the {clean}"

        # def _normalize_auditor_name(name: str) -> str:
        #     """Add 'the' article but preserve original casing from API."""
        #     if not name:
        #         return ""
        #     clean = name.strip()
        #     # Just add "the" - don't modify casing at all
        #     return clean if clean.lower().startswith("the ") else f"the {clean}"

        # Don't use _title_with_article, just add "The" prefix
        # def _add_article_the(name: str) -> str:
        #     """Add 'The' article but preserve original casing from API."""
        #     if not name:
        #         return ""
        #     clean = name.strip()
        #     return clean if clean.lower().startswith("the ") else f"The {clean}"

        # def _no_article(name: str) -> str:
        #     """Return the name without 'The' article."""
        #     if not name:
        #         return ""
        #     clean = name.strip()
        #     # Remove "The " if it exists at the beginning
        #     if clean.lower().startswith("the "):
        #         return clean[4:].strip()
        #     return clean
        # Get RAW names from FAC (exactly as stored in database)
        raw_auditee = gen[0].get("auditee_name") or req.auditee_name
        raw_auditor = fac_defaults.get("auditor_name") or ""
        #recipient = _title_with_article(req.recipient_name or req.auditee_name)
        #recipient = _add_article_the(req.recipient_name or req.auditee_name)
        # Now add articles while preserving casing
        # For address block - use WITHOUT "The"
        # address_recipient = _no_article(raw_auditee.upper() if raw_auditee else "")
        # narrative_recipient = _add_article_the(raw_auditee.upper() if raw_auditee else "")
        # auditor = _normalize_auditor_name(raw_auditor.upper() if raw_auditor else "")
        #recipient = _add_article_the(raw_auditee.upper() if raw_auditee else "")
        #auditor = _normalize_auditor_name(raw_auditor.upper() if raw_auditor else "")

        # ✅ NEW CODE - Use standard case everywhere, no "The" article:
        recipient_formatted = _format_name_standard_case(raw_auditee)
        auditor_formatted = raw_auditor
        header_overrides = {
            # recipient & period end
            "recipient_name": recipient_formatted,
            "period_end_text": req.fy_end_text or fac_defaults.get("period_end_text") or mdl_model.get("period_end_text"),

            # address (title case street + city, uppercase state, keep zip as-is)
            "street_address": _title_case(req.street_address or fac_defaults.get("street_address")),
            "city": _title_case(req.city or fac_defaults.get("city")),
            "state": (req.state or fac_defaults.get("state") or "").upper(),
            "zip_code": req.zip_code or fac_defaults.get("zip_code") or "",

            # auditor
            #"auditor_name": _normalize_auditor_name(req.auditor_name or fac_defaults.get("auditor_name") or ""),
            "auditor_name": auditor_formatted,  # use normalized name with "the" article
            "auditee_name": recipient_formatted,
            # POC (title case name + title)
            "poc_name": _title_case(req.poc_name or fac_defaults.get("poc_name")),
            "poc_title": req.poc_title or fac_defaults.get("poc_title"),
        }

        # apply non-empty values only
        for k, v in header_overrides.items():
            if v:
                mdl_model[k] = v
        # ADD THIS DEBUG LOGGING:
        logging.info(f"🔍 After header overrides, mdl_model auditor_name: {mdl_model.get('auditor_name')}")
        logging.info(f"🔍 After header overrides, mdl_model auditee_name: {mdl_model.get('auditee_name')}")
        logging.info(f"🔍 After header overrides, mdl_model recipient_name: {mdl_model.get('recipient_name')}")
        # ------------- sensible defaults for things the caller omitted -------------
        # Treasury listings: if not provided, use the SLFRF + common Treasury programs for demo
        if not req.treasury_listings:
            req.treasury_listings = ["21.027", "21.023", "21.026"]

        # Template defaults if not provided
        template_path = _none_if_placeholder(req.template_path) or "templates/MDL_Template_Data_Mapping_Comments.docx"
        aln_xlsx      = _none_if_placeholder(req.aln_reference_xlsx) or "templates/Additional_Reference_Documentation_MDLs.xlsx"

        # Pass mapping workbook path into the model builder via existing parameter if you support it
        # (If build_mdl_model_from_fac already accepts aln_reference_xlsx, we've passed it above.)

        # Destination folder defaults
        dest_folder = _str_or_default(req.dest_path, f"mdl/{req.audit_year}/").lstrip("/")

        # 4) Build DOCX (unchanged except variable names)
        try:
            data = build_docx_from_template(mdl_model, template_path=template_path)

            # ✅ Post-process the generated docx bytes AFTER everything else
            data = postprocess_docx(data, mdl_model)

        except HTTPException as e:
            return {"ok": False, "message": f"Template error: {e.detail}"}
        except Exception as e:
            return {"ok": False, "message": f"Unexpected template error: {e}"}

        # 5) Upload (unchanged)
        base = f"MDL-{sanitize(req.auditee_name)}-{sanitize(req.ein)}-{req.audit_year}.docx"
        blob_name = f"{dest_folder}{base}" if dest_folder else base
        url = upload_and_sas(AZURE_CONTAINER, blob_name, data) if AZURE_CONN_STR else save_local_and_url(blob_name, data)

        return {"ok": True, "url": url, "blob_path": f"{AZURE_CONTAINER}/{blob_name}"}

    except HTTPException as e:
        return JSONResponse(status_code=200, content={"ok": False, "message": f"{e.status_code}: {e.detail}"})
    except Exception as e:
        logger.exception("Unhandled error")
        return JSONResponse(status_code=200, content={"ok": False, "message": str(e)})


@app.post("/build-mdl")
def build_mdl(req: BuildRequest):
    return build_mdl_docx_auto(req)


@app.get("/local/{path:path}")
def get_local_file(path: str):
    full = os.path.join(Config.LOCAL_SAVE_DIR, path)
    if not os.path.isfile(full):
        raise HTTPException(404, "Not found")
    return FileResponse(
        full,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# ============================================================
# CLI
# ============================================================

if __name__ == "__main__":
    import sys
    
    if len(sys.argv) >= 4:
        gen = MDLGenerator()
        result = gen.generate_from_fac(
            auditee_name=sys.argv[1],
            ein=sys.argv[2],
            audit_year=int(sys.argv[3]),
        )
        if result.get("ok"):
            print(f"✓ Generated: {result.get('url')}")
            print(f"  Report: {result.get('report_id')}")
            print(f"  Findings: {result.get('findings_count')}")
        else:
            print(f"✗ Error: {result.get('message')}")
            sys.exit(1)
    else:
        import uvicorn
        print("Starting MDL Generator API on port 8000...")
        uvicorn.run(app, host="0.0.0.0", port=8000)
