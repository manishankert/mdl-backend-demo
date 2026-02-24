# routes/docx_routes.py
import os
import re
import base64
import logging
from io import BytesIO
from datetime import datetime
from typing import Dict, Any

from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse, JSONResponse
from docx import Document

from config import (
    AZURE_CONTAINER,
    AZURE_CONN_STR,
    LOCAL_SAVE_DIR,
    MDL_TEMPLATE_PATH,
    TREASURY_PROGRAMS,
)
from models.schemas import (
    BuildDocx,
    BuildFromFAC,
    BuildByReport,
    BuildByReportTemplated,
    BuildAuto,
)
from utils.text_utils import sanitize, title_case, format_name_standard_case
from services.storage import upload_and_sas, save_local_and_url, parse_conn_str
from services.fac_api import fac_get, or_param, aln_overrides_from_summary, from_fac_general
from services.html_converter import html_to_docx_bytes
from services.mdl_builder import build_mdl_model_from_fac, render_mdl_html
from services.template_processor import build_docx_from_template
from services.document_editor import postprocess_docx

logging.basicConfig(level=logging.INFO)

router = APIRouter()

# --- placeholder cleanup helpers ---
PLACEHOLDER_RE = re.compile(r"^\s*\$\{[^}]+\}\s*$")


def none_if_placeholder(v):
    """Return None if value looks like an unresolved ${var} placeholder."""
    return None if isinstance(v, str) and PLACEHOLDER_RE.match(v.strip()) else v


def str_or_default(v, default=""):
    """If v is placeholder/blank/None return default, else v."""
    v = none_if_placeholder(v)
    if isinstance(v, str) and v.strip():
        return v
    return default


@router.get("/healthz")
def healthz():
    logging.info(f"Incoming payload to healthz endpoint")
    return {"ok": True, "time": datetime.utcnow().isoformat()}


@router.post("/echo")
def echo(payload: Dict[str, Any]):
    logging.info(f"Incoming payload to echo endpoint: {payload}")
    return {"received": payload, "ts": datetime.utcnow().isoformat()}


@router.get("/debug/env")
def debug_env():
    key = os.getenv("FAC_API_KEY") or ""
    masked = (key[:4] + "..." + key[-2:]) if key else None
    return {"fac_api_key_present": bool(key), "fac_api_key_masked": masked}


@router.get("/debug/storage")
def debug_storage():
    info = parse_conn_str(AZURE_CONN_STR) if AZURE_CONN_STR else {}
    return {"using_storage": bool(AZURE_CONN_STR), "account": info.get("AccountName"), "blob_endpoint": info.get("BlobEndpoint")}


@router.get("/debug/sas")
def debug_sas():
    if not AZURE_CONN_STR:
        raise HTTPException(400, "Set AZURE_STORAGE_CONNECTION_STRING to test SAS.")
    url = upload_and_sas(AZURE_CONTAINER, "debug/hello.txt", b"hi", ttl_minutes=5)
    return {"url": url}


@router.get("/local/{path:path}")
def get_local_file(path: str):
    full = os.path.join(LOCAL_SAVE_DIR, path)
    if not os.path.isfile(full):
        raise HTTPException(404, "Not found")
    return FileResponse(
        full,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@router.post("/build-docx-demo")
def build_docx_demo():
    document = Document()
    document.add_heading("Hello from the DOCX demo", level=1)
    document.add_paragraph("If you can read this, your write/upload pipeline is good.")
    bio = BytesIO()
    document.save(bio)
    data = bio.getvalue()
    blob_name = "demo/hello.docx"
    url = upload_and_sas(AZURE_CONTAINER, blob_name, data) if AZURE_CONN_STR else save_local_and_url(blob_name, data)
    return {"url": url, "size_bytes": len(data)}


@router.post("/build-docx")
def build_docx(req: BuildDocx):
    html_str = (req.body_html or "").strip()
    if (not html_str) and req.body_html_b64:
        try:
            html_str = base64.b64decode(req.body_html_b64).decode("utf-8", errors="ignore")
        except Exception:
            raise HTTPException(400, "Invalid base64 in body_html_b64")
    if not html_str:
        raise HTTPException(400, "body_html (or body_html_b64) is required")

    data = html_to_docx_bytes(html_str, force_basic=False)
    folder = (req.dest_path or "").lstrip("/")
    base = req.filename or f"MDL-{sanitize(req.auditee_name)}-{sanitize(req.ein)}-{req.audit_year}.docx"
    blob_name = f"{folder}{base}" if folder else base
    url = upload_and_sas(AZURE_CONTAINER, blob_name, data) if AZURE_CONN_STR else save_local_and_url(blob_name, data)
    return {"url": url, "blob_path": f"{AZURE_CONTAINER}/{blob_name}", "size_bytes": len(data)}


@router.post("/build-docx-from-fac")
def build_docx_from_fac_route(req: BuildFromFAC):
    model = build_mdl_model_from_fac(
        auditee_name=req.auditee_name,
        ein=req.ein,
        audit_year=req.audit_year,
        fac_general=req.fac_general,
        fac_findings=req.fac_findings,
        fac_findings_text=req.fac_findings_text,
        fac_caps=req.fac_caps,
        federal_awards=req.federal_awards,
        only_flagged=False,
        max_refs=25,
        include_no_qc_line=True,
    )
    html_str = render_mdl_html(model)
    data = html_to_docx_bytes(html_str, force_basic=True)

    folder = (req.dest_path or "").lstrip("/")
    base = req.filename or f"MDL-Preview-{sanitize(req.auditee_name)}-{sanitize(req.ein)}-{req.audit_year}.docx"
    blob_name = f"{folder}{base}" if folder else base
    url = upload_and_sas(AZURE_CONTAINER, blob_name, data) if AZURE_CONN_STR else save_local_and_url(blob_name, data)
    return {"url": url, "blob_path": f"{AZURE_CONTAINER}/{blob_name}", "size_bytes": len(data)}


@router.post("/build-docx-by-report")
def build_docx_by_report(req: BuildByReport):
    fac_general = fac_get("general", {
        "report_id": f"eq.{req.report_id}",
        "select": "report_id,fac_accepted_date",
        "limit": 1
    })

    findings_row_limit = max(req.max_refs * 20, 500)
    findings_params = {
        "report_id": f"eq.{req.report_id}",
        "select": "reference_number,award_reference,type_requirement,"
                  "is_material_weakness,is_significant_deficiency,is_questioned_costs,"
                  "is_modified_opinion,is_other_findings,is_other_matters,is_repeat_finding",
        "order": "reference_number.asc",
        "limit": str(findings_row_limit)
    }
    if req.only_flagged:
        flagged = [
            "is_material_weakness", "is_significant_deficiency", "is_questioned_costs",
            "is_modified_opinion", "is_other_findings", "is_other_matters", "is_repeat_finding"
        ]
        findings_params["or"] = "(" + ",".join([f"{f}.eq.true" for f in flagged]) + ")"
    fac_findings = fac_get("findings", findings_params)

    seen_refs = set()
    refs = []
    for row in fac_findings:
        ref = row.get("reference_number")
        if ref and ref not in seen_refs:
            seen_refs.add(ref)
            refs.append(ref)
    refs = refs[: req.max_refs]

    if refs:
        fac_findings_text = fac_get("findings_text", {
            "report_id": f"eq.{req.report_id}",
            "select": "finding_ref_number,finding_text",
            "order": "finding_ref_number.asc",
            "limit": str(len(refs)),
            "or": or_param("finding_ref_number", refs)
        })
        fac_caps = fac_get("corrective_action_plans", {
            "report_id": f"eq.{req.report_id}",
            "select": "finding_ref_number,planned_action",
            "order": "finding_ref_number.asc",
            "limit": str(len(refs)),
            "or": or_param("finding_ref_number", refs)
        })
    else:
        fac_findings_text, fac_caps = [], []

    federal_awards = []
    if req.include_awards:
        federal_awards = fac_get("federal_awards", {
            "report_id": f"eq.{req.report_id}",
            "select": "award_reference,federal_program_name",
            "order": "award_reference.asc",
            "limit": "200"
        })

    model = build_mdl_model_from_fac(
        auditee_name=req.auditee_name,
        ein=req.ein,
        audit_year=req.audit_year,
        fac_general=fac_general,
        fac_findings=fac_findings,
        fac_findings_text=fac_findings_text,
        fac_caps=fac_caps,
        federal_awards=federal_awards,
        only_flagged=req.only_flagged,
        max_refs=req.max_refs,
        include_no_qc_line=True,
    )
    html_str = render_mdl_html(model)
    data = html_to_docx_bytes(html_str, force_basic=True)

    folder = (req.dest_path or "").lstrip("/")
    base = f"MDL-Preview-{sanitize(req.auditee_name)}-{sanitize(req.ein)}-{req.audit_year}.docx"
    blob_name = f"{folder}{base}" if folder else base
    url = upload_and_sas(AZURE_CONTAINER, blob_name, data) if AZURE_CONN_STR else save_local_and_url(blob_name, data)
    return {"url": url, "blob_path": f"{AZURE_CONTAINER}/{blob_name}", "size_bytes": len(data)}


@router.post("/build-mdl-docx-by-report-templated")
def build_mdl_docx_by_report_templated(req: BuildByReportTemplated):
    try:
        # 1) General
        fac_general = fac_get("general", {
            "report_id": f"eq.{req.report_id}",
            "select": "report_id,fac_accepted_date",
            "limit": 1
        })

        # 2) Findings
        findings_row_limit = max(req.max_refs * 20, 500)
        findings_params = {
            "report_id": f"eq.{req.report_id}",
            "select": "reference_number,award_reference,type_requirement,"
                      "is_material_weakness,is_significant_deficiency,is_questioned_costs,"
                      "is_modified_opinion,is_other_findings,is_other_matters,is_repeat_finding",
            "order": "reference_number.asc",
            "limit": str(findings_row_limit)
        }
        if req.only_flagged:
            flagged = [
                "is_material_weakness", "is_significant_deficiency", "is_questioned_costs",
                "is_modified_opinion", "is_other_findings", "is_other_matters", "is_repeat_finding"
            ]
            findings_params["or"] = "(" + ",".join([f"{f}.eq.true" for f in flagged]) + ")"
        fac_findings = fac_get("findings", findings_params)

        # 3) refs - deduplicate (API returns one row per finding-per-award)
        seen_refs = set()
        refs = []
        for row in fac_findings:
            ref = row.get("reference_number")
            if ref and ref not in seen_refs:
                seen_refs.add(ref)
                refs.append(ref)
        refs = refs[: req.max_refs]

        # 4) texts & CAPs
        if refs:
            fac_findings_text = fac_get("findings_text", {
                "report_id": f"eq.{req.report_id}",
                "select": "finding_ref_number,finding_text",
                "order": "finding_ref_number.asc",
                "limit": str(len(refs)),
                "or": or_param("finding_ref_number", refs)
            })
            fac_caps = fac_get("corrective_action_plans", {
                "report_id": f"eq.{req.report_id}",
                "select": "finding_ref_number,planned_action",
                "order": "finding_ref_number.asc",
                "limit": str(len(refs)),
                "or": or_param("finding_ref_number", refs)
            })
        else:
            fac_findings_text, fac_caps = [], []

        # 5) awards
        federal_awards = []
        if req.include_awards:
            federal_awards = fac_get("federal_awards", {
                "report_id": f"eq.{req.report_id}",
                "select": "award_reference,federal_program_name",
                "order": "award_reference.asc",
                "limit": "200"
            })

        # 6) model
        mdl_model = build_mdl_model_from_fac(
            auditee_name=req.auditee_name,
            ein=req.ein,
            audit_year=req.audit_year,
            fac_general=fac_general,
            fac_findings=fac_findings,
            fac_findings_text=fac_findings_text,
            fac_caps=fac_caps,
            federal_awards=federal_awards,
            only_flagged=req.only_flagged,
            max_refs=req.max_refs,
            include_no_qc_line=True,
        )

        # 7) header autofill (if provided)
        if req.fy_end_text:
            mdl_model["period_end_text"] = req.fy_end_text
        if req.auditor_name:
            mdl_model["auditor_name"] = req.auditor_name
        if req.recipient_name:
            mdl_model["auditee_name"] = req.recipient_name
        if req.street_address:
            mdl_model["street_address"] = req.street_address
        if req.city:
            mdl_model["city"] = req.city
        if req.state:
            mdl_model["state"] = req.state
        if req.zip_code:
            mdl_model["zip_code"] = req.zip_code
        if req.poc_name:
            mdl_model["poc_name"] = req.poc_name
        if req.poc_title:
            mdl_model["poc_title"] = req.poc_title

        # 8) template path
        template_path = req.template_path or MDL_TEMPLATE_PATH
        if not template_path:
            return JSONResponse(status_code=200, content={"ok": False, "message": "Template path not provided (template_path or MDL_TEMPLATE_PATH)."})

        # 9) build docx
        try:
            data = build_docx_from_template(mdl_model, template_path=template_path)
        except HTTPException as e:
            return JSONResponse(status_code=200, content={"ok": False, "message": f"Template error: {e.detail}"})
        except Exception as e:
            return JSONResponse(status_code=200, content={"ok": False, "message": f"Unexpected template error: {e}"})

        # 10) upload
        folder = (req.dest_path or "").lstrip("/")
        base = f"MDL-{sanitize(req.auditee_name)}-{sanitize(req.ein)}-{req.audit_year}.docx"
        blob_name = f"{folder}{base}" if folder else base
        url = upload_and_sas(AZURE_CONTAINER, blob_name, data) if AZURE_CONN_STR else save_local_and_url(blob_name, data)
        return {"ok": True, "url": url, "blob_path": f"{AZURE_CONTAINER}/{blob_name}", "size_bytes": len(data)}

    except HTTPException as e:
        # Convert any other HTTPException into a soft error so AnythingLLM can surface the message
        return JSONResponse(status_code=200, content={"ok": False, "message": f"{e.status_code}: {e.detail}"})
    except Exception as e:
        return JSONResponse(status_code=200, content={"ok": False, "message": f"Unhandled error: {e}"})


@router.post("/build-mdl-docx-auto")
def build_mdl_docx_auto(req: BuildAuto):
    try:
        # 1a) Fetch ALL audit records for this EIN to find the latest one with a valid auditee_name
        # First, get multiple recent records (not just limit 1) so we can find one with auditee_name
        all_audits = fac_get("general", {
            "auditee_ein": f"eq.{req.ein}",
            "select": "report_id, audit_year, fac_accepted_date, auditee_address_line_1, auditee_city, auditee_state, auditee_zip, auditor_firm_name, fy_end_date, auditee_contact_name, auditee_contact_title, auditee_name",
            "order": "audit_year.desc,fac_accepted_date.desc",
            "limit": 10  # Get multiple records to find one with valid auditee_name
        })

        if not all_audits:
            return {"ok": False, "message": f"No FAC records found for EIN {req.ein}."}

        # Log all available audit years for debugging
        available_years = [str(a.get("audit_year")) for a in all_audits]
        logging.info(f"Available audit years for EIN {req.ein}: {', '.join(available_years)}")

        # Find the latest record that has a non-empty auditee_name
        gen_latest = None
        auditee_name_from_latest = ""
        for audit_record in all_audits:
            candidate_name = (audit_record.get("auditee_name") or "").strip()
            if candidate_name:
                gen_latest = audit_record
                auditee_name_from_latest = candidate_name
                break

        # If no record has auditee_name, use the first record anyway
        if not gen_latest:
            gen_latest = all_audits[0]
            logging.warning(f"No FAC records with valid auditee_name found for EIN {req.ein}, using first record")

        latest_year = gen_latest.get("audit_year")
        logging.info(f"Latest audit year WITH auditee_name for EIN {req.ein}: {latest_year} (input year: {req.audit_year})")

        # Get auditee_name from latest year (primary source) or fall back to request
        effective_auditee_name = auditee_name_from_latest or req.auditee_name or ""

        if not effective_auditee_name:
            return {"ok": False, "message": f"Could not determine auditee name for EIN {req.ein}."}

        # Get POC (Point of Contact) from latest year as well
        poc_name_from_latest = (gen_latest.get("auditee_contact_name") or "").strip()
        poc_title_from_latest = (gen_latest.get("auditee_contact_title") or "").strip()

        logging.info(f"Using auditee_name from latest FAC year ({latest_year}): {effective_auditee_name}")
        logging.info(f"Using POC from latest FAC year ({latest_year}): {poc_name_from_latest} ({poc_title_from_latest})")

        # 1b) Find report_id for the INPUT year (for findings data AND all other info except auditee_name)
        gen = fac_get("general", {
            "audit_year": f"eq.{req.audit_year}",
            "auditee_ein": f"eq.{req.ein}",
            "select": "report_id, fac_accepted_date, auditee_address_line_1, auditee_city, auditee_state, auditee_zip, auditor_firm_name, fy_end_date, auditee_contact_name,auditee_contact_title, auditee_name",
            "order": "fac_accepted_date.desc",
            "limit": 1
        })
        if not gen:
            return {"ok": False, "message": f"No FAC report found for EIN {req.ein} in {req.audit_year}."}

        report_id = gen[0]["report_id"]
        logging.info(f"Found report_id {report_id} for EIN {req.ein} in {req.audit_year}")
        try:
            aln_by_award, aln_by_finding = aln_overrides_from_summary(report_id)
        except Exception:
            aln_by_award, aln_by_finding = {}, {}
        logging.info(f"ALN overrides loaded: {len(aln_by_award)} awards, {len(aln_by_finding)} findings")
        logging.info("aln_by_award")
        logging.info(aln_by_award)
        logging.info("aln_by_finding")
        logging.info(aln_by_finding)

        # 2) Fetch findings / finding text / CAPs (ALWAYS initialize lists)
        # NOTE: The FAC API returns one row per finding-per-award combination,
        # so a single finding can produce 10+ rows.  Use a high row limit
        # to ensure we capture all findings; the model builder deduplicates
        # by unique reference_number and then applies max_refs.
        effective_max_refs = req.max_refs or 15
        findings_row_limit = max(effective_max_refs * 20, 500)  # high limit; deduplicate after
        findings_params = {
            "report_id": f"eq.{report_id}",
            "select": ("reference_number,award_reference,type_requirement,"
                       "is_material_weakness,is_significant_deficiency,is_questioned_costs,"
                       "is_modified_opinion,is_other_findings,is_other_matters,is_repeat_finding"),
            "order": "reference_number.asc",
            "limit": str(findings_row_limit),
        }
        if req.only_flagged:
            flagged = [
                "is_material_weakness", "is_significant_deficiency", "is_questioned_costs",
                "is_modified_opinion", "is_other_findings", "is_other_matters", "is_repeat_finding",
            ]
            findings_params["or"] = "(" + ",".join(f"{f}.eq.true" for f in flagged) + ")"

        try:
            fac_findings = fac_get("findings", findings_params) or []
        except Exception:
            fac_findings = []

        # Deduplicate refs to unique finding numbers (API returns one row per finding-per-award)
        seen_refs = set()
        refs = []
        for r in fac_findings:
            ref = r.get("reference_number")
            if ref and ref not in seen_refs:
                seen_refs.add(ref)
                refs.append(ref)
        refs = refs[:effective_max_refs]

        # Always define these
        if refs:
            try:
                fac_findings_text = fac_get("findings_text", {
                    "report_id": f"eq.{report_id}",
                    "select": "finding_ref_number,finding_text",
                    "order": "finding_ref_number.asc",
                    "limit": str(len(refs)),
                    "or": or_param("finding_ref_number", refs),
                }) or []
            except Exception:
                fac_findings_text = []
            try:
                fac_caps = fac_get("corrective_action_plans", {
                    "report_id": f"eq.{report_id}",
                    "select": "finding_ref_number,planned_action",
                    "order": "finding_ref_number.asc",
                    "limit": str(len(refs)),
                    "or": or_param("finding_ref_number", refs),
                }) or []
            except Exception:
                fac_caps = []
        else:
            fac_findings_text, fac_caps = [], []

        # Awards (optional, still safe)
        # NOTE: The `assistance_listing` column only exists for newer FAC data.
        # For older data (e.g. 2022), it doesn't exist and the query fails.
        # Try with it first, then retry without it on failure.
        federal_awards = []
        if req.include_awards:
            try:
                federal_awards = fac_get("federal_awards", {
                    "report_id": f"eq.{report_id}",
                    "select": "award_reference,federal_program_name,assistance_listing",
                    "order": "award_reference.asc",
                    "limit": "200",
                }) or []
            except Exception:
                try:
                    logging.info("Retrying federal_awards without assistance_listing column")
                    federal_awards = fac_get("federal_awards", {
                        "report_id": f"eq.{report_id}",
                        "select": "award_reference,federal_program_name",
                        "order": "award_reference.asc",
                        "limit": "200",
                    }) or []
                except Exception:
                    federal_awards = []

        # APPLY ALN OVERRIDES BEFORE BUILDING MODEL
        if federal_awards and aln_by_award:
            logging.info(f"Applying ALN overrides from {len(aln_by_award)} award mappings")
            for a in federal_awards:
                ar = (a.get("award_reference") or "").strip()
                current_aln = (a.get("assistance_listing") or "").strip()

                logging.info(f"   Award {ar}: current ALN = '{current_aln}'")

                # If no ALN or it's 'Unknown', try to get from override
                if (not current_aln or current_aln == "Unknown") and ar in aln_by_award:
                    a["assistance_listing"] = aln_by_award[ar]
                    logging.info(f"   Applied override: {ar} -> {aln_by_award[ar]}")
                elif ar in aln_by_award:
                    # Even if there's an ALN, if override exists and differs, consider using it
                    override_aln = aln_by_award[ar]
                    if override_aln != current_aln:
                        logging.info(f"   ALN mismatch for {ar}: '{current_aln}' vs override '{override_aln}'")

        # HARDCODED ALN FIX
        for a in federal_awards:
            aln = (a.get("assistance_listing") or "").strip()
            if aln in TREASURY_PROGRAMS:
                a["federal_program_name"] = TREASURY_PROGRAMS[aln]
                logging.info(f"Set program name for {aln}")

        template_path = none_if_placeholder(req.template_path) or "templates/MDL_Template_Data_Mapping_Comments.docx"
        aln_xlsx = none_if_placeholder(req.aln_reference_xlsx) or "templates/Additional_Reference_Documentation_MDLs.xlsx"

        # ---------- NEW: build the model -------------
        mdl_model = build_mdl_model_from_fac(
            auditee_name=effective_auditee_name,
            ein=req.ein,
            audit_year=req.audit_year,
            fac_general=gen,
            fac_findings=fac_findings,
            fac_findings_text=fac_findings_text,
            fac_caps=fac_caps,
            federal_awards=federal_awards,
            only_flagged=req.only_flagged,
            max_refs=req.max_refs,
            include_no_qc_line=True,
            treasury_listings=req.treasury_listings or ["21.027", "21.023", "21.026"],
            aln_reference_xlsx=aln_xlsx,
            aln_overrides_by_finding=aln_by_finding,
        )

        # ---------- NEW: enrich headers from FAC + defaults ----------
        # Use INPUT year (gen) for: address, auditor, period_end
        # Use LATEST year for: auditee_name, recipient_name, poc_name, poc_title
        fac_defaults = from_fac_general(gen)  # Use input year for address, auditor, period_end

        # auditee_name, recipient_name, and POC come from the LATEST FAC year
        raw_auditee = effective_auditee_name  # From latest year
        raw_auditor = fac_defaults.get("auditor_name") or ""  # From input year
        raw_poc_name = poc_name_from_latest or req.poc_name or ""  # From latest year
        raw_poc_title = poc_title_from_latest or req.poc_title or ""  # From latest year

        # Critical logging: Verify auditee_name and POC source
        input_year_auditee = gen[0].get("auditee_name") or "(empty)"
        input_year_poc_name = gen[0].get("auditee_contact_name") or "(empty)"
        input_year_poc_title = gen[0].get("auditee_contact_title") or "(empty)"
        logging.info(f"=== AUDITEE NAME & POC SOURCE VERIFICATION ===")
        logging.info(f"  Input year ({req.audit_year}) auditee_name: {input_year_auditee}")
        logging.info(f"  Latest year ({latest_year}) auditee_name: {auditee_name_from_latest}")
        logging.info(f"  USING auditee_name (from latest year): {raw_auditee}")
        logging.info(f"  ---")
        logging.info(f"  Input year ({req.audit_year}) POC: {input_year_poc_name} ({input_year_poc_title})")
        logging.info(f"  Latest year ({latest_year}) POC: {poc_name_from_latest} ({poc_title_from_latest})")
        logging.info(f"  USING POC (from latest year): {raw_poc_name} ({raw_poc_title})")
        logging.info(f"===============================================")

        # NEW CODE - Use standard case everywhere, no "The" article:
        recipient_formatted = format_name_standard_case(raw_auditee)
        auditor_formatted = raw_auditor
        header_overrides = {
            # recipient & period end
            "recipient_name": recipient_formatted,
            "period_end_text": req.fy_end_text or fac_defaults.get("period_end_text") or mdl_model.get("period_end_text"),

            # address (title case street + city, uppercase state, keep zip as-is)
            "street_address": title_case(req.street_address or fac_defaults.get("street_address")),
            "city": title_case(req.city or fac_defaults.get("city")),
            "state": (req.state or fac_defaults.get("state") or "").upper(),
            "zip_code": req.zip_code or fac_defaults.get("zip_code") or "",

            # auditor
            "auditor_name": auditor_formatted,  # use normalized name with "the" article
            "auditee_name": recipient_formatted,
            # POC (title case name + title) - from LATEST year
            "poc_name": title_case(raw_poc_name),
            "poc_title": raw_poc_title,
        }

        # apply non-empty values only
        for k, v in header_overrides.items():
            if v:
                mdl_model[k] = v

        # ADD THIS DEBUG LOGGING:
        logging.info(f"After header overrides, mdl_model auditor_name: {mdl_model.get('auditor_name')}")
        logging.info(f"After header overrides, mdl_model auditee_name: {mdl_model.get('auditee_name')}")
        logging.info(f"After header overrides, mdl_model recipient_name: {mdl_model.get('recipient_name')}")

        # ------------- sensible defaults for things the caller omitted -------------
        # Treasury listings: if not provided, use the SLFRF + common Treasury programs for demo
        if not req.treasury_listings:
            req.treasury_listings = ["21.027", "21.023", "21.026"]

        # Template defaults if not provided
        template_path = none_if_placeholder(req.template_path) or "templates/MDL_Template_Data_Mapping_Comments.docx"
        aln_xlsx = none_if_placeholder(req.aln_reference_xlsx) or "templates/Additional_Reference_Documentation_MDLs.xlsx"

        # Destination folder defaults
        dest_folder = str_or_default(req.dest_path, f"mdl/{req.audit_year}/").lstrip("/")

        # 4) Build DOCX (unchanged except variable names)
        try:
            data = build_docx_from_template(mdl_model, template_path=template_path)

            # Post-process the generated docx bytes AFTER everything else
            data = postprocess_docx(data, mdl_model)

        except HTTPException as e:
            return {"ok": False, "message": f"Template error: {e.detail}"}
        except Exception as e:
            return {"ok": False, "message": f"Unexpected template error: {e}"}

        # 5) Upload (unchanged)
        base = f"MDL-{sanitize(effective_auditee_name)}-{sanitize(req.ein)}-{req.audit_year}.docx"
        blob_name = f"{dest_folder}{base}" if dest_folder else base
        url = upload_and_sas(AZURE_CONTAINER, blob_name, data) if AZURE_CONN_STR else save_local_and_url(blob_name, data)

        return {"ok": True, "url": url, "blob_path": f"{AZURE_CONTAINER}/{blob_name}"}

    except HTTPException as e:
        return JSONResponse(status_code=200, content={"ok": False, "message": f"{e.status_code}: {e.detail}"})
    except Exception as e:
        return JSONResponse(status_code=200, content={"ok": False, "message": f"Unhandled error: {e}"})
